# -*- coding: utf-8 -*-
"""
국회 의사중계 자막 추출기 v16.4
- 참조 코드 기반 안정적 구조
- PyQt6 모던 UI
- 추가 기능: 타임스탬프, 실시간 저장, 검색, 테마, URL 히스토리, SRT/VTT 내보내기, 통계, 키워드 하이라이트, 세션 저장
- v16.1: 생중계 xcgcd 자동 감지 기능 추가
- v16.2: 코드 품질 개선 (원래 URL 복귀, xcgcd 유효성 검증, 자동 셀렉터 감지 통합, 헤드리스 타임아웃 증가)
- v16.3: 스레드 안전성 강화, 키워드 설정 영구 저장, 종료 처리 개선
"""

import sys
import os
import time
import threading
import queue
import re
import json
import logging
from datetime import datetime, timedelta
from pathlib import Path

# HiDPI 지원 - PyQt6 임포트 전에 설정 필요
os.environ['QT_ENABLE_HIGHDPI_SCALING'] = '1'
os.environ['QT_AUTO_SCREEN_SCALE_FACTOR'] = '1'

# ============================================================
# 로깅 설정
# ============================================================

def setup_logging():
    """로깅 시스템 초기화 - 파일 및 콘솔 출력"""
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    
    log_file = log_dir / f"subtitle_{datetime.now().strftime('%Y%m%d')}.log"
    
    # 로거 설정
    logger = logging.getLogger("SubtitleExtractor")
    logger.setLevel(logging.DEBUG)
    
    # 이미 핸들러가 있으면 추가하지 않음
    if logger.handlers:
        return logger
    
    # 파일 핸들러 (DEBUG 레벨)
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_format = logging.Formatter(
        '%(asctime)s [%(levelname)s] %(funcName)s: %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(file_format)
    
    # 콘솔 핸들러 (INFO 레벨)
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_format = logging.Formatter('[%(levelname)s] %(message)s')
    console_handler.setFormatter(console_format)
    
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    return logger


# 전역 로거 초기화
logger = setup_logging()

try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QLabel, QPushButton, QLineEdit, QTextEdit, QComboBox, QCheckBox,
        QFrame, QProgressBar, QMessageBox, QFileDialog,
        QGroupBox, QGridLayout, QDialog, QDialogButtonBox, QListWidget,
        QSplitter, QMenu, QMenuBar, QInputDialog, QSystemTrayIcon, QAbstractItemView
    )
    from PyQt6.QtCore import Qt, QTimer, QSettings
    from PyQt6.QtGui import (
        QFont, QColor, QTextCursor, QTextCharFormat, QAction,
        QShortcut, QKeySequence, QIcon
    )
except ImportError:
    print("PyQt6 필요: pip install PyQt6")
    sys.exit(1)

try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.common.exceptions import WebDriverException, NoSuchElementException, StaleElementReferenceException
    from selenium.webdriver.chrome.options import Options
except ImportError:
    print("selenium 필요: pip install selenium")
    sys.exit(1)


# ============================================================
# 상수 정의
# ============================================================

class Config:
    """프로그램 설정 상수"""
    VERSION = "16.5"  # 스레드 안전성 전면 강화, 비동기 파일 저장, 헤드리스 포트 동적 할당
    APP_NAME = "국회 의사중계 자막 추출기"
    
    # 타이밍 상수 (초)
    SUBTITLE_FINALIZE_DELAY = 2.0      # 자막 확정까지 대기 시간
    FINALIZE_CHECK_INTERVAL = 500      # 자막 확정 체크 간격 (ms)
    QUEUE_PROCESS_INTERVAL = 100       # 메시지 큐 처리 간격 (ms)
    STATS_UPDATE_INTERVAL = 1000       # 통계 업데이트 간격 (ms)
    SUBTITLE_CHECK_INTERVAL = 0.2      # 자막 확인 간격 (초)
    THREAD_STOP_TIMEOUT = 3            # 스레드 종료 대기 시간 (초)
    PAGE_LOAD_WAIT = 3                 # 페이지 로딩 대기 시간 (초)
    WEBDRIVER_WAIT_TIMEOUT = 20        # WebDriver 대기 타임아웃 (초)
    SCRIPT_DELAY = 0.5                 # 스크립트 실행 후 대기 (초)
    
    # 자동 백업
    AUTO_BACKUP_INTERVAL = 300000      # 5분 (ms)
    MAX_BACKUP_COUNT = 10
    MAX_URL_HISTORY = 50               # URL 히스토리 최대 개수
    
    # 경로
    LOG_DIR = "logs"
    SESSION_DIR = "sessions"
    REALTIME_DIR = "realtime_output"
    BACKUP_DIR = "backups"
    PRESET_FILE = "committee_presets.json"
    
    # 기본 CSS 선택자
    DEFAULT_SELECTORS = [
        "#viewSubtit .incont",
        "#viewSubtit",
        ".subtitle_area"
    ]
    
    # 기본 URL
    DEFAULT_URL = "https://assembly.webcast.go.kr/main/player.asp"
    
    # 상임위원회 기본 프리셋 (v16.0 기준 동작하는 xcode 값)
    # xcode: 위원회(채널) 구분 고정 값
    # xcgcd: 해당 회의의 고유 방송 ID (매 회의마다 변경)
    DEFAULT_COMMITTEE_PRESETS = {
        "본회의": "https://assembly.webcast.go.kr/main/player.asp",
        "국회운영위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=24",
        "법제사법위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=25",
        "기획재정위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=38",
        "교육위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=58",
        "과학기술정보방송통신위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=56",
        "외교통일위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=48",
        "국방위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=37",
        "행정안전위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=45",
        "문화체육관광위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=59",
        "농림축산식품해양수산위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=53",
        "산업통상자원중소벤처기업위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=55",
        "보건복지위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=33",
        "환경노동위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=62",
        "국토교통위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=54",
        "정보위원회": "https://assembly.webcast.go.kr/main/player.asp",
        "여성가족위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=36",
        "예산결산특별위원회": "https://assembly.webcast.go.kr/main/player.asp?xcode=21",
    }
    
    # 상임위원회 xcode 값 매핑 (v16.0 기준 동작하는 값)
    COMMITTEE_XCODE_MAP = {
        "본회의": None,  # 본회의는 xcode 없이 접근
        "국회운영위원회": 24,
        "법제사법위원회": 25,
        "기획재정위원회": 38,
        "교육위원회": 58,
        "과학기술정보방송통신위원회": 56,
        "외교통일위원회": 48,
        "국방위원회": 37,
        "행정안전위원회": 45,
        "문화체육관광위원회": 59,
        "농림축산식품해양수산위원회": 53,
        "산업통상자원중소벤처기업위원회": 55,
        "보건복지위원회": 33,
        "환경노동위원회": 62,
        "국토교통위원회": 54,
        "정보위원회": None,
        "여성가족위원회": 36,
        "예산결산특별위원회": 21,
    }
    
    # 상임위원회 약칭 매핑 (사이트 내 표기 포함)
    COMMITTEE_ABBREVIATIONS = {
        # 기본 약칭
        "운영위": "운영위원회",
        "법사위": "법제사법위원회",
        "정무위": "정무위원회",
        "기재위": "기획재정위원회",
        "재경위": "기획재정위원회",  # 사이트 내 타이틀 표기
        "교육위": "교육위원회",
        "과방위": "과학기술정보방송통신위원회",
        "외통위": "외교통일위원회",
        "국방위": "국방위원회",
        "행안위": "행정안전위원회",
        "문체위": "문화체육관광위원회",
        "농해수위": "농림축산식품해양수산위원회",
        "산자위": "산업통상자원중소벤처기업위원회",
        "산자중기위": "산업통상자원중소벤처기업위원회",
        "복지위": "보건복지위원회",
        "환노위": "환경노동위원회",
        "기후노동위": "환경노동위원회",  # 사이트 내 타이틀 표기
        "국토위": "국토교통위원회",
        "여가위": "여성가족위원회",
        "성평등가족위": "여성가족위원회",  # 사이트 내 타이틀 표기
        "예결위": "예산결산특별위원회",
        "특별위": "특별위원회",
    }
    
    # 폰트 설정
    DEFAULT_FONT_SIZE = 14
    MIN_FONT_SIZE = 10
    MAX_FONT_SIZE = 24


# ============================================================
# 테마 정의
# ============================================================

DARK_THEME = """
/* ===== 다크 테마 - 프리미엄 모던 디자인 ===== */
QMainWindow, QWidget {
    background-color: #0d1117;
    color: #e6edf3;
    font-family: '맑은 고딕', 'Malgun Gothic', 'Segoe UI', sans-serif;
}

/* 라벨 */
QLabel { 
    color: #e6edf3; 
}
QLabel#headerLabel {
    font-size: 20px;
    font-weight: bold;
    color: #58a6ff;
    padding: 12px;
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 rgba(88, 166, 255, 0.1), stop:1 rgba(136, 146, 255, 0.1));
    border-radius: 10px;
}

/* 버튼 기본 스타일 */
QPushButton {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #21262d, stop:1 #161b22);
    color: #e6edf3;
    border: 1px solid #30363d;
    padding: 10px 22px;
    border-radius: 8px;
    font-weight: bold;
    font-size: 12px;
}
QPushButton:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #30363d, stop:1 #21262d);
    border-color: #58a6ff;
}
QPushButton:pressed {
    background: #161b22;
    border-color: #58a6ff;
}
QPushButton:disabled { 
    background-color: #21262d; 
    color: #484f58;
    border-color: #21262d;
}

/* 시작 버튼 - 그린 그라데이션 */
QPushButton#startBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #238636, stop:1 #1a7f37);
    border: 1px solid #2ea043;
    font-size: 14px;
    min-width: 130px;
    padding: 12px 24px;
}
QPushButton#startBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2ea043, stop:1 #238636);
    border-color: #3fb950;
}
QPushButton#startBtn:pressed {
    background: #1a7f37;
}

/* 중지 버튼 - 레드 그라데이션 */
QPushButton#stopBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #da3633, stop:1 #b62324);
    border: 1px solid #f85149;
    font-size: 14px;
    min-width: 130px;
    padding: 12px 24px;
}
QPushButton#stopBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f85149, stop:1 #da3633);
    border-color: #ff7b72;
}
QPushButton#stopBtn:pressed {
    background: #b62324;
}

/* 입력 필드 */
QLineEdit, QComboBox {
    background-color: #161b22;
    border: 1px solid #30363d;
    border-radius: 8px;
    padding: 10px 14px;
    color: #e6edf3;
    font-size: 13px;
    selection-background-color: #264f78;
}
QLineEdit:focus, QComboBox:focus {
    border: 1px solid #58a6ff;
    background-color: #0d1117;
}
QLineEdit:hover, QComboBox:hover {
    border-color: #484f58;
}
QComboBox::drop-down {
    border: none;
    padding-right: 10px;
}
QComboBox::down-arrow {
    width: 12px;
    height: 12px;
}
QComboBox QAbstractItemView {
    background-color: #161b22;
    color: #e6edf3;
    border: 1px solid #30363d;
    border-radius: 8px;
    selection-background-color: #21262d;
    selection-color: #58a6ff;
    padding: 5px;
    outline: none;
}

/* 텍스트 편집 영역 */
QTextEdit {
    background-color: #0d1117;
    border: 1px solid #30363d;
    border-radius: 12px;
    padding: 16px;
    color: #e6edf3;
    font-size: 14px;
    line-height: 1.6;
    selection-background-color: #264f78;
}
QTextEdit:focus {
    border-color: #58a6ff;
}

/* 그룹박스 - 카드 스타일 */
QGroupBox {
    background-color: rgba(22, 27, 34, 0.6);
    border: 1px solid #30363d;
    border-radius: 12px;
    margin-top: 20px;
    padding: 20px 15px 15px 15px;
    font-weight: bold;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 18px;
    padding: 0 10px;
    color: #58a6ff;
    font-size: 13px;
    background-color: #0d1117;
    border-radius: 4px;
}

/* 체크박스 */
QCheckBox {
    spacing: 10px;
    color: #e6edf3;
}
QCheckBox::indicator {
    width: 18px;
    height: 18px;
    border-radius: 4px;
    border: 1px solid #30363d;
    background: #161b22;
}
QCheckBox::indicator:checked {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #58a6ff, stop:1 #8892ff);
    border-color: #58a6ff;
}
QCheckBox::indicator:hover {
    border-color: #58a6ff;
}
QCheckBox::indicator:checked:hover {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #79b8ff, stop:1 #a0aaff);
}

/* 프로그레스바 */
QProgressBar {
    background-color: #21262d;
    border-radius: 6px;
    height: 8px;
    text-align: center;
    border: none;
}
QProgressBar::chunk { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #58a6ff, stop:0.5 #8892ff, stop:1 #a371f7);
    border-radius: 6px; 
}

/* 메뉴바 */
QMenuBar { 
    background-color: #161b22; 
    color: #e6edf3;
    padding: 6px;
    border-bottom: 1px solid #21262d;
}
QMenuBar::item { 
    padding: 8px 16px;
    border-radius: 6px;
}
QMenuBar::item:selected { 
    background-color: #21262d; 
}

/* 메뉴 */
QMenu { 
    background-color: #161b22; 
    color: #e6edf3; 
    border: 1px solid #30363d;
    border-radius: 10px;
    padding: 6px;
}
QMenu::item { 
    padding: 10px 28px;
    border-radius: 6px;
}
QMenu::item:selected { 
    background-color: #21262d;
    color: #58a6ff;
}
QMenu::separator {
    height: 1px;
    background-color: #30363d;
    margin: 6px 10px;
}

/* 스크롤바 */
QScrollBar:vertical {
    background: #0d1117;
    width: 12px;
    border-radius: 6px;
    margin: 4px 2px;
}
QScrollBar::handle:vertical {
    background: #30363d;
    border-radius: 5px;
    min-height: 40px;
}
QScrollBar::handle:vertical:hover {
    background: #484f58;
}
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
    height: 0px;
}
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
    background: none;
}

QScrollBar:horizontal {
    background: #0d1117;
    height: 12px;
    border-radius: 6px;
    margin: 2px 4px;
}
QScrollBar::handle:horizontal {
    background: #30363d;
    border-radius: 5px;
    min-width: 40px;
}
QScrollBar::handle:horizontal:hover {
    background: #484f58;
}

/* 스플리터 */
QSplitter::handle {
    background: #30363d;
    border-radius: 2px;
}
QSplitter::handle:hover {
    background: #58a6ff;
}

/* 프레임 */
QFrame {
    border: none;
}
"""

LIGHT_THEME = """
/* ===== 라이트 테마 - 클린 모던 디자인 ===== */
QMainWindow, QWidget {
    background-color: #ffffff;
    color: #24292f;
    font-family: '맑은 고딕', 'Malgun Gothic', 'Segoe UI', sans-serif;
}

/* 라벨 */
QLabel { 
    color: #24292f; 
}
QLabel#headerLabel {
    font-size: 20px;
    font-weight: bold;
    color: #0969da;
    padding: 12px;
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 rgba(9, 105, 218, 0.08), stop:1 rgba(130, 80, 223, 0.08));
    border-radius: 10px;
}

/* 버튼 기본 스타일 */
QPushButton {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f6f8fa, stop:1 #eaeef2);
    color: #24292f;
    border: 1px solid #d0d7de;
    padding: 10px 22px;
    border-radius: 8px;
    font-weight: bold;
    font-size: 12px;
}
QPushButton:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f3f4f6, stop:1 #ebecf0);
    border-color: #afb8c1;
}
QPushButton:pressed {
    background: #eaeef2;
    border-color: #afb8c1;
}
QPushButton:disabled { 
    background-color: #f6f8fa; 
    color: #8c959f;
    border-color: #d0d7de;
}

/* 시작 버튼 - 그린 그라데이션 */
QPushButton#startBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2da44e, stop:1 #1a7f37);
    color: white;
    border: 1px solid #1a7f37;
    font-size: 14px;
    min-width: 130px;
    padding: 12px 24px;
}
QPushButton#startBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3fb950, stop:1 #2da44e);
    border-color: #2da44e;
}
QPushButton#startBtn:pressed {
    background: #1a7f37;
}

/* 중지 버튼 - 레드 그라데이션 */
QPushButton#stopBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #cf222e, stop:1 #a40e26);
    color: white;
    border: 1px solid #a40e26;
    font-size: 14px;
    min-width: 130px;
    padding: 12px 24px;
}
QPushButton#stopBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #fa4549, stop:1 #cf222e);
    border-color: #cf222e;
}
QPushButton#stopBtn:pressed {
    background: #a40e26;
}

/* 입력 필드 */
QLineEdit, QComboBox {
    background-color: #ffffff;
    border: 1px solid #d0d7de;
    border-radius: 8px;
    padding: 10px 14px;
    color: #24292f;
    font-size: 13px;
    selection-background-color: #ddf4ff;
}
QLineEdit:focus, QComboBox:focus {
    border: 1px solid #0969da;
    background-color: #ffffff;
}
QLineEdit:hover, QComboBox:hover {
    border-color: #afb8c1;
}
QComboBox::drop-down {
    border: none;
    padding-right: 10px;
}
QComboBox::down-arrow {
    width: 12px;
    height: 12px;
}
QComboBox QAbstractItemView {
    background-color: #ffffff;
    color: #24292f;
    border: 1px solid #d0d7de;
    border-radius: 8px;
    selection-background-color: #ddf4ff;
    selection-color: #0969da;
    padding: 5px;
    outline: none;
}

/* 텍스트 편집 영역 */
QTextEdit {
    background-color: #ffffff;
    border: 1px solid #d0d7de;
    border-radius: 12px;
    padding: 16px;
    color: #24292f;
    font-size: 14px;
    line-height: 1.6;
    selection-background-color: #ddf4ff;
}
QTextEdit:focus {
    border-color: #0969da;
}

/* 그룹박스 - 카드 스타일 */
QGroupBox {
    background-color: #f6f8fa;
    border: 1px solid #d0d7de;
    border-radius: 12px;
    margin-top: 20px;
    padding: 20px 15px 15px 15px;
    font-weight: bold;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 18px;
    padding: 0 10px;
    color: #0969da;
    font-size: 13px;
    background-color: #ffffff;
    border-radius: 4px;
}

/* 체크박스 */
QCheckBox {
    spacing: 10px;
    color: #24292f;
}
QCheckBox::indicator {
    width: 18px;
    height: 18px;
    border-radius: 4px;
    border: 1px solid #d0d7de;
    background: #ffffff;
}
QCheckBox::indicator:checked {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #0969da, stop:1 #8250df);
    border-color: #0969da;
}
QCheckBox::indicator:hover {
    border-color: #0969da;
}
QCheckBox::indicator:checked:hover {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #218bff, stop:1 #a475f9);
}

/* 프로그레스바 */
QProgressBar {
    background-color: #eaeef2;
    border-radius: 6px;
    height: 8px;
    text-align: center;
    border: none;
}
QProgressBar::chunk { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #0969da, stop:0.5 #8250df, stop:1 #bf3989);
    border-radius: 6px; 
}

/* 메뉴바 */
QMenuBar { 
    background-color: #f6f8fa; 
    color: #24292f;
    padding: 6px;
    border-bottom: 1px solid #d0d7de;
}
QMenuBar::item { 
    padding: 8px 16px;
    border-radius: 6px;
}
QMenuBar::item:selected { 
    background-color: #eaeef2; 
}

/* 메뉴 */
QMenu { 
    background-color: #ffffff; 
    color: #24292f; 
    border: 1px solid #d0d7de;
    border-radius: 10px;
    padding: 6px;
}
QMenu::item { 
    padding: 10px 28px;
    border-radius: 6px;
}
QMenu::item:selected { 
    background-color: #ddf4ff;
    color: #0969da;
}
QMenu::separator {
    height: 1px;
    background-color: #d0d7de;
    margin: 6px 10px;
}

/* 스크롤바 */
QScrollBar:vertical {
    background: #f6f8fa;
    width: 12px;
    border-radius: 6px;
    margin: 4px 2px;
}
QScrollBar::handle:vertical {
    background: #afb8c1;
    border-radius: 5px;
    min-height: 40px;
}
QScrollBar::handle:vertical:hover {
    background: #8c959f;
}
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
    height: 0px;
}
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
    background: none;
}

QScrollBar:horizontal {
    background: #f6f8fa;
    height: 12px;
    border-radius: 6px;
    margin: 2px 4px;
}
QScrollBar::handle:horizontal {
    background: #afb8c1;
    border-radius: 5px;
    min-width: 40px;
}
QScrollBar::handle:horizontal:hover {
    background: #8c959f;
}

/* 스플리터 */
QSplitter::handle {
    background: #d0d7de;
    border-radius: 2px;
}
QSplitter::handle:hover {
    background: #0969da;
}

/* 프레임 */
QFrame {
    border: none;
}
"""


# ============================================================
# 토스트 알림 위젯
# ============================================================

class ToastWidget(QFrame):
    """비차단 토스트 알림 위젯 - 스택 처리 지원, 모던 디자인"""
    
    def __init__(self, parent, message: str, duration: int = 3000, 
                 toast_type: str = "info", y_offset: int = 10, on_close=None):
        super().__init__(parent)
        self.setObjectName("toastWidget")
        self.on_close = on_close
        
        # 개선된 색상 팔레트 (더 세련된 컬러)
        colors = {
            "info": ("#58a6ff", "#161b22", "#21262d"),      # 블루 - 정보
            "success": ("#3fb950", "#0d1117", "#1c2128"),   # 그린 - 성공
            "warning": ("#d29922", "#1c1a14", "#2d2a21"),   # 옐로우 - 경고
            "error": ("#f85149", "#1c1418", "#2d2128")      # 레드 - 오류
        }
        accent_color, bg_color, border_bg = colors.get(toast_type, colors["info"])
        
        self.setStyleSheet(f"""
            QFrame#toastWidget {{
                background-color: {bg_color};
                border: 1px solid {accent_color};
                border-left: 4px solid {accent_color};
                border-radius: 10px;
                padding: 12px;
            }}
            QLabel {{
                color: #e6edf3;
                font-size: 13px;
                font-weight: 500;
                background: transparent;
            }}
            QLabel#toastIcon {{
                font-size: 16px;
                min-width: 24px;
            }}
            QLabel#toastMessage {{
                color: #e6edf3;
                padding-left: 8px;
            }}
        """)
        
        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(10)
        
        # 아이콘 (개선된 이모지)
        icons = {"info": "ℹ️", "success": "✅", "warning": "⚠️", "error": "❌"}
        icon_label = QLabel(icons.get(toast_type, "ℹ️"))
        icon_label.setObjectName("toastIcon")
        layout.addWidget(icon_label)
        
        # 메시지
        msg_label = QLabel(message)
        msg_label.setObjectName("toastMessage")
        msg_label.setWordWrap(True)
        layout.addWidget(msg_label, 1)
        
        # 위치 설정 (부모 위젯 상단 중앙 + y_offset)
        self.adjustSize()
        if parent:
            parent_rect = parent.rect()
            x = (parent_rect.width() - self.width()) // 2
            self.move(x, y_offset)
        
        self.show()
        self.raise_()
        
        # 자동 사라짐
        QTimer.singleShot(duration, self._fade_out)
    
    def _fade_out(self):
        """토스트 사라지기"""
        if self.on_close:
            self.on_close(self)
        self.hide()
        self.deleteLater()


# ============================================================
# 자막 데이터 클래스
# ============================================================

class SubtitleEntry:
    """자막 항목 - 타입 힌트 포함"""
    
    def __init__(self, text: str, timestamp=None):
        self.text: str = text
        self.timestamp: datetime = timestamp or datetime.now()
        self.start_time: datetime = None  # SRT용
        self.end_time: datetime = None    # SRT용
    
    def to_dict(self) -> dict:
        """딕셔너리로 변환 (세션 저장용)"""
        return {
            'text': self.text,
            'timestamp': self.timestamp.isoformat(),
            'start_time': self.start_time.isoformat() if self.start_time else None,
            'end_time': self.end_time.isoformat() if self.end_time else None,
        }
    
    @classmethod
    def from_dict(cls, data: dict) -> 'SubtitleEntry':
        """딕셔너리에서 생성 - 방어 코드 포함"""
        # 필수 필드 검증
        text = data.get('text', '')
        timestamp_str = data.get('timestamp')
        
        if not text:
            raise ValueError("자막 텍스트가 비어있습니다")
        if not timestamp_str:
            raise ValueError("타임스탬프가 없습니다")
        
        entry = cls(text)
        entry.timestamp = datetime.fromisoformat(timestamp_str)
        if data.get('start_time'):
            entry.start_time = datetime.fromisoformat(data['start_time'])
        if data.get('end_time'):
            entry.end_time = datetime.fromisoformat(data['end_time'])
        return entry


# ============================================================
# 메인 윈도우
# ============================================================

class MainWindow(QMainWindow):
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{Config.APP_NAME} v{Config.VERSION}")
        self.setMinimumSize(1100, 750)
        self.resize(1200, 800)
        
        # 설정
        self.settings = QSettings("AssemblySubtitle", "Extractor")
        self.is_dark_theme = self.settings.value("dark_theme", True, type=bool)
        self.font_size = self.settings.value("font_size", Config.DEFAULT_FONT_SIZE, type=int)
        self.minimize_to_tray = self.settings.value("minimize_to_tray", False, type=bool)
        
        # 메시지 큐
        self.message_queue = queue.Queue()
        
        # 상태
        self.worker = None
        self.driver = None
        self.is_running = False
        self.stop_event = threading.Event()  # 스레드 안전한 종료 시그널
        self.subtitle_lock = threading.Lock()  # 자막 리스트 접근 동기화
        self.start_time = None
        self.last_subtitle = ""
        
        # 자막 데이터 (타임스탬프 포함)
        self.subtitles = []  # List[SubtitleEntry]
        # 저장된 키워드 로드
        saved_keywords = self.settings.value("highlight_keywords", "")
        self.keywords = [k.strip() for k in saved_keywords.split(",") if k.strip()] if saved_keywords else []
        saved_alert = self.settings.value("alert_keywords", "")
        self.alert_keywords = [k.strip() for k in saved_alert.split(",") if k.strip()] if saved_alert else []
        self.last_update_time = 0  # 마지막 자막 업데이트 시간
        
        # 토스트 스택 관리
        self.active_toasts = []  # 현재 표시 중인 토스트 목록
        
        # 실시간 저장
        self.realtime_file = None
        
        # URL 히스토리
        self.url_history = self._load_url_history()
        
        # UI 생성
        self._create_menu()
        self._create_ui()
        self._apply_theme()
        self._setup_shortcuts()
        
        # 타이머
        self.queue_timer = QTimer(self)
        self.queue_timer.timeout.connect(self._process_message_queue)
        self.queue_timer.start(Config.QUEUE_PROCESS_INTERVAL)
        
        self.stats_timer = QTimer(self)
        self.stats_timer.timeout.connect(self._update_stats)
        
        # 자막 확정 타이머
        self.finalize_timer = QTimer(self)
        self.finalize_timer.timeout.connect(self._check_finalize)
        
        # 자동 백업 타이머
        self.backup_timer = QTimer(self)
        self.backup_timer.timeout.connect(self._auto_backup)
        
        # 디렉토리 생성
        Path(Config.SESSION_DIR).mkdir(exist_ok=True)
        Path(Config.REALTIME_DIR).mkdir(exist_ok=True)
        Path(Config.BACKUP_DIR).mkdir(exist_ok=True)
        
        # 창 위치/크기 복원
        geometry = self.settings.value("geometry")
        if geometry:
            self.restoreGeometry(geometry)
        state = self.settings.value("windowState")
        if state:
            self.restoreState(state)
        
        # 시스템 트레이 설정
        self._setup_tray()
    
    def _setup_tray(self):
        """시스템 트레이 아이콘 설정"""
        self.tray_icon = QSystemTrayIcon(self)
        self.tray_icon.setToolTip(f"{Config.APP_NAME} v{Config.VERSION}")
        
        # 기본 아이콘 (앱 아이콘이 없으면 기본 아이콘 사용)
        app_icon = self.windowIcon()
        if not app_icon.isNull():
            self.tray_icon.setIcon(app_icon)
        else:
            # 기본 아이콘 사용
            self.tray_icon.setIcon(self.style().standardIcon(self.style().StandardPixmap.SP_ComputerIcon))
        
        # 트레이 메뉴
        tray_menu = QMenu()
        
        show_action = QAction("🏛️ 창 보이기", self)
        show_action.triggered.connect(self._show_from_tray)
        tray_menu.addAction(show_action)
        
        tray_menu.addSeparator()
        
        # 추출 상태 표시
        self.tray_status_action = QAction("⚪ 대기 중", self)
        self.tray_status_action.setEnabled(False)
        tray_menu.addAction(self.tray_status_action)
        
        tray_menu.addSeparator()
        
        start_action = QAction("▶ 시작", self)
        start_action.triggered.connect(self._start)
        tray_menu.addAction(start_action)
        
        stop_action = QAction("⏹ 중지", self)
        stop_action.triggered.connect(self._stop)
        tray_menu.addAction(stop_action)
        
        tray_menu.addSeparator()
        
        quit_action = QAction("❌ 종료", self)
        quit_action.triggered.connect(self._quit_from_tray)
        tray_menu.addAction(quit_action)
        
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.activated.connect(self._tray_activated)
        self.tray_icon.show()
    
    def _show_from_tray(self):
        """트레이에서 창 복원"""
        self.showNormal()
        self.activateWindow()
        self.raise_()
    
    def _quit_from_tray(self):
        """트레이에서 완전 종료"""
        self.minimize_to_tray = False  # 트레이 최소화 비활성화
        self.close()
    
    def _tray_activated(self, reason):
        """트레이 아이콘 클릭 처리"""
        if reason == QSystemTrayIcon.ActivationReason.DoubleClick:
            self._show_from_tray()
    
    def _update_tray_status(self, status: str):
        """트레이 상태 업데이트"""
        if hasattr(self, 'tray_status_action'):
            self.tray_status_action.setText(status)
    
    def _create_menu(self):
        menubar = self.menuBar()
        
        # 파일 메뉴
        file_menu = menubar.addMenu("파일")
        
        save_txt = QAction("TXT 저장", self)
        save_txt.setShortcut("Ctrl+S")
        save_txt.triggered.connect(self._save_txt)
        file_menu.addAction(save_txt)
        
        save_srt = QAction("SRT 저장", self)
        save_srt.triggered.connect(self._save_srt)
        file_menu.addAction(save_srt)
        
        save_vtt = QAction("VTT 저장", self)
        save_vtt.triggered.connect(self._save_vtt)
        file_menu.addAction(save_vtt)
        
        save_docx = QAction("DOCX 저장 (Word)", self)
        save_docx.triggered.connect(self._save_docx)
        file_menu.addAction(save_docx)
        
        save_hwp = QAction("HWP 저장 (한글)", self)
        save_hwp.triggered.connect(self._save_hwp)
        file_menu.addAction(save_hwp)
        
        save_rtf = QAction("RTF 저장", self)
        save_rtf.triggered.connect(self._save_rtf)
        file_menu.addAction(save_rtf)
        
        file_menu.addSeparator()
        
        save_session = QAction("세션 저장", self)
        save_session.setShortcut("Ctrl+Shift+S")
        save_session.triggered.connect(self._save_session)
        file_menu.addAction(save_session)
        
        load_session = QAction("세션 불러오기", self)
        load_session.setShortcut("Ctrl+O")
        load_session.triggered.connect(self._load_session)
        file_menu.addAction(load_session)
        
        file_menu.addSeparator()
        
        export_stats = QAction("📊 통계 내보내기", self)
        export_stats.triggered.connect(self._export_stats)
        file_menu.addAction(export_stats)
        
        file_menu.addSeparator()
        
        exit_action = QAction("종료", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # 편집 메뉴
        edit_menu = menubar.addMenu("편집")
        
        search_action = QAction("검색", self)
        search_action.setShortcut("Ctrl+F")
        search_action.triggered.connect(self._show_search)
        edit_menu.addAction(search_action)
        
        keyword_action = QAction("하이라이트 키워드 설정", self)
        keyword_action.triggered.connect(self._set_keywords)
        edit_menu.addAction(keyword_action)
        
        alert_keyword_action = QAction("알림 키워드 설정", self)
        alert_keyword_action.setToolTip("특정 키워드 감지 시 알림을 표시합니다")
        alert_keyword_action.triggered.connect(self._set_alert_keywords)
        edit_menu.addAction(alert_keyword_action)
        
        edit_menu.addSeparator()
        
        edit_subtitle_action = QAction("✏️ 자막 편집", self)
        edit_subtitle_action.setShortcut("Ctrl+E")
        edit_subtitle_action.setToolTip("선택한 자막을 편집합니다")
        edit_subtitle_action.triggered.connect(self._edit_subtitle)
        edit_menu.addAction(edit_subtitle_action)
        
        delete_subtitle_action = QAction("🗑️ 자막 삭제", self)
        delete_subtitle_action.setShortcut("Delete")
        delete_subtitle_action.setToolTip("선택한 자막을 삭제합니다")
        delete_subtitle_action.triggered.connect(self._delete_subtitle)
        edit_menu.addAction(delete_subtitle_action)
        
        edit_menu.addSeparator()
        
        copy_action = QAction("클립보드 복사", self)
        copy_action.setShortcut("Ctrl+C")
        copy_action.setToolTip("전체 자막을 클립보드에 복사합니다")
        copy_action.triggered.connect(self._copy_to_clipboard)
        edit_menu.addAction(copy_action)
        
        clear_action = QAction("내용 지우기", self)
        clear_action.setToolTip("모든 자막 내용을 삭제합니다")
        clear_action.triggered.connect(self._clear_text)
        edit_menu.addAction(clear_action)
        
        # 보기 메뉴
        view_menu = menubar.addMenu("보기")
        
        self.theme_action = QAction("라이트 테마" if self.is_dark_theme else "다크 테마", self)
        self.theme_action.setShortcut("Ctrl+T")
        self.theme_action.triggered.connect(self._toggle_theme)
        view_menu.addAction(self.theme_action)
        
        self.timestamp_action = QAction("타임스탬프 표시", self)
        self.timestamp_action.setCheckable(True)
        self.timestamp_action.setChecked(True)
        self.timestamp_action.triggered.connect(self._refresh_text)
        view_menu.addAction(self.timestamp_action)
        
        self.tray_action = QAction("🔽 트레이로 최소화", self)
        self.tray_action.setCheckable(True)
        self.tray_action.setChecked(self.minimize_to_tray)
        self.tray_action.setToolTip("활성화 시 창을 닫으면 트레이로 최소화됩니다")
        self.tray_action.triggered.connect(self._toggle_tray_option)
        view_menu.addAction(self.tray_action)
        
        view_menu.addSeparator()
        
        # 글자 크기 서브메뉴
        font_menu = view_menu.addMenu("📝 글자 크기")
        for size in [12, 14, 16, 18, 20, 22, 24]:
            font_action = QAction(f"{size}pt", self)
            font_action.triggered.connect(lambda checked, s=size: self._set_font_size(s))
            font_menu.addAction(font_action)
        
        view_menu.addSeparator()
        
        font_increase = QAction("글자 크기 키우기", self)
        font_increase.setShortcut("Ctrl++")
        font_increase.triggered.connect(lambda: self._adjust_font_size(2))
        view_menu.addAction(font_increase)
        
        font_decrease = QAction("글자 크기 줄이기", self)
        font_decrease.setShortcut("Ctrl+-")
        font_decrease.triggered.connect(lambda: self._adjust_font_size(-2))
        view_menu.addAction(font_decrease)
        
        # 도움말 메뉴
        help_menu = menubar.addMenu("도움말")
        
        guide_action = QAction("사용법 가이드", self)
        guide_action.setShortcut("F1")
        guide_action.setToolTip("프로그램 사용법 안내")
        guide_action.triggered.connect(self._show_guide)
        help_menu.addAction(guide_action)
        
        features_action = QAction("기능 소개", self)
        features_action.setToolTip("프로그램의 주요 기능 소개")
        features_action.triggered.connect(self._show_features)
        help_menu.addAction(features_action)
        
        shortcuts_action = QAction("⌨️ 키보드 단축키", self)
        shortcuts_action.setToolTip("사용 가능한 키보드 단축키 목록")
        shortcuts_action.triggered.connect(self._show_shortcuts)
        help_menu.addAction(shortcuts_action)
        
        help_menu.addSeparator()
        
        about_action = QAction("정보", self)
        about_action.setToolTip("프로그램 정보 및 버전")
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)
    
    def _create_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        
        layout = QVBoxLayout(central)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # === 헤더 ===
        header = QLabel("🏛️ 국회 의사중계 자막 추출기")
        header.setObjectName("headerLabel")
        header.setFont(QFont("맑은 고딕", 18, QFont.Weight.Bold))
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(header)
        
        # === URL/설정 영역 ===
        settings_group = QGroupBox("⚙️ 설정")
        settings_layout = QGridLayout(settings_group)
        settings_layout.setSpacing(10)
        
        # URL
        url_label = QLabel("📌 URL:")
        url_label.setToolTip("국회 의사중계 웹사이트 URL을 입력하세요")
        settings_layout.addWidget(url_label, 0, 0)
        
        url_layout = QHBoxLayout()
        self.url_combo = QComboBox()
        self.url_combo.setEditable(True)
        self.url_combo.setToolTip("국회 의사중계 웹사이트 URL\n최근 사용한 URL이 자동 저장됩니다\n태그가 있으면 [태그] 형태로 표시됩니다")
        
        # URL 히스토리 로드 및 콤보박스 초기화
        for url, tag in self.url_history.items():
            if tag:
                self.url_combo.addItem(f"[{tag}] {url}", url)
            else:
                self.url_combo.addItem(url, url)
        
        if not self.url_history:
            self.url_combo.addItem("https://assembly.webcast.go.kr/main/player.asp")
        
        url_layout.addWidget(self.url_combo, 1)
        
        # 태그 버튼
        self.tag_btn = QPushButton("🏷️ 태그")
        self.tag_btn.setToolTip("현재 URL에 태그 추가/편집\n예: 본회의, 법사위, 상임위")
        self.tag_btn.setFixedWidth(85)
        self.tag_btn.clicked.connect(self._edit_url_tag)
        url_layout.addWidget(self.tag_btn)
        
        # 상임위원회 프리셋 버튼
        self.preset_btn = QPushButton("📋 상임위")
        self.preset_btn.setToolTip("상임위원회 프리셋 선택\n빠른 URL 입력을 위한 기능")
        self.preset_btn.setFixedWidth(110)
        
        # 프리셋 메뉴 생성
        self.preset_menu = QMenu(self)
        self._load_committee_presets()
        self._build_preset_menu()
        self.preset_btn.setMenu(self.preset_menu)
        url_layout.addWidget(self.preset_btn)
        
        settings_layout.addLayout(url_layout, 0, 1)
        
        # 선택자
        selector_label = QLabel("🔍 선택자:")
        selector_label.setToolTip("자막 요소의 CSS 선택자")
        settings_layout.addWidget(selector_label, 1, 0)
        self.selector_combo = QComboBox()
        self.selector_combo.setEditable(True)
        self.selector_combo.addItems(["#viewSubtit .incont", "#viewSubtit", ".subtitle_area"])
        self.selector_combo.setToolTip("자막 텍스트가 표시되는 HTML 요소의 CSS 선택자\n기본값을 사용하거나 직접 입력하세요")
        settings_layout.addWidget(self.selector_combo, 1, 1)
        
        # 옵션
        options_layout = QHBoxLayout()
        options_layout.setSpacing(20)
        
        self.auto_scroll_check = QCheckBox("📜 자동 스크롤")
        self.auto_scroll_check.setChecked(True)
        self.auto_scroll_check.setToolTip("새 자막이 추가될 때 자동으로 맨 아래로 스크롤합니다")
        
        self.realtime_save_check = QCheckBox("💾 실시간 저장")
        self.realtime_save_check.setChecked(False)
        self.realtime_save_check.setToolTip("자막이 확정될 때마다 자동으로 파일에 저장합니다\n저장 위치: realtime_output 폴더")
        
        self.headless_check = QCheckBox("🔇 헤드리스 모드 (인터넷창 숨김)")
        self.headless_check.setChecked(False)
        self.headless_check.setToolTip("Chrome 브라우저 창을 숨기고 백그라운드에서 실행합니다.\n자막 추출 중 다른 작업을 할 수 있습니다.")
        options_layout.addWidget(self.auto_scroll_check)
        options_layout.addWidget(self.realtime_save_check)
        options_layout.addWidget(self.headless_check)
        options_layout.addStretch()
        settings_layout.addLayout(options_layout, 2, 0, 1, 2)
        
        layout.addWidget(settings_group)
        
        # === 컨트롤 버튼 ===
        btn_layout = QHBoxLayout()
        
        self.start_btn = QPushButton("▶  시작")
        self.start_btn.setObjectName("startBtn")
        self.start_btn.clicked.connect(self._start)
        
        self.stop_btn = QPushButton("⏹  중지")
        self.stop_btn.setObjectName("stopBtn")
        self.stop_btn.setEnabled(False)
        self.stop_btn.clicked.connect(self._stop)
        
        btn_layout.addWidget(self.start_btn)
        btn_layout.addWidget(self.stop_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        # === 진행 표시 ===
        self.progress = QProgressBar()
        self.progress.setMaximum(0)
        self.progress.hide()
        layout.addWidget(self.progress)
        
        # === 메인 콘텐츠 (자막 + 통계) ===
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # 자막 텍스트
        self.subtitle_text = QTextEdit()
        self.subtitle_text.setReadOnly(True)
        splitter.addWidget(self.subtitle_text)
        
        # 통계 패널 - 모던 디자인
        stats_group = QGroupBox("📊 통계")
        stats_layout = QVBoxLayout(stats_group)
        stats_layout.setSpacing(12)
        stats_layout.setContentsMargins(12, 20, 12, 12)
        
        self.stat_time = QLabel("⏱️ 실행 시간: 00:00:00")
        self.stat_chars = QLabel("📝 글자 수: 0")
        self.stat_words = QLabel("📖 단어 수: 0")
        self.stat_sents = QLabel("💬 문장 수: 0")
        self.stat_cpm = QLabel("⚡ 분당 글자: 0")
        
        stat_labels = [self.stat_time, self.stat_chars, self.stat_words, self.stat_sents, self.stat_cpm]
        for label in stat_labels:
            label.setFont(QFont("맑은 고딕", 10))
            label.setStyleSheet("""
                padding: 6px 8px;
                border-radius: 6px;
                background-color: rgba(88, 166, 255, 0.08);
            """)
            stats_layout.addWidget(label)
        
        stats_layout.addStretch()
        stats_group.setFixedWidth(220)
        splitter.addWidget(stats_group)
        
        splitter.setSizes([860, 220])
        layout.addWidget(splitter)
        
        # === 검색바 (숨김) - 개선된 디자인 ===
        self.search_frame = QFrame()
        self.search_frame.setStyleSheet("""
            QFrame {
                background-color: rgba(88, 166, 255, 0.05);
                border: 1px solid rgba(88, 166, 255, 0.2);
                border-radius: 10px;
                padding: 5px;
            }
        """)
        search_layout = QHBoxLayout(self.search_frame)
        search_layout.setContentsMargins(12, 8, 12, 8)
        search_layout.setSpacing(8)
        
        # 검색 아이콘 라벨
        search_icon = QLabel("🔍")
        search_icon.setStyleSheet("background: transparent; border: none;")
        search_layout.addWidget(search_icon)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("검색어 입력...")
        self.search_input.setStyleSheet("""
            QLineEdit {
                background-color: transparent;
                border: none;
                padding: 6px;
                font-size: 13px;
            }
        """)
        self.search_input.returnPressed.connect(self._do_search)
        search_layout.addWidget(self.search_input, 1)
        
        self.search_count = QLabel("")
        self.search_count.setStyleSheet("""
            color: #8b949e;
            font-size: 12px;
            background: transparent;
            border: none;
            padding: 0 8px;
        """)
        search_layout.addWidget(self.search_count)
        
        # 검색 버튼 스타일
        search_btn_style = """
            QPushButton {
                background-color: rgba(88, 166, 255, 0.1);
                border: 1px solid rgba(88, 166, 255, 0.3);
                border-radius: 6px;
                padding: 6px;
                font-size: 12px;
                min-width: 32px;
            }
            QPushButton:hover {
                background-color: rgba(88, 166, 255, 0.2);
                border-color: #58a6ff;
            }
        """
        
        search_prev = QPushButton("◀")
        search_prev.setStyleSheet(search_btn_style)
        search_prev.setFixedWidth(36)
        search_prev.clicked.connect(lambda: self._nav_search(-1))
        
        search_next = QPushButton("▶")
        search_next.setStyleSheet(search_btn_style)
        search_next.setFixedWidth(36)
        search_next.clicked.connect(lambda: self._nav_search(1))
        
        search_close = QPushButton("✕")
        search_close.setStyleSheet("""
            QPushButton {
                background-color: rgba(248, 81, 73, 0.1);
                border: 1px solid rgba(248, 81, 73, 0.3);
                border-radius: 6px;
                padding: 6px;
                font-size: 12px;
                min-width: 32px;
            }
            QPushButton:hover {
                background-color: rgba(248, 81, 73, 0.2);
                border-color: #f85149;
            }
        """)
        search_close.setFixedWidth(36)
        search_close.clicked.connect(self._hide_search)
        
        search_layout.addWidget(search_prev)
        search_layout.addWidget(search_next)
        search_layout.addWidget(search_close)
        
        # 검색바 초기에 숨김 상태
        self.search_frame.hide()
        layout.addWidget(self.search_frame)
        
        # === 상태바 - 모던 디자인 ===
        status_frame = QFrame()
        status_frame.setStyleSheet("""
            QFrame {
                background-color: rgba(48, 54, 61, 0.3);
                border-radius: 8px;
                padding: 4px;
            }
        """)
        status_layout = QHBoxLayout(status_frame)
        status_layout.setContentsMargins(12, 8, 12, 8)
        
        self.status_label = QLabel("⚪ 대기 중")
        self.status_label.setStyleSheet("background: transparent; border: none; font-weight: 500;")
        
        self.count_label = QLabel("📝 0문장 | 0자")
        self.count_label.setStyleSheet("color: #8b949e; background: transparent; border: none;")
        
        status_layout.addWidget(self.status_label)
        status_layout.addStretch()
        status_layout.addWidget(self.count_label)
        
        layout.addWidget(status_frame)
        
        # 검색 상태
        self.search_matches = []
        self.search_idx = 0
        
        # 저장된 폰트 크기 적용
        self._set_font_size(self.font_size)
    
    def _apply_theme(self):
        self.setStyleSheet(DARK_THEME if self.is_dark_theme else LIGHT_THEME)
        self.theme_action.setText("라이트 테마" if self.is_dark_theme else "다크 테마")
        
        # 테마에 따른 동적 스타일 업데이트
        if self.is_dark_theme:
            stat_bg = "rgba(88, 166, 255, 0.08)"
            search_bg = "rgba(88, 166, 255, 0.05)"
            search_border = "rgba(88, 166, 255, 0.2)"
            status_bg = "rgba(48, 54, 61, 0.3)"
            count_color = "#8b949e"
        else:
            stat_bg = "rgba(9, 105, 218, 0.06)"
            search_bg = "rgba(9, 105, 218, 0.04)"
            search_border = "rgba(9, 105, 218, 0.15)"
            status_bg = "rgba(208, 215, 222, 0.4)"
            count_color = "#57606a"
        
        # 통계 라벨 스타일 업데이트
        try:
            stat_labels = [self.stat_time, self.stat_chars, self.stat_words, self.stat_sents, self.stat_cpm]
            for label in stat_labels:
                label.setStyleSheet(f"""
                    padding: 6px 8px;
                    border-radius: 6px;
                    background-color: {stat_bg};
                """)
            
            # 검색바 스타일 업데이트
            self.search_frame.setStyleSheet(f"""
                QFrame {{
                    background-color: {search_bg};
                    border: 1px solid {search_border};
                    border-radius: 10px;
                    padding: 5px;
                }}
            """)
            
            # 상태바 카운트 라벨 색상 업데이트
            self.count_label.setStyleSheet(f"color: {count_color}; background: transparent; border: none;")
        except AttributeError:
            # UI가 아직 완전히 초기화되지 않은 경우
            pass
    
    def _toggle_theme(self):
        self.is_dark_theme = not self.is_dark_theme
        self.settings.setValue("dark_theme", self.is_dark_theme)
        self._apply_theme()
    
    def _toggle_tray_option(self):
        """트레이 최소화 옵션 토글"""
        self.minimize_to_tray = self.tray_action.isChecked()
        self.settings.setValue("minimize_to_tray", self.minimize_to_tray)
        if self.minimize_to_tray:
            self._show_toast("창을 닫으면 트레이로 최소화됩니다.", "info")
        else:
            self._show_toast("창을 닫으면 프로그램이 종료됩니다.", "info")
    
    def _setup_shortcuts(self):
        QShortcut(QKeySequence("F5"), self, self._start)
        QShortcut(QKeySequence("Escape"), self, self._stop)
        QShortcut(QKeySequence("F3"), self, lambda: self._nav_search(1))
        QShortcut(QKeySequence("Shift+F3"), self, lambda: self._nav_search(-1))
    
    # ========== 토스트 및 상태 표시 ==========
    
    def _show_toast(self, message: str, toast_type: str = "info", duration: int = 3000) -> None:
        """토스트 알림 표시 - 스택 처리로 겹침 방지"""
        # 만료된 토스트 정리
        self.active_toasts = [t for t in self.active_toasts if t.isVisible()]
        
        # 새 토스트 y 위치 계산 (기존 토스트 아래에 배치)
        y_offset = 10
        for toast in self.active_toasts:
            y_offset += toast.height() + 5
        
        # 토스트 제거 콜백
        def remove_toast(t):
            if t in self.active_toasts:
                self.active_toasts.remove(t)
        
        # 토스트 생성
        toast = ToastWidget(
            self.centralWidget(), message, duration, toast_type,
            y_offset=y_offset, on_close=remove_toast
        )
        self.active_toasts.append(toast)
    
    def _set_status(self, text: str, status_type: str = "info"):
        """상태 표시 (아이콘 + 색상)"""
        icons = {
            "info": "ℹ️",
            "success": "✅",
            "warning": "⚠️",
            "error": "❌",
            "running": "🔄"
        }
        colors = {
            "info": "#4fc3f7",
            "success": "#4caf50",
            "warning": "#ff9800",
            "error": "#f44336",
            "running": "#ab47bc"
        }
        icon = icons.get(status_type, "")
        color = colors.get(status_type, "#eaeaea")
        self.status_label.setText(f"{icon} {text}"[:100])
        self.status_label.setStyleSheet(f"color: {color};")
    
    def _update_count_label(self):
        """자막 카운트 라벨 업데이트"""
        count = len(self.subtitles)
        chars = sum(len(s.text) for s in self.subtitles)
        self.count_label.setText(f"📝 {count}문장 | {chars:,}자")
    
    # ========== 글자 크기 ==========
    
    def _set_font_size(self, size: int):
        """자막 영역 폰트 크기 변경"""
        size = max(Config.MIN_FONT_SIZE, min(size, Config.MAX_FONT_SIZE))
        self.font_size = size
        font = self.subtitle_text.font()
        font.setPointSize(size)
        self.subtitle_text.setFont(font)
        self.settings.setValue("font_size", size)
    
    def _adjust_font_size(self, delta: int):
        """폰트 크기 조절"""
        self._set_font_size(self.font_size + delta)
    
    # ========== URL 히스토리 (태그 지원) ==========
    
    def _load_url_history(self):
        """URL 히스토리 로드 - {url: tag} 형태"""
        try:
            if Path("url_history.json").exists():
                with open("url_history.json", 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # dict 형태인지 확인 (새로운 형식)
                    if isinstance(data, dict):
                        return data
                    # 이전 list 형태면 dict로 변환
                    elif isinstance(data, list):
                        return {url: "" for url in data}
        except Exception as e:
            logger.warning(f"URL 히스토리 로드 오류: {e}")
        return {}
    
    def _save_url_history(self):
        """URL 히스토리 저장"""
        try:
            if not isinstance(self.url_history, dict):
                self.url_history = {}
            with open("url_history.json", 'w', encoding='utf-8') as f:
                json.dump(self.url_history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"URL 히스토리 저장 오류: {e}")
    
    def _add_to_history(self, url, tag=""):
        """URL 히스토리에 추가 (자동 태그 매칭)"""
        if not isinstance(self.url_history, dict):
            self.url_history = {}
        
        # 태그가 없으면 자동 감지
        if not tag:
            # 1. 이미 저장된 태그가 있는지 확인
            if url in self.url_history and self.url_history[url]:
                tag = self.url_history[url]
            else:
                # 2. 프리셋/약칭에서 매칭 확인
                tag = self._autodetect_tag(url)
        
        self.url_history[url] = tag
        
        # 히스토리 크기 제한
        if len(self.url_history) > Config.MAX_URL_HISTORY:
            # 가장 오래된 항목 삭제 (dict는 삽입 순서 유지)
            oldest_key = next(iter(self.url_history))
            del self.url_history[oldest_key]
        
        self._save_url_history()
        self._refresh_url_combo()

    def _autodetect_tag(self, url):
        """URL을 기반으로 위원회 이름/약칭 자동 감지"""
        # 1. 정확한 URL 매칭 확인 (프리셋)
        for name, preset_url in self.committee_presets.items():
            if url == preset_url:
                # 약칭이 있으면 약칭 사용 (더 짧고 보기 좋음)
                for abbr, full_name in Config.COMMITTEE_ABBREVIATIONS.items():
                    if full_name == name:
                        return abbr
                return name
        
        # 2. xcode 파라미터 매칭 (숫자 또는 문자열 xcode 모두 지원)
        import re
        match = re.search(r'xcode=([^&]+)', url)
        if match:
            xcode = match.group(1)
            # 프리셋에서 해당 xcode를 가진 URL 찾기
            for name, preset_url in self.committee_presets.items():
                if f"xcode={xcode}" in preset_url:
                    # 약칭 리턴
                    for abbr, full_name in Config.COMMITTEE_ABBREVIATIONS.items():
                        if full_name == name:
                            return abbr
                    return name
                    
        return ""
    
    def _refresh_url_combo(self):
        """URL 콤보박스 새로고침"""
        current_text = self.url_combo.currentText()
        self.url_combo.clear()
        
        for url, tag in self.url_history.items():
            if tag:
                self.url_combo.addItem(f"[{tag}] {url}", url)
            else:
                self.url_combo.addItem(url, url)
        
        # 기본 URL 추가
        if not self.url_history:
            self.url_combo.addItem("https://assembly.webcast.go.kr/main/player.asp")
        
        # 이전 텍스트 복원
        if current_text:
            self.url_combo.setCurrentText(current_text)
    
    def _get_current_url(self):
        """현재 선택된 URL 반환 (태그 제거)"""
        # userData에 원본 URL이 저장되어 있으면 그것 사용
        data = self.url_combo.currentData()
        if data:
            return data
        
        # 없으면 텍스트에서 추출
        text = self.url_combo.currentText()
        if text.startswith("[") and "] " in text:
            return text.split("] ", 1)[1]
        return text
    
    def _edit_url_tag(self):
        """현재 URL의 태그 편집"""
        url = self._get_current_url()
        if not url or not url.startswith(('http://', 'https://')):
            QMessageBox.warning(self, "알림", "태그를 지정할 URL을 먼저 선택하세요.")
            return
        
        current_tag = self.url_history.get(url, "")
        tag, ok = QInputDialog.getText(
            self, "URL 태그 설정",
            f"URL: {url[:50]}...\n\n태그 입력 (예: 본회의, 법사위, 상임위):",
            text=current_tag
        )
        
        if ok:
            self._add_to_history(url, tag.strip())
            QMessageBox.information(self, "성공", f"태그가 설정되었습니다: [{tag}]" if tag else "태그가 제거되었습니다.")
    
    # ========== 상임위 프리셋 ==========
    
    def _load_committee_presets(self):
        """프리셋 파일에서 로드 (없으면 기본값 사용)"""
        self.committee_presets = dict(Config.DEFAULT_COMMITTEE_PRESETS)
        self.custom_presets = {}
        
        try:
            if Path(Config.PRESET_FILE).exists():
                with open(Config.PRESET_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if "presets" in data:
                        self.committee_presets.update(data["presets"])
                    if "custom" in data:
                        self.custom_presets = data["custom"]
        except Exception as e:
            logger.warning(f"프리셋 로드 오류: {e}")
    
    def _save_committee_presets(self):
        """프리셋을 파일에 저장"""
        try:
            data = {
                "presets": self.committee_presets,
                "custom": self.custom_presets
            }
            with open(Config.PRESET_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"프리셋 저장 오류: {e}")
    
    def _build_preset_menu(self):
        """프리셋 메뉴 구성"""
        self.preset_menu.clear()
        
        # 기본 상임위원회
        for name, url in self.committee_presets.items():
            action = QAction(name, self)
            action.setData(url)
            action.triggered.connect(lambda checked, u=url, n=name: self._select_preset(u, n))
            self.preset_menu.addAction(action)
        
        # 사용자 정의 프리셋이 있으면 구분선 추가
        if self.custom_presets:
            self.preset_menu.addSeparator()
            self.preset_menu.addAction("── 사용자 정의 ──").setEnabled(False)
            
            for name, url in self.custom_presets.items():
                action = QAction(f"⭐ {name}", self)
                action.setData(url)
                action.triggered.connect(lambda checked, u=url, n=name: self._select_preset(u, n))
                self.preset_menu.addAction(action)
        
        # 관리 메뉴
        self.preset_menu.addSeparator()
        add_action = QAction("➕ 프리셋 추가...", self)
        add_action.triggered.connect(self._add_custom_preset)
        self.preset_menu.addAction(add_action)
        
        edit_action = QAction("✏️ 프리셋 관리...", self)
        edit_action.triggered.connect(self._manage_presets)
        self.preset_menu.addAction(edit_action)
        
        self.preset_menu.addSeparator()
        export_action = QAction("📤 프리셋 내보내기...", self)
        export_action.triggered.connect(self._export_presets)
        self.preset_menu.addAction(export_action)
        
        import_action = QAction("📥 프리셋 가져오기...", self)
        import_action.triggered.connect(self._import_presets)
        self.preset_menu.addAction(import_action)
    
    def _select_preset(self, url, name):
        """프리셋 선택 시 URL 설정"""
        self.url_combo.setCurrentText(url)
        self._show_toast(f"'{name}' 선택됨", "success", 1500)
    
    def _add_custom_preset(self):
        """사용자 정의 프리셋 추가"""
        name, ok = QInputDialog.getText(
            self, "프리셋 추가",
            "프리셋 이름을 입력하세요:"
        )
        if not ok or not name.strip():
            return
        
        name = name.strip()
        current_url = self._get_current_url()
        
        url, ok = QInputDialog.getText(
            self, "프리셋 URL",
            f"'{name}' 프리셋의 URL을 입력하세요:",
            text=current_url if current_url.startswith('http') else Config.DEFAULT_URL
        )
        
        if ok and url.strip():
            self.custom_presets[name] = url.strip()
            self._save_committee_presets()
            self._build_preset_menu()
            self._show_toast(f"프리셋 '{name}' 추가됨", "success")
    
    def _manage_presets(self):
        """프리셋 관리 대화상자"""
        if not self.custom_presets:
            QMessageBox.information(
                self, "프리셋 관리",
                "사용자 정의 프리셋이 없습니다.\n\n"
                "'➕ 프리셋 추가'를 통해 새 프리셋을 추가하세요."
            )
            return
        
        # 작업 선택
        names = list(self.custom_presets.keys())
        actions = ["수정", "삭제"]
        action, ok = QInputDialog.getItem(
            self, "프리셋 관리",
            "작업을 선택하세요:",
            actions, 0, False
        )
        
        if not ok:
            return
        
        # 프리셋 선택
        name, ok = QInputDialog.getItem(
            self, f"프리셋 {action}",
            f"{action}할 프리셋을 선택하세요:",
            names, 0, False
        )
        
        if not ok or not name:
            return
        
        if action == "삭제":
            reply = QMessageBox.question(
                self, "확인",
                f"'{name}' 프리셋을 삭제하시겠습니까?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                del self.custom_presets[name]
                self._save_committee_presets()
                self._build_preset_menu()
                self._show_toast(f"프리셋 '{name}' 삭제됨", "warning")
        
        elif action == "수정":
            # 이름 수정
            new_name, ok = QInputDialog.getText(
                self, "프리셋 이름 수정",
                f"'{name}' 프리셋의 새 이름:",
                text=name
            )
            if not ok:
                return
            
            # URL 수정
            current_url = self.custom_presets[name]
            new_url, ok = QInputDialog.getText(
                self, "프리셋 URL 수정",
                f"'{new_name}' 프리셋의 URL:",
                text=current_url
            )
            if not ok:
                return
            
            # 기존 프리셋 삭제 후 새로 추가 (이름이 변경되었을 수 있으므로)
            del self.custom_presets[name]
            self.custom_presets[new_name.strip()] = new_url.strip()
            self._save_committee_presets()
            self._build_preset_menu()
            self._show_toast(f"프리셋 '{new_name}' 수정됨", "success")
    
    def _export_presets(self):
        """프리셋을 파일로 내보내기"""
        filename = f"presets_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        path, _ = QFileDialog.getSaveFileName(
            self, "프리셋 내보내기", filename, "JSON 파일 (*.json)"
        )
        
        if path:
            try:
                data = {
                    "version": Config.VERSION,
                    "exported": datetime.now().isoformat(),
                    "committee": self.committee_presets,
                    "custom": self.custom_presets
                }
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                total = len(self.committee_presets) + len(self.custom_presets)
                self._show_toast(f"프리셋 {total}개 내보내기 완료!", "success")
                logger.info(f"프리셋 내보내기 완료: {path}")
            except Exception as e:
                QMessageBox.critical(self, "오류", f"프리셋 내보내기 실패: {e}")
    
    def _import_presets(self):
        """파일에서 프리셋 가져오기"""
        path, _ = QFileDialog.getOpenFileName(
            self, "프리셋 가져오기", "", "JSON 파일 (*.json)"
        )
        
        if path:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                imported_count = 0
                
                # 사용자 정의 프리셋 가져오기 (기존 것에 추가)
                if "custom" in data and isinstance(data["custom"], dict):
                    for name, url in data["custom"].items():
                        if name not in self.custom_presets:
                            self.custom_presets[name] = url
                            imported_count += 1
                
                if imported_count > 0:
                    self._save_committee_presets()
                    self._build_preset_menu()
                    self._show_toast(f"프리셋 {imported_count}개 가져오기 완료!", "success")
                else:
                    self._show_toast("가져올 새 프리셋이 없습니다", "info")
                
                logger.info(f"프리셋 가져오기 완료: {path}, {imported_count}개")
            except json.JSONDecodeError:
                QMessageBox.warning(self, "오류", "잘못된 JSON 파일 형식입니다.")
            except Exception as e:
                QMessageBox.critical(self, "오류", f"프리셋 가져오기 실패: {e}")
    
    # ========== 추출 제어 ==========
    
    def _start(self):
        if self.is_running:
            return
        
        url = self._get_current_url().strip()
        selector = self.selector_combo.currentText().strip()
        
        if not url or not selector:
            QMessageBox.warning(self, "오류", "URL과 선택자를 입력하세요.")
            return
        
        if not url.startswith(('http://', 'https://')):
            QMessageBox.warning(self, "오류", "올바른 URL을 입력하세요.")
            return
        
        try:
            # 히스토리에 추가
            self._add_to_history(url)
            
            # 초기화
            self.subtitle_text.clear()
            self.subtitles = []
            self.last_subtitle = ""
            self.start_time = time.time()
            
            # 큐 비우기
            while not self.message_queue.empty():
                try:
                    self.message_queue.get_nowait()
                except queue.Empty:
                    break
            
            # 실시간 저장 설정 (예외 처리 추가)
            self.realtime_file = None
            if self.realtime_save_check.isChecked():
                try:
                    Path(Config.REALTIME_DIR).mkdir(exist_ok=True)
                    filename = f"{Config.REALTIME_DIR}/자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                    self.realtime_file = open(filename, 'w', encoding='utf-8')
                    self._set_status(f"실시간 저장: {filename}", "success")
                except Exception as e:
                    logger.error(f"실시간 저장 파일 생성 오류: {e}")
            
            # UI 업데이트
            self.is_running = True
            self.stop_event.clear()  # 종료 이벤트 초기화
            self.start_btn.setEnabled(False)
            self.stop_btn.setEnabled(True)
            self.url_combo.setEnabled(False)
            self.selector_combo.setEnabled(False)
            self.progress.show()
            
            self._set_status("Chrome 브라우저 시작 중...", "running")
            self._update_tray_status("🟢 추출 중")
            
            # UI 값을 시작 시점에 복사 (스레드 안전성)
            headless = self.headless_check.isChecked()
            
            # 워커 시작
            self.worker = threading.Thread(
                target=self._extraction_worker,
                args=(url, selector, headless),
                daemon=True
            )
            self.worker.start()
            
            # 타이머 시작
            self.stats_timer.start(Config.STATS_UPDATE_INTERVAL)
            self.backup_timer.start(Config.AUTO_BACKUP_INTERVAL)
        
        except Exception as e:
            logger.exception(f"시작 오류: {e}")
            self._reset_ui()
            QMessageBox.critical(self, "오류", f"시작 중 오류 발생: {e}")
    
    def _stop(self):
        if not self.is_running:
            return
        
        self.is_running = False
        self.stop_event.set()  # 워커 스레드에 종료 신호
        self._set_status("중지 중...", "warning")
        
        # 워커 스레드 종료 대기 (최대 3초)
        if self.worker and self.worker.is_alive():
            self.worker.join(timeout=Config.THREAD_STOP_TIMEOUT)
            if self.worker.is_alive():
                logger.warning("워커 스레드가 시간 내에 종료되지 않음")
        
        # 마지막 자막 저장
        if self.last_subtitle:
            self._finalize_subtitle(self.last_subtitle)
            self.last_subtitle = ""
        
        # 실시간 저장 종료
        if self.realtime_file:
            try:
                self.realtime_file.close()
            except Exception as e:
                logger.debug(f"파일 닫기 오류: {e}")
            self.realtime_file = None
        
        # WebDriver 종료 (중복 종료 방지)
        if self.driver:
            try:
                driver = self.driver
                self.driver = None  # 먼저 None 설정
                driver.quit()
            except Exception as e:
                logger.debug(f"WebDriver 종료 중 오류 (무시됨): {e}")
        
        self._reset_ui()
        self._set_status("중지됨", "warning")
        self._update_tray_status("⚪ 대기 중")
    
    def _reset_ui(self):
        self.is_running = False
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.url_combo.setEnabled(True)
        self.selector_combo.setEnabled(True)
        self.progress.hide()
        self.stats_timer.stop()
        self.backup_timer.stop()
    
    # ========== 워커 스레드 ==========
    
    def _detect_live_broadcast(self, driver, original_url: str) -> str:
        """현재 진행 중인 생중계의 xcgcd를 자동 감지
        
        Args:
            driver: Selenium WebDriver
            original_url: 원래 요청된 URL
            
        Returns:
            str: 감지된 xcgcd를 포함한 URL, 감지 실패 시 원래 URL 반환
        """
        try:
            # 원래 URL에 이미 xcgcd가 있으면 그대로 사용
            if 'xcgcd=' in original_url:
                logger.info(f"URL에 이미 xcgcd 포함됨: {original_url}")
                return original_url
            
            self.message_queue.put(("status", "🔍 현재 생중계 감지 중..."))
            
            # 방법 1: 현재 페이지의 JavaScript 변수에서 xcgcd 가져오기
            scripts = [
                # 전역 변수에서 xcgcd 찾기
                "return typeof xcgcd !== 'undefined' ? xcgcd : null;",
                "return typeof XCGCD !== 'undefined' ? XCGCD : null;",
                "return window.xcgcd || null;",
                "return window.XCGCD || null;",
                # URL 파라미터에서 추출
                "return new URLSearchParams(window.location.search).get('xcgcd');",
                # 현재 스트림 정보에서 추출
                "if(typeof streamInfo !== 'undefined' && streamInfo.xcgcd) return streamInfo.xcgcd; return null;",
                # 플레이어 정보에서 추출
                "if(typeof playerConfig !== 'undefined' && playerConfig.xcgcd) return playerConfig.xcgcd; return null;",
            ]
            
            xcgcd = None
            for script in scripts:
                try:
                    result = driver.execute_script(script)
                    if result:
                        xcgcd = str(result)
                        logger.info(f"JavaScript에서 xcgcd 발견: {xcgcd}")
                        break
                except Exception as e:
                    logger.debug(f"Script 실행 오류: {e}")
            
            # 방법 2: URL이 리다이렉트 되었는지 확인
            if not xcgcd:
                current_url = driver.current_url
                if 'xcgcd=' in current_url:
                    match = re.search(r'xcgcd=([^&]+)', current_url)
                    if match:
                        xcgcd = match.group(1)
                        logger.info(f"리다이렉트된 URL에서 xcgcd 발견: {xcgcd}")
            
            # 방법 3: 페이지 내 생중계 목록에서 현재 방송 찾기
            if not xcgcd:
                try:
                    # 현재 진행 중인 방송 링크 찾기
                    live_scripts = [
                        # 오늘의 생중계 목록에서 첫 번째 방송 링크 가져오기
                        """
                        var links = document.querySelectorAll('a[href*="xcgcd="]');
                        for(var i=0; i<links.length; i++) {
                            var href = links[i].getAttribute('href');
                            if(href && href.includes('xcgcd=')) {
                                var match = href.match(/xcgcd=([^&]+)/);
                                if(match) return match[1];
                            }
                        }
                        return null;
                        """,
                        # 생중계 플레이어 iframe에서 xcgcd 추출
                        """
                        var iframe = document.querySelector('iframe[src*="xcgcd="]');
                        if(iframe) {
                            var src = iframe.getAttribute('src');
                            var match = src.match(/xcgcd=([^&]+)/);
                            if(match) return match[1];
                        }
                        return null;
                        """,
                        # 숨겨진 input 필드에서 찾기
                        """
                        var input = document.querySelector('input[name="xcgcd"], input#xcgcd');
                        if(input) return input.value;
                        return null;
                        """,
                    ]
                    
                    for script in live_scripts:
                        result = driver.execute_script(script)
                        if result:
                            xcgcd = str(result)
                            logger.info(f"페이지 요소에서 xcgcd 발견: {xcgcd}")
                            break
                except Exception as e:
                    logger.debug(f"생중계 목록 파싱 오류: {e}")
            
            # 방법 4: 메인 페이지에서 오늘의 생중계 정보 가져오기
            if not xcgcd:
                navigated_to_main = False
                try:
                    # xcode 추출 (숫자 또는 문자열 xcode 모두 지원)
                    xcode_match = re.search(r'xcode=([^&]+)', original_url)
                    if xcode_match:
                        xcode = xcode_match.group(1)
                        
                        # 메인 페이지로 이동해서 해당 xcode의 생중계 정보 찾기
                        main_url = "https://assembly.webcast.go.kr/main/"
                        driver.get(main_url)
                        navigated_to_main = True
                        time.sleep(3)
                        
                        # 해당 xcode의 생중계 링크 찾기
                        script = f"""
                        var links = document.querySelectorAll('a[href*="xcode={xcode}"][href*="xcgcd="]');
                        for(var i=0; i<links.length; i++) {{
                            var href = links[i].getAttribute('href');
                            var match = href.match(/xcgcd=([^&]+)/);
                            if(match) return match[1];
                        }}
                        return null;
                        """
                        result = driver.execute_script(script)
                        if result:
                            xcgcd = str(result)
                            logger.info(f"메인 페이지에서 xcgcd 발견: {xcgcd}")
                except Exception as e:
                    logger.debug(f"메인 페이지 조회 오류: {e}")
                finally:
                    # 메인 페이지로 이동했지만 xcgcd를 못 찾은 경우 원래 URL로 복귀
                    if navigated_to_main and not xcgcd:
                        try:
                            driver.get(original_url)
                            time.sleep(2)
                            logger.info(f"원래 URL로 복귀: {original_url}")
                        except Exception as e:
                            logger.debug(f"원래 URL 복귀 실패: {e}")
            
            # xcgcd를 찾았으면 URL 업데이트 (유효성 검증 포함)
            if xcgcd and len(xcgcd) >= 10:  # 최소 길이 검증 (유효한 xcgcd는 보통 20자 이상)
                if '?' in original_url:
                    new_url = f"{original_url}&xcgcd={xcgcd}"
                else:
                    new_url = f"{original_url}?xcgcd={xcgcd}"
                
                display_xcgcd = xcgcd[:15] + '...' if len(xcgcd) > 15 else xcgcd
                self.message_queue.put(("status", f"✅ 생중계 감지 성공! (xcgcd={display_xcgcd})"))
                logger.info(f"생중계 URL 업데이트: {new_url}")
                return new_url
            else:
                self.message_queue.put(("status", "⚠️ 생중계 정보를 찾을 수 없음 - 기본 URL 사용"))
                logger.warning("생중계 xcgcd를 찾을 수 없음, 원래 URL 사용")
                return original_url
                
        except Exception as e:
            logger.error(f"생중계 감지 오류: {e}")
            self.message_queue.put(("status", f"⚠️ 생중계 감지 실패: {e}"))
            return original_url
    
    def _extraction_worker(self, url, selector, headless):
        """자막 추출 워커 스레드
        
        Args:
            url: 대상 웹사이트 URL
            selector: 자막 요소 CSS 선택자
            headless: 헤드리스 모드 여부 (시작 시점에 복사됨)
        """
        driver = None
        
        try:
            options = Options()
            options.add_argument("--log-level=3")
            options.add_argument("--disable-blink-features=AutomationControlled")
            options.add_experimental_option("excludeSwitches", ["enable-logging", "enable-automation"])
            options.add_experimental_option('useAutomationExtension', False)
            
            # 헤드리스 모드
            if headless:
                options.add_argument("--headless=new")
                options.add_argument("--window-size=1280,720")
                options.add_argument("--disable-gpu")
                options.add_argument("--no-sandbox")
                options.add_argument("--disable-dev-shm-usage")
                options.add_argument("--disable-setuid-sandbox")
                options.add_argument("--remote-debugging-port=0")  # 동적 포트 할당 (다중 인스턴스 충돌 방지)
                options.add_argument("--enable-javascript")
                self.message_queue.put(("status", "헤드리스 모드로 시작 중..."))
            
            # WebDriver 초기화 재시도 (최대 3회)
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    driver = webdriver.Chrome(options=options)
                    self.driver = driver
                    self.message_queue.put(("status", "Chrome 시작 완료"))
                    break
                except Exception as e:
                    if attempt < max_retries - 1:
                        self.message_queue.put(("status", f"Chrome 시작 재시도 ({attempt+2}/{max_retries})..."))
                        time.sleep(2)
                    else:
                        self.message_queue.put(("error", f"Chrome 오류 ({max_retries}회 시도 실패): {e}"))
                        return
            
            self.message_queue.put(("status", "페이지 로딩 중..."))
            driver.get(url)
            time.sleep(5)  # 초기 로딩 대기 (v16.0과 동일하게 5초로 복원)
            
            # 현재 생중계 xcgcd 자동 감지 (URL에 xcgcd가 없는 경우에만)
            # xcgcd가 이미 URL에 있으면 자동 감지 건너뛰기
            if 'xcgcd=' not in url:
                detected_url = self._detect_live_broadcast(driver, url)
                # 감지된 URL이 다르면 페이지 다시 로드
                if detected_url != url and 'xcgcd=' in detected_url:
                    self.message_queue.put(("status", "🔄 생중계 URL로 이동 중..."))
                    driver.get(detected_url)
                    time.sleep(5)
            
            self.message_queue.put(("status", "AI 자막 활성화 중..."))
            self._activate_subtitle(driver)
            
            self.message_queue.put(("status", "자막 요소 검색 중..."))
            
            # 헤드리스 모드에서는 타임아웃 증가 (JavaScript 로딩 지연 대비)
            timeout = 30 if headless else 20
            wait = WebDriverWait(driver, timeout)
            
            # _find_subtitle_selector를 활용하여 자막 요소 자동 감지
            detected_selector = self._find_subtitle_selector(driver)
            selectors_to_try = [selector, detected_selector, "#viewSubtit .incont", "#viewSubtit", ".subtitle_area"]
            # 중복 제거
            selectors_to_try = list(dict.fromkeys(selectors_to_try))
            
            found = False
            working_selector = selector  # 실제 동작하는 셀렉터 저장
            for sel in selectors_to_try:
                try:
                    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, sel)))
                    self.message_queue.put(("status", f"자막 요소 찾음: {sel}"))
                    found = True
                    working_selector = sel
                    break
                except Exception:
                    continue
            
            if not found:
                # 사용자에게 직접 링크를 가져오라고 안내
                self.message_queue.put((
                    "subtitle_not_found",
                    "자막 요소를 찾을 수 없습니다.\n\n"
                    "현재 진행 중인 생중계가 없거나, URL이 올바르지 않을 수 있습니다.\n\n"
                    "📌 해결 방법:\n"
                    "1. https://assembly.webcast.go.kr 사이트에서\n"
                    "2. 현재 진행 중인 중계를 직접 클릭하여\n"
                    "3. 해당 페이지의 URL을 복사해서 입력해주세요.\n\n"
                    "💡 Tip: URL에 'xcgcd=' 파라미터가 포함된 링크가 필요합니다."
                ))
                return
            
            # 찾은 셀렉터를 사용하여 모니터링
            selector = working_selector
            
            self.message_queue.put(("status", "자막 모니터링 중"))
            
            last_check = time.time()
            consecutive_errors = 0  # 네트워크 오류 재연결 카운터
            max_consecutive_errors = 5  # 최대 연속 오류 횟수
            empty_text_count = 0  # 빈 텍스트 연속 카운터
            max_empty_text = 30  # 30회 연속 빈 텍스트 (약 6초) 시 경고
            
            # stop_event 사용으로 더 빠른 종료 응답
            while not self.stop_event.is_set():
                try:
                    now = time.time()
                    if now - last_check >= 0.2:
                        try:
                            text = driver.find_element(By.CSS_SELECTOR, selector).text.strip()
                            consecutive_errors = 0  # 성공 시 카운터 리셋
                        except (NoSuchElementException, StaleElementReferenceException):
                            text = ""
                        
                        text = self._clean_text(text)
                        
                        # 빈 텍스트 연속 감지 (방송 종료 또는 자막 없음)
                        if not text:
                            empty_text_count += 1
                            if empty_text_count == max_empty_text:
                                self.message_queue.put(("status", "⏸️ 자막 대기 중... (방송 확인 필요)"))
                        else:
                            empty_text_count = 0  # 리셋
                        
                        # 텍스트 정규화 비교 (공백/특수문자 차이 무시)
                        normalized_text = re.sub(r'\s+', ' ', text).strip()
                        normalized_last = re.sub(r'\s+', ' ', self.last_subtitle).strip() if self.last_subtitle else ""
                        
                        if text and normalized_text != normalized_last:
                            self.message_queue.put(("preview", text))
                        
                        last_check = now
                    
                    # stop_event 대기 (0.05초, 즉시 응답 가능)
                    self.stop_event.wait(timeout=0.05)
                    
                except WebDriverException as e:
                    # 네트워크/브라우저 오류 - 재연결 시도
                    consecutive_errors += 1
                    if not self.stop_event.is_set():
                        logger.warning(f"브라우저 오류 ({consecutive_errors}/{max_consecutive_errors}): {e}")
                        self.message_queue.put(("status", f"⚠️ 연결 오류 ({consecutive_errors}/{max_consecutive_errors})"))
                        
                        if consecutive_errors >= max_consecutive_errors:
                            self.message_queue.put(("error", f"네트워크 오류로 중단됨: {consecutive_errors}회 연속 실패"))
                            break
                        
                        # 재연결 대기
                        time.sleep(2)
                        try:
                            driver.refresh()
                            time.sleep(3)
                            self._activate_subtitle(driver)
                            self.message_queue.put(("status", "✅ 재연결 성공"))
                        except Exception:
                            pass
                            
                except Exception as e:
                    if not self.stop_event.is_set():
                        logger.warning(f"모니터링 중 오류: {e}")
                    time.sleep(0.5)
        
        except Exception as e:
            if not self.stop_event.is_set():
                self.message_queue.put(("error", str(e)))
        
        finally:
            if driver:
                try:
                    driver.quit()
                except Exception as e:
                    logger.debug(f"워커 WebDriver 종료 중 오류: {e}")
                self.driver = None
            self.message_queue.put(("finished", ""))
    
    def _activate_subtitle(self, driver) -> bool:
        """자막 레이어 활성화 - 다양한 방법 시도
        
        Returns:
            bool: 활성화 성공 여부
        """
        activation_scripts = [
            # 방법 1: layerSubtit 함수 호출
            "if(typeof layerSubtit==='function'){layerSubtit(); return true;} return false;",
            # 방법 2: 자막 버튼 클릭
            "var btn=document.querySelector('.btn_subtit'); if(btn){btn.click(); return true;} return false;",
            "var btn=document.querySelector('#btnSubtit'); if(btn){btn.click(); return true;} return false;",
            # 방법 3: AI 자막 버튼
            "var btn=document.querySelector('[data-action=\\'subtitle\\']'); if(btn){btn.click(); return true;} return false;",
            # 방법 4: 자막 레이어 직접 표시
            "var layer=document.querySelector('#viewSubtit'); if(layer){layer.style.display='block'; return true;} return false;",
        ]
        
        activated = False
        for script in activation_scripts:
            try:
                result = driver.execute_script(script)
                if result:
                    logger.info(f"자막 활성화 성공: {script[:50]}...")
                    activated = True
                time.sleep(0.5)
            except Exception as e:
                logger.debug(f"자막 활성화 스크립트 실패: {e}")
        
        # 추가 대기 - 자막 레이어 로딩
        time.sleep(2.0)
        
        # 활성화 검증: 자막 요소가 실제로 보이는지 확인
        try:
            elem = driver.find_element(By.CSS_SELECTOR, "#viewSubtit")
            is_visible = driver.execute_script(
                "return arguments[0].offsetParent !== null && getComputedStyle(arguments[0]).display !== 'none';",
                elem
            )
            if is_visible:
                logger.info("자막 레이어 활성화 확인됨")
                return True
            else:
                logger.warning("자막 레이어가 숨겨져 있음")
        except Exception as e:
            logger.debug(f"자막 레이어 확인 실패: {e}")
        
        return activated
    
    def _find_subtitle_selector(self, driver) -> str:
        """사용 가능한 자막 셀렉터 자동 감지
        
        Returns:
            str: 찾은 셀렉터 또는 기본 셀렉터
        """
        # 우선순위대로 셀렉터 확인
        selectors = [
            "#viewSubtit .incont",
            "#viewSubtit span",
            "#viewSubtit",
            ".subtitle_area",
            ".ai_subtitle",
            "[class*='subtitle']",
        ]
        
        for sel in selectors:
            try:
                elements = driver.find_elements(By.CSS_SELECTOR, sel)
                for elem in elements:
                    text = elem.text.strip()
                    if text and len(text) > 2:  # 의미 있는 텍스트가 있는 경우
                        logger.info(f"자막 셀렉터 발견: {sel}")
                        return sel
            except Exception:
                continue
        
        return "#viewSubtit .incont"  # 기본값
    
    def _clean_text(self, text: str) -> str:
        """자막 텍스트 정리"""
        if not text:
            return ""
        # 년도 제거
        text = re.sub(r'\b\d{4}년\b', '', text)
        # 특수 문자 정리
        text = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', text)  # Zero-width 문자 제거
        # 연속 공백 정리
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    # ========== 메시지 큐 처리 ==========
    
    def _process_message_queue(self):
        """메시지 큐 처리 (100ms마다 호출) - 예외 처리 강화"""
        try:
            # 최대 10개 메시지까지 처리 (무한 루프 방지)
            for _ in range(10):
                try:
                    msg_type, data = self.message_queue.get_nowait()
                    self._handle_message(msg_type, data)
                except queue.Empty:
                    break
        except Exception as e:
            logger.error(f"큐 처리 오류: {e}")
    
    def _handle_message(self, msg_type, data):
        """개별 메시지 처리"""
        try:
            if msg_type == "status":
                self.status_label.setText(str(data)[:200])
            
            elif msg_type == "preview":
                self._process_raw_text(data)
            
            elif msg_type == "error":
                self.progress.hide()
                self._reset_ui()
                QMessageBox.critical(self, "오류", str(data))
            
            elif msg_type == "finished":
                if self.last_subtitle:
                    self._finalize_subtitle(self.last_subtitle)
                    self.last_subtitle = ""
                
                self._refresh_text()
                self._reset_ui()
                
                # 스레드 안전하게 통계 계산
                with self.subtitle_lock:
                    subtitle_count = len(self.subtitles)
                    total_chars = sum(len(s.text) for s in self.subtitles)
                self.status_label.setText(f"완료 - {subtitle_count}문장, {total_chars:,}자")
            
            elif msg_type == "subtitle_not_found":
                # 자막 요소를 찾지 못했을 때 사용자 안내
                self.progress.hide()
                self._reset_ui()
                self._update_tray_status("⚪ 대기 중")
                
                # 상세 안내 다이얼로그
                msg_box = QMessageBox(self)
                msg_box.setWindowTitle("자막을 찾을 수 없습니다")
                msg_box.setIcon(QMessageBox.Icon.Warning)
                msg_box.setText(str(data))
                msg_box.setStandardButtons(
                    QMessageBox.StandardButton.Open | 
                    QMessageBox.StandardButton.Ok
                )
                msg_box.button(QMessageBox.StandardButton.Open).setText("🌐 사이트 열기")
                msg_box.button(QMessageBox.StandardButton.Ok).setText("확인")
                
                result = msg_box.exec()
                
                # 사이트 열기 버튼 클릭 시 브라우저에서 열기
                if result == QMessageBox.StandardButton.Open:
                    import webbrowser
                    webbrowser.open("https://assembly.webcast.go.kr")
        
        except Exception as e:
            logger.error(f"메시지 처리 오류 ({msg_type}): {e}")
    
    def _process_raw_text(self, raw):
        if not raw or raw == self.last_subtitle:
            return
        
        # 시간 기반 확정: 자막이 변경될 때마다 시간 기록
        # 2초 동안 변경 없으면 확정
        self.last_update_time = time.time()
        self.last_subtitle = raw
        
        # 확정 타이머 시작 (아직 안돌고 있으면)
        if not self.finalize_timer.isActive():
            self.finalize_timer.start(Config.FINALIZE_CHECK_INTERVAL)
        
        # 화면에는 현재 자막만 표시 (누적 없음, 미리보기만)
        self._update_preview(raw)
    
    def _check_finalize(self):
        """2초 동안 변경 없으면 현재 자막 확정"""
        if not self.last_subtitle:
            return
        
        elapsed = time.time() - self.last_update_time
        if elapsed >= 2.0:  # 2초 동안 변경 없음
            self._finalize_subtitle(self.last_subtitle)
            self.last_subtitle = ""
            self.finalize_timer.stop()
            self._refresh_text()
    
    def _update_preview(self, raw: str) -> None:
        """미리보기 업데이트 - 확정 자막 + 현재 자막 모두 표시"""
        self._render_subtitles(current_raw=raw)
    
    def _render_subtitles(self, current_raw: str = None) -> None:
        """자막 렌더링 공통 메소드 - 중복 코드 제거
        
        Args:
            current_raw: 현재 진행 중인 자막 (None이면 확정 자막만 표시)
        """
        self.subtitle_text.clear()
        cursor = self.subtitle_text.textCursor()
        
        show_ts = self.timestamp_action.isChecked()
        
        # 확정된 자막들 표시 (스레드 안전하게 복사하여 순회)
        with self.subtitle_lock:
            subtitles_copy = list(self.subtitles)
        for i, entry in enumerate(subtitles_copy):
            if i > 0:
                cursor.insertText("\n\n")
            
            # 타임스탬프
            if show_ts:
                fmt = QTextCharFormat()
                fmt.setForeground(QColor("#888888"))
                cursor.insertText(f"[{entry.timestamp.strftime('%H:%M:%S')}] ", fmt)
            
            # 텍스트 (키워드 하이라이트)
            self._insert_highlighted_text(cursor, entry.text)
        
        # 현재 진행 중인 자막 (회색, 진행 중 표시)
        if current_raw:
            if self.subtitles:
                cursor.insertText("\n\n")
            
            fmt = QTextCharFormat()
            fmt.setForeground(QColor("#aaaaaa"))
            cursor.insertText("⏳ ", fmt)
            cursor.insertText(current_raw)
        
        if self.auto_scroll_check.isChecked():
            self.subtitle_text.moveCursor(QTextCursor.MoveOperation.End)
    
    def _refresh_text(self) -> None:
        """확정된 자막만 표시 (진행 중인 자막 없음)"""
        self._render_subtitles(current_raw=None)
    
    def _finalize_subtitle(self, text: str) -> None:
        """자막 확정 - 이전 확정 자막과 겹치는 부분 제거 후 이어붙이기"""
        if not text:
            return
        
        with self.subtitle_lock:
            # 이전에 확정된 자막이 있으면, 겹치는 부분 제거
            if self.subtitles:
                last_text = self.subtitles[-1].text
                
                # 새 텍스트가 이전 텍스트로 시작하면 -> 새로운 부분만 추가
                if text.startswith(last_text):
                    new_part = text[len(last_text):].strip()
                    if new_part:
                        # 이전 자막에 이어붙이기
                        self.subtitles[-1].text = text
                        self.subtitles[-1].end_time = datetime.now()
                        
                        # 실시간 저장 (덮어쓰기 대신 업데이트)
                        if self.realtime_file:
                            try:
                                timestamp = self.subtitles[-1].timestamp.strftime('%H:%M:%S')
                                self.realtime_file.write(f"[{timestamp}] +{new_part}\n")
                                self.realtime_file.flush()
                            except IOError as e:
                                logger.warning(f"실시간 저장 쓰기 오류: {e}")
                    return
                
                # 텍스트 겹침 확인 (앞부분이 잘려서 시작하는 경우)
                for overlap_len in range(min(len(last_text), len(text)), 10, -1):
                    if last_text.endswith(text[:overlap_len]):
                        # 겹치는 부분을 제외하고 새 부분만 이어붙이기
                        new_part = text[overlap_len:].strip()
                        if new_part:
                            self.subtitles[-1].text += " " + new_part
                            self.subtitles[-1].end_time = datetime.now()
                            
                            if self.realtime_file:
                                try:
                                    self.realtime_file.write(f"+ {new_part}\n")
                                    self.realtime_file.flush()
                                except IOError as e:
                                    logger.warning(f"실시간 저장 쓰기 오류: {e}")
                        return
            
            # 첫 번째 자막 또는 완전히 새로운 문장
            entry = SubtitleEntry(text)
            entry.start_time = datetime.now()
            entry.end_time = datetime.now()
            
            self.subtitles.append(entry)
        
        # 키워드 알림 확인 (lock 외부)
        self._check_keyword_alert(text)
        
        # 카운트 라벨 업데이트
        self._update_count_label()
        
        # 실시간 저장
        if self.realtime_file:
            try:
                timestamp = entry.timestamp.strftime('%H:%M:%S')
                self.realtime_file.write(f"[{timestamp}] {text}\n")
                self.realtime_file.flush()
            except IOError as e:
                logger.warning(f"실시간 저장 쓰기 오류: {e}")
        
        self._refresh_text()

    
    def _insert_highlighted_text(self, cursor, text):
        """텍스트에서 키워드만 하이라이트"""
        # 빈 키워드 필터링
        valid_keywords = [k.strip() for k in self.keywords if k and k.strip()]
        
        if not valid_keywords:
            cursor.insertText(text)
            return
        
        # 키워드 하이라이트 (해당 단어만)
        pattern = '|'.join(re.escape(k) for k in valid_keywords)
        parts = re.split(f'({pattern})', text, flags=re.IGNORECASE)
        
        # 키워드 소문자 리스트 (비교용)
        keywords_lower = [k.lower() for k in valid_keywords]
        
        for part in parts:
            if not part:  # 빈 문자열 건너뛰기
                continue
            
            if part.lower() in keywords_lower:
                # 키워드만 하이라이트 (눈에 띄는 스타일)
                fmt = QTextCharFormat()
                fmt.setBackground(QColor("#ffd700"))  # 골드 배경
                fmt.setForeground(QColor("#000000"))  # 검정 글자
                fmt.setFontWeight(QFont.Weight.Bold)  # 볼드
                cursor.insertText(part, fmt)
            else:
                # 일반 텍스트 (기본 포맷으로)
                fmt = QTextCharFormat()
                cursor.insertText(part, fmt)
    
    # ========== 통계 ==========
    
    def _update_stats(self):
        if self.start_time:
            elapsed = int(time.time() - self.start_time)
            h, r = divmod(elapsed, 3600)
            m, s = divmod(r, 60)
            self.stat_time.setText(f"⏱️ 실행 시간: {h:02d}:{m:02d}:{s:02d}")
            
            # 스레드 안전하게 자막 접근
            with self.subtitle_lock:
                subtitles_snapshot = list(self.subtitles)
            
            total_chars = sum(len(s.text) for s in subtitles_snapshot)
            total_words = sum(len(s.text.split()) for s in subtitles_snapshot)
            
            self.stat_chars.setText(f"📝 글자 수: {total_chars:,}")
            self.stat_words.setText(f"📖 단어 수: {total_words:,}")
            self.stat_sents.setText(f"💬 문장 수: {len(subtitles_snapshot)}")
            
            if elapsed > 0:
                cpm = int(total_chars / (elapsed / 60))
                self.stat_cpm.setText(f"⚡ 분당 글자: {cpm}")
    
    # ========== 검색 ==========
    
    def _show_search(self):
        self.search_frame.show()
        self.search_input.setFocus()
        self.search_input.selectAll()
    
    def _hide_search(self):
        self.search_frame.hide()
        self._refresh_text()
    
    def _do_search(self):
        query = self.search_input.text()
        if not query:
            return
        
        text = self.subtitle_text.toPlainText()
        self.search_matches = []
        
        start = 0
        while True:
            idx = text.lower().find(query.lower(), start)
            if idx == -1:
                break
            self.search_matches.append(idx)
            start = idx + 1
        
        self.search_idx = 0
        self.search_count.setText(f"{len(self.search_matches)}개")
        
        if self.search_matches:
            self._highlight_search(0)
    
    def _nav_search(self, delta):
        if not self.search_matches:
            return
        
        self.search_idx = (self.search_idx + delta) % len(self.search_matches)
        self._highlight_search(self.search_idx)
    
    def _highlight_search(self, idx):
        if not self.search_matches:
            return
        
        pos = self.search_matches[idx]
        query = self.search_input.text()
        
        cursor = self.subtitle_text.textCursor()
        cursor.setPosition(pos)
        cursor.setPosition(pos + len(query), QTextCursor.MoveMode.KeepAnchor)
        
        self.subtitle_text.setTextCursor(cursor)
        self.subtitle_text.ensureCursorVisible()
        
        self.search_count.setText(f"{idx + 1}/{len(self.search_matches)}")
    
    # ========== 키워드 ==========
    
    def _set_keywords(self):
        """하이라이트 키워드 설정"""
        current = ", ".join(self.keywords)
        text, ok = QInputDialog.getText(
            self, "하이라이트 키워드 설정",
            "하이라이트할 키워드 (쉼표로 구분):",
            text=current
        )
        
        if ok:
            self.keywords = [k.strip() for k in text.split(",") if k.strip()]
            # 설정 저장
            self.settings.setValue("highlight_keywords", ", ".join(self.keywords))
            self._refresh_text()
            self._show_toast(f"하이라이트 키워드 {len(self.keywords)}개 설정됨", "success")
    
    def _set_alert_keywords(self):
        """알림 키워드 설정 - 해당 키워드 감지 시 토스트 알림"""
        current = ", ".join(self.alert_keywords)
        text, ok = QInputDialog.getText(
            self, "알림 키워드 설정",
            "알림을 받을 키워드 (쉼표로 구분):\n예: 법안, 의결, 통과",
            text=current
        )
        
        if ok:
            self.alert_keywords = [k.strip() for k in text.split(",") if k.strip()]
            # 설정 저장
            self.settings.setValue("alert_keywords", ", ".join(self.alert_keywords))
            self._show_toast(f"알림 키워드 {len(self.alert_keywords)}개 설정됨", "success")
    
    def _check_keyword_alert(self, text: str):
        """키워드 포함 시 알림 표시"""
        if not self.alert_keywords:
            return
        
        for keyword in self.alert_keywords:
            if keyword and keyword.lower() in text.lower():
                self._show_toast(f"🔔 키워드 감지: {keyword}", "warning", 5000)
                break  # 한 번만 알림
    
    # ========== 자동 백업 ==========
    
    def _auto_backup(self):
        """자동 백업 실행"""
        if not self.subtitles:
            return
        
        try:
            backup_dir = Path(Config.BACKUP_DIR)
            backup_dir.mkdir(exist_ok=True)
            
            # 백업 파일 생성
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_file = backup_dir / f"backup_{timestamp}.json"
            
            # 스레드 안전하게 자막 복사
            with self.subtitle_lock:
                subtitles_copy = [s.to_dict() for s in self.subtitles]
            
            data = {
                'version': Config.VERSION,
                'created': datetime.now().isoformat(),
                'url': self._get_current_url(),
                'subtitles': subtitles_copy
            }
            
            with open(backup_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            # 오래된 백업 삭제 (최대 개수 유지)
            self._cleanup_old_backups()
            
            logger.info(f"자동 백업 완료: {backup_file}")
            
        except Exception as e:
            logger.error(f"자동 백업 오류: {e}")
    
    def _cleanup_old_backups(self):
        """오래된 백업 파일 정리"""
        try:
            backup_dir = Path(Config.BACKUP_DIR)
            backups = sorted(backup_dir.glob("backup_*.json"), reverse=True)
            
            # 최대 개수 초과분 삭제
            for old_backup in backups[Config.MAX_BACKUP_COUNT:]:
                old_backup.unlink()
                logger.debug(f"오래된 백업 삭제: {old_backup}")
        except Exception as e:
            logger.warning(f"백업 정리 중 오류: {e}")
    
    # ========== 파일 저장 ==========
    
    def _save_in_background(self, save_func, path: str, success_msg: str, error_prefix: str):
        """백그라운드에서 파일 저장 (자막 수집 중단 없이)
        
        Args:
            save_func: 실제 저장을 수행하는 함수 (path를 인자로 받음)
            path: 저장할 파일 경로
            success_msg: 성공 시 토스트 메시지
            error_prefix: 실패 시 에러 메시지 접두어
        """
        def background_save():
            try:
                save_func(path)
                # UI 스레드에서 토스트 표시 (QTimer.singleShot 사용)
                QTimer.singleShot(0, lambda: self._show_toast(success_msg, "success"))
            except Exception as e:
                logger.error(f"{error_prefix}: {e}")
                QTimer.singleShot(0, lambda: self._show_toast(f"{error_prefix}: {e}", "error", 5000))
        
        # 백그라운드 스레드에서 저장 실행
        save_thread = threading.Thread(target=background_save, daemon=True)
        save_thread.start()
        
        # 저장 시작 알림 (즉시)
        self._show_toast(f"💾 저장 중... ({Path(path).name})", "info", 1500)
    
    def _get_accumulated_text(self):
        with self.subtitle_lock:
            return "\n".join(s.text for s in self.subtitles)
    
    def _export_stats(self):
        """자막 통계 내보내기"""
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "내보낼 내용이 없습니다.")
            return
        
        # 스레드 안전하게 자막 스냅샷 생성
        with self.subtitle_lock:
            subtitles_snapshot = list(self.subtitles)
        
        # 통계 계산
        total_chars = sum(len(s.text) for s in subtitles_snapshot)
        total_words = sum(len(s.text.split()) for s in subtitles_snapshot)
        
        # 시간대별 통계
        hour_counts = {}
        for entry in subtitles_snapshot:
            hour = entry.timestamp.hour
            hour_counts[hour] = hour_counts.get(hour, 0) + 1
        
        # 가장 긴/짧은 문장
        longest = max(subtitles_snapshot, key=lambda s: len(s.text))
        shortest = min(subtitles_snapshot, key=lambda s: len(s.text))
        
        # 파일 저장
        filename = f"자막통계_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        path, _ = QFileDialog.getSaveFileName(self, "통계 내보내기", filename, "텍스트 (*.txt)")
        
        if path:
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write("=" * 50 + "\n")
                    f.write("        🏛️ 국회 자막 통계 보고서\n")
                    f.write("=" * 50 + "\n\n")
                    
                    f.write(f"생성 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}\n\n")
                    
                    f.write("📊 기본 통계\n")
                    f.write("-" * 30 + "\n")
                    f.write(f"  총 문장 수: {len(subtitles_snapshot):,}개\n")
                    f.write(f"  총 글자 수: {total_chars:,}자\n")
                    f.write(f"  총 단어 수: {total_words:,}개\n")
                    f.write(f"  평균 문장 길이: {total_chars/len(subtitles_snapshot):.1f}자\n")
                    f.write(f"  평균 단어 수: {total_words/len(subtitles_snapshot):.1f}개\n\n")
                    
                    f.write("📏 문장 분석\n")
                    f.write("-" * 30 + "\n")
                    f.write(f"  가장 긴 문장: {len(longest.text)}자\n")
                    f.write(f"    \"{longest.text[:50]}{'...' if len(longest.text) > 50 else ''}\"\n")
                    f.write(f"  가장 짧은 문장: {len(shortest.text)}자\n")
                    f.write(f"    \"{shortest.text}\"\n\n")
                    
                    if hour_counts:
                        f.write("⏰ 시간대별 분포\n")
                        f.write("-" * 30 + "\n")
                        for h in sorted(hour_counts.keys()):
                            bar = "█" * min(hour_counts[h] // 2, 20)
                            f.write(f"  {h:02d}시: {bar} {hour_counts[h]}개\n")
                        f.write("\n")
                    
                    # 하이라이트 키워드 통계
                    if self.keywords:
                        f.write("🔍 키워드 빈도\n")
                        f.write("-" * 30 + "\n")
                        all_text = " ".join(s.text for s in subtitles_snapshot).lower()
                        for kw in self.keywords:
                            count = all_text.count(kw.lower())
                            if count > 0:
                                f.write(f"  {kw}: {count}회\n")
                        f.write("\n")
                    
                    f.write("=" * 50 + "\n")
                
                self._show_toast("통계 내보내기 완료!", "success")
            except Exception as e:
                QMessageBox.critical(self, "오류", f"통계 저장 실패: {e}")
    
    def _save_txt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        path, _ = QFileDialog.getSaveFileName(self, "TXT 저장", filename, "텍스트 (*.txt)")
        
        if path:
            # 스레드 안전하게 자막 스냅샷 생성 (UI 스레드에서)
            with self.subtitle_lock:
                subtitles_snapshot = list(self.subtitles)
            
            # 실제 저장 함수 정의
            def do_save(filepath):
                with open(filepath, 'w', encoding='utf-8') as f:
                    for entry in subtitles_snapshot:
                        f.write(f"[{entry.timestamp.strftime('%H:%M:%S')}] {entry.text}\n")
            
            # 백그라운드에서 저장 실행
            self._save_in_background(do_save, path, "TXT 저장 완료!", "TXT 저장 실패")
    
    def _save_srt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.srt"
        path, _ = QFileDialog.getSaveFileName(self, "SRT 저장", filename, "SubRip (*.srt)")
        
        if path:
            # 스레드 안전하게 자막 스냅샷 생성 (UI 스레드에서)
            with self.subtitle_lock:
                subtitles_snapshot = list(self.subtitles)
            
            def do_save(filepath):
                with open(filepath, 'w', encoding='utf-8') as f:
                    for i, entry in enumerate(subtitles_snapshot, 1):
                        if entry.start_time and entry.end_time:
                            start = entry.start_time.strftime('%H:%M:%S,000')
                            end = entry.end_time.strftime('%H:%M:%S,000')
                        else:
                            start = entry.timestamp.strftime('%H:%M:%S,000')
                            end_time = entry.timestamp + timedelta(seconds=3)
                            end = end_time.strftime('%H:%M:%S,000')
                        f.write(f"{i}\n{start} --> {end}\n{entry.text}\n\n")
            
            self._save_in_background(do_save, path, "SRT 저장 완료!", "SRT 저장 실패")
    
    def _save_vtt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.vtt"
        path, _ = QFileDialog.getSaveFileName(self, "VTT 저장", filename, "WebVTT (*.vtt)")
        
        if path:
            with self.subtitle_lock:
                subtitles_snapshot = list(self.subtitles)
            
            def do_save(filepath):
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write("WEBVTT\n\n")
                    for i, entry in enumerate(subtitles_snapshot, 1):
                        if entry.start_time and entry.end_time:
                            start = entry.start_time.strftime('%H:%M:%S.000')
                            end = entry.end_time.strftime('%H:%M:%S.000')
                        else:
                            start = entry.timestamp.strftime('%H:%M:%S.000')
                            end_time = entry.timestamp + timedelta(seconds=3)
                            end = end_time.strftime('%H:%M:%S.000')
                        f.write(f"{i}\n{start} --> {end}\n{entry.text}\n\n")
            
            self._save_in_background(do_save, path, "VTT 저장 완료!", "VTT 저장 실패")
    
    def _save_docx(self):
        """DOCX (Word) 파일로 저장"""
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.warning(
                self, "라이브러리 필요",
                "DOCX 저장을 위해 python-docx 라이브러리가 필요합니다.\n\n"
                "설치: pip install python-docx"
            )
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path, _ = QFileDialog.getSaveFileName(self, "DOCX 저장", filename, "Word 문서 (*.docx)")
        
        if path:
            try:
                doc = Document()
                
                # 제목
                title = doc.add_heading("국회 의사중계 자막", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 생성 일시
                doc.add_paragraph(f"생성 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}")
                doc.add_paragraph()
                
                # 스레드 안전하게 자막 스냅샷 생성
                with self.subtitle_lock:
                    subtitles_snapshot = list(self.subtitles)
                
                # 자막 내용
                for entry in subtitles_snapshot:
                    timestamp = entry.timestamp.strftime('%H:%M:%S')
                    p = doc.add_paragraph()
                    run = p.add_run(f"[{timestamp}] ")
                    run.font.size = Pt(9)
                    run.font.color.rgb = None  # 기본 색상
                    p.add_run(entry.text)
                
                # 통계
                doc.add_paragraph()
                total_chars = sum(len(s.text) for s in subtitles_snapshot)
                doc.add_paragraph(f"총 {len(subtitles_snapshot)}문장, {total_chars:,}자")
                
                doc.save(path)
                QMessageBox.information(self, "성공", f"DOCX 저장 완료!\n\n파일: {path}")
            
            except Exception as e:
                QMessageBox.critical(self, "오류", f"DOCX 저장 실패: {e}")
    
    def _save_hwp(self):
        """HWP 파일로 저장 (Hancom Office 필요)"""
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        try:
            import win32com.client
        except ImportError:
            QMessageBox.warning(
                self, "라이브러리 필요",
                "HWP 저장을 위해 pywin32 라이브러리가 필요합니다.\n\n"
                "설치: pip install pywin32\n\n"
                "RTF 형식으로 저장을 시도합니다."
            )
            self._save_rtf()
            return
        
        hwp = None
        
        try:
            # win32com.client.dynamic.Dispatch 사용으로 캐시 문제 회피
            hwp = win32com.client.dynamic.Dispatch("HWPFrame.HwpObject")
            hwp.XHwpWindows.Item(0).Visible = True
            hwp.RegisterModule("FilePathCheckDLL", "SecurityModule")  # 보안 모듈 등록 시도
            
            # 새 문서 생성
            hwp.HAction.Run("FileNew")
            
            # 제목 입력
            hwp.HAction.Run("CharShapeBold")
            hwp.HAction.Run("ParagraphShapeAlignCenter")
            hwp.HAction.GetDefault("InsertText", hwp.HParameterSet.HInsertText.HSet)
            hwp.HParameterSet.HInsertText.Text = "국회 의사중계 자막\r\n"
            hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)
            
            hwp.HAction.Run("CharShapeBold")
            hwp.HAction.Run("ParagraphShapeAlignLeft")
            
            # 생성 일시
            hwp.HParameterSet.HInsertText.Text = f"생성 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}\r\n\r\n"
            hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)
            
            # 스레드 안전하게 자막 스냅샷 생성
            with self.subtitle_lock:
                subtitles_snapshot = list(self.subtitles)
            
            # 자막 내용
            for entry in subtitles_snapshot:
                timestamp = entry.timestamp.strftime('%H:%M:%S')
                hwp.HParameterSet.HInsertText.Text = f"[{timestamp}] {entry.text}\r\n"
                hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)
            
            # 통계
            total_chars = sum(len(s.text) for s in subtitles_snapshot)
            hwp.HParameterSet.HInsertText.Text = f"\r\n총 {len(subtitles_snapshot)}문장, {total_chars:,}자\r\n"
            hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)
            
            # 저장 대화상자
            filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.hwp"
            path, _ = QFileDialog.getSaveFileName(self, "HWP 저장", filename, "HWP 문서 (*.hwp)")
            
            if path:
                # FileSaveAs_S 액션 사용
                hwp.HAction.GetDefault("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
                hwp.HParameterSet.HFileOpenSave.filename = path
                hwp.HParameterSet.HFileOpenSave.Format = "HWP"
                hwp.HAction.Execute("FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet)
                
                QMessageBox.information(self, "성공", f"HWP 저장 완료!\n\n파일: {path}")
            
        except Exception as e:
            error_msg = str(e).lower()
            logger.error(f"HWP 저장 실패: {e}")
            
            # 권한 문제 힌트 제공
            if "access denied" in error_msg or "권한" in str(e):
                advice = "\n\n관리자 권한으로 실행하거나 한글 프로그램을 먼저 실행해 보세요."
            elif "server execution failed" in error_msg:
                advice = "\n\n한글 프로그램이 응답하지 않습니다. 한글을 종료하고 다시 시도하세요."
            else:
                advice = ""
                
            QMessageBox.warning(
                self, "HWP 저장 실패", 
                f"한글 파일 저장 중 오류가 발생했습니다: {e}{advice}\n\n"
                "대신 RTF 형식으로 저장을 시도합니다."
            )
            self._save_rtf()
            
        finally:
            if hwp:
                try:
                    hwp.Quit()
                except Exception:
                    pass
    
    def _rtf_encode(self, text: str) -> str:
        """유니코드 문자를 RTF 형식으로 인코딩 (특수문자 이스케이프)"""
        result = []
        for char in text:
            if char == '\\':
                result.append("\\\\")
            elif char == '{':
                result.append("\\{")
            elif char == '}':
                result.append("\\}")
            else:
                result.append(char)
        return ''.join(result)
    
    def _save_rtf(self):
        """RTF 파일로 저장 (HWP에서 열기 가능)"""
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.rtf"
        path, _ = QFileDialog.getSaveFileName(self, "RTF 저장", filename, "RTF 문서 (*.rtf)")
        
        if path:
            try:
                # 스레드 안전하게 자막 스냅샷 생성
                with self.subtitle_lock:
                    subtitles_snapshot = list(self.subtitles)
                
                with open(path, 'wb') as f:  # 바이너리 모드로 변경
                    # RTF 헤더 (유니코드 지원)
                    f.write(b"{\\rtf1\\ansi\\ansicpg949\\deff0")
                    f.write(b"{\\fonttbl{\\f0\\fnil\\fcharset129 \\'b8\\'c0\\'c0\\'ba \\'b0\\'ed\\'b5\\'f1;}}")
                    f.write(b"{\\colortbl;\\red0\\green0\\blue0;\\red128\\green128\\blue128;}")
                    f.write(b"\n")
                    
                    # 제목
                    title = self._rtf_encode("국회 의사중계 자막")
                    f.write(b"\\pard\\qc\\b\\fs28 ")
                    f.write(title.encode('cp949', errors='replace'))
                    f.write(b"\\b0\\par\n")
                    
                    # 생성 일시
                    date_str = self._rtf_encode(f"생성 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}")
                    f.write(b"\\pard\\ql\\fs20 ")
                    f.write(date_str.encode('cp949', errors='replace'))
                    f.write(b"\\par\\par\n")
                    
                    # 자막 내용
                    for entry in subtitles_snapshot:
                        timestamp = entry.timestamp.strftime('%H:%M:%S')
                        text = self._rtf_encode(entry.text)
                        f.write(b"\\cf2[")
                        f.write(timestamp.encode('cp949', errors='replace'))
                        f.write(b"]\\cf1 ")
                        f.write(text.encode('cp949', errors='replace'))
                        f.write(b"\\par\n")
                    
                    # 통계
                    total_chars = sum(len(s.text) for s in subtitles_snapshot)
                    stats = self._rtf_encode(f"총 {len(subtitles_snapshot)}문장, {total_chars:,}자")
                    f.write(b"\\par\\fs18 ")
                    f.write(stats.encode('cp949', errors='replace'))
                    f.write(b"\\par}")
                
                QMessageBox.information(self, "성공", f"RTF 저장 완료!\n\n파일: {path}\n\n이 파일은 한글(HWP)에서 열 수 있습니다.")
            
            except Exception as e:
                QMessageBox.critical(self, "오류", f"RTF 저장 실패: {e}")
    
    # ========== 세션 ==========
    
    def _save_session(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"{Config.SESSION_DIR}/세션_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        path, _ = QFileDialog.getSaveFileName(self, "세션 저장", filename, "JSON (*.json)")
        
        if path:
            try:
                # 디렉토리 존재 확인
                Path(path).parent.mkdir(parents=True, exist_ok=True)
                
                # 스레드 안전하게 자막 복사
                with self.subtitle_lock:
                    subtitles_copy = [s.to_dict() for s in self.subtitles]
                
                data = {
                    'version': Config.VERSION,
                    'created': datetime.now().isoformat(),
                    'url': self._get_current_url(),
                    'subtitles': subtitles_copy
                }
                
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                self._show_toast("세션 저장 완료!", "success")
            except Exception as e:
                logger.error(f"세션 저장 오류: {e}")
                QMessageBox.critical(self, "오류", f"세션 저장 실패: {e}")
    
    def _load_session(self):
        path, _ = QFileDialog.getOpenFileName(self, "세션 불러오기", f"{Config.SESSION_DIR}/", "JSON (*.json)")
        
        if path:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # 세션 버전 호환성 확인
                session_version = data.get('version', 'unknown')
                if session_version != Config.VERSION:
                    reply = QMessageBox.question(
                        self, "버전 불일치",
                        f"세션 버전({session_version})이 현재 버전({Config.VERSION})과 다릅니다.\n"
                        "계속 불러오시겠습니까?",
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                    )
                    if reply == QMessageBox.StandardButton.No:
                        return
                
                # 스레드 안전하게 자막 교체
                new_subtitles = [
                    SubtitleEntry.from_dict(item) 
                    for item in data.get('subtitles', [])
                ]
                with self.subtitle_lock:
                    self.subtitles = new_subtitles
                
                if data.get('url'):
                    self.url_combo.setCurrentText(data['url'])
                
                self._refresh_text()
                self._update_count_label()
                self._show_toast(f"세션 불러오기 완료! {len(self.subtitles)}개 문장", "success")
            
            except Exception as e:
                QMessageBox.critical(self, "오류", f"불러오기 실패: {e}")
    
    # ========== 유틸리티 ==========
    
    def _copy_to_clipboard(self):
        text = self._get_accumulated_text()
        if not text:
            self._show_toast("복사할 내용이 없습니다.", "warning")
            return
        
        QApplication.clipboard().setText(text)
        self._show_toast(f"클립보드에 복사되었습니다. ({len(text):,}자)", "success")
    
    def _clear_text(self):
        if not self.subtitles:
            return
        
        reply = QMessageBox.question(
            self, "확인", "모든 내용을 지우시겠습니까?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            with self.subtitle_lock:
                self.subtitles = []
            self.last_subtitle = ""
            self.subtitle_text.clear()
            self.status_label.setText("내용 삭제됨")
    
    def _edit_subtitle(self):
        """선택한 자막 편집"""
        if not self.subtitles:
            self._show_toast("편집할 자막이 없습니다.", "warning")
            return
        
        # 자막 목록 다이얼로그
        dialog = QDialog(self)
        dialog.setWindowTitle("자막 편집")
        dialog.setMinimumSize(600, 400)
        
        layout = QVBoxLayout(dialog)
        
        # 안내 라벨
        info_label = QLabel("편집할 자막을 선택하세요:")
        layout.addWidget(info_label)
        
        # 자막 목록
        list_widget = QListWidget()
        for i, entry in enumerate(self.subtitles):
            timestamp = entry.timestamp.strftime('%H:%M:%S')
            text_preview = entry.text[:50] + "..." if len(entry.text) > 50 else entry.text
            list_widget.addItem(f"[{timestamp}] {text_preview}")
        layout.addWidget(list_widget)
        
        # 편집 영역
        edit_label = QLabel("자막 내용:")
        layout.addWidget(edit_label)
        
        edit_text = QTextEdit()
        edit_text.setMaximumHeight(100)
        layout.addWidget(edit_text)
        
        # 선택 시 내용 로드
        def on_selection_changed():
            idx = list_widget.currentRow()
            if 0 <= idx < len(self.subtitles):
                edit_text.setText(self.subtitles[idx].text)
        
        list_widget.currentRowChanged.connect(on_selection_changed)
        
        # 버튼
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Save | QDialogButtonBox.StandardButton.Cancel
        )
        
        def save_edit():
            idx = list_widget.currentRow()
            if 0 <= idx < len(self.subtitles):
                new_text = edit_text.toPlainText().strip()
                if new_text:
                    with self.subtitle_lock:
                        self.subtitles[idx].text = new_text
                    self._refresh_text()
                    self._show_toast("자막이 수정되었습니다.", "success")
                    dialog.accept()
                else:
                    QMessageBox.warning(dialog, "알림", "자막 내용을 입력해주세요.")
            else:
                QMessageBox.warning(dialog, "알림", "편집할 자막을 선택해주세요.")
        
        buttons.accepted.connect(save_edit)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        
        # 첫 번째 항목 선택
        if list_widget.count() > 0:
            list_widget.setCurrentRow(0)
        
        dialog.exec()
    
    def _delete_subtitle(self):
        """선택한 자막 삭제"""
        if not self.subtitles:
            self._show_toast("삭제할 자막이 없습니다.", "warning")
            return
        
        # 자막 목록 다이얼로그
        dialog = QDialog(self)
        dialog.setWindowTitle("자막 삭제")
        dialog.setMinimumSize(600, 400)
        
        layout = QVBoxLayout(dialog)
        
        # 안내 라벨
        info_label = QLabel("삭제할 자막을 선택하세요 (다중 선택 가능):")
        layout.addWidget(info_label)
        
        # 자막 목록 (다중 선택 가능)
        list_widget = QListWidget()
        list_widget.setSelectionMode(QAbstractItemView.SelectionMode.MultiSelection)
        
        for i, entry in enumerate(self.subtitles):
            timestamp = entry.timestamp.strftime('%H:%M:%S')
            text_preview = entry.text[:50] + "..." if len(entry.text) > 50 else entry.text
            list_widget.addItem(f"[{timestamp}] {text_preview}")
        layout.addWidget(list_widget)
        
        # 버튼
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.button(QDialogButtonBox.StandardButton.Ok).setText("삭제")
        
        def delete_selected():
            selected_rows = sorted([i.row() for i in list_widget.selectedIndexes()], reverse=True)
            if not selected_rows:
                QMessageBox.warning(dialog, "알림", "삭제할 자막을 선택해주세요.")
                return
            
            reply = QMessageBox.question(
                dialog, "확인",
                f"선택한 {len(selected_rows)}개의 자막을 삭제하시겠습니까?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                with self.subtitle_lock:
                    for row in selected_rows:
                        del self.subtitles[row]
                self._refresh_text()
                self._update_count_label()
                self._show_toast(f"{len(selected_rows)}개 자막이 삭제되었습니다.", "success")
                dialog.accept()
        
        buttons.accepted.connect(delete_selected)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        
        dialog.exec()
    
    # ========== 도움말 ==========
    
    def _show_guide(self):
        """사용법 가이드 표시"""
        guide = """
<h2>🏛️ 사용법 가이드</h2>

<h3>📋 기본 사용법</h3>
<ol>
<li><b>URL 입력</b> - 국회 의사중계 페이지 URL을 입력합니다</li>
<li><b>선택자 확인</b> - 기본값을 사용하거나 수정합니다</li>
<li><b>옵션 설정</b>
    <ul>
    <li>자동 스크롤: 새 자막 자동 따라가기</li>
    <li>실시간 저장: 자막 실시간 파일 저장</li>
    <li>헤드리스 모드: 브라우저 창 숨기고 실행</li>
    </ul>
</li>
<li><b>시작</b> 버튼 클릭 (또는 F5)</li>
<li>자막 추출 완료 후 <b>파일 저장</b></li>
</ol>

<h3>⌨️ 주요 단축키</h3>
<table>
<tr><td><b>F5</b></td><td>시작</td></tr>
<tr><td><b>Escape</b></td><td>중지</td></tr>
<tr><td><b>Ctrl+F</b></td><td>검색</td></tr>
<tr><td><b>F3</b></td><td>다음 검색</td></tr>
<tr><td><b>Ctrl+T</b></td><td>테마 전환</td></tr>
<tr><td><b>Ctrl+S</b></td><td>TXT 저장</td></tr>
</table>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("사용법 가이드")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(guide)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()
    
    def _show_shortcuts(self):
        """키보드 단축키 목록 표시"""
        shortcuts = """
<h2>⌨️ 키보드 단축키</h2>

<h3>📋 기본 조작</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>단축키</th><th>기능</th></tr>
<tr><td><b>F5</b></td><td>추출 시작</td></tr>
<tr><td><b>Escape</b></td><td>추출 중지</td></tr>
<tr><td><b>Ctrl+Q</b></td><td>프로그램 종료</td></tr>
</table>

<h3>🔍 검색 및 편집</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>단축키</th><th>기능</th></tr>
<tr><td><b>Ctrl+F</b></td><td>검색창 열기</td></tr>
<tr><td><b>F3</b></td><td>다음 검색 결과</td></tr>
<tr><td><b>Shift+F3</b></td><td>이전 검색 결과</td></tr>
<tr><td><b>Ctrl+E</b></td><td>자막 편집</td></tr>
<tr><td><b>Delete</b></td><td>자막 삭제</td></tr>
<tr><td><b>Ctrl+C</b></td><td>클립보드 복사</td></tr>
</table>

<h3>💾 저장</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>단축키</th><th>기능</th></tr>
<tr><td><b>Ctrl+S</b></td><td>TXT 저장</td></tr>
<tr><td><b>Ctrl+Shift+S</b></td><td>세션 저장</td></tr>
<tr><td><b>Ctrl+O</b></td><td>세션 불러오기</td></tr>
</table>

<h3>🎨 보기</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>단축키</th><th>기능</th></tr>
<tr><td><b>Ctrl+T</b></td><td>테마 전환</td></tr>
<tr><td><b>Ctrl++</b></td><td>글자 크기 키우기</td></tr>
<tr><td><b>Ctrl+-</b></td><td>글자 크기 줄이기</td></tr>
<tr><td><b>F1</b></td><td>사용법 가이드</td></tr>
</table>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("키보드 단축키")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(shortcuts)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()
    
    def _show_features(self):
        """기능 소개 표시"""
        features = """
<h2>✨ 기능 소개</h2>

<h3>🎯 실시간 자막 추출</h3>
<p>국회 의사중계 웹사이트의 AI 자막을 실시간으로 캡처합니다.<br>
2초 동안 자막이 변경되지 않으면 자동으로 확정됩니다.</p>

<h3>💾 다양한 저장 형식</h3>
<ul>
<li><b>TXT</b> - 일반 텍스트</li>
<li><b>SRT</b> - 자막 파일 형식</li>
<li><b>VTT</b> - WebVTT 자막 형식</li>
<li><b>DOCX</b> - Word 문서</li>
</ul>

<h3>🔍 검색 및 하이라이트</h3>
<ul>
<li><b>실시간 검색</b> - Ctrl+F로 자막 내 텍스트 검색</li>
<li><b>키워드 하이라이트</b> - 특정 단어 강조</li>
</ul>

<h3>⚙️ 헤드리스 모드 (인터넷창 숨김)</h3>
<p>브라우저 창을 숨기고 백그라운드에서 실행합니다.<br>
자막 추출 중 다른 작업을 할 수 있습니다.</p>

<h3>📊 통계 패널</h3>
<p>실행 시간, 글자 수, 단어 수, 분당 글자 수를 표시합니다.</p>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("기능 소개")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(features)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()
    
    def _show_about(self):
        """프로그램 정보 표시"""
        about = f"""
<h2>🏛️ 국회 의사중계 자막 추출기</h2>
<p><b>버전:</b> {Config.VERSION}</p>
<p><b>설명:</b> 국회 의사중계 웹사이트에서 실시간 AI 자막을<br>
자동으로 추출하고 저장하는 프로그램입니다.</p>

<h3>📦 필요 라이브러리</h3>
<ul>
<li>PyQt6</li>
<li>selenium</li>
<li>python-docx (DOCX 저장용)</li>
</ul>

<p><b>© 2024-2025</b></p>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("정보")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(about)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()
    
    def closeEvent(self, event):
        # 트레이 최소화 모드
        if self.minimize_to_tray and self.tray_icon.isVisible():
            self.hide()
            self.tray_icon.showMessage(
                Config.APP_NAME,
                "프로그램이 트레이로 최소화되었습니다.\n트레이 아이콘을 더블클릭하여 다시 열 수 있습니다.",
                QSystemTrayIcon.MessageIcon.Information,
                2000
            )
            event.ignore()
            return
        
        # 저장하지 않은 자막이 있으면 확인
        if self.subtitles:
            reply = QMessageBox.question(
                self, "종료 확인",
                f"저장하지 않은 자막 {len(self.subtitles)}개가 있습니다.\n\n"
                "저장하시겠습니까?",
                QMessageBox.StandardButton.Save | 
                QMessageBox.StandardButton.Discard | 
                QMessageBox.StandardButton.Cancel
            )
            if reply == QMessageBox.StandardButton.Cancel:
                event.ignore()
                return
            elif reply == QMessageBox.StandardButton.Save:
                # 저장 대화상자에서 취소하면 종료도 취소
                filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                path, _ = QFileDialog.getSaveFileName(
                    self, "TXT 저장", filename, "텍스트 (*.txt)"
                )
                if not path:  # 사용자가 취소한 경우
                    event.ignore()
                    return
                # 파일 저장
                try:
                    with open(path, 'w', encoding='utf-8') as f:
                        for entry in self.subtitles:
                            f.write(f"[{entry.timestamp.strftime('%H:%M:%S')}] {entry.text}\n")
                except Exception as e:
                    QMessageBox.critical(self, "오류", f"저장 실패: {e}")
                    event.ignore()
                    return
        
        # 추출 중이면 확인
        if self.is_running:
            reply = QMessageBox.question(
                self, "종료", "추출 중입니다. 종료하시겠습니까?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                event.ignore()
                return
        
        # 창 위치/크기 저장
        self.settings.setValue("geometry", self.saveGeometry())
        self.settings.setValue("windowState", self.saveState())
        
        # 타이머 정리
        self.queue_timer.stop()
        self.stats_timer.stop()
        self.finalize_timer.stop()
        self.backup_timer.stop()
        
        # 워커 스레드 종료
        self.is_running = False
        self.stop_event.set()
        if self.worker and self.worker.is_alive():
            self.worker.join(timeout=Config.THREAD_STOP_TIMEOUT)
        
        if self.realtime_file:
            try:
                self.realtime_file.close()
            except Exception as e:
                logger.debug(f"파일 닫기 오류: {e}")
        if self.driver:
            try:
                self.driver.quit()
            except Exception as e:
                logger.debug(f"WebDriver 종료 오류: {e}")
        
        logger.info("프로그램 종료")
        event.accept()


def main():
    """메인 함수 - 예외 처리 강화"""
    try:
        # IDLE 호환성을 위한 이벤트 루프 체크
        app = QApplication.instance()
        if app is None:
            app = QApplication(sys.argv)
        
        app.setStyle("Fusion")
        app.setFont(QFont("맑은 고딕", 10))
        
        window = MainWindow()
        window.show()
        
        # IDLE에서 실행 시 exec() 대신 processEvents 사용
        if hasattr(sys, 'ps1'):  # IDLE/인터프리터 환경
            while window.isVisible():
                app.processEvents()
                time.sleep(0.01)
        else:
            sys.exit(app.exec())
    
    except Exception as e:
        logger.exception(f"프로그램 오류: {e}")
        QMessageBox.critical(None, "오류", f"프로그램 실행 중 오류 발생:\n{e}")


if __name__ == '__main__':
    main()
