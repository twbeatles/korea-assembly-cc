# -*- coding: utf-8 -*-

import os
import time
import threading
import queue
import re
import json
from urllib.request import Request, urlopen
from urllib.error import HTTPError, URLError
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from PyQt6.QtWidgets import (
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QLabel,
    QPushButton,
    QLineEdit,
    QTextEdit,
    QComboBox,
    QCheckBox,
    QApplication,
    QFrame,
    QProgressBar,
    QMessageBox,
    QFileDialog,
    QGroupBox,
    QGridLayout,
    QDialog,
    QDialogButtonBox,
    QListWidget,
    QSplitter,
    QMenu,
    QMenuBar,
    QInputDialog,
    QSystemTrayIcon,
    QAbstractItemView,
    QTreeWidget,
    QTreeWidgetItem,
    QToolBar,
    QSizePolicy,
)
from PyQt6.QtCore import Qt, QTimer, QSettings
from PyQt6.QtGui import (
    QFont,
    QColor,
    QTextCursor,
    QTextCharFormat,
    QAction,
    QShortcut,
    QKeySequence,
    QIcon,
)

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import (
    WebDriverException,
    NoSuchElementException,
    StaleElementReferenceException,
)
from selenium.webdriver.chrome.options import Options

from core.config import Config
from core.logging_utils import logger
from core.models import SubtitleEntry
from ui.dialogs import LiveBroadcastDialog
from ui.themes import DARK_THEME, LIGHT_THEME
from ui.widgets import CollapsibleGroupBox, ToastWidget
from core import utils
from core.subtitle_processor import SubtitleProcessor

try:
    from database import DatabaseManager

    DB_AVAILABLE = True
except ImportError:
    DB_AVAILABLE = False
    logger.warning(
        "database.py ëª¨ë“ˆì„ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤. ë°ì´í„°ë² ì´ìŠ¤ ê¸°ëŠ¥ì€ ë¹„í™œì„±í™”ë©ë‹ˆë‹¤."
    )


class RecoverableWebDriverError(RuntimeError):
    """ì›¹ë“œë¼ì´ë²„ ì¬ì—°ê²°ë¡œ ë³µêµ¬ ê°€ëŠ¥í•œ ì˜¤ë¥˜"""


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{Config.APP_NAME} v{Config.VERSION}")
        self.setMinimumSize(1100, 750)
        self.resize(1200, 800)

        # ì„¤ì •
        self.settings = QSettings("AssemblySubtitle", "Extractor")
        self.is_dark_theme = self.settings.value("dark_theme", True, type=bool)
        self.font_size = self.settings.value(
            "font_size", Config.DEFAULT_FONT_SIZE, type=int
        )
        self.minimize_to_tray = self.settings.value(
            "minimize_to_tray", False, type=bool
        )

        # ë©”ì‹œì§€ í
        self.message_queue = queue.Queue()

        # ìƒíƒœ
        self.worker = None
        self.driver = None
        self.is_running = False
        self.stop_event = threading.Event()  # ìŠ¤ë ˆë“œ ì•ˆì „í•œ ì¢…ë£Œ ì‹œê·¸ë„
        self.subtitle_lock = threading.Lock()  # ìë§‰ ë¦¬ìŠ¤íŠ¸ ì ‘ê·¼ ë™ê¸°í™”
        self._auto_backup_lock = threading.Lock()  # ìë™ ë°±ì—… ì¤‘ë³µ ì‹¤í–‰ ë°©ì§€
        self.start_time = None
        self.last_subtitle = ""
        self._last_raw_text = ""
        self._last_processed_raw = ""
        self._stream_start_time = None

        # [NEW] ê¸€ë¡œë²Œ íˆìŠ¤í† ë¦¬ + Suffix ë§¤ì¹­ (#GlobalHistorySuffix)
        self._confirmed_compact = ""  # í™•ì •ëœ ëª¨ë“  í…ìŠ¤íŠ¸ (compact, ê³µë°± ì œê±°)
        self._trailing_suffix = ""  # íˆìŠ¤í† ë¦¬ì˜ ë§ˆì§€ë§‰ Nì (suffix ë§¤ì¹­ìš©)
        self._suffix_length = 50  # suffix ê¸¸ì´
        self._preview_desync_count = 0
        self._preview_ambiguous_skip_count = 0
        self._last_good_raw_compact = ""
        self._preview_resync_threshold = 10
        self._preview_ambiguous_resync_threshold = 6

        # ìë§‰ ë°ì´í„° (íƒ€ì„ìŠ¤íƒ¬í”„ í¬í•¨)
        self.subtitles = []  # List[SubtitleEntry]

        # [Fix] ìë§‰ ì¤‘ë³µ ë°©ì§€ í”„ë¡œì„¸ì„œ (#SubtitleRepetitionFix)
        self.subtitle_processor = SubtitleProcessor()

        # ì €ì¥ëœ í‚¤ì›Œë“œ ë¡œë“œ
        saved_keywords = self.settings.value("highlight_keywords", "")
        if not saved_keywords:
            legacy_keywords = self.settings.value("keywords", "", type=str)
            if legacy_keywords:
                saved_keywords = legacy_keywords
                self.settings.setValue("highlight_keywords", legacy_keywords)
                self.settings.remove("keywords")
        self.keywords = (
            [k.strip() for k in saved_keywords.split(",") if k.strip()]
            if saved_keywords
            else []
        )
        saved_alert = self.settings.value("alert_keywords", "")
        self.alert_keywords = (
            [k.strip() for k in saved_alert.split(",") if k.strip()]
            if saved_alert
            else []
        )
        self._alert_keywords_cache = []
        self._rebuild_alert_keyword_cache(
            self.alert_keywords, update_settings=False
        )
        self.last_update_time = 0  # ë§ˆì§€ë§‰ ìë§‰ ì—…ë°ì´íŠ¸ ì‹œê°„

        # ì„±ëŠ¥ ìµœì í™”: QTextCharFormat ìºì‹±
        self._highlight_fmt = QTextCharFormat()
        self._highlight_fmt.setBackground(QColor("#ffd700"))  # ê³¨ë“œ ë°°ê²½
        self._highlight_fmt.setForeground(QColor("#000000"))  # ê²€ì • ê¸€ì
        self._highlight_fmt.setFontWeight(QFont.Weight.Bold)
        self._normal_fmt = QTextCharFormat()
        self._timestamp_fmt = QTextCharFormat()
        self._timestamp_fmt.setForeground(QColor("#888888"))

        # ì„±ëŠ¥ ìµœì í™”: í‚¤ì›Œë“œ íŒ¨í„´ ìºì‹±
        self._keyword_pattern = None  # ì»´íŒŒì¼ëœ ì •ê·œì‹ íŒ¨í„´
        self._keywords_lower_set = set()  # ë¹ ë¥¸ ê²€ìƒ‰ìš© set
        self._rebuild_keyword_cache(self.keywords, update_settings=False, refresh=False)

        # ì„±ëŠ¥ ìµœì í™”: í†µê³„ ìºì‹±
        self._cached_total_chars = 0
        self._cached_total_words = 0

        # ë Œë”ë§ ìƒíƒœ ìºì‹±
        self._last_rendered_count = 0
        self._last_rendered_last_text = ""
        self._last_render_offset = 0
        self._last_render_show_ts = None

        # í† ìŠ¤íŠ¸ ìŠ¤íƒ ê´€ë¦¬
        self.active_toasts = []  # í˜„ì¬ í‘œì‹œ ì¤‘ì¸ í† ìŠ¤íŠ¸ ëª©ë¡

        # ì‹¤ì‹œê°„ ì €ì¥
        self.realtime_file = None
        self._realtime_error_count = 0  # ì‹¤ì‹œê°„ ì €ì¥ ì—°ì† ì˜¤ë¥˜ ì¹´ìš´íŠ¸ (#3)

        # ì¤‘ì§€ ì‹œ ë¸Œë¼ìš°ì € ìœ ì§€ìš© (ì¢…ë£Œ ì‹œ ì •ë¦¬)
        self._detached_drivers = []
        self._detached_drivers_lock = threading.Lock()
        self._last_subtitle_frame_path = ()

        # ì—°ê²° ìƒíƒœ ëª¨ë‹ˆí„°ë§ (#30)
        self.connection_status = "disconnected"  # connected, disconnected, reconnecting
        self.last_ping_time = 0
        self.ping_latency = 0  # ms

        # ìë™ ì¬ì—°ê²° (#31)
        self.reconnect_attempts = 0
        self.auto_reconnect_enabled = self.settings.value(
            "auto_reconnect", True, type=bool
        )
        self.current_url = ""  # í˜„ì¬ ì—°ê²° ì¤‘ì¸ URL ì €ì¥

        # ìŠ¤ë§ˆíŠ¸ ìŠ¤í¬ë¡¤ ìƒíƒœ (ì‚¬ìš©ìê°€ ìœ„ë¡œ ìŠ¤í¬ë¡¤í•˜ë©´ ìë™ ìŠ¤í¬ë¡¤ ì¼ì‹œ ì¤‘ì§€)
        self._user_scrolled_up = False
        self._is_stopping = False
        self._last_status_message = ""
        self._session_save_in_progress = False
        self._session_load_in_progress = False
        self._db_history_dialog_state = None

        # URL íˆìŠ¤í† ë¦¬
        self.url_history = self._load_url_history()

        # UI ìƒì„±
        self._create_menu()
        self._create_ui()
        self._apply_theme()
        self._setup_shortcuts()

        # íƒ€ì´ë¨¸
        self.queue_timer = QTimer(self)
        self.queue_timer.timeout.connect(self._process_message_queue)
        self.queue_timer.start(Config.QUEUE_PROCESS_INTERVAL)

        self.stats_timer = QTimer(self)
        self.stats_timer.timeout.connect(self._update_stats)

        # ìë§‰ í™•ì • íƒ€ì´ë¨¸
        self.finalize_timer = QTimer(self)
        self.finalize_timer.timeout.connect(self._check_finalize)

        # ìë™ ë°±ì—… íƒ€ì´ë¨¸
        self.backup_timer = QTimer(self)
        self.backup_timer.timeout.connect(self._auto_backup)

        # ë””ë ‰í† ë¦¬ ìƒì„±
        Path(Config.SESSION_DIR).mkdir(exist_ok=True)
        Path(Config.REALTIME_DIR).mkdir(exist_ok=True)
        Path(Config.BACKUP_DIR).mkdir(exist_ok=True)

        # ë°ì´í„°ë² ì´ìŠ¤ ë§¤ë‹ˆì € ì´ˆê¸°í™” (#26)
        self.db = None
        self._db_tasks_inflight = set()
        if DB_AVAILABLE:
            try:
                self.db = DatabaseManager(Config.DATABASE_PATH)
            except Exception as e:
                logger.error(f"ë°ì´í„°ë² ì´ìŠ¤ ì´ˆê¸°í™” ì‹¤íŒ¨: {e}")

        # ì°½ ìœ„ì¹˜/í¬ê¸° ë³µì›
        geometry = self.settings.value("geometry")
        if geometry:
            self.restoreGeometry(geometry)
        state = self.settings.value("windowState")
        if state:
            self.restoreState(state)

        # ì‹œìŠ¤í…œ íŠ¸ë ˆì´ ì„¤ì •
        self._setup_tray()

    def _rebuild_stats_cache(self) -> None:
        """í˜„ì¬ ìë§‰ ë¦¬ìŠ¤íŠ¸ ê¸°ë°˜ìœ¼ë¡œ í†µê³„ ìºì‹œë¥¼ ì¬ê³„ì‚°.

        - ì„¸ì…˜ ë¡œë“œ/ë³‘í•©/ì‚­ì œ/í¸ì§‘ ë“±ìœ¼ë¡œ ìë§‰ ë¦¬ìŠ¤íŠ¸ê°€ ë°”ë€ŒëŠ” ê²½ìš° í˜¸ì¶œí•œë‹¤.
        - ìºì‹œëŠ” UI í†µê³„/ì¹´ìš´íŠ¸ ë¼ë²¨ì—ì„œ ì¬ê³„ì‚° ë¹„ìš©ì„ ì¤„ì´ê¸° ìœ„í•´ ì‚¬ìš©ëœë‹¤.
        """
        with self.subtitle_lock:
            self._cached_total_chars = sum(s.char_count for s in self.subtitles)
            self._cached_total_words = sum(s.word_count for s in self.subtitles)

    def _setup_tray(self):
        """ì‹œìŠ¤í…œ íŠ¸ë ˆì´ ì•„ì´ì½˜ ì„¤ì •"""
        self.tray_icon = QSystemTrayIcon(self)
        self.tray_icon.setToolTip(f"{Config.APP_NAME} v{Config.VERSION}")

        # ê¸°ë³¸ ì•„ì´ì½˜ (ì•± ì•„ì´ì½˜ì´ ì—†ìœ¼ë©´ ê¸°ë³¸ ì•„ì´ì½˜ ì‚¬ìš©)
        app_icon = self.windowIcon()
        if not app_icon.isNull():
            self.tray_icon.setIcon(app_icon)
        else:
            # ê¸°ë³¸ ì•„ì´ì½˜ ì‚¬ìš©
            self.tray_icon.setIcon(
                self.style().standardIcon(self.style().StandardPixmap.SP_ComputerIcon)
            )

        # íŠ¸ë ˆì´ ë©”ë‰´
        tray_menu = QMenu()

        show_action = QAction("ğŸ›ï¸ ì°½ ë³´ì´ê¸°", self)
        show_action.triggered.connect(self._show_from_tray)
        tray_menu.addAction(show_action)

        tray_menu.addSeparator()

        # ì¶”ì¶œ ìƒíƒœ í‘œì‹œ
        self.tray_status_action = QAction("âšª ëŒ€ê¸° ì¤‘", self)
        self.tray_status_action.setEnabled(False)
        tray_menu.addAction(self.tray_status_action)

        tray_menu.addSeparator()

        start_action = QAction("â–¶ ì‹œì‘", self)
        start_action.triggered.connect(self._start)
        tray_menu.addAction(start_action)

        stop_action = QAction("â¹ ì¤‘ì§€", self)
        stop_action.triggered.connect(self._stop)
        tray_menu.addAction(stop_action)

        tray_menu.addSeparator()

        quit_action = QAction("âŒ ì¢…ë£Œ", self)
        quit_action.triggered.connect(self._quit_from_tray)
        tray_menu.addAction(quit_action)

        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.activated.connect(self._tray_activated)
        self.tray_icon.show()

    def _show_from_tray(self):
        """íŠ¸ë ˆì´ì—ì„œ ì°½ ë³µì›"""
        self.showNormal()
        self.activateWindow()
        self.raise_()

    def _quit_from_tray(self):
        """íŠ¸ë ˆì´ì—ì„œ ì™„ì „ ì¢…ë£Œ"""
        self.minimize_to_tray = False  # íŠ¸ë ˆì´ ìµœì†Œí™” ë¹„í™œì„±í™”
        self.close()

    def _tray_activated(self, reason):
        """íŠ¸ë ˆì´ ì•„ì´ì½˜ í´ë¦­ ì²˜ë¦¬"""
        if reason == QSystemTrayIcon.ActivationReason.DoubleClick:
            self._show_from_tray()

    def _update_tray_status(self, status: str):
        """íŠ¸ë ˆì´ ìƒíƒœ ì—…ë°ì´íŠ¸"""
        if hasattr(self, "tray_status_action"):
            self.tray_status_action.setText(status)

    def _create_menu(self):
        menubar = self.menuBar()

        # íŒŒì¼ ë©”ë‰´
        file_menu = menubar.addMenu("íŒŒì¼")

        save_txt = QAction("TXT ì €ì¥", self)
        save_txt.setShortcut("Ctrl+S")
        save_txt.triggered.connect(self._save_txt)
        file_menu.addAction(save_txt)

        save_srt = QAction("SRT ì €ì¥", self)
        save_srt.triggered.connect(self._save_srt)
        file_menu.addAction(save_srt)

        save_vtt = QAction("VTT ì €ì¥", self)
        save_vtt.triggered.connect(self._save_vtt)
        file_menu.addAction(save_vtt)

        save_docx = QAction("DOCX ì €ì¥ (Word)", self)
        save_docx.triggered.connect(self._save_docx)
        file_menu.addAction(save_docx)

        save_hwp = QAction("HWP ì €ì¥ (í•œê¸€)", self)
        save_hwp.triggered.connect(self._save_hwp)
        file_menu.addAction(save_hwp)

        save_rtf = QAction("RTF ì €ì¥", self)
        save_rtf.triggered.connect(self._save_rtf)
        file_menu.addAction(save_rtf)

        file_menu.addSeparator()

        save_session = QAction("ì„¸ì…˜ ì €ì¥", self)
        save_session.setShortcut("Ctrl+Shift+S")
        save_session.triggered.connect(self._save_session)
        file_menu.addAction(save_session)

        load_session = QAction("ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸°", self)
        load_session.setShortcut("Ctrl+O")
        load_session.triggered.connect(self._load_session)
        file_menu.addAction(load_session)

        file_menu.addSeparator()

        export_stats = QAction("ğŸ“Š í†µê³„ ë‚´ë³´ë‚´ê¸°", self)
        export_stats.triggered.connect(self._export_stats)
        file_menu.addAction(export_stats)

        file_menu.addSeparator()

        exit_action = QAction("ì¢…ë£Œ", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # í¸ì§‘ ë©”ë‰´
        edit_menu = menubar.addMenu("í¸ì§‘")

        search_action = QAction("ê²€ìƒ‰", self)
        search_action.setShortcut("Ctrl+F")
        search_action.triggered.connect(self._show_search)
        edit_menu.addAction(search_action)

        keyword_action = QAction("í•˜ì´ë¼ì´íŠ¸ í‚¤ì›Œë“œ ì„¤ì •", self)
        keyword_action.triggered.connect(self._set_keywords)
        edit_menu.addAction(keyword_action)

        alert_keyword_action = QAction("ì•Œë¦¼ í‚¤ì›Œë“œ ì„¤ì •", self)
        alert_keyword_action.setToolTip("íŠ¹ì • í‚¤ì›Œë“œ ê°ì§€ ì‹œ ì•Œë¦¼ì„ í‘œì‹œí•©ë‹ˆë‹¤")
        alert_keyword_action.triggered.connect(self._set_alert_keywords)
        edit_menu.addAction(alert_keyword_action)

        edit_menu.addSeparator()

        edit_subtitle_action = QAction("âœï¸ ìë§‰ í¸ì§‘", self)
        edit_subtitle_action.setShortcut("Ctrl+E")
        edit_subtitle_action.setToolTip("ì„ íƒí•œ ìë§‰ì„ í¸ì§‘í•©ë‹ˆë‹¤")
        edit_subtitle_action.triggered.connect(self._edit_subtitle)
        edit_menu.addAction(edit_subtitle_action)

        delete_subtitle_action = QAction("ğŸ—‘ï¸ ìë§‰ ì‚­ì œ", self)
        delete_subtitle_action.setShortcut("Delete")
        delete_subtitle_action.setToolTip("ì„ íƒí•œ ìë§‰ì„ ì‚­ì œí•©ë‹ˆë‹¤")
        delete_subtitle_action.triggered.connect(self._delete_subtitle)
        edit_menu.addAction(delete_subtitle_action)

        edit_menu.addSeparator()

        copy_action = QAction("í´ë¦½ë³´ë“œ ë³µì‚¬", self)
        # copy_action.setShortcut("Ctrl+C") # í…ìŠ¤íŠ¸ ì„ íƒ ë³µì‚¬ ì¶©ëŒ ë°©ì§€
        copy_action.setToolTip("ì „ì²´ ìë§‰ì„ í´ë¦½ë³´ë“œì— ë³µì‚¬í•©ë‹ˆë‹¤")
        copy_action.triggered.connect(self._copy_to_clipboard)
        edit_menu.addAction(copy_action)

        clear_action = QAction("ë‚´ìš© ì§€ìš°ê¸°", self)
        clear_action.setToolTip("ëª¨ë“  ìë§‰ ë‚´ìš©ì„ ì‚­ì œí•©ë‹ˆë‹¤")
        clear_action.triggered.connect(self._clear_text)
        edit_menu.addAction(clear_action)

        # ë³´ê¸° ë©”ë‰´
        view_menu = menubar.addMenu("ë³´ê¸°")

        self.theme_action = QAction(
            "ë¼ì´íŠ¸ í…Œë§ˆ" if self.is_dark_theme else "ë‹¤í¬ í…Œë§ˆ", self
        )
        self.theme_action.setShortcut("Ctrl+T")
        self.theme_action.triggered.connect(self._toggle_theme)
        view_menu.addAction(self.theme_action)

        self.timestamp_action = QAction("íƒ€ì„ìŠ¤íƒ¬í”„ í‘œì‹œ", self)
        self.timestamp_action.setCheckable(True)
        self.timestamp_action.setChecked(True)
        self.timestamp_action.triggered.connect(self._refresh_text_full)
        view_menu.addAction(self.timestamp_action)

        self.tray_action = QAction("ğŸ”½ íŠ¸ë ˆì´ë¡œ ìµœì†Œí™”", self)
        self.tray_action.setCheckable(True)
        self.tray_action.setChecked(self.minimize_to_tray)
        self.tray_action.setToolTip("í™œì„±í™” ì‹œ ì°½ì„ ë‹«ìœ¼ë©´ íŠ¸ë ˆì´ë¡œ ìµœì†Œí™”ë©ë‹ˆë‹¤")
        self.tray_action.triggered.connect(self._toggle_tray_option)
        view_menu.addAction(self.tray_action)

        view_menu.addSeparator()

        # ê¸€ì í¬ê¸° ì„œë¸Œë©”ë‰´
        font_menu = view_menu.addMenu("ğŸ“ ê¸€ì í¬ê¸°")
        for size in [12, 14, 16, 18, 20, 22, 24]:
            font_action = QAction(f"{size}pt", self)
            font_action.triggered.connect(
                lambda checked, s=size: self._set_font_size(s)
            )
            font_menu.addAction(font_action)

        view_menu.addSeparator()

        font_increase = QAction("ê¸€ì í¬ê¸° í‚¤ìš°ê¸°", self)
        font_increase.setShortcut("Ctrl++")
        font_increase.triggered.connect(lambda: self._adjust_font_size(2))
        view_menu.addAction(font_increase)

        font_decrease = QAction("ê¸€ì í¬ê¸° ì¤„ì´ê¸°", self)
        font_decrease.setShortcut("Ctrl+-")
        font_decrease.triggered.connect(lambda: self._adjust_font_size(-2))
        view_menu.addAction(font_decrease)

        # ë„êµ¬ ë©”ë‰´ (#20)
        tools_menu = menubar.addMenu("ë„êµ¬")

        merge_action = QAction("ğŸ“ ìë§‰ ë³‘í•©...", self)
        merge_action.setShortcut("Ctrl+Shift+M")
        merge_action.setToolTip("ì—¬ëŸ¬ ì„¸ì…˜ íŒŒì¼ì„ í•˜ë‚˜ë¡œ ë³‘í•©í•©ë‹ˆë‹¤")
        merge_action.triggered.connect(self._show_merge_dialog)
        tools_menu.addAction(merge_action)

        # ì¤„ë„˜ê¹€ ì •ë¦¬
        clean_newlines_action = QAction("ì¤„ë„˜ê¹€ ì •ë¦¬", self)
        clean_newlines_action.setShortcut("Ctrl+Shift+L")
        clean_newlines_action.setToolTip(
            "ë¬¸ì¥ ë¶€í˜¸ë¡œ ëë‚˜ì§€ ì•ŠëŠ” ì¤„ì„ ë³‘í•©í•˜ì—¬ ë¬¸ì¥ì„ ì •ë¦¬í•©ë‹ˆë‹¤"
        )
        clean_newlines_action.triggered.connect(self._clean_newlines)
        tools_menu.addAction(clean_newlines_action)

        # ë°ì´í„°ë² ì´ìŠ¤ ë©”ë‰´ (#26)
        db_menu = menubar.addMenu("ë°ì´í„°ë² ì´ìŠ¤")

        db_history_action = QAction("ğŸ“‹ ì„¸ì…˜ íˆìŠ¤í† ë¦¬", self)
        db_history_action.setToolTip("ì €ì¥ëœ ëª¨ë“  ì„¸ì…˜ ëª©ë¡ì„ í™•ì¸í•©ë‹ˆë‹¤")
        db_history_action.triggered.connect(self._show_db_history)
        db_menu.addAction(db_history_action)

        db_search_action = QAction("ğŸ” ìë§‰ ê²€ìƒ‰", self)
        db_search_action.setToolTip("ëª¨ë“  ì„¸ì…˜ì—ì„œ í‚¤ì›Œë“œë¥¼ ê²€ìƒ‰í•©ë‹ˆë‹¤")
        db_search_action.triggered.connect(self._show_db_search)
        db_menu.addAction(db_search_action)

        db_menu.addSeparator()

        db_stats_action = QAction("ğŸ“Š ì „ì²´ í†µê³„", self)
        db_stats_action.setToolTip("ë°ì´í„°ë² ì´ìŠ¤ ì „ì²´ í†µê³„ë¥¼ í™•ì¸í•©ë‹ˆë‹¤")
        db_stats_action.triggered.connect(self._show_db_stats)
        db_menu.addAction(db_stats_action)

        # ë„ì›€ë§ ë©”ë‰´
        help_menu = menubar.addMenu("ë„ì›€ë§")

        guide_action = QAction("ì‚¬ìš©ë²• ê°€ì´ë“œ", self)
        guide_action.setShortcut("F1")
        guide_action.setToolTip("í”„ë¡œê·¸ë¨ ì‚¬ìš©ë²• ì•ˆë‚´")
        guide_action.triggered.connect(self._show_guide)
        help_menu.addAction(guide_action)

        features_action = QAction("ê¸°ëŠ¥ ì†Œê°œ", self)
        features_action.setToolTip("í”„ë¡œê·¸ë¨ì˜ ì£¼ìš” ê¸°ëŠ¥ ì†Œê°œ")
        features_action.triggered.connect(self._show_features)
        help_menu.addAction(features_action)

        shortcuts_action = QAction("âŒ¨ï¸ í‚¤ë³´ë“œ ë‹¨ì¶•í‚¤", self)
        shortcuts_action.setToolTip("ì‚¬ìš© ê°€ëŠ¥í•œ í‚¤ë³´ë“œ ë‹¨ì¶•í‚¤ ëª©ë¡")
        shortcuts_action.triggered.connect(self._show_shortcuts)
        help_menu.addAction(shortcuts_action)

        help_menu.addSeparator()

        about_action = QAction("ì •ë³´", self)
        about_action.setToolTip("í”„ë¡œê·¸ë¨ ì •ë³´ ë° ë²„ì „")
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)

    def _create_ui(self):
        central = QWidget()
        self.setCentralWidget(central)

        layout = QVBoxLayout(central)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # === í—¤ë” ===
        header = QLabel("ğŸ›ï¸ êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ìë§‰ ì¶”ì¶œê¸°")
        header.setObjectName("headerLabel")
        header.setFont(QFont("ë§‘ì€ ê³ ë”•", 18, QFont.Weight.Bold))
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        # === ìƒë‹¨ ì˜ì—­ ì»¨í…Œì´ë„ˆ ===
        self.top_header_container = QWidget()
        top_header_layout = QVBoxLayout(self.top_header_container)
        top_header_layout.setContentsMargins(0, 0, 0, 0)
        top_header_layout.setSpacing(15)
        top_header_layout.addWidget(header)

        # === í€µ ì•¡ì…˜ íˆ´ë°” ===
        toolbar_frame = QFrame()
        toolbar_frame.setStyleSheet("""
            QFrame {
                background-color: rgba(88, 166, 255, 0.05);
                border: 1px solid rgba(88, 166, 255, 0.15);
                border-radius: 10px;
                padding: 4px;
            }
        """)
        toolbar_layout = QHBoxLayout(toolbar_frame)
        toolbar_layout.setContentsMargins(12, 8, 12, 8)
        toolbar_layout.setSpacing(8)

        # íˆ´ë°” ë²„íŠ¼ ìŠ¤íƒ€ì¼
        toolbar_btn_style = """
            QPushButton {
                background-color: rgba(88, 166, 255, 0.1);
                border: 1px solid rgba(88, 166, 255, 0.2);
                border-radius: 8px;
                padding: 8px 16px;
                font-size: 12px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: rgba(88, 166, 255, 0.2);
                border-color: rgba(88, 166, 255, 0.4);
            }
            QPushButton:pressed {
                background-color: rgba(88, 166, 255, 0.3);
            }
        """

        # ë¹ ë¥¸ ì €ì¥ ë²„íŠ¼
        quick_save_btn = QPushButton("ğŸ’¾ ë¹ ë¥¸ ì €ì¥")
        quick_save_btn.setStyleSheet(toolbar_btn_style)
        quick_save_btn.setToolTip("í˜„ì¬ ìë§‰ì„ TXTë¡œ ë¹ ë¥´ê²Œ ì €ì¥ (Ctrl+S)")
        quick_save_btn.clicked.connect(self._save_txt)

        # ê²€ìƒ‰ ë²„íŠ¼
        search_btn = QPushButton("ğŸ” ê²€ìƒ‰")
        search_btn.setStyleSheet(toolbar_btn_style)
        search_btn.setToolTip("ìë§‰ ë‚´ í‚¤ì›Œë“œ ê²€ìƒ‰ (Ctrl+F)")
        search_btn.clicked.connect(self._show_search)

        # í´ë¦½ë³´ë“œ ë³µì‚¬ ë²„íŠ¼
        copy_btn = QPushButton("ğŸ“‹ ë³µì‚¬")
        copy_btn.setStyleSheet(toolbar_btn_style)
        copy_btn.setToolTip("ì „ì²´ ìë§‰ì„ í´ë¦½ë³´ë“œì— ë³µì‚¬")
        copy_btn.clicked.connect(self._copy_to_clipboard)

        # ìë§‰ ì¤„ë„˜ê¹€ ì •ë¦¬ ë²„íŠ¼ (New)
        clean_btn = QPushButton("âœ¨ ì¤„ë„˜ê¹€ ì •ë¦¬")
        clean_btn.setStyleSheet(toolbar_btn_style)
        clean_btn.setToolTip(
            "ë¬¸ì¥ ë¶€í˜¸ë¡œ ëŠì–´ì§€ì§€ ì•Šì€ ì¤„ì„ ìë™ìœ¼ë¡œ ë³‘í•© (Ctrl+Shift+L)"
        )
        clean_btn.clicked.connect(self._clean_newlines)

        # ìë§‰ ì§€ìš°ê¸° ë²„íŠ¼
        clear_btn = QPushButton("ğŸ—‘ï¸ ì§€ìš°ê¸°")
        clear_btn.setStyleSheet(toolbar_btn_style)
        clear_btn.setToolTip("í˜„ì¬ ìë§‰ ëª©ë¡ ì´ˆê¸°í™”")
        clear_btn.clicked.connect(self._clear_subtitles)

        toolbar_layout.addWidget(quick_save_btn)
        toolbar_layout.addWidget(search_btn)
        toolbar_layout.addWidget(copy_btn)
        toolbar_layout.addWidget(clean_btn)
        toolbar_layout.addWidget(clear_btn)
        toolbar_layout.addStretch()

        # í…Œë§ˆ í† ê¸€ ë²„íŠ¼
        self.theme_toggle_btn = QPushButton("ğŸŒ™" if self.is_dark_theme else "â˜€ï¸")
        self.theme_toggle_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                font-size: 18px;
                padding: 4px 8px;
            }
            QPushButton:hover {
                background-color: rgba(88, 166, 255, 0.1);
                border-radius: 6px;
            }
        """)
        self.theme_toggle_btn.setToolTip("í…Œë§ˆ ì „í™˜")
        self.theme_toggle_btn.clicked.connect(self._toggle_theme_from_button)
        toolbar_layout.addWidget(self.theme_toggle_btn)

        top_header_layout.addWidget(toolbar_frame)
        layout.addWidget(self.top_header_container)

        # === URL/ì„¤ì • ì˜ì—­ (ì ‘ê¸°/í¼ì¹˜ê¸° ê°€ëŠ¥) ===
        self.settings_group = CollapsibleGroupBox("âš™ï¸ ì„¤ì •")
        settings_layout = QGridLayout(self.settings_group)
        settings_layout.setSpacing(10)

        # URL
        url_label = QLabel("ğŸ“Œ URL:")
        url_label.setToolTip("êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ì›¹ì‚¬ì´íŠ¸ URLì„ ì…ë ¥í•˜ì„¸ìš”")
        settings_layout.addWidget(url_label, 0, 0)

        url_layout = QHBoxLayout()
        self.url_combo = QComboBox()
        self.url_combo.setEditable(True)
        self.url_combo.setToolTip(
            "êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ì›¹ì‚¬ì´íŠ¸ URL\nìµœê·¼ ì‚¬ìš©í•œ URLì´ ìë™ ì €ì¥ë©ë‹ˆë‹¤\níƒœê·¸ê°€ ìˆìœ¼ë©´ [íƒœê·¸] í˜•íƒœë¡œ í‘œì‹œë©ë‹ˆë‹¤"
        )

        # URL íˆìŠ¤í† ë¦¬ ë¡œë“œ ë° ì½¤ë³´ë°•ìŠ¤ ì´ˆê¸°í™”
        for url, tag in self.url_history.items():
            if tag:
                self.url_combo.addItem(f"[{tag}] {url}", url)
            else:
                self.url_combo.addItem(url, url)

        if not self.url_history:
            self.url_combo.addItem("https://assembly.webcast.go.kr/main/player.asp")

        url_layout.addWidget(self.url_combo, 1)

        # íƒœê·¸ ë²„íŠ¼
        self.tag_btn = QPushButton("ğŸ·ï¸ íƒœê·¸")
        self.tag_btn.setToolTip("í˜„ì¬ URLì— íƒœê·¸ ì¶”ê°€/í¸ì§‘\nì˜ˆ: ë³¸íšŒì˜, ë²•ì‚¬ìœ„, ìƒì„ìœ„")
        self.tag_btn.setFixedWidth(90)
        self.tag_btn.clicked.connect(self._edit_url_tag)
        url_layout.addWidget(self.tag_btn)

        # ìƒì„ìœ„ì›íšŒ í”„ë¦¬ì…‹ ë²„íŠ¼
        self.preset_btn = QPushButton("ğŸ“‹ ìƒì„ìœ„")
        self.preset_btn.setToolTip("ìƒì„ìœ„ì›íšŒ í”„ë¦¬ì…‹ ì„ íƒ\në¹ ë¥¸ URL ì…ë ¥ì„ ìœ„í•œ ê¸°ëŠ¥")
        self.preset_btn.setFixedWidth(120)

        # í”„ë¦¬ì…‹ ë©”ë‰´ ìƒì„±
        self.preset_menu = QMenu(self)
        self._load_committee_presets()
        self._build_preset_menu()
        self.preset_btn.setMenu(self.preset_menu)
        url_layout.addWidget(self.preset_btn)

        settings_layout.addLayout(url_layout, 0, 1)

        # ìƒì¤‘ê³„ ëª©ë¡ ë²„íŠ¼ (ì¶”ê°€)
        self.live_btn = QPushButton("ğŸ“¡ ìƒì¤‘ê³„ ëª©ë¡")
        self.live_btn.setToolTip("í˜„ì¬ ì§„í–‰ ì¤‘ì¸ ìƒì¤‘ê³„ ëª©ë¡ì„ í™•ì¸í•˜ê³  ì„ íƒí•©ë‹ˆë‹¤")
        self.live_btn.setFixedWidth(140)
        self.live_btn.clicked.connect(self._show_live_dialog)
        url_layout.addWidget(self.live_btn)

        # ì„ íƒì
        selector_label = QLabel("ğŸ” ì„ íƒì:")
        selector_label.setToolTip("ìë§‰ ìš”ì†Œì˜ CSS ì„ íƒì")
        settings_layout.addWidget(selector_label, 1, 0)
        self.selector_combo = QComboBox()
        self.selector_combo.setEditable(True)
        self.selector_combo.addItems(Config.DEFAULT_SELECTORS)
        self.selector_combo.setToolTip(
            "ìë§‰ í…ìŠ¤íŠ¸ê°€ í‘œì‹œë˜ëŠ” HTML ìš”ì†Œì˜ CSS ì„ íƒì\nê¸°ë³¸ê°’ì„ ì‚¬ìš©í•˜ê±°ë‚˜ ì§ì ‘ ì…ë ¥í•˜ì„¸ìš”"
        )
        settings_layout.addWidget(self.selector_combo, 1, 1)

        # ì˜µì…˜
        options_layout = QHBoxLayout()
        options_layout.setSpacing(20)

        self.auto_scroll_check = QCheckBox("ğŸ“œ ìë™ ìŠ¤í¬ë¡¤")
        self.auto_scroll_check.setChecked(True)
        self.auto_scroll_check.setToolTip(
            "ìƒˆ ìë§‰ì´ ì¶”ê°€ë  ë•Œ ìë™ìœ¼ë¡œ ë§¨ ì•„ë˜ë¡œ ìŠ¤í¬ë¡¤í•©ë‹ˆë‹¤"
        )

        self.realtime_save_check = QCheckBox("ğŸ’¾ ì‹¤ì‹œê°„ ì €ì¥")
        self.realtime_save_check.setChecked(False)
        self.realtime_save_check.setToolTip(
            "ìë§‰ì´ í™•ì •ë  ë•Œë§ˆë‹¤ ìë™ìœ¼ë¡œ íŒŒì¼ì— ì €ì¥í•©ë‹ˆë‹¤\nì €ì¥ ìœ„ì¹˜: realtime_output í´ë”"
        )

        self.headless_check = QCheckBox("ğŸ”‡ í—¤ë“œë¦¬ìŠ¤ ëª¨ë“œ (ì¸í„°ë„·ì°½ ìˆ¨ê¹€)")
        self.headless_check.setChecked(False)
        self.headless_check.setToolTip(
            "Chrome ë¸Œë¼ìš°ì € ì°½ì„ ìˆ¨ê¸°ê³  ë°±ê·¸ë¼ìš´ë“œì—ì„œ ì‹¤í–‰í•©ë‹ˆë‹¤.\nìë§‰ ì¶”ì¶œ ì¤‘ ë‹¤ë¥¸ ì‘ì—…ì„ í•  ìˆ˜ ìˆìŠµë‹ˆë‹¤."
        )
        options_layout.addWidget(self.auto_scroll_check)
        options_layout.addWidget(self.realtime_save_check)
        options_layout.addWidget(self.headless_check)
        options_layout.addStretch()
        settings_layout.addLayout(options_layout, 2, 0, 1, 2)

        # í‚¤ì›Œë“œ ì•Œë¦¼ (ì¶”ê°€)
        keyword_layout = QHBoxLayout()
        keyword_label = QLabel("âœ¨ í•˜ì´ë¼ì´íŠ¸ í‚¤ì›Œë“œ:")
        keyword_label.setToolTip("ê°•ì¡° í‘œì‹œí•  í‚¤ì›Œë“œ (ì‰¼í‘œë¡œ êµ¬ë¶„)")
        self.keyword_input = QLineEdit()
        self.keyword_input.setPlaceholderText("ì˜ˆ: ì˜ˆì‚°, ê²½ì œ, ì˜ì›ë‹˜ (ì‰¼í‘œ êµ¬ë¶„)")
        self.keyword_input.setToolTip(
            "ìë§‰ì— í•´ë‹¹ í‚¤ì›Œë“œê°€ ë“±ì¥í•˜ë©´ ê°•ì¡° í‘œì‹œë©ë‹ˆë‹¤.\nì—¬ëŸ¬ í‚¤ì›Œë“œëŠ” ì‰¼í‘œ(,)ë¡œ êµ¬ë¶„í•˜ì„¸ìš”."
        )
        # í‚¤ì›Œë“œ ë³€ê²½ ì‹œ ìºì‹œ ì—…ë°ì´íŠ¸ (ë””ë°”ìš´ì‹± ì ìš©)
        self.keyword_input.textChanged.connect(self._update_keyword_cache)

        # í‚¤ì›Œë“œ ì´ˆê¸°ê°’ ë¡œë“œ
        saved_keywords = self.settings.value("highlight_keywords", "", type=str)
        self.keyword_input.setText(saved_keywords)
        self._perform_keyword_cache_update()  # ì´ˆê¸° ìºì‹œ ë¹Œë“œ

        keyword_layout.addWidget(keyword_label)
        keyword_layout.addWidget(self.keyword_input)
        settings_layout.addLayout(keyword_layout, 3, 0, 1, 2)

        layout.addWidget(self.settings_group)

        # === ì»¨íŠ¸ë¡¤ ë²„íŠ¼ ===
        btn_layout = QHBoxLayout()

        self.start_btn = QPushButton("â–¶  ì‹œì‘")
        self.start_btn.setObjectName("startBtn")
        self.start_btn.clicked.connect(self._start)

        self.stop_btn = QPushButton("â¹  ì¤‘ì§€")
        self.stop_btn.setObjectName("stopBtn")
        self.stop_btn.setEnabled(False)
        self.stop_btn.clicked.connect(self._stop)

        btn_layout.addWidget(self.start_btn)
        btn_layout.addWidget(self.stop_btn)
        btn_layout.addStretch()

        # ìƒë‹¨ ì ‘ê¸°/í¼ì¹˜ê¸° ë²„íŠ¼
        self.toggle_header_btn = QPushButton("ğŸ”¼ ìƒë‹¨ ì ‘ê¸°")
        self.toggle_header_btn.setToolTip(
            "ìƒë‹¨ íƒ€ì´í‹€ê³¼ íˆ´ë°”ë¥¼ ìˆ¨ê²¨ ìë§‰ ì˜ì—­ì„ ë„“í™ë‹ˆë‹¤"
        )
        self.toggle_header_btn.setFixedWidth(120)
        self.toggle_header_btn.setStyleSheet("""
            QPushButton {
                background-color: rgba(88, 166, 255, 0.1);
                border: 1px solid rgba(88, 166, 255, 0.3);
                border-radius: 6px;
                padding: 6px 12px;
            }
            QPushButton:hover {
                background-color: rgba(88, 166, 255, 0.2);
            }
        """)
        self.toggle_header_btn.clicked.connect(self._toggle_top_header)
        btn_layout.addWidget(self.toggle_header_btn)

        layout.addLayout(btn_layout)

        # === ì§„í–‰ í‘œì‹œ ===
        self.progress = QProgressBar()
        self.progress.setMaximum(0)
        self.progress.hide()
        layout.addWidget(self.progress)

        # === ë©”ì¸ ì½˜í…ì¸  (ìë§‰ + í†µê³„) ===
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)

        # ìë§‰ ì˜ì—­ ì»¨í…Œì´ë„ˆ (ìë§‰ í…ìŠ¤íŠ¸ + í†µê³„ í† ê¸€ ë²„íŠ¼)
        subtitle_container = QWidget()
        subtitle_layout = QVBoxLayout(subtitle_container)
        subtitle_layout.setContentsMargins(0, 0, 0, 0)
        subtitle_layout.setSpacing(8)

        # ìë§‰ í…ìŠ¤íŠ¸
        self.subtitle_text = QTextEdit()
        self.subtitle_text.setReadOnly(True)
        self.subtitle_text.setUndoRedoEnabled(False)

        # ìŠ¤ë§ˆíŠ¸ ìŠ¤í¬ë¡¤: ìŠ¤í¬ë¡¤ë°” ê°’ ë³€ê²½ ì‹œ ì‚¬ìš©ì ìŠ¤í¬ë¡¤ ê°ì§€
        self.subtitle_text.verticalScrollBar().valueChanged.connect(
            self._on_scroll_changed
        )

        subtitle_layout.addWidget(self.subtitle_text)

        # ì‹¤ì‹œê°„ ë¯¸ë¦¬ë³´ê¸° ì˜ì—­ (ë³„ë„ í‘œì‹œ)
        self.preview_frame = QFrame()
        preview_layout = QHBoxLayout(self.preview_frame)
        preview_layout.setContentsMargins(8, 6, 8, 6)
        preview_layout.setSpacing(8)

        preview_title = QLabel("â³ ë¯¸ë¦¬ë³´ê¸°")
        preview_title.setFixedWidth(90)
        preview_title.setStyleSheet(
            "background: transparent; border: none; font-weight: 600;"
        )

        self.preview_label = QLabel("")
        self.preview_label.setWordWrap(True)
        self.preview_label.setStyleSheet("background: transparent; border: none;")

        preview_layout.addWidget(preview_title)
        preview_layout.addWidget(self.preview_label, 1)

        self.preview_frame.hide()
        subtitle_layout.addWidget(self.preview_frame)

        # í†µê³„ í† ê¸€ ë²„íŠ¼ (ìë§‰ ì˜ì—­ í•˜ë‹¨)
        self.toggle_stats_btn = QPushButton("ğŸ“Š í†µê³„ ìˆ¨ê¸°ê¸°")
        self.toggle_stats_btn.setToolTip("í†µê³„ íŒ¨ë„ ìˆ¨ê¸°ê¸°/ë³´ì´ê¸°")
        self.toggle_stats_btn.setFixedHeight(28)
        self.toggle_stats_btn.setStyleSheet("""
            QPushButton {
                background-color: rgba(88, 166, 255, 0.1);
                border: 1px solid rgba(88, 166, 255, 0.3);
                border-radius: 6px;
                padding: 4px 12px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: rgba(88, 166, 255, 0.2);
            }
        """)
        self.toggle_stats_btn.clicked.connect(self._toggle_stats_panel)
        subtitle_layout.addWidget(self.toggle_stats_btn)

        self.main_splitter.addWidget(subtitle_container)

        # í†µê³„ íŒ¨ë„ - ëª¨ë˜ ë””ìì¸
        self.stats_group = QGroupBox("ğŸ“Š í†µê³„")
        stats_layout = QVBoxLayout(self.stats_group)
        stats_layout.setSpacing(12)
        stats_layout.setContentsMargins(12, 20, 12, 12)

        self.stat_time = QLabel("â±ï¸ ì‹¤í–‰ ì‹œê°„: 00:00:00")
        self.stat_chars = QLabel("ğŸ“ ê¸€ì ìˆ˜: 0")
        self.stat_words = QLabel("ğŸ“– ë‹¨ì–´ ìˆ˜: 0")
        self.stat_sents = QLabel("ğŸ’¬ ë¬¸ì¥ ìˆ˜: 0")
        self.stat_cpm = QLabel("âš¡ ë¶„ë‹¹ ê¸€ì: 0")

        stat_labels = [
            self.stat_time,
            self.stat_chars,
            self.stat_words,
            self.stat_sents,
            self.stat_cpm,
        ]
        for label in stat_labels:
            label.setFont(QFont("ë§‘ì€ ê³ ë”•", 10))
            label.setStyleSheet("""
                padding: 6px 8px;
                border-radius: 6px;
                background-color: rgba(88, 166, 255, 0.08);
            """)
            stats_layout.addWidget(label)

        stats_layout.addStretch()
        self.stats_group.setFixedWidth(220)
        self.main_splitter.addWidget(self.stats_group)

        self.main_splitter.setSizes([860, 220])
        layout.addWidget(self.main_splitter)

        # === "ìµœì‹  ìë§‰" í”Œë¡œíŒ… ë²„íŠ¼ (ìŠ¤ë§ˆíŠ¸ ìŠ¤í¬ë¡¤ìš©) ===
        self.scroll_to_bottom_btn = QPushButton("â¬‡ï¸ ìµœì‹  ìë§‰")
        self.scroll_to_bottom_btn.setToolTip("ìµœì‹  ìë§‰ìœ¼ë¡œ ì´ë™í•˜ê³  ìë™ ìŠ¤í¬ë¡¤ ì¬ê°œ")
        self.scroll_to_bottom_btn.setStyleSheet("""
            QPushButton {
                background-color: rgba(88, 166, 255, 0.9);
                color: white;
                border: none;
                border-radius: 16px;
                padding: 8px 16px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: rgba(88, 166, 255, 1.0);
            }
        """)
        self.scroll_to_bottom_btn.clicked.connect(self._scroll_to_bottom)
        self.scroll_to_bottom_btn.hide()  # ì´ˆê¸°ì—ëŠ” ìˆ¨ê¹€
        layout.addWidget(self.scroll_to_bottom_btn)

        # === ê²€ìƒ‰ë°” (ìˆ¨ê¹€) - ê°œì„ ëœ ë””ìì¸ ===
        self.search_frame = QFrame()
        self.search_frame.setStyleSheet("""
            QFrame {
                background-color: rgba(88, 166, 255, 0.05);
                border: 1px solid rgba(88, 166, 255, 0.2);
                border-radius: 10px;
                padding: 5px;
            }
        """)
        search_layout = QHBoxLayout(self.search_frame)
        search_layout.setContentsMargins(12, 8, 12, 8)
        search_layout.setSpacing(8)

        # ê²€ìƒ‰ ì•„ì´ì½˜ ë¼ë²¨
        search_icon = QLabel("ğŸ”")
        search_icon.setStyleSheet("background: transparent; border: none;")
        search_layout.addWidget(search_icon)

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("ê²€ìƒ‰ì–´ ì…ë ¥...")
        self.search_input.setStyleSheet("""
            QLineEdit {
                background-color: transparent;
                border: none;
                padding: 6px;
                font-size: 13px;
            }
        """)
        self.search_input.returnPressed.connect(self._do_search)
        search_layout.addWidget(self.search_input, 1)

        self.search_count = QLabel("")
        self.search_count.setStyleSheet("""
            color: #8b949e;
            font-size: 12px;
            background: transparent;
            border: none;
            padding: 0 8px;
        """)
        search_layout.addWidget(self.search_count)

        # ê²€ìƒ‰ ë²„íŠ¼ ìŠ¤íƒ€ì¼
        search_btn_style = """
            QPushButton {
                background-color: rgba(88, 166, 255, 0.1);
                border: 1px solid rgba(88, 166, 255, 0.3);
                border-radius: 6px;
                padding: 6px;
                font-size: 12px;
                min-width: 32px;
            }
            QPushButton:hover {
                background-color: rgba(88, 166, 255, 0.2);
                border-color: #58a6ff;
            }
        """

        search_prev = QPushButton("â—€")
        search_prev.setStyleSheet(search_btn_style)
        search_prev.setFixedWidth(36)
        search_prev.clicked.connect(lambda: self._nav_search(-1))

        search_next = QPushButton("â–¶")
        search_next.setStyleSheet(search_btn_style)
        search_next.setFixedWidth(36)
        search_next.clicked.connect(lambda: self._nav_search(1))

        search_close = QPushButton("âœ•")
        search_close.setStyleSheet("""
            QPushButton {
                background-color: rgba(248, 81, 73, 0.1);
                border: 1px solid rgba(248, 81, 73, 0.3);
                border-radius: 6px;
                padding: 6px;
                font-size: 12px;
                min-width: 32px;
            }
            QPushButton:hover {
                background-color: rgba(248, 81, 73, 0.2);
                border-color: #f85149;
            }
        """)
        search_close.setFixedWidth(36)
        search_close.clicked.connect(self._hide_search)

        search_layout.addWidget(search_prev)
        search_layout.addWidget(search_next)
        search_layout.addWidget(search_close)

        # ê²€ìƒ‰ë°” ì´ˆê¸°ì— ìˆ¨ê¹€ ìƒíƒœ
        self.search_frame.hide()
        layout.addWidget(self.search_frame)

        # === ìƒíƒœë°” - ëª¨ë˜ ë””ìì¸ ===
        # === ìƒíƒœë°” - ëª¨ë˜ ë””ìì¸ ===
        status_frame = QFrame()
        status_frame.setFixedHeight(48)  # ë†’ì´ ê³ ì •
        status_frame.setStyleSheet("""
            QFrame {
                background-color: rgba(48, 54, 61, 0.3);
                border-radius: 10px;
                padding: 4px;
            }
        """)
        status_layout = QHBoxLayout(status_frame)
        status_layout.setContentsMargins(16, 4, 16, 4)  # ìƒí•˜ ì—¬ë°± ì¶•ì†Œ
        status_layout.setSpacing(12)

        # ìƒíƒœ í…ìŠ¤íŠ¸
        self.status_label = QLabel("âšª ëŒ€ê¸° ì¤‘")
        self.status_label.setStyleSheet(
            "background: transparent; border: none; font-weight: 600; font-size: 13px;"
        )

        # ì—°ê²° ìƒíƒœ ì¸ë””ì¼€ì´í„° (#30) - ê°œì„ ëœ ë””ìì¸
        self.connection_indicator = QLabel("âš«")
        self.connection_indicator.setToolTip("ì—°ê²° ìƒíƒœ: ëŒ€ê¸° ì¤‘")
        self.connection_indicator.setStyleSheet("""
            background: transparent; 
            border: none; 
            font-size: 14px;
            padding: 2px 8px;
            border-radius: 4px;
        """)

        # êµ¬ë¶„ì„ 
        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.VLine)
        separator.setStyleSheet(
            "background-color: rgba(88, 166, 255, 0.3); max-width: 1px;"
        )

        # ì¹´ìš´íŠ¸ ë¼ë²¨ - ê°œì„ ëœ ìŠ¤íƒ€ì¼
        self.count_label = QLabel("ğŸ“ 0ë¬¸ì¥ | 0ì")
        self.count_label.setStyleSheet("""
            color: #8b949e; 
            background: rgba(88, 166, 255, 0.08); 
            border: none;
            border-radius: 6px;
            padding: 4px 10px;
            font-weight: 500;
        """)

        status_layout.addWidget(self.status_label)
        status_layout.addStretch()
        status_layout.addWidget(self.connection_indicator)
        status_layout.addWidget(separator)
        status_layout.addWidget(self.count_label)

        layout.addWidget(status_frame)

        # ê²€ìƒ‰ ìƒíƒœ
        self.search_matches = []
        self.search_idx = 0

        # ì €ì¥ëœ í°íŠ¸ í¬ê¸° ì ìš©
        self._set_font_size(self.font_size)

    def _apply_theme(self):
        self.setStyleSheet(DARK_THEME if self.is_dark_theme else LIGHT_THEME)
        self.theme_action.setText("ë¼ì´íŠ¸ í…Œë§ˆ" if self.is_dark_theme else "ë‹¤í¬ í…Œë§ˆ")

        # í…Œë§ˆì— ë”°ë¥¸ ë™ì  ìŠ¤íƒ€ì¼ ì—…ë°ì´íŠ¸
        if self.is_dark_theme:
            stat_bg = "rgba(88, 166, 255, 0.08)"
            search_bg = "rgba(88, 166, 255, 0.05)"
            search_border = "rgba(88, 166, 255, 0.2)"
            status_bg = "rgba(48, 54, 61, 0.3)"
            count_color = "#8b949e"
            preview_bg = "rgba(88, 166, 255, 0.08)"
            preview_border = "rgba(88, 166, 255, 0.25)"
        else:
            stat_bg = "rgba(9, 105, 218, 0.06)"
            search_bg = "rgba(9, 105, 218, 0.04)"
            search_border = "rgba(9, 105, 218, 0.15)"
            status_bg = "rgba(208, 215, 222, 0.4)"
            count_color = "#57606a"
            preview_bg = "rgba(9, 105, 218, 0.05)"
            preview_border = "rgba(9, 105, 218, 0.15)"

        # í†µê³„ ë¼ë²¨ ìŠ¤íƒ€ì¼ ì—…ë°ì´íŠ¸
        try:
            stat_labels = [
                self.stat_time,
                self.stat_chars,
                self.stat_words,
                self.stat_sents,
                self.stat_cpm,
            ]
            for label in stat_labels:
                label.setStyleSheet(f"""
                    padding: 6px 8px;
                    border-radius: 6px;
                    background-color: {stat_bg};
                """)

            # ê²€ìƒ‰ë°” ìŠ¤íƒ€ì¼ ì—…ë°ì´íŠ¸
            self.search_frame.setStyleSheet(f"""
                QFrame {{
                    background-color: {search_bg};
                    border: 1px solid {search_border};
                    border-radius: 10px;
                    padding: 5px;
                }}
            """)

            # ìƒíƒœë°” ì¹´ìš´íŠ¸ ë¼ë²¨ ìƒ‰ìƒ ì—…ë°ì´íŠ¸
            self.count_label.setStyleSheet(
                f"color: {count_color}; background: transparent; border: none;"
            )

            # ë¯¸ë¦¬ë³´ê¸° ì˜ì—­ ìŠ¤íƒ€ì¼ ì—…ë°ì´íŠ¸
            if hasattr(self, "preview_frame"):
                self.preview_frame.setStyleSheet(
                    f"background-color: {preview_bg}; border: 1px solid {preview_border}; "
                    "border-radius: 8px;"
                )
        except AttributeError:
            # UIê°€ ì•„ì§ ì™„ì „íˆ ì´ˆê¸°í™”ë˜ì§€ ì•Šì€ ê²½ìš°
            pass

    def _toggle_theme(self):
        self.is_dark_theme = not self.is_dark_theme
        self.settings.setValue("dark_theme", self.is_dark_theme)
        self._apply_theme()

    def _toggle_tray_option(self):
        """íŠ¸ë ˆì´ ìµœì†Œí™” ì˜µì…˜ í† ê¸€"""
        self.minimize_to_tray = self.tray_action.isChecked()
        self.settings.setValue("minimize_to_tray", self.minimize_to_tray)
        if self.minimize_to_tray:
            self._show_toast("ì°½ì„ ë‹«ìœ¼ë©´ íŠ¸ë ˆì´ë¡œ ìµœì†Œí™”ë©ë‹ˆë‹¤.", "info")
        else:
            self._show_toast("ì°½ì„ ë‹«ìœ¼ë©´ í”„ë¡œê·¸ë¨ì´ ì¢…ë£Œë©ë‹ˆë‹¤.", "info")

    def _setup_shortcuts(self):
        QShortcut(QKeySequence("F5"), self, self._start)
        QShortcut(QKeySequence("Escape"), self, self._stop)
        QShortcut(QKeySequence("F3"), self, lambda: self._nav_search(1))
        QShortcut(QKeySequence("Shift+F3"), self, lambda: self._nav_search(-1))

    # ========== í† ìŠ¤íŠ¸ ë° ìƒíƒœ í‘œì‹œ ==========

    def _show_toast(
        self, message: str, toast_type: str = "info", duration: int = 3000
    ) -> None:
        """í† ìŠ¤íŠ¸ ì•Œë¦¼ í‘œì‹œ - ìŠ¤íƒ ì²˜ë¦¬ë¡œ ê²¹ì¹¨ ë°©ì§€"""
        # ë§Œë£Œëœ í† ìŠ¤íŠ¸ ì •ë¦¬
        self.active_toasts = [t for t in self.active_toasts if t.isVisible()]

        # ìƒˆ í† ìŠ¤íŠ¸ y ìœ„ì¹˜ ê³„ì‚° (ê¸°ì¡´ í† ìŠ¤íŠ¸ ì•„ë˜ì— ë°°ì¹˜)
        y_offset = 10
        for toast in self.active_toasts:
            y_offset += toast.height() + 5

        # í† ìŠ¤íŠ¸ ì œê±° ì½œë°±
        def remove_toast(t):
            if t in self.active_toasts:
                self.active_toasts.remove(t)

        # í† ìŠ¤íŠ¸ ìƒì„±
        toast = ToastWidget(
            self.centralWidget(),
            message,
            duration,
            toast_type,
            y_offset=y_offset,
            on_close=remove_toast,
        )
        self.active_toasts.append(toast)

    def _set_status(self, text: str, status_type: str = "info"):
        """ìƒíƒœ í‘œì‹œ (ì•„ì´ì½˜ + ìƒ‰ìƒ)"""
        icons = {
            "info": "â„¹ï¸",
            "success": "âœ…",
            "warning": "âš ï¸",
            "error": "âŒ",
            "running": "ğŸ”„",
        }
        colors = {
            "info": "#4fc3f7",
            "success": "#4caf50",
            "warning": "#ff9800",
            "error": "#f44336",
            "running": "#ab47bc",
        }
        icon = icons.get(status_type, "")
        color = colors.get(status_type, "#eaeaea")
        rendered = f"{icon} {text}"[:100]
        self.status_label.setText(rendered)
        self.status_label.setStyleSheet(f"color: {color};")
        self._last_status_message = rendered

    def _update_count_label(self):
        """ìë§‰ ì¹´ìš´íŠ¸ ë¼ë²¨ ì—…ë°ì´íŠ¸"""
        with self.subtitle_lock:
            count = len(self.subtitles)
        chars = self._cached_total_chars
        self.count_label.setText(f"ğŸ“ {count}ë¬¸ì¥ | {chars:,}ì")

    def _update_connection_status(self, status: str, latency: int = None):
        """ì—°ê²° ìƒíƒœ ì¸ë””ì¼€ì´í„° ì—…ë°ì´íŠ¸ (#30)

        Args:
            status: 'connected', 'disconnected', 'reconnecting'
            latency: ì‘ë‹µ ì‹œê°„ (ms), ì—°ê²°ëœ ê²½ìš°ì—ë§Œ
        """
        self.connection_status = status

        # ìƒíƒœë³„ ì•„ì´ì½˜ê³¼ íˆ´íŒ
        status_config = {
            "connected": ("ğŸŸ¢", "#4caf50", "ì—°ê²°ë¨"),
            "disconnected": ("ğŸ”´", "#f44336", "ì—°ê²° ëŠê¹€"),
            "reconnecting": ("ğŸŸ¡", "#ff9800", "ì¬ì—°ê²° ì¤‘..."),
        }

        icon, color, text = status_config.get(status, ("âš«", "#888", "ì•Œ ìˆ˜ ì—†ìŒ"))

        # ë ˆì´í„´ì‹œê°€ ìˆìœ¼ë©´ íˆ´íŒì— í‘œì‹œ
        if latency is not None and status == "connected":
            self.ping_latency = latency
            tooltip = f"ì—°ê²° ìƒíƒœ: {text} ({latency}ms)"
        elif status == "reconnecting":
            tooltip = f"ì—°ê²° ìƒíƒœ: {text} (ì‹œë„ {self.reconnect_attempts}/{Config.MAX_RECONNECT_ATTEMPTS})"
        else:
            tooltip = f"ì—°ê²° ìƒíƒœ: {text}"

        self.connection_indicator.setText(icon)
        self.connection_indicator.setToolTip(tooltip)
        self.connection_indicator.setStyleSheet(
            f"background: transparent; border: none; font-size: 12px; color: {color};"
        )

    # ========== ê¸€ì í¬ê¸° ==========

    def _set_font_size(self, size: int):
        """ìë§‰ ì˜ì—­ í°íŠ¸ í¬ê¸° ë³€ê²½"""
        size = max(Config.MIN_FONT_SIZE, min(size, Config.MAX_FONT_SIZE))
        self.font_size = size
        font = self.subtitle_text.font()
        font.setPointSize(size)
        self.subtitle_text.setFont(font)
        self.settings.setValue("font_size", size)

    def _adjust_font_size(self, delta: int):
        """í°íŠ¸ í¬ê¸° ì¡°ì ˆ"""
        self._set_font_size(self.font_size + delta)

    # ========== URL íˆìŠ¤í† ë¦¬ (íƒœê·¸ ì§€ì›) ==========

    def _load_url_history(self):
        """URL íˆìŠ¤í† ë¦¬ ë¡œë“œ - {url: tag} í˜•íƒœ"""
        try:
            if Path(Config.URL_HISTORY_FILE).exists():
                with open(Config.URL_HISTORY_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    # dict í˜•íƒœì¸ì§€ í™•ì¸ (ìƒˆë¡œìš´ í˜•ì‹)
                    if isinstance(data, dict):
                        return data
                    # ì´ì „ list í˜•íƒœë©´ dictë¡œ ë³€í™˜
                    elif isinstance(data, list):
                        return {url: "" for url in data}
        except Exception as e:
            logger.warning(f"URL íˆìŠ¤í† ë¦¬ ë¡œë“œ ì˜¤ë¥˜: {e}")
        return {}

    def _save_url_history(self):
        """URL íˆìŠ¤í† ë¦¬ ì €ì¥"""
        try:
            if not isinstance(self.url_history, dict):
                self.url_history = {}
            utils.atomic_write_json(
                Config.URL_HISTORY_FILE,
                self.url_history,
                ensure_ascii=False,
                indent=2,
            )
        except Exception as e:
            logger.warning(f"URL íˆìŠ¤í† ë¦¬ ì €ì¥ ì˜¤ë¥˜: {e}")

    def _add_to_history(self, url, tag=""):
        """URL íˆìŠ¤í† ë¦¬ì— ì¶”ê°€ (ìë™ íƒœê·¸ ë§¤ì¹­)"""
        if not isinstance(self.url_history, dict):
            self.url_history = {}

        # íƒœê·¸ê°€ ì—†ìœ¼ë©´ ìë™ ê°ì§€
        if not tag:
            # 1. ì´ë¯¸ ì €ì¥ëœ íƒœê·¸ê°€ ìˆëŠ”ì§€ í™•ì¸
            if url in self.url_history and self.url_history[url]:
                tag = self.url_history[url]
            else:
                # 2. í”„ë¦¬ì…‹/ì•½ì¹­ì—ì„œ ë§¤ì¹­ í™•ì¸
                tag = self._autodetect_tag(url)

        self.url_history[url] = tag

        # íˆìŠ¤í† ë¦¬ í¬ê¸° ì œí•œ
        if len(self.url_history) > Config.MAX_URL_HISTORY:
            # ê°€ì¥ ì˜¤ë˜ëœ í•­ëª© ì‚­ì œ (dictëŠ” ì‚½ì… ìˆœì„œ ìœ ì§€)
            oldest_key = next(iter(self.url_history))
            del self.url_history[oldest_key]

        self._save_url_history()
        self._refresh_url_combo()

    def _autodetect_tag(self, url):
        """URLì„ ê¸°ë°˜ìœ¼ë¡œ ìœ„ì›íšŒ ì´ë¦„/ì•½ì¹­ ìë™ ê°ì§€"""
        # 1. ì •í™•í•œ URL ë§¤ì¹­ í™•ì¸ (í”„ë¦¬ì…‹)
        for name, preset_url in self.committee_presets.items():
            if url == preset_url:
                # ì•½ì¹­ì´ ìˆìœ¼ë©´ ì•½ì¹­ ì‚¬ìš© (ë” ì§§ê³  ë³´ê¸° ì¢‹ìŒ)
                for abbr, full_name in Config.COMMITTEE_ABBREVIATIONS.items():
                    if full_name == name:
                        return abbr
                return name

        # 2. xcode íŒŒë¼ë¯¸í„° ë§¤ì¹­ (ìˆ«ì ë˜ëŠ” ë¬¸ìì—´ xcode ëª¨ë‘ ì§€ì›)
        import re

        match = re.search(r"xcode=([^&]+)", url)
        if match:
            xcode = match.group(1)
            # í”„ë¦¬ì…‹ì—ì„œ í•´ë‹¹ xcodeë¥¼ ê°€ì§„ URL ì°¾ê¸°
            for name, preset_url in self.committee_presets.items():
                if f"xcode={xcode}" in preset_url:
                    # ì•½ì¹­ ë¦¬í„´
                    for abbr, full_name in Config.COMMITTEE_ABBREVIATIONS.items():
                        if full_name == name:
                            return abbr
                    return name

        return ""

    def _refresh_url_combo(self):
        """URL ì½¤ë³´ë°•ìŠ¤ ìƒˆë¡œê³ ì¹¨"""
        current_text = self.url_combo.currentText()
        self.url_combo.clear()

        for url, tag in self.url_history.items():
            if tag:
                self.url_combo.addItem(f"[{tag}] {url}", url)
            else:
                self.url_combo.addItem(url, url)

        # ê¸°ë³¸ URL ì¶”ê°€
        if not self.url_history:
            self.url_combo.addItem("https://assembly.webcast.go.kr/main/player.asp")

        # ì´ì „ í…ìŠ¤íŠ¸ ë³µì›
        if current_text:
            self.url_combo.setCurrentText(current_text)

    def _get_current_url(self):
        """í˜„ì¬ ì„ íƒëœ URL ë°˜í™˜ (íƒœê·¸ ì œê±°)"""
        text = self.url_combo.currentText().strip()
        text_url = text
        if text.startswith("[") and "] " in text:
            text_url = text.split("] ", 1)[1].strip()

        data = self.url_combo.currentData()
        if data:
            data_url = str(data).strip()
            if text_url and text_url != data_url:
                return text_url
            return data_url

        return text_url

    def _edit_url_tag(self):
        """í˜„ì¬ URLì˜ íƒœê·¸ í¸ì§‘"""
        url = self._get_current_url()
        if not url or not url.startswith(("http://", "https://")):
            QMessageBox.warning(self, "ì•Œë¦¼", "íƒœê·¸ë¥¼ ì§€ì •í•  URLì„ ë¨¼ì € ì„ íƒí•˜ì„¸ìš”.")
            return

        current_tag = self.url_history.get(url, "")
        tag, ok = QInputDialog.getText(
            self,
            "URL íƒœê·¸ ì„¤ì •",
            f"URL: {url[:50]}...\n\níƒœê·¸ ì…ë ¥ (ì˜ˆ: ë³¸íšŒì˜, ë²•ì‚¬ìœ„, ìƒì„ìœ„):",
            text=current_tag,
        )

        if ok:
            self._add_to_history(url, tag.strip())
            QMessageBox.information(
                self,
                "ì„±ê³µ",
                f"íƒœê·¸ê°€ ì„¤ì •ë˜ì—ˆìŠµë‹ˆë‹¤: [{tag}]" if tag else "íƒœê·¸ê°€ ì œê±°ë˜ì—ˆìŠµë‹ˆë‹¤.",
            )

    # ========== ìƒì„ìœ„ í”„ë¦¬ì…‹ ==========

    def _load_committee_presets(self):
        """í”„ë¦¬ì…‹ íŒŒì¼ì—ì„œ ë¡œë“œ (ì—†ìœ¼ë©´ ê¸°ë³¸ê°’ ì‚¬ìš©)"""
        self.committee_presets = dict(Config.DEFAULT_COMMITTEE_PRESETS)
        self.custom_presets = {}

        try:
            if Path(Config.PRESET_FILE).exists():
                with open(Config.PRESET_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if "presets" in data:
                        self.committee_presets.update(data["presets"])
                    if "custom" in data:
                        self.custom_presets = data["custom"]
        except Exception as e:
            logger.warning(f"í”„ë¦¬ì…‹ ë¡œë“œ ì˜¤ë¥˜: {e}")

    def _save_committee_presets(self):
        """í”„ë¦¬ì…‹ì„ íŒŒì¼ì— ì €ì¥"""
        try:
            data = {"presets": self.committee_presets, "custom": self.custom_presets}
            utils.atomic_write_json(
                Config.PRESET_FILE,
                data,
                ensure_ascii=False,
                indent=2,
            )
        except Exception as e:
            logger.warning(f"í”„ë¦¬ì…‹ ì €ì¥ ì˜¤ë¥˜: {e}")

    def _build_preset_menu(self):
        """í”„ë¦¬ì…‹ ë©”ë‰´ êµ¬ì„±"""
        self.preset_menu.clear()

        # ê¸°ë³¸ ìƒì„ìœ„ì›íšŒ
        for name, url in self.committee_presets.items():
            action = QAction(name, self)
            action.setData(url)
            action.triggered.connect(
                lambda checked, u=url, n=name: self._select_preset(u, n)
            )
            self.preset_menu.addAction(action)

        # ì‚¬ìš©ì ì •ì˜ í”„ë¦¬ì…‹ì´ ìˆìœ¼ë©´ êµ¬ë¶„ì„  ì¶”ê°€
        if self.custom_presets:
            self.preset_menu.addSeparator()
            self.preset_menu.addAction("â”€â”€ ì‚¬ìš©ì ì •ì˜ â”€â”€").setEnabled(False)

            for name, url in self.custom_presets.items():
                action = QAction(f"â­ {name}", self)
                action.setData(url)
                action.triggered.connect(
                    lambda checked, u=url, n=name: self._select_preset(u, n)
                )
                self.preset_menu.addAction(action)

        # ê´€ë¦¬ ë©”ë‰´
        self.preset_menu.addSeparator()
        add_action = QAction("â• í”„ë¦¬ì…‹ ì¶”ê°€...", self)
        add_action.triggered.connect(self._add_custom_preset)
        self.preset_menu.addAction(add_action)

        edit_action = QAction("âœï¸ í”„ë¦¬ì…‹ ê´€ë¦¬...", self)
        edit_action.triggered.connect(self._manage_presets)
        self.preset_menu.addAction(edit_action)

        self.preset_menu.addSeparator()
        export_action = QAction("ğŸ“¤ í”„ë¦¬ì…‹ ë‚´ë³´ë‚´ê¸°...", self)
        export_action.triggered.connect(self._export_presets)
        self.preset_menu.addAction(export_action)

        import_action = QAction("ğŸ“¥ í”„ë¦¬ì…‹ ê°€ì ¸ì˜¤ê¸°...", self)
        import_action.triggered.connect(self._import_presets)
        self.preset_menu.addAction(import_action)

    def _select_preset(self, url, name):
        """í”„ë¦¬ì…‹ ì„ íƒ ì‹œ URL ì„¤ì •"""
        self.url_combo.setCurrentText(url)
        self._show_toast(f"'{name}' ì„ íƒë¨", "success", 1500)

    def _add_custom_preset(self):
        """ì‚¬ìš©ì ì •ì˜ í”„ë¦¬ì…‹ ì¶”ê°€"""
        name, ok = QInputDialog.getText(
            self, "í”„ë¦¬ì…‹ ì¶”ê°€", "í”„ë¦¬ì…‹ ì´ë¦„ì„ ì…ë ¥í•˜ì„¸ìš”:"
        )
        if not ok or not name.strip():
            return

        name = name.strip()
        current_url = self._get_current_url()

        url, ok = QInputDialog.getText(
            self,
            "í”„ë¦¬ì…‹ URL",
            f"'{name}' í”„ë¦¬ì…‹ì˜ URLì„ ì…ë ¥í•˜ì„¸ìš”:",
            text=current_url if current_url.startswith("http") else Config.DEFAULT_URL,
        )

        if ok and url.strip():
            self.custom_presets[name] = url.strip()
            self._save_committee_presets()
            self._build_preset_menu()
            self._show_toast(f"í”„ë¦¬ì…‹ '{name}' ì¶”ê°€ë¨", "success")

    def _manage_presets(self):
        """í”„ë¦¬ì…‹ ê´€ë¦¬ ëŒ€í™”ìƒì"""
        if not self.custom_presets:
            QMessageBox.information(
                self,
                "í”„ë¦¬ì…‹ ê´€ë¦¬",
                "ì‚¬ìš©ì ì •ì˜ í”„ë¦¬ì…‹ì´ ì—†ìŠµë‹ˆë‹¤.\n\n"
                "'â• í”„ë¦¬ì…‹ ì¶”ê°€'ë¥¼ í†µí•´ ìƒˆ í”„ë¦¬ì…‹ì„ ì¶”ê°€í•˜ì„¸ìš”.",
            )
            return

        # ì‘ì—… ì„ íƒ
        names = list(self.custom_presets.keys())
        actions = ["ìˆ˜ì •", "ì‚­ì œ"]
        action, ok = QInputDialog.getItem(
            self, "í”„ë¦¬ì…‹ ê´€ë¦¬", "ì‘ì—…ì„ ì„ íƒí•˜ì„¸ìš”:", actions, 0, False
        )

        if not ok:
            return

        # í”„ë¦¬ì…‹ ì„ íƒ
        name, ok = QInputDialog.getItem(
            self,
            f"í”„ë¦¬ì…‹ {action}",
            f"{action}í•  í”„ë¦¬ì…‹ì„ ì„ íƒí•˜ì„¸ìš”:",
            names,
            0,
            False,
        )

        if not ok or not name:
            return

        if action == "ì‚­ì œ":
            reply = QMessageBox.question(
                self,
                "í™•ì¸",
                f"'{name}' í”„ë¦¬ì…‹ì„ ì‚­ì œí•˜ì‹œê² ìŠµë‹ˆê¹Œ?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if reply == QMessageBox.StandardButton.Yes:
                del self.custom_presets[name]
                self._save_committee_presets()
                self._build_preset_menu()
                self._show_toast(f"í”„ë¦¬ì…‹ '{name}' ì‚­ì œë¨", "warning")

        elif action == "ìˆ˜ì •":
            # ì´ë¦„ ìˆ˜ì •
            new_name, ok = QInputDialog.getText(
                self, "í”„ë¦¬ì…‹ ì´ë¦„ ìˆ˜ì •", f"'{name}' í”„ë¦¬ì…‹ì˜ ìƒˆ ì´ë¦„:", text=name
            )
            if not ok:
                return

            # URL ìˆ˜ì •
            current_url = self.custom_presets[name]
            new_url, ok = QInputDialog.getText(
                self, "í”„ë¦¬ì…‹ URL ìˆ˜ì •", f"'{new_name}' í”„ë¦¬ì…‹ì˜ URL:", text=current_url
            )
            if not ok:
                return

            # ê¸°ì¡´ í”„ë¦¬ì…‹ ì‚­ì œ í›„ ìƒˆë¡œ ì¶”ê°€ (ì´ë¦„ì´ ë³€ê²½ë˜ì—ˆì„ ìˆ˜ ìˆìœ¼ë¯€ë¡œ)
            del self.custom_presets[name]
            self.custom_presets[new_name.strip()] = new_url.strip()
            self._save_committee_presets()
            self._build_preset_menu()
            self._show_toast(f"í”„ë¦¬ì…‹ '{new_name}' ìˆ˜ì •ë¨", "success")

    def _export_presets(self):
        """í”„ë¦¬ì…‹ì„ íŒŒì¼ë¡œ ë‚´ë³´ë‚´ê¸°"""
        filename = f"presets_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        path, _ = QFileDialog.getSaveFileName(
            self, "í”„ë¦¬ì…‹ ë‚´ë³´ë‚´ê¸°", filename, "JSON íŒŒì¼ (*.json)"
        )

        if path:
            try:
                data = {
                    "version": Config.VERSION,
                    "exported": datetime.now().isoformat(),
                    "committee": self.committee_presets,
                    "custom": self.custom_presets,
                }
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)

                total = len(self.committee_presets) + len(self.custom_presets)
                self._show_toast(f"í”„ë¦¬ì…‹ {total}ê°œ ë‚´ë³´ë‚´ê¸° ì™„ë£Œ!", "success")
                logger.info(f"í”„ë¦¬ì…‹ ë‚´ë³´ë‚´ê¸° ì™„ë£Œ: {path}")
            except Exception as e:
                QMessageBox.critical(self, "ì˜¤ë¥˜", f"í”„ë¦¬ì…‹ ë‚´ë³´ë‚´ê¸° ì‹¤íŒ¨: {e}")

    def _import_presets(self):
        """íŒŒì¼ì—ì„œ í”„ë¦¬ì…‹ ê°€ì ¸ì˜¤ê¸°"""
        path, _ = QFileDialog.getOpenFileName(
            self, "í”„ë¦¬ì…‹ ê°€ì ¸ì˜¤ê¸°", "", "JSON íŒŒì¼ (*.json)"
        )

        if path:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)

                imported_count = 0

                # ì‚¬ìš©ì ì •ì˜ í”„ë¦¬ì…‹ ê°€ì ¸ì˜¤ê¸° (ê¸°ì¡´ ê²ƒì— ì¶”ê°€)
                if "custom" in data and isinstance(data["custom"], dict):
                    for name, url in data["custom"].items():
                        if name not in self.custom_presets:
                            self.custom_presets[name] = url
                            imported_count += 1

                if imported_count > 0:
                    self._save_committee_presets()
                    self._build_preset_menu()
                    self._show_toast(
                        f"í”„ë¦¬ì…‹ {imported_count}ê°œ ê°€ì ¸ì˜¤ê¸° ì™„ë£Œ!", "success"
                    )
                else:
                    self._show_toast("ê°€ì ¸ì˜¬ ìƒˆ í”„ë¦¬ì…‹ì´ ì—†ìŠµë‹ˆë‹¤", "info")

                logger.info(f"í”„ë¦¬ì…‹ ê°€ì ¸ì˜¤ê¸° ì™„ë£Œ: {path}, {imported_count}ê°œ")
            except json.JSONDecodeError:
                QMessageBox.warning(self, "ì˜¤ë¥˜", "ì˜ëª»ëœ JSON íŒŒì¼ í˜•ì‹ì…ë‹ˆë‹¤.")
            except Exception as e:
                QMessageBox.critical(self, "ì˜¤ë¥˜", f"í”„ë¦¬ì…‹ ê°€ì ¸ì˜¤ê¸° ì‹¤íŒ¨: {e}")

    # ========== ì¶”ì¶œ ì œì–´ ==========

    def _start(self):
        if self.is_running:
            return

        url = self._get_current_url().strip()
        selector = self.selector_combo.currentText().strip()

        if not url or not selector:
            QMessageBox.warning(self, "ì˜¤ë¥˜", "URLê³¼ ì„ íƒìë¥¼ ì…ë ¥í•˜ì„¸ìš”.")
            return

        if not url.startswith(("http://", "https://")):
            QMessageBox.warning(self, "ì˜¤ë¥˜", "ì˜¬ë°”ë¥¸ URLì„ ì…ë ¥í•˜ì„¸ìš”.")
            return

        try:
            # íˆìŠ¤í† ë¦¬ì— ì¶”ê°€
            self._add_to_history(url)
            self.current_url = url

            # ì´ˆê¸°í™”
            self.subtitle_text.clear()
            with self.subtitle_lock:
                self.subtitles = []
            self._cached_total_chars = 0
            self._cached_total_words = 0
            self._last_rendered_count = 0
            self._last_rendered_last_text = ""
            self._last_render_offset = 0
            self._last_render_show_ts = None
            self._last_printed_ts = None
            self._update_count_label()

            self.last_subtitle = ""
            self.last_update_time = 0
            self._last_raw_text = ""
            self._last_processed_raw = ""
            self._stream_start_time = None
            self._confirmed_compact = ""
            self._trailing_suffix = ""
            self._preview_desync_count = 0
            self._preview_ambiguous_skip_count = 0
            self._last_good_raw_compact = ""
            self._is_stopping = False
            self.finalize_timer.stop()
            self._clear_preview()
            self.start_time = time.time()

            # í ë¹„ìš°ê¸°
            self._clear_message_queue()

            # ì‹¤ì‹œê°„ ì €ì¥ ì„¤ì • (ì˜ˆì™¸ ì²˜ë¦¬ ì¶”ê°€)
            self.realtime_file = None
            if self.realtime_save_check.isChecked():
                try:
                    Path(Config.REALTIME_DIR).mkdir(exist_ok=True)
                    filename = f"{Config.REALTIME_DIR}/ìë§‰_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                    self.realtime_file = open(
                        filename, "w", encoding="utf-8-sig"
                    )  # BOM í¬í•¨ (#12)
                    self._set_status(f"ì‹¤ì‹œê°„ ì €ì¥: {filename}", "success")
                except Exception as e:
                    logger.error(f"ì‹¤ì‹œê°„ ì €ì¥ íŒŒì¼ ìƒì„± ì˜¤ë¥˜: {e}")

            # UI ì—…ë°ì´íŠ¸
            self.is_running = True
            self.stop_event.clear()  # ì¢…ë£Œ ì´ë²¤íŠ¸ ì´ˆê¸°í™”
            self.start_btn.setEnabled(False)
            self.stop_btn.setEnabled(True)
            self.url_combo.setEnabled(False)
            self.selector_combo.setEnabled(False)
            self.progress.show()

            self._set_status("Chrome ë¸Œë¼ìš°ì € ì‹œì‘ ì¤‘...", "running")
            self._update_tray_status("ğŸŸ¢ ì¶”ì¶œ ì¤‘")

            # UI ê°’ì„ ì‹œì‘ ì‹œì ì— ë³µì‚¬ (ìŠ¤ë ˆë“œ ì•ˆì „ì„±)
            headless = self.headless_check.isChecked()

            # ì›Œì»¤ ì‹œì‘
            self.worker = threading.Thread(
                target=self._extraction_worker,
                args=(url, selector, headless),
                daemon=True,
            )
            self.worker.start()

            # íƒ€ì´ë¨¸ ì‹œì‘
            self.stats_timer.start(Config.STATS_UPDATE_INTERVAL)
            self.backup_timer.start(Config.AUTO_BACKUP_INTERVAL)

            # ì‹œì‘ ì‹œ ìë™ìœ¼ë¡œ ìƒë‹¨ UI ì ‘ê¸° (ì‚¬ìš©ì ìš”ì²­)
            if self.top_header_container.isVisible():
                self._toggle_top_header()

        except Exception as e:
            logger.exception(f"ì‹œì‘ ì˜¤ë¥˜: {e}")
            # íŒŒì¼ í•¸ë“¤ ì •ë¦¬
            if self.realtime_file:
                try:
                    self.realtime_file.close()
                except Exception as close_err:
                    logger.debug(f"ì‹¤ì‹œê°„ ì €ì¥ íŒŒì¼ ë‹«ê¸° ì˜¤ë¥˜: {close_err}")
                self.realtime_file = None
            self._reset_ui()
            QMessageBox.critical(self, "ì˜¤ë¥˜", f"ì‹œì‘ ì¤‘ ì˜¤ë¥˜ ë°œìƒ: {e}")

    def _stop(self, for_app_exit: bool = False):
        if not self.is_running:
            return

        try:
            self.is_running = False
            self.stop_event.set()  # ì›Œì»¤ ìŠ¤ë ˆë“œì— ì¢…ë£Œ ì‹ í˜¸
            self._set_status("ì¤‘ì§€ ì¤‘...", "warning")
            self._is_stopping = True

            # í™•ì • íƒ€ì´ë¨¸ ì¤‘ì§€
            self.finalize_timer.stop()

            # ì¢…ë£Œ ì§ì „ í preview ì†Œì§„ + ë§ˆì§€ë§‰ ìë§‰ í™•ì •
            self._drain_pending_previews(requeue_others=True)
            self._finalize_pending_subtitle()

            force_driver_quit = for_app_exit or (not Config.KEEP_BROWSER_ON_STOP)

            # ì•± ì¢…ë£Œ ë˜ëŠ” ë¸Œë¼ìš°ì € ë¯¸ìœ ì§€ ì„¤ì •ì¼ ë•ŒëŠ” ì›Œì»¤ ëŒ€ê¸° ì „ì— ë“œë¼ì´ë²„ë¥¼ ë¨¼ì € ì •ë¦¬
            if force_driver_quit and self.driver:
                driver = self.driver
                self.driver = None
                self._force_quit_driver_with_timeout(
                    driver, timeout=2.0, source="stop_initial"
                )

            worker_stopped = self._wait_worker_shutdown(
                timeout=Config.THREAD_STOP_TIMEOUT
            )

            # ìˆ˜ë™ ì¤‘ì§€ + ë¸Œë¼ìš°ì € ìœ ì§€ ëª¨ë“œì—ì„œ ì¢…ë£Œê°€ ì§€ì—°ë˜ë©´ 1íšŒ ì—ìŠ¤ì»¬ë ˆì´ì…˜
            if not worker_stopped and not force_driver_quit and self.driver:
                logger.warning("ì›Œì»¤ ìŠ¤ë ˆë“œ ì¢…ë£Œ ì§€ì—° ê°ì§€ - ë“œë¼ì´ë²„ ê°•ì œ ì¢…ë£Œ í›„ ì¬ëŒ€ê¸°")
                driver = self.driver
                self.driver = None
                self._force_quit_driver_with_timeout(
                    driver, timeout=2.0, source="stop_escalation"
                )
                worker_stopped = self._wait_worker_shutdown(timeout=1.0)

            if not worker_stopped:
                logger.warning("ì›Œì»¤ ìŠ¤ë ˆë“œê°€ ì‹œê°„ ë‚´ì— ì¢…ë£Œë˜ì§€ ì•ŠìŒ(ì¢…ë£Œ ê³„ì† ì§„í–‰)")

            # ì¢…ë£Œ í›„ì—ë„ ë‚¨ì•„ìˆë˜ preview ì²˜ë¦¬
            self._drain_pending_previews(requeue_others=True)
            self._finalize_pending_subtitle()
            self._clear_preview()

            # ì‹¤ì‹œê°„ ì €ì¥ ì¢…ë£Œ
            if self.realtime_file:
                try:
                    self.realtime_file.close()
                except Exception as e:
                    logger.debug(f"íŒŒì¼ ë‹«ê¸° ì˜¤ë¥˜: {e}")
                self.realtime_file = None

            # ì¤‘ì§€ ì´í›„ ë‚¨ì•„ìˆì„ ìˆ˜ ìˆëŠ” detached WebDriver ì •ë¦¬
            self._cleanup_detached_drivers_with_timeout(timeout=2.0)

            self._clear_message_queue()
            self._reset_ui()
            self._set_status("ì¤‘ì§€ë¨", "warning")
            self._update_tray_status("âšª ëŒ€ê¸° ì¤‘")
        except Exception as e:
            logger.error(f"ì¤‘ì§€ ì¤‘ ì˜¤ë¥˜ ë°œìƒ: {e}")
            self._reset_ui()
        finally:
            self._is_stopping = False

    def _force_quit_driver_with_timeout(
        self, driver, timeout: float = 2.0, source: str = "shutdown"
    ) -> bool:
        """WebDriver quit()ë¥¼ íƒ€ì„ë°•ìŠ¤ë¡œ ì‹¤í–‰í•´ UI ìŠ¤ë ˆë“œ ë¬´ê¸°í•œ ëŒ€ê¸°ë¥¼ ë°©ì§€í•œë‹¤."""
        if not driver:
            return True

        done = threading.Event()
        error_holder = {"error": None}

        def _quit_driver():
            try:
                driver.quit()
            except Exception as e:
                error_holder["error"] = e
            finally:
                done.set()

        threading.Thread(
            target=_quit_driver,
            daemon=True,
            name=f"DriverQuitThread-{source}",
        ).start()

        if not done.wait(timeout=timeout):
            logger.warning(
                "WebDriver ì¢…ë£Œ íƒ€ì„ì•„ì›ƒ (source=%s, timeout=%.1fs)",
                source,
                timeout,
            )
            return False

        if error_holder["error"] is not None:
            logger.debug(
                "WebDriver ì¢…ë£Œ ì˜¤ë¥˜ (source=%s): %s",
                source,
                error_holder["error"],
            )
            return False

        return True

    def _wait_worker_shutdown(self, timeout: float) -> bool:
        """ì›Œì»¤ ìŠ¤ë ˆë“œ ì¢…ë£Œë¥¼ ì œí•œ ì‹œê°„ ë™ì•ˆ ëŒ€ê¸°í•œë‹¤."""
        if not self.worker or not self.worker.is_alive():
            return True
        self.worker.join(timeout=timeout)
        return not self.worker.is_alive()

    def _cleanup_detached_drivers_with_timeout(self, timeout: float = 2.0) -> None:
        """ë¶„ë¦¬ëœ(detached) ë“œë¼ì´ë²„ë¥¼ íƒ€ì„ë°•ìŠ¤ë¡œ ì¢…ë£Œí•œë‹¤."""
        with self._detached_drivers_lock:
            detached_drivers = list(self._detached_drivers)
            self._detached_drivers.clear()

        for idx, drv in enumerate(detached_drivers, start=1):
            self._force_quit_driver_with_timeout(
                drv,
                timeout=timeout,
                source=f"detached_{idx}",
            )

    def _reset_ui(self):
        self.is_running = False
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.url_combo.setEnabled(True)
        self.selector_combo.setEnabled(True)
        self.progress.hide()
        self.stats_timer.stop()
        self.backup_timer.stop()
        self.finalize_timer.stop()

    # ========== ì›Œì»¤ ìŠ¤ë ˆë“œ ==========

    def _get_query_param(self, url: str, name: str) -> str:
        """URL ì¿¼ë¦¬ íŒŒë¼ë¯¸í„° ê°’ ì¶”ì¶œ (ì—†ìœ¼ë©´ ë¹ˆ ë¬¸ìì—´)"""
        match = re.search(r"(?:^|[?&])" + re.escape(name) + r"=([^&]*)", url)
        return match.group(1) if match else ""

    def _set_query_param(self, url: str, name: str, value: str) -> str:
        """URL ì¿¼ë¦¬ íŒŒë¼ë¯¸í„° ì„¤ì •/êµì²´"""
        base_url = url.strip().rstrip("&")
        pattern = re.compile(r"([?&])" + re.escape(name) + r"=[^&]*")
        if pattern.search(base_url):
            return pattern.sub(
                lambda m: f"{m.group(1)}{name}={value}", base_url, count=1
            )
        sep = "&" if "?" in base_url else "?"
        return f"{base_url}{sep}{name}={value}"

    def _fetch_live_list(self):
        """êµ­íšŒ ìƒì¤‘ê³„ ëª©ë¡ APIì—ì„œ í˜„ì¬ ë°©ì†¡ ëª©ë¡ ê°€ì ¸ì˜¤ê¸°"""
        api_url = f"https://assembly.webcast.go.kr/main/service/live_list.asp?vv={int(time.time())}"
        try:
            req = Request(api_url, headers={"User-Agent": "Mozilla/5.0"})
            with urlopen(req, timeout=10) as response:
                payload = response.read().decode("utf-8", errors="replace")
            data = json.loads(payload)
            if isinstance(data, dict):
                return data.get("xlist", [])
        except (HTTPError, URLError, json.JSONDecodeError) as e:
            logger.debug(f"live_list API ì˜¤ë¥˜: {e}")
        except Exception as e:
            logger.debug(f"live_list ì²˜ë¦¬ ì˜¤ë¥˜: {e}")
        return []

    def _show_live_dialog(self):
        """ìƒì¤‘ê³„ ëª©ë¡ ë‹¤ì´ì–¼ë¡œê·¸ í‘œì‹œ"""
        dialog = LiveBroadcastDialog(self)
        dialog.finished.connect(dialog.deleteLater)
        if dialog.exec():
            data = dialog.selected_broadcast
            if data:
                # ì„ íƒëœ ë°©ì†¡ ì •ë³´ë¡œ URL ìƒì„±
                # player.asp?xcode={xcode}&xcgcd={xcgcd} í˜•ì‹ì´ ê°€ì¥ ì•ˆì •ì ì„
                xcode = str(data.get("xcode", "")).strip()
                xcgcd = str(data.get("xcgcd", "")).strip()

                if xcode and xcgcd:
                    # ê¸°ë³¸ URL
                    base_url = "https://assembly.webcast.go.kr/main/player.asp"
                    new_url = f"{base_url}?xcode={xcode}&xcgcd={xcgcd}"

                    # ì½¤ë³´ë°•ìŠ¤ì— ì„¤ì •
                    name = data.get("xname", "").strip()
                    self._add_to_history(new_url, name)
                    idx = self.url_combo.findData(new_url)
                    if idx >= 0:
                        self.url_combo.setCurrentIndex(idx)
                    else:
                        self.url_combo.setEditText(new_url)

                    # íƒœê·¸ ìë™ ì„¤ì • (ë°©ì†¡ëª…) - _add_to_historyì—ì„œ ì²˜ë¦¬ë¨
                    # ë°”ë¡œ ì‹œì‘í• ì§€ ë¬¼ì–´ë³´ëŠ” ê²ƒë„ ì¢‹ì§€ë§Œ, ì¼ë‹¨ URLë§Œ ì±„ì›Œì¤Œ
                    # Toast ì•Œë¦¼
                    self._show_toast(
                        f"ë°©ì†¡ì´ ì„ íƒë˜ì—ˆìŠµë‹ˆë‹¤:\n{name}", toast_type="success"
                    )

    def _resolve_live_url_from_list(self, original_url: str, target_xcode: str) -> str:
        """live_list APIë¡œ xcgcd/xcodeë¥¼ ë³´ì™„í•˜ì—¬ URL ìƒì„±"""
        broadcasts = self._fetch_live_list()
        if not broadcasts:
            return original_url

        current_xcgcd = self._get_query_param(original_url, "xcgcd").strip()
        target_norm = target_xcode.strip().upper() if target_xcode else ""

        # xcgcdê°€ ìˆìœ¼ë©´ ìœ íš¨í•œ ê²ƒìœ¼ë¡œ ê°„ì£¼ (xstat ìƒíƒœì™€ ë¬´ê´€í•˜ê²Œ ì‹œë„)
        def is_valid_broadcast(item):
            return bool(str(item.get("xcgcd", "")).strip())

        if current_xcgcd and not target_norm:
            for bc in broadcasts:
                bc_xcgcd = str(bc.get("xcgcd", "")).strip()
                if bc_xcgcd and bc_xcgcd == current_xcgcd:
                    bc_xcode = str(bc.get("xcode", "")).strip()
                    if bc_xcode:
                        new_url = self._set_query_param(original_url, "xcode", bc_xcode)
                        logger.info(f"live_listë¡œ xcode ë³´ì™„: xcode={bc_xcode}")
                        return new_url

        if target_norm:
            for bc in broadcasts:
                bc_xcode = str(bc.get("xcode", "")).strip()
                # xcodeê°€ ì¼ì¹˜í•˜ê³  xcgcdê°€ ìˆìœ¼ë©´ ì‚¬ìš© (xstat ì¡°ê±´ ì™„í™”)
                if bc_xcode.upper() == target_norm and is_valid_broadcast(bc):
                    xcgcd = str(bc.get("xcgcd", "")).strip()
                    new_url = self._set_query_param(original_url, "xcgcd", xcgcd)
                    if not self._get_query_param(new_url, "xcode"):
                        new_url = self._set_query_param(new_url, "xcode", bc_xcode)
                    logger.info(f"live_list ë§¤ì¹­ ì„±ê³µ: xcode={bc_xcode}, xcgcd={xcgcd}")
                    return new_url
            logger.warning(f"live_listì—ì„œ xcode={target_norm} ìƒì¤‘ê³„ ë¯¸ë°œê²¬")
        else:
            for bc in broadcasts:
                if is_valid_broadcast(bc):
                    xcgcd = str(bc.get("xcgcd", "")).strip()
                    new_url = self._set_query_param(original_url, "xcgcd", xcgcd)
                    if not self._get_query_param(new_url, "xcode"):
                        bc_xcode = str(bc.get("xcode", "")).strip()
                        if bc_xcode:
                            new_url = self._set_query_param(new_url, "xcode", bc_xcode)
                    logger.info(f"live_list ì²« ìƒì¤‘ê³„ ì‚¬ìš©: xcgcd={xcgcd}")
                    return new_url

        return original_url

    def _detect_live_broadcast(self, driver, original_url: str) -> str:
        """í˜„ì¬ ì§„í–‰ ì¤‘ì¸ ìƒì¤‘ê³„ì˜ xcgcdë¥¼ ìë™ ê°ì§€

        Args:
            driver: Selenium WebDriver
            original_url: ì›ë˜ ìš”ì²­ëœ URL

        Returns:
            str: ê°ì§€ëœ xcgcdë¥¼ í¬í•¨í•œ URL, ê°ì§€ ì‹¤íŒ¨ ì‹œ ì›ë˜ URL ë°˜í™˜
        """
        try:
            existing_xcgcd = self._get_query_param(original_url, "xcgcd").strip()
            existing_xcode = self._get_query_param(original_url, "xcode").strip()

            # ì›ë˜ URLì— xcode/xcgcdê°€ ëª¨ë‘ ìˆìœ¼ë©´ ê·¸ëŒ€ë¡œ ì‚¬ìš©
            if existing_xcgcd and existing_xcode:
                logger.info(f"URLì— ì´ë¯¸ xcode/xcgcd í¬í•¨ë¨: {original_url}")
                return original_url

            self.message_queue.put(("status", "ğŸ” í˜„ì¬ ìƒì¤‘ê³„ ê°ì§€ ì¤‘..."))

            # xcode ì¶”ì¶œ (ëª¨ë“  ë°©ë²•ì—ì„œ ê³µí†µìœ¼ë¡œ ì‚¬ìš©)
            target_xcode = existing_xcode or None
            target_xcode_norm = target_xcode.upper() if target_xcode else None
            logger.info(f"xcgcd íƒìƒ‰ ì‹œì‘ - target_xcode: {target_xcode}")

            # ë°©ë²• 0: live_list APIë¡œ ìƒì¤‘ê³„ ì •ë³´ í™•ì¸ (ì‚¬ì´íŠ¸ êµ¬ì¡° ê¸°ë°˜)
            resolved_url = self._resolve_live_url_from_list(original_url, target_xcode)
            if resolved_url != original_url and self._get_query_param(
                resolved_url, "xcgcd"
            ):
                logger.info(f"live_list ê¸°ë°˜ URL ê°ì§€ ì„±ê³µ: {resolved_url}")
                return resolved_url

            # xcgcdì—ì„œ xcode ì¶”ì¶œí•˜ëŠ” í—¬í¼ í•¨ìˆ˜
            # xcgcd í˜•ì‹: DCM0000XX... ì—¬ê¸°ì„œ XXê°€ xcode ë¶€ë¶„ (ì˜ˆ: IO, 25 ë“±)
            def extract_xcode_from_xcgcd(xcgcd_val):
                """xcgcd ê°’ì—ì„œ xcode ë¶€ë¶„ ì¶”ì¶œ ì‹œë„"""
                if not xcgcd_val:
                    return None
                # DCM0000 ì ‘ë‘ì‚¬ ì´í›„ ë¶€ë¶„ ì¶”ì¶œ ì‹œë„
                # ì˜ˆ: DCM0000IO224310401 -> IO
                # ì˜ˆ: DCM000025224310401 -> 25
                match = re.search(r"DCM0000([A-Za-z0-9]+)", xcgcd_val)
                if match:
                    code = match.group(1)
                    # ìˆ«ì+ë‚˜ë¨¸ì§€ íŒ¨í„´ (ì˜ˆ: 25224310401 -> 25)
                    num_match = re.match(r"^(\d{2})", code)
                    if num_match:
                        return num_match.group(1)
                    # ë¬¸ì+ë‚˜ë¨¸ì§€ íŒ¨í„´ (ì˜ˆ: IO224310401 -> IO)
                    alpha_match = re.match(r"^([A-Za-z]+)", code)
                    if alpha_match:
                        return alpha_match.group(1)
                return None

            # ë°©ë²• 1: í˜„ì¬ í˜ì´ì§€ì˜ JavaScript ë³€ìˆ˜ì—ì„œ xcgcd ê°€ì ¸ì˜¤ê¸°
            scripts = [
                # ì „ì—­ ë³€ìˆ˜ì—ì„œ xcgcd ì°¾ê¸°
                "return typeof xcgcd !== 'undefined' ? xcgcd : null;",
                "return typeof XCGCD !== 'undefined' ? XCGCD : null;",
                "return window.xcgcd || null;",
                "return window.XCGCD || null;",
                # URL íŒŒë¼ë¯¸í„°ì—ì„œ ì¶”ì¶œ
                "return new URLSearchParams(window.location.search).get('xcgcd');",
                # í˜„ì¬ ìŠ¤íŠ¸ë¦¼ ì •ë³´ì—ì„œ ì¶”ì¶œ
                "if(typeof streamInfo !== 'undefined' && streamInfo.xcgcd) return streamInfo.xcgcd; return null;",
                # í”Œë ˆì´ì–´ ì •ë³´ì—ì„œ ì¶”ì¶œ
                "if(typeof playerConfig !== 'undefined' && playerConfig.xcgcd) return playerConfig.xcgcd; return null;",
            ]

            xcgcd = None
            for script in scripts:
                try:
                    result = driver.execute_script(script)
                    if result:
                        found_xcgcd = str(result)
                        # target_xcodeê°€ ìˆìœ¼ë©´ xcgcdì˜ xcode ë¶€ë¶„ ê²€ì¦
                        if target_xcode:
                            found_xcode = extract_xcode_from_xcgcd(found_xcgcd)
                            if (
                                found_xcode
                                and target_xcode_norm
                                and found_xcode.upper() != target_xcode_norm
                            ):
                                logger.warning(
                                    f"JavaScript xcgcdì˜ xcode({found_xcode})ê°€ target({target_xcode})ì™€ ë¶ˆì¼ì¹˜ - ë¬´ì‹œ"
                                )
                                continue
                        xcgcd = found_xcgcd
                        logger.info(f"JavaScriptì—ì„œ xcgcd ë°œê²¬: {xcgcd}")
                        break
                except Exception as e:
                    logger.debug(f"Script ì‹¤í–‰ ì˜¤ë¥˜: {e}")

            # ë°©ë²• 2: URLì´ ë¦¬ë‹¤ì´ë ‰íŠ¸ ë˜ì—ˆëŠ”ì§€ í™•ì¸
            if not xcgcd:
                current_url = driver.current_url
                found_xcgcd = self._get_query_param(current_url, "xcgcd").strip()
                if found_xcgcd:
                    # target_xcode ê²€ì¦
                    if target_xcode:
                        found_xcode = extract_xcode_from_xcgcd(found_xcgcd)
                        if (
                            found_xcode
                            and target_xcode_norm
                            and found_xcode.upper() != target_xcode_norm
                        ):
                            logger.warning(
                                f"ë¦¬ë‹¤ì´ë ‰íŠ¸ëœ xcgcdì˜ xcode({found_xcode})ê°€ target({target_xcode})ì™€ ë¶ˆì¼ì¹˜ - ë¬´ì‹œ"
                            )
                        else:
                            xcgcd = found_xcgcd
                            logger.info(f"ë¦¬ë‹¤ì´ë ‰íŠ¸ëœ URLì—ì„œ xcgcd ë°œê²¬: {xcgcd}")
                    else:
                        xcgcd = found_xcgcd
                        logger.info(f"ë¦¬ë‹¤ì´ë ‰íŠ¸ëœ URLì—ì„œ xcgcd ë°œê²¬: {xcgcd}")

            # ë°©ë²• 3: í˜ì´ì§€ ë‚´ ìƒì¤‘ê³„ ëª©ë¡ì—ì„œ í˜„ì¬ ë°©ì†¡ ì°¾ê¸°
            # ì£¼ì˜: ì´ í˜ì´ì§€ì— ì—¬ëŸ¬ ë°©ì†¡ ë§í¬ê°€ ìˆì„ ìˆ˜ ìˆìœ¼ë¯€ë¡œ xcode ê²€ì¦ í•„ìš”
            # (target_xcodeëŠ” ì´ë¯¸ ìœ„ì—ì„œ ì¶”ì¶œë¨)
            if not xcgcd:
                try:
                    if target_xcode:
                        # target_xcodeê°€ ìˆìœ¼ë©´ í•´ë‹¹ xcodeê°€ í¬í•¨ëœ ë§í¬ë§Œ ê²€ìƒ‰
                        script = f"""
                        var links = document.querySelectorAll('a[href*="xcode={target_xcode}"][href*="xcgcd="]');
                        for(var i=0; i<links.length; i++) {{
                            var href = links[i].getAttribute('href');
                            var match = href.match(/xcgcd=([^&]+)/);
                            if(match) return match[1];
                        }}
                        return null;
                        """
                        result = driver.execute_script(script)
                        if result:
                            xcgcd = str(result)
                            logger.info(
                                f"í˜ì´ì§€ ìš”ì†Œì—ì„œ xcode={target_xcode} ë§¤ì¹­ xcgcd ë°œê²¬: {xcgcd}"
                            )
                    else:
                        # target_xcodeê°€ ì—†ìœ¼ë©´(ë³¸íšŒì˜ ë“±) ê¸°ì¡´ ë¡œì§ ì‚¬ìš©
                        live_scripts = [
                            """
                            var links = document.querySelectorAll('a[href*="xcgcd="]');
                            for(var i=0; i<links.length; i++) {
                                var href = links[i].getAttribute('href');
                                if(href && href.includes('xcgcd=')) {
                                    var match = href.match(/xcgcd=([^&]+)/);
                                    if(match) return match[1];
                                }
                            }
                            return null;
                            """,
                            """
                            var iframe = document.querySelector('iframe[src*="xcgcd="]');
                            if(iframe) {
                                var src = iframe.getAttribute('src');
                                var match = src.match(/xcgcd=([^&]+)/);
                                if(match) return match[1];
                            }
                            return null;
                            """,
                            """
                            var input = document.querySelector('input[name="xcgcd"], input#xcgcd');
                            if(input) return input.value;
                            return null;
                            """,
                        ]

                        for script in live_scripts:
                            result = driver.execute_script(script)
                            if result:
                                xcgcd = str(result)
                                logger.info(f"í˜ì´ì§€ ìš”ì†Œì—ì„œ xcgcd ë°œê²¬: {xcgcd}")
                                break
                except Exception as e:
                    logger.debug(f"ìƒì¤‘ê³„ ëª©ë¡ íŒŒì‹± ì˜¤ë¥˜: {e}")

            # ë°©ë²• 4: ë©”ì¸ í˜ì´ì§€ì—ì„œ ì˜¤ëŠ˜ì˜ ìƒì¤‘ê³„ ì •ë³´ ê°€ì ¸ì˜¤ê¸° (ê°œì„ ë¨)
            # (target_xcodeëŠ” ì´ë¯¸ ìœ„ì—ì„œ ì¶”ì¶œë¨)
            if not xcgcd:
                navigated_to_main = False
                try:
                    # ë©”ì¸ í˜ì´ì§€ë¡œ ì´ë™
                    main_url = "https://assembly.webcast.go.kr/main/"
                    self.message_queue.put(
                        ("status", "ğŸ” ë©”ì¸ í˜ì´ì§€ì—ì„œ ìƒì¤‘ê³„ ëª©ë¡ í™•ì¸ ì¤‘...")
                    )
                    driver.get(main_url)
                    navigated_to_main = True

                    # ë™ì  ì½˜í…ì¸  ë¡œë”© ëŒ€ê¸° (ìµœëŒ€ 10ì´ˆ)
                    try:
                        wait = WebDriverWait(driver, 10)
                        wait.until(
                            EC.presence_of_element_located(
                                (By.CSS_SELECTOR, 'a[href*="xcgcd="]')
                            )
                        )
                    except Exception:
                        # íƒ€ì„ì•„ì›ƒ ì‹œ ê¸°ë³¸ ëŒ€ê¸° (ì¢…ë£Œ ì‹ í˜¸ì— ì¦‰ì‹œ ë°˜ì‘)
                        self.stop_event.wait(timeout=3)

                    # ëª¨ë“  ìƒì¤‘ê³„ ë§í¬ ìˆ˜ì§‘
                    all_broadcasts_script = """
                    var results = [];
                    var links = document.querySelectorAll('a[href*="xcgcd="]');
                    for(var i=0; i<links.length; i++) {
                        var href = links[i].getAttribute('href');
                        var text = links[i].innerText || links[i].textContent || '';
                        if(href && href.includes('xcgcd=')) {
                            var xcgcdMatch = href.match(/xcgcd=([^&]+)/);
                            var xcodeMatch = href.match(/xcode=([^&]+)/);
                            if(xcgcdMatch) {
                                results.push({
                                    xcgcd: xcgcdMatch[1],
                                    xcode: xcodeMatch ? xcodeMatch[1] : null,
                                    text: text.trim()
                                });
                            }
                        }
                    }
                    return JSON.stringify(results);
                    """

                    broadcasts_json = driver.execute_script(all_broadcasts_script)
                    broadcasts = json.loads(broadcasts_json) if broadcasts_json else []
                    logger.info(f"ë°œê²¬ëœ ìƒì¤‘ê³„ ëª©ë¡: {len(broadcasts)}ê°œ")

                    if broadcasts:
                        # target_xcodeê°€ ìˆìœ¼ë©´ í•´ë‹¹ xcode ë§¤ì¹­ ìš°ì„ 
                        if target_xcode:
                            for bc in broadcasts:
                                bc_xcode = str(bc.get("xcode", "")).strip()
                                if (
                                    target_xcode_norm
                                    and bc_xcode.upper() == target_xcode_norm
                                ):
                                    xcgcd = bc["xcgcd"]
                                    logger.info(
                                        f"xcode={target_xcode} ë§¤ì¹­ ì„±ê³µ: xcgcd={xcgcd}"
                                    )
                                    break
                            # target_xcodeê°€ ìˆëŠ”ë° ë§¤ì¹­ ì‹¤íŒ¨í•˜ë©´ xcgcdë¥¼ ì„¤ì •í•˜ì§€ ì•ŠìŒ
                            if not xcgcd:
                                logger.warning(
                                    f"xcode={target_xcode}ì— í•´ë‹¹í•˜ëŠ” ìƒì¤‘ê³„ë¥¼ ì°¾ì§€ ëª»í•¨"
                                )
                        else:
                            # target_xcodeê°€ ì—†ëŠ” ê²½ìš°(ë³¸íšŒì˜ ë“±)ì—ë§Œ ì²« ë²ˆì§¸ ìƒì¤‘ê³„ ì‚¬ìš©
                            xcgcd = broadcasts[0]["xcgcd"]
                            first_bc = broadcasts[0]
                            logger.info(
                                f"ì²« ë²ˆì§¸ ìƒì¤‘ê³„ ì‚¬ìš©: xcgcd={xcgcd}, text={first_bc.get('text', '')[:30]}"
                            )

                except Exception as e:
                    logger.debug(f"ë©”ì¸ í˜ì´ì§€ ì¡°íšŒ ì˜¤ë¥˜: {e}")

            # ë°©ë²• 5: ë©”ì¸ í˜ì´ì§€ ë¦¬ë‹¤ì´ë ‰íŠ¸ ì‹œ í™”ë©´ì˜ 'ìƒì¤‘ê³„' ë²„íŠ¼ ìë™ í´ë¦­ (ê°œì„ ë¨)
            if not xcgcd and target_xcode:
                # í˜„ì¬ URLì´ ë©”ì¸ í˜ì´ì§€ì¸ì§€ í™•ì¸ (player.aspê°€ ì•„ë‹˜)
                current_url = driver.current_url
                if "/main/player.asp" not in current_url:
                    try:
                        self.message_queue.put(
                            (
                                "status",
                                f"ğŸ–±ï¸ ë©”ì¸ í™”ë©´ì—ì„œ xcode={target_xcode} ë²„íŠ¼ íƒìƒ‰ ì¤‘...",
                            )
                        )
                        logger.info(
                            f"ë©”ì¸ í˜ì´ì§€ ë¦¬ë‹¤ì´ë ‰íŠ¸ ê°ì§€ - ë²„íŠ¼ í´ë¦­ ì‹œë„ (xcode={target_xcode})"
                        )

                        # ë™ì  ì½˜í…ì¸  ë¡œë”© ëŒ€ê¸° (ìµœëŒ€ 10ì´ˆ) - onair í´ë˜ìŠ¤ê°€ ë‚˜íƒ€ë‚  ë•Œê¹Œì§€
                        try:
                            WebDriverWait(driver, 10).until(
                                EC.presence_of_element_located(
                                    (By.CSS_SELECTOR, ".onair")
                                )
                            )
                        except Exception:
                            logger.debug(
                                "onair ìš”ì†Œ ëŒ€ê¸° íƒ€ì„ì•„ì›ƒ (ìƒì¤‘ê³„ê°€ ì—†ê±°ë‚˜ ë¡œë”© ì§€ì—°)"
                            )

                        # 1. onair ë²„íŠ¼ ì°¾ê¸° (xcode ë§¤ì¹­)
                        # ë‹¤ì–‘í•œ ì„ íƒì ì‹œë„ - ë” í¬ê´„ì ìœ¼ë¡œ
                        selectors = [
                            f'a.onair[href*="xcode={target_xcode}"]',
                            f'a.btn[href*="xcode={target_xcode}"]',  # onair í´ë˜ìŠ¤ê°€ ì—†ì„ ìˆ˜ë„ ìˆìŒ
                            f'div.onair a[href*="xcode={target_xcode}"]',
                            f'a[href*="xcode={target_xcode}"]:has(.icon_onair)',
                            f'a[href*="xcode={target_xcode}"]',  # ìµœí›„ì˜ ìˆ˜ë‹¨: ê·¸ëƒ¥ ë§í¬ ì°¾ê¸°
                        ]

                        btn = None
                        for sel in selectors:
                            try:
                                elems = driver.find_elements(By.CSS_SELECTOR, sel)
                                # onair í´ë˜ìŠ¤ê°€ ìˆëŠ” ìš”ì†Œ ìš°ì„  ì„ íƒ
                                for elem in elems:
                                    if "onair" in elem.get_attribute(
                                        "class"
                                    ) or elem.find_elements(By.CSS_SELECTOR, ".onair"):
                                        btn = elem
                                        break

                                if btn:
                                    break

                                # onairê°€ ì—†ë”ë¼ë„ ì²« ë²ˆì§¸ ìš”ì†Œ ì„ íƒ (í´ë¦­í•´ë³´ëŠ” ê²ƒì´ ë‚˜ìŒ)
                                if elems and not btn:
                                    btn = elems[0]
                                    break
                            except Exception:
                                continue

                        if btn:
                            # 2. ìŠ¤í¬ë¡¤í•˜ì—¬ ìš”ì†Œ ë³´ì´ê²Œ í•˜ê¸°
                            driver.execute_script(
                                "arguments[0].scrollIntoView({block: 'center'});", btn
                            )
                            # ìŠ¤í¬ë¡¤ í›„ ì•ˆì •í™” ëŒ€ê¸° (ì¢…ë£Œ ì‹ í˜¸ì— ì¦‰ì‹œ ë°˜ì‘)
                            self.stop_event.wait(timeout=1.0)

                            # 3. í´ë¦­ (JavaScript ì‚¬ìš©ì´ ë” ì•ˆì •ì )
                            driver.execute_script("arguments[0].click();", btn)
                            logger.info(
                                f"ë©”ì¸ í˜ì´ì§€ì—ì„œ ìƒì¤‘ê³„ ë²„íŠ¼ ìë™ í´ë¦­ ì„±ê³µ: xcode={target_xcode}"
                            )
                            self.message_queue.put(
                                ("status", "âœ… ìƒì¤‘ê³„ ë²„íŠ¼ ìë™ í´ë¦­ ì„±ê³µ")
                            )

                            # 4. í˜ì´ì§€ ì „í™˜ ëŒ€ê¸°
                            try:
                                WebDriverWait(driver, 5).until(
                                    lambda d: "player.asp" in d.current_url
                                )
                                # í˜ì´ì§€ ì „í™˜ í›„ URL ë°˜í™˜
                                return driver.current_url
                            except Exception:
                                logger.warning("ë²„íŠ¼ í´ë¦­ í›„ í˜ì´ì§€ ì „í™˜ íƒ€ì„ì•„ì›ƒ")
                        else:
                            logger.warning(
                                f"ë©”ì¸ í˜ì´ì§€ì—ì„œ xcode={target_xcode} ìƒì¤‘ê³„ ë²„íŠ¼ì„ ì°¾ì„ ìˆ˜ ì—†ìŒ"
                            )

                    except Exception as e:
                        logger.warning(f"ìƒì¤‘ê³„ ë²„íŠ¼ í´ë¦­ ë¡œì§ ì‹¤íŒ¨: {e}")
                    finally:
                        # ë²„íŠ¼ í´ë¦­ë„ ì‹¤íŒ¨í•œ ê²½ìš° ì›ë˜ URLë¡œ ë³µê·€
                        # ë‹¨, ë²„íŠ¼ í´ë¦­ìœ¼ë¡œ í˜ì´ì§€ê°€ ì´ë™í–ˆë‹¤ë©´ ë³µê·€í•˜ì§€ ì•ŠìŒ
                        if (
                            navigated_to_main
                            and not xcgcd
                            and "/main/player.asp" not in driver.current_url
                        ):
                            try:
                                # ì´ë¯¸ ë²„íŠ¼ í´ë¦­ ì‹œë„ë¡œ URLì´ ë°”ë€Œì—ˆì„ ìˆ˜ ìˆìœ¼ë¯€ë¡œ ì²´í¬
                                if original_url not in driver.current_url:
                                    driver.get(original_url)
                                    self.stop_event.wait(timeout=2)
                                    logger.info(f"ì›ë˜ URLë¡œ ë³µê·€: {original_url}")
                            except Exception as e:
                                logger.debug(f"ì›ë˜ URL ë³µê·€ ì‹¤íŒ¨: {e}")

            # xcgcdë¥¼ ì°¾ì•˜ìœ¼ë©´ URL ì—…ë°ì´íŠ¸ (ìœ íš¨ì„± ê²€ì¦ í¬í•¨)
            if (
                xcgcd and len(xcgcd) >= 10
            ):  # ìµœì†Œ ê¸¸ì´ ê²€ì¦ (ìœ íš¨í•œ xcgcdëŠ” ë³´í†µ 20ì ì´ìƒ)
                new_url = self._set_query_param(original_url, "xcgcd", xcgcd)
                if not self._get_query_param(new_url, "xcode"):
                    inferred_xcode = target_xcode or extract_xcode_from_xcgcd(xcgcd)
                    if inferred_xcode:
                        new_url = self._set_query_param(
                            new_url, "xcode", inferred_xcode
                        )

                display_xcgcd = xcgcd[:15] + "..." if len(xcgcd) > 15 else xcgcd
                self.message_queue.put(
                    ("status", f"âœ… ìƒì¤‘ê³„ ê°ì§€ ì„±ê³µ! (xcgcd={display_xcgcd})")
                )
                logger.info(f"ìƒì¤‘ê³„ URL ì—…ë°ì´íŠ¸: {new_url}")
                return new_url
            else:
                # target_xcode ì •ë³´ë¥¼ í¬í•¨í•˜ì—¬ ë” êµ¬ì²´ì ì¸ ë©”ì‹œì§€ í‘œì‹œ
                target_xcode = (
                    self._get_query_param(original_url, "xcode").strip() or None
                )

                if target_xcode:
                    self.message_queue.put(
                        (
                            "status",
                            f"âš ï¸ xcode={target_xcode} ìœ„ì›íšŒì˜ ì§„í–‰ ì¤‘ì¸ ìƒì¤‘ê³„ë¥¼ ì°¾ì„ ìˆ˜ ì—†ìŒ",
                        )
                    )
                    logger.warning(
                        f"xcode={target_xcode}ì— í•´ë‹¹í•˜ëŠ” ìƒì¤‘ê³„ê°€ ì—†ìŒ, ì›ë˜ URL ì‚¬ìš©"
                    )
                else:
                    self.message_queue.put(
                        ("status", "âš ï¸ ìƒì¤‘ê³„ ì •ë³´ë¥¼ ì°¾ì„ ìˆ˜ ì—†ìŒ - ê¸°ë³¸ URL ì‚¬ìš©")
                    )
                    logger.warning("ìƒì¤‘ê³„ xcgcdë¥¼ ì°¾ì„ ìˆ˜ ì—†ìŒ, ì›ë˜ URL ì‚¬ìš©")
                return original_url

        except Exception as e:
            logger.error(f"ìƒì¤‘ê³„ ê°ì§€ ì˜¤ë¥˜: {e}")
            self.message_queue.put(("status", f"âš ï¸ ìƒì¤‘ê³„ ê°ì§€ ì‹¤íŒ¨: {e}"))
            return original_url

    def _get_reconnect_delay(self, attempt: int) -> float:
        """ì§€ìˆ˜ ë°±ì˜¤í”„ ê¸°ë°˜ ì¬ì—°ê²° ëŒ€ê¸° ì‹œê°„(ì´ˆ) ê³„ì‚°"""
        if attempt <= 0:
            return 0.0
        delay = Config.RECONNECT_BASE_DELAY * (2 ** (attempt - 1))
        return min(delay, Config.RECONNECT_MAX_DELAY)

    def _is_recoverable_webdriver_error(self, error: Exception) -> bool:
        """ì¬ì—°ê²°ë¡œ ë³µêµ¬ ê°€ëŠ¥í•œ ì›¹ë“œë¼ì´ë²„ ì˜¤ë¥˜ì¸ì§€ íŒë‹¨"""
        msg = str(error).lower()
        markers = [
            "invalid session",
            "no such execution context",
            "chrome not reachable",
            "disconnected",
            "target closed",
            "session deleted",
            "connection reset",
            "connection refused",
            "web view not found",
        ]
        return any(marker in msg for marker in markers)

    def _ping_driver(self, driver):
        """ì›¹ë“œë¼ì´ë²„ ì‘ë‹µ ì‹œê°„ì„ ì¸¡ì • (ms). ì‹¤íŒ¨ ì‹œ None ë°˜í™˜."""
        start = time.time()
        try:
            driver.execute_script("return 1")
        except Exception:
            return None
        return int((time.time() - start) * 1000)

    def _extraction_worker(self, url, selector, headless):
        """ìë§‰ ì¶”ì¶œ ì›Œì»¤ ìŠ¤ë ˆë“œ (Legacy Logic Restoration)"""
        driver = None

        try:
            options = Options()
            options.add_argument("--log-level=3")
            options.add_argument("--disable-blink-features=AutomationControlled")
            options.add_experimental_option(
                "excludeSwitches", ["enable-logging", "enable-automation"]
            )
            options.add_experimental_option("useAutomationExtension", False)

            # í—¤ë“œë¦¬ìŠ¤ ëª¨ë“œ
            if headless:
                options.add_argument("--headless=new")
                options.add_argument("--window-size=1280,720")
                self.message_queue.put(("status", "í—¤ë“œë¦¬ìŠ¤ ëª¨ë“œë¡œ ì‹œì‘ ì¤‘..."))

            try:
                driver = webdriver.Chrome(options=options)
                self.driver = driver
                self.message_queue.put(("status", "Chrome ì‹œì‘ ì™„ë£Œ"))
            except Exception as e:
                self.message_queue.put(("error", f"Chrome ì˜¤ë¥˜: {e}"))
                return

            self.message_queue.put(("status", "í˜ì´ì§€ ë¡œë”© ì¤‘..."))
            driver.get(url)
            self.stop_event.wait(timeout=3)

            self.message_queue.put(("status", "AI ìë§‰ í™œì„±í™” ì¤‘..."))
            self._activate_subtitle(driver)

            self.message_queue.put(("status", "ìë§‰ ìš”ì†Œ ê²€ìƒ‰ ì¤‘..."))
            wait = WebDriverWait(driver, 20)

            found = False
            selector_candidates = self._build_subtitle_selector_candidates(selector)
            active_selector = ""
            for sel in selector_candidates:
                try:
                    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, sel)))
                    self.message_queue.put(("status", f"ìë§‰ ìš”ì†Œ ì°¾ìŒ: {sel}"))
                    active_selector = sel
                    found = True
                    break
                except Exception:
                    continue

            if not found:
                detected_selector = self._find_subtitle_selector(driver)
                if detected_selector:
                    selector_candidates = self._build_subtitle_selector_candidates(
                        detected_selector, selector_candidates
                    )
                    active_selector = selector_candidates[0]
                    self.message_queue.put(
                        ("status", f"ìë§‰ ìš”ì†Œ ìë™ ê°ì§€: {active_selector}")
                    )
                    found = True

            if not found:
                self.message_queue.put(("error", "ìë§‰ ìš”ì†Œë¥¼ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤."))
                return

            self.message_queue.put(("status", "ìë§‰ ëª¨ë‹ˆí„°ë§ ì¤‘"))

            # MutationObserver ì£¼ì… (í•˜ì´ë¸Œë¦¬ë“œ ì•„í‚¤í…ì²˜)
            observer_active, observer_frame_path = self._inject_mutation_observer(
                driver, ",".join(selector_candidates)
            )
            observer_retry_interval = 3.0
            last_observer_retry = time.time()
            last_selector_refresh = time.time()

            last_check = time.time()
            last_connection_check = time.time()
            # [Fix] ì§€ì—­ ë³€ìˆ˜ë¡œ ë³€ê²½ - ìŠ¤ë ˆë“œ ì•ˆì „ì„± í™•ë³´ (Race condition ë°©ì§€)
            worker_last_raw_text = ""
            worker_last_raw_compact = ""
            reconnect_attempt = 0

            # stop_event ì‚¬ìš©ìœ¼ë¡œ ë” ë¹ ë¥¸ ì¢…ë£Œ ì‘ë‹µ
            while not self.stop_event.is_set():
                try:
                    now = time.time()

                    # ì—°ê²° ìƒíƒœ ëª¨ë‹ˆí„°ë§ (#5) - 5ì´ˆë§ˆë‹¤ ì²´í¬
                    if now - last_connection_check >= 5.0:
                        ping_time = self._ping_driver(driver)
                        if ping_time is not None:
                            self.message_queue.put(
                                (
                                    "connection_status",
                                    {"status": "connected", "latency": ping_time},
                                )
                            )
                            reconnect_attempt = 0  # ì—°ê²° ì„±ê³µ ì‹œ ì¬ì—°ê²° íšŸìˆ˜ ì´ˆê¸°í™”
                        else:
                            self.message_queue.put(
                                ("connection_status", {"status": "disconnected"})
                            )
                        last_connection_check = now

                    if now - last_check >= 0.2:
                        changes_processed = False

                        # 1ë‹¨ê³„: MutationObserver ë²„í¼ì—ì„œ ìˆ˜ì§‘ (ì´ë²¤íŠ¸ ê¸°ë°˜)
                        if observer_active:
                            observer_changes = self._collect_observer_changes(
                                driver, observer_frame_path
                            )
                            if observer_changes is None:
                                # Observerê°€ ì£½ì—ˆìœ¼ë©´ ë¹„í™œì„±í™” í›„ í´ë§ fallback
                                observer_active = False
                                logger.warning("MutationObserver ë¹„í™œì„±í™”, í´ë§ fallback")
                            elif observer_changes:
                                for change_text in observer_changes:
                                    # í´ë¦¬ì–´ ë§ˆì»¤ ê°ì§€ (ë°œì–¸ì ì „í™˜)
                                    if change_text == "__SUBTITLE_CLEARED__":
                                        self.message_queue.put(
                                            ("subtitle_reset", "observer_cleared")
                                        )
                                        worker_last_raw_text = ""
                                        worker_last_raw_compact = ""
                                        continue
                                    c_text = utils.clean_text_display(change_text)
                                    c_compact = utils.compact_subtitle_text(c_text)
                                    if (
                                        c_text
                                        and c_compact
                                        and c_compact != worker_last_raw_compact
                                    ):
                                        worker_last_raw_text = c_text
                                        worker_last_raw_compact = c_compact
                                        self.message_queue.put(("preview", c_text))
                                        changes_processed = True

                        # 2ë‹¨ê³„: Observer ê²°ê³¼ê°€ ì—†ìœ¼ë©´ ê¸°ì¡´ í´ë§ fallback
                        if not changes_processed:
                            text, matched_selector, selector_found = (
                                self._read_subtitle_text_by_selectors(
                                    driver, selector_candidates
                                )
                            )
                            if selector_found:
                                reconnect_attempt = 0
                                if matched_selector and matched_selector != active_selector:
                                    active_selector = matched_selector
                                    selector_candidates = (
                                        self._build_subtitle_selector_candidates(
                                            active_selector, selector_candidates
                                        )
                                    )
                            elif now - last_selector_refresh >= 5.0:
                                # ì…€ë ‰í„° ë³€í™” ëŒ€ì‘: ì£¼ê¸°ì  ìë™ ì¬íƒìƒ‰
                                detected_selector = self._find_subtitle_selector(driver)
                                if detected_selector:
                                    selector_candidates = (
                                        self._build_subtitle_selector_candidates(
                                            detected_selector, selector_candidates
                                        )
                                    )
                                    active_selector = selector_candidates[0]
                                    logger.info(
                                        "ìë§‰ ì…€ë ‰í„° ìë™ ì „í™˜: %s", active_selector
                                    )
                                last_selector_refresh = now

                            if (
                                not observer_active
                                and now - last_observer_retry >= observer_retry_interval
                            ):
                                observer_active, observer_frame_path = self._inject_mutation_observer(
                                    driver, ",".join(selector_candidates)
                                )
                                last_observer_retry = now

                            text = utils.clean_text_display(text)
                            text_compact = utils.compact_subtitle_text(text)

                            if (
                                text
                                and text_compact
                                and text_compact != worker_last_raw_compact
                            ):
                                worker_last_raw_text = text
                                worker_last_raw_compact = text_compact
                                self.message_queue.put(("preview", text))
                            elif (
                                not text
                                and selector_found
                                and worker_last_raw_compact
                            ):
                                # í´ë§ì—ì„œë„ ë¹ˆ í…ìŠ¤íŠ¸ ê°ì§€ (ë°œì–¸ì ì „í™˜)
                                self.message_queue.put(
                                    ("subtitle_reset", "polling_cleared")
                                )
                                worker_last_raw_text = ""
                                worker_last_raw_compact = ""

                        last_check = now

                    # stop_event ëŒ€ê¸° (0.05ì´ˆ, ì¦‰ì‹œ ì‘ë‹µ ê°€ëŠ¥)
                    self.stop_event.wait(timeout=0.05)

                except Exception as e:
                    if self.stop_event.is_set():
                        break

                    # ìë™ ì¬ì—°ê²° ë¡œì§ (#4)
                    if (
                        self.auto_reconnect_enabled
                        and self._is_recoverable_webdriver_error(e)
                    ):
                        reconnect_attempt += 1
                        if reconnect_attempt <= Config.MAX_RECONNECT_ATTEMPTS:
                            delay = self._get_reconnect_delay(reconnect_attempt)
                            self.message_queue.put(
                                (
                                    "reconnecting",
                                    {
                                        "attempt": reconnect_attempt,
                                        "max_attempts": Config.MAX_RECONNECT_ATTEMPTS,
                                        "delay": delay,
                                    },
                                )
                            )
                            logger.warning(
                                f"WebDriver ì—°ê²° ì˜¤ë¥˜, {delay}ì´ˆ í›„ ì¬ì—°ê²° ì‹œë„ ({reconnect_attempt}/{Config.MAX_RECONNECT_ATTEMPTS})"
                            )

                            # ê¸°ì¡´ ë“œë¼ì´ë²„ ì •ë¦¬
                            if driver:
                                try:
                                    driver.quit()
                                except Exception as quit_err:
                                    logger.debug(
                                        f"ë“œë¼ì´ë²„ ì¢…ë£Œ ì‹¤íŒ¨, detached ëª©ë¡ì— ì¶”ê°€: {quit_err}"
                                    )
                                    with self._detached_drivers_lock:
                                        self._detached_drivers.append(driver)

                            # ëŒ€ê¸° í›„ ì¬ì—°ê²° (ì¢…ë£Œ ì‹ í˜¸ì— ì¦‰ì‹œ ë°˜ì‘)
                            if self.stop_event.wait(timeout=delay):
                                break

                            try:
                                driver = webdriver.Chrome(options=options)
                                self.driver = driver
                                driver.get(url)
                                if self.stop_event.wait(timeout=2):
                                    break
                                self._activate_subtitle(driver)
                                self.message_queue.put(
                                    (
                                        "status",
                                        f"âœ… ì¬ì—°ê²° ì„±ê³µ (ì‹œë„ {reconnect_attempt})",
                                    )
                                )
                                self.message_queue.put(
                                    ("connection_status", {"status": "connected"})
                                )
                                last_check = time.time()
                                last_connection_check = time.time()
                                detected_selector = self._find_subtitle_selector(driver)
                                if detected_selector:
                                    selector_candidates = (
                                        self._build_subtitle_selector_candidates(
                                            detected_selector, selector_candidates
                                        )
                                    )
                                    active_selector = selector_candidates[0]
                                # ì¬ì—°ê²° í›„ MutationObserver ì¬ì£¼ì…
                                observer_active, observer_frame_path = self._inject_mutation_observer(
                                    driver, ",".join(selector_candidates)
                                )
                                last_observer_retry = time.time()
                                continue
                            except Exception as reconnect_error:
                                logger.error(f"ì¬ì—°ê²° ì‹¤íŒ¨: {reconnect_error}")
                        else:
                            self.message_queue.put(
                                (
                                    "error",
                                    f"ìµœëŒ€ ì¬ì—°ê²° ì‹œë„ íšŸìˆ˜({Config.MAX_RECONNECT_ATTEMPTS}) ì´ˆê³¼",
                                )
                            )
                            break
                    else:
                        # ë³µêµ¬ ë¶ˆê°€ëŠ¥í•œ ì˜¤ë¥˜
                        logger.warning(f"ëª¨ë‹ˆí„°ë§ ì¤‘ ì˜¤ë¥˜: {e}")
                        self.stop_event.wait(timeout=0.5)

        except Exception as e:
            if not self.stop_event.is_set():
                self.message_queue.put(("error", str(e)))

        finally:
            if driver:
                try:
                    driver.quit()
                except Exception as e:
                    logger.debug(f"WebDriver ì¢…ë£Œ ì˜¤ë¥˜: {e}")
                self.driver = None
            self.message_queue.put(("finished", ""))

    def _build_subtitle_selector_candidates(
        self, primary_selector: str, extras: list[str] | None = None
    ) -> list[str]:
        """ìš°ì„ ìˆœìœ„ê°€ ë°˜ì˜ëœ ìë§‰ CSS ì…€ë ‰í„° í›„ë³´ ëª©ë¡ì„ ìƒì„±í•œë‹¤."""
        candidates = []

        def _add(sel: str) -> None:
            if not isinstance(sel, str):
                return
            norm = sel.strip()
            if not norm or norm in candidates:
                return
            candidates.append(norm)

        _add(primary_selector or "")
        for sel in [
            "#viewSubtit .smi_word:last-child",
            "#viewSubtit .smi_word",
            "#viewSubtit .incont",
            "#viewSubtit span",
            "#viewSubtit",
            ".subtitle_area",
            ".ai_subtitle",
            "[class*='subtitle']",
        ]:
            _add(sel)

        if extras:
            for sel in extras:
                _add(sel)

        # ì»¨í…Œì´ë„ˆí˜• ì…€ë ‰í„°ë³´ë‹¤ ì‹¤ì œ ìë§‰ ë¼ì¸(.smi_word)ì„ ìš°ì„ í•œë‹¤.
        broad_selectors = {
            "#viewSubtit .incont",
            "#viewSubtit",
            ".subtitle_area",
            ".ai_subtitle",
            "[class*='subtitle']",
        }
        priority = {
            "#viewSubtit .smi_word:last-child": 0,
            "#viewSubtit .smi_word": 1,
            "#viewSubtit span": 2,
            "#viewSubtit .incont": 7,
            "#viewSubtit": 8,
            ".subtitle_area": 9,
            ".ai_subtitle": 10,
            "[class*='subtitle']": 11,
        }
        primary_norm = (primary_selector or "").strip()

        order_map = {sel: idx for idx, sel in enumerate(candidates)}

        def _weight(sel: str) -> tuple[int, int]:
            original_idx = order_map.get(sel, 999)
            if sel in priority:
                return priority[sel], original_idx
            if sel == primary_norm and sel not in broad_selectors:
                return 3, original_idx
            if sel in broad_selectors:
                return 12, original_idx
            return 4, original_idx

        return sorted(candidates, key=_weight)

    def _join_stream_text(self, base: str, addition: str) -> str:
        """ìŠ¤íŠ¸ë¦¬ë° í…ìŠ¤íŠ¸ë¥¼ ê³µë°±/ë¬¸ì¥ë¶€í˜¸ë¥¼ ë³´ì¡´í•´ ê²°í•©í•œë‹¤."""
        left = str(base or "")
        right = str(addition or "")
        if not left:
            return right.strip()
        if not right:
            return left.strip()

        left = left.rstrip()
        right = right.lstrip()
        if not left:
            return right
        if not right:
            return left

        no_space_before = set(".,!?;:)]}%\"'â€â€™â€¦")
        no_space_after = set("([{<\"'â€œâ€˜")
        if right[0] in no_space_before or left[-1] in no_space_after:
            return left + right
        return left + " " + right

    def _read_subtitle_text_by_selectors(
        self,
        driver,
        selectors: list[str],
        preferred_frame_path: tuple[int, ...] = (),
    ) -> tuple[str, str, bool]:
        """ì—¬ëŸ¬ ì…€ë ‰í„°ë¥¼ ìˆœì°¨ ì‹œë„í•´ ìë§‰ í…ìŠ¤íŠ¸ë¥¼ ì½ëŠ”ë‹¤.

        Returns:
            (text, matched_selector, found_element)
        """
        def _read_in_current_context() -> tuple[str, str, bool]:
            for sel in selectors:
                try:
                    element = driver.find_element(By.CSS_SELECTOR, sel)
                except (NoSuchElementException, StaleElementReferenceException):
                    continue
                except Exception as e:
                    logger.debug("ì…€ë ‰í„° ì¡°íšŒ ì˜¤ë¥˜ (%s): %s", sel, e)
                    continue

                try:
                    text = (element.text or "").strip()
                except StaleElementReferenceException:
                    continue
                except Exception:
                    text = ""
                return text, sel, True
            return "", "", False

        # ì„ í˜¸ í”„ë ˆì„(Observerê°€ ì„¤ì¹˜ëœ í”„ë ˆì„)ì„ ìš°ì„  í™•ì¸
        if preferred_frame_path:
            try:
                if self._switch_to_frame_path(driver, preferred_frame_path):
                    result = _read_in_current_context()
                    if result[2]:
                        self._last_subtitle_frame_path = preferred_frame_path
                        return result
            finally:
                driver.switch_to.default_content()

        # ê¸°ë³¸ ë¬¸ì„œ í™•ì¸
        try:
            driver.switch_to.default_content()
        except Exception:
            pass
        result = _read_in_current_context()
        if result[2]:
            self._last_subtitle_frame_path = ()
            return result

        # ì¤‘ì²© iframe/frame ìˆœíšŒ í™•ì¸
        for frame_path in self._iter_frame_paths(driver, max_depth=3, max_frames=60):
            try:
                if self._switch_to_frame_path(driver, frame_path):
                    result = _read_in_current_context()
                    if result[2]:
                        self._last_subtitle_frame_path = frame_path
                        return result
            finally:
                try:
                    driver.switch_to.default_content()
                except Exception:
                    pass

        return "", "", False

    def _switch_to_frame_path(self, driver, frame_path: tuple[int, ...]) -> bool:
        """frame index ê²½ë¡œë¡œ ì´ë™í•œë‹¤. ì‹¤íŒ¨ ì‹œ False."""
        try:
            driver.switch_to.default_content()
        except Exception:
            return False

        for idx in frame_path:
            try:
                frames = driver.find_elements(By.CSS_SELECTOR, "iframe,frame")
                if idx < 0 or idx >= len(frames):
                    return False
                driver.switch_to.frame(frames[idx])
            except Exception:
                return False
        return True

    def _iter_frame_paths(
        self, driver, max_depth: int = 3, max_frames: int = 60
    ) -> list[tuple[int, ...]]:
        """ì¤‘ì²© iframe/frame ê²½ë¡œ ëª©ë¡ì„ ë°˜í™˜í•œë‹¤."""
        paths: list[tuple[int, ...]] = []

        def _walk(path: tuple[int, ...], depth: int) -> None:
            if len(paths) >= max_frames or depth > max_depth:
                return
            if not self._switch_to_frame_path(driver, path):
                return
            try:
                frames = driver.find_elements(By.CSS_SELECTOR, "iframe,frame")
            except Exception:
                return

            for idx in range(len(frames)):
                child = path + (idx,)
                paths.append(child)
                if len(paths) >= max_frames:
                    return
                _walk(child, depth + 1)

        try:
            _walk((), 0)
        finally:
            try:
                driver.switch_to.default_content()
            except Exception:
                pass

        return paths

    def _inject_mutation_observer_here(
        self, driver, selector: str, allow_poll_fallback: bool = False
    ) -> bool:
        """í˜„ì¬ ë¬¸ë§¥(í˜„ì¬ frame)ì—ì„œ Observerë¥¼ ì£¼ì…í•œë‹¤."""
        default_selector = (
            "#viewSubtit .smi_word:last-child, #viewSubtit .smi_word, "
            "#viewSubtit .incont, #viewSubtit, .subtitle_area"
        )
        safe_selector = (
            selector if isinstance(selector, str) and selector.strip() else default_selector
        )
        result = driver.execute_script(
            """
            return (function(selectorArg, allowPollFallbackArg) {
                if (window.__subtitleObserver) {
                    try { window.__subtitleObserver.disconnect(); } catch(e) {}
                }
                if (window.__subtitlePollTimer) {
                    try { clearInterval(window.__subtitlePollTimer); } catch(e) {}
                    window.__subtitlePollTimer = null;
                }
                window.__subtitleBuffer = [];
                window.__subtitleLastText = '';
                window.__subtitleLastEmitTs = 0;

                var rawSelector = (typeof selectorArg === 'string') ? selectorArg : '';
                var allowPollFallback = !!allowPollFallbackArg;
                var selectors = rawSelector
                    .split(',')
                    .map(function(s) { return (s || '').trim(); })
                    .filter(function(s) { return s.length > 0; });
                if (!selectors.length) {
                    selectors = [
                        '#viewSubtit .smi_word:last-child',
                        '#viewSubtit .smi_word',
                        '#viewSubtit .incont',
                        '#viewSubtit',
                        '.subtitle_area',
                        '.ai_subtitle',
                        "[class*='subtitle']"
                    ];
                }

                var target = null;
                for (var i = 0; i < selectors.length; i++) {
                    try {
                        target = document.querySelector(selectors[i]);
                    } catch (e) {
                        target = null;
                    }
                    if (target) break;
                }

                function normalizeText(text) {
                    return String(text || '').replace(/\\s+/g, ' ').trim();
                }

                function isLikelySubtitleText(text) {
                    if (!text) return false;
                    if (text.length < 3 || text.length > 320) return false;
                    if (!/[ê°€-í£A-Za-z]/.test(text)) return false;
                    if (/^[\\d\\s:.,\\-_/()%]+$/.test(text)) return false;
                    return true;
                }

                function pickBestMutationText(mutations) {
                    var bestText = '';
                    var bestScore = -1;
                    for (var i = 0; i < mutations.length; i++) {
                        var m = mutations[i];
                        var node = m && m.target ? m.target : null;
                        var el = null;
                        if (node && node.nodeType === 1) el = node;
                        else if (node && node.parentElement) el = node.parentElement;
                        if (!el) continue;
                        if (el.tagName === 'SCRIPT' || el.tagName === 'STYLE') continue;
                        if (typeof el.closest === 'function') {
                            var bad = el.closest('script,style,head,noscript');
                            if (bad) continue;
                        }

                        var text = normalizeText(el.innerText || el.textContent || '');
                        if (!isLikelySubtitleText(text)) continue;

                        if (text.length > 120) {
                            var lines = String(el.innerText || '').split('\\n')
                                .map(function(v) { return normalizeText(v); })
                                .filter(function(v) { return !!v; });
                            if (lines.length) {
                                var tail = lines[lines.length - 1];
                                if (isLikelySubtitleText(tail)) text = tail;
                            }
                        }

                        var score = 0;
                        try {
                            var idClass = ((el.id || '') + ' ' + (el.className || '')).toLowerCase();
                            if (/subtit|subtitle|caption|script|stt|transcript|incont|viewsubtit/.test(idClass)) score += 6;
                            var rect = el.getBoundingClientRect ? el.getBoundingClientRect() : null;
                            if (rect && rect.width > 0 && rect.height > 0) score += 2;
                            if (rect && rect.bottom >= (window.innerHeight * 0.35)) score += 1;
                        } catch (e) {}

                        score += Math.min(4, Math.floor(text.length / 25));

                        if (score > bestScore) {
                            bestScore = score;
                            bestText = text;
                        }
                    }
                    return bestText;
                }

                if (target) {
                    window.__subtitleObserver = new MutationObserver(function() {
                        try {
                            var text = target.innerText || target.textContent || '';
                            text = normalizeText(text);
                            if (!text && window.__subtitleLastText) {
                                window.__subtitleBuffer.push('__SUBTITLE_CLEARED__');
                                window.__subtitleLastText = '';
                                return;
                            }
                            if (text && text !== window.__subtitleLastText) {
                                window.__subtitleLastText = text;
                                window.__subtitleBuffer.push(text);
                                if (window.__subtitleBuffer.length > 100) {
                                    window.__subtitleBuffer = window.__subtitleBuffer.slice(-50);
                                }
                            }
                        } catch (e) {}
                    });

                    window.__subtitleObserver.observe(target, {
                        childList: true,
                        subtree: true,
                        characterData: true,
                        attributes: true
                    });
                    return true;
                }

                var root = document.body || document.documentElement;
                if (!root || !allowPollFallback) return false;

                // íƒ€ê²Ÿì„ ëª» ì°¾ì€ ê²½ìš°: ì£¼ê¸°ì  selector ìŠ¤ìº”ìœ¼ë¡œ Observer ë²„í¼ ë¸Œë¦¬ì§€
                window.__subtitlePollTimer = setInterval(function() {
                    try {
                        var now = Date.now();
                        if (now - (window.__subtitleLastEmitTs || 0) < 100) {
                            return;
                        }
                        var liveTarget = null;
                        for (var i = 0; i < selectors.length; i++) {
                            try {
                                liveTarget = document.querySelector(selectors[i]);
                            } catch (e) {
                                liveTarget = null;
                            }
                            if (liveTarget) break;
                        }
                        if (!liveTarget) {
                            return;
                        }

                        var text = normalizeText(liveTarget.innerText || liveTarget.textContent || '');
                        if (!text && window.__subtitleLastText) {
                            window.__subtitleBuffer.push('__SUBTITLE_CLEARED__');
                            window.__subtitleLastText = '';
                            window.__subtitleLastEmitTs = now;
                            return;
                        }
                        if (!text || !isLikelySubtitleText(text)) {
                            return;
                        }
                        if (text && text !== window.__subtitleLastText) {
                            window.__subtitleLastText = text;
                            window.__subtitleLastEmitTs = now;
                            window.__subtitleBuffer.push(text);
                            if (window.__subtitleBuffer.length > 100) {
                                window.__subtitleBuffer = window.__subtitleBuffer.slice(-50);
                            }
                        }
                    } catch (e) {
                    }
                }, 180);
                return true;
            })(arguments[0], arguments[1]);
            """,
            safe_selector,
            allow_poll_fallback,
        )
        return bool(result)

    def _inject_mutation_observer(self, driver, selector: str) -> tuple[bool, tuple[int, ...]]:
        """MutationObserverë¥¼ í˜ì´ì§€ì— ì£¼ì…í•˜ì—¬ ìë§‰ ë³€ê²½ì„ ì´ë²¤íŠ¸ ê¸°ë°˜ìœ¼ë¡œ ìº¡ì²˜í•œë‹¤.

        Returns:
            (ì£¼ì… ì„±ê³µ ì—¬ë¶€, observer frame ê²½ë¡œ)
        """
        try:
            safe_selector = selector if isinstance(selector, str) else ""
            priority_paths: list[tuple[int, ...]] = []
            last_path = getattr(self, "_last_subtitle_frame_path", ())
            if isinstance(last_path, tuple):
                priority_paths.append(last_path)
            priority_paths.append(())
            for p in self._iter_frame_paths(driver, max_depth=3, max_frames=60):
                if p not in priority_paths:
                    priority_paths.append(p)

            # 1) íƒ€ê²Ÿ ê¸°ë°˜ Observer ìš°ì„  ì‹œë„
            for frame_path in priority_paths:
                if not self._switch_to_frame_path(driver, frame_path):
                    continue
                if self._inject_mutation_observer_here(
                    driver, safe_selector, allow_poll_fallback=False
                ):
                    location = "default" if frame_path == () else f"frame={frame_path}"
                    logger.info(
                        "MutationObserver ì£¼ì… ì„±ê³µ: %s (%s)", location, safe_selector
                    )
                    return True, frame_path

            # 2) íƒ€ê²Ÿ ë¯¸íƒìƒ‰ ì‹œ JS í´ë§ ë¸Œë¦¬ì§€ fallback
            for frame_path in priority_paths:
                if not self._switch_to_frame_path(driver, frame_path):
                    continue
                if self._inject_mutation_observer_here(
                    driver, safe_selector, allow_poll_fallback=True
                ):
                    location = "default" if frame_path == () else f"frame={frame_path}"
                    logger.info(
                        "MutationObserver í´ë§ ë¸Œë¦¬ì§€ í™œì„±í™”: %s (%s)",
                        location,
                        safe_selector,
                    )
                    return True, frame_path

            logger.warning("MutationObserver ì£¼ì… ì‹¤íŒ¨: ëŒ€ìƒ ìš”ì†Œ ì—†ìŒ (%s)", safe_selector)
            return False, ()
        except Exception as e:
            logger.warning("MutationObserver ì£¼ì… ì˜¤ë¥˜: %s", e)
            return False, ()
        finally:
            try:
                driver.switch_to.default_content()
            except Exception:
                pass

    def _collect_observer_changes(
        self, driver, frame_path: tuple[int, ...] = ()
    ) -> list | None:
        """MutationObserver ë²„í¼ì—ì„œ ë³€ê²½ëœ í…ìŠ¤íŠ¸ë¥¼ ìˆ˜ì§‘í•œë‹¤.

        Returns:
            list: ë³€ê²½ëœ í…ìŠ¤íŠ¸ ëª©ë¡ (ë¹„ì–´ìˆì„ ìˆ˜ ìˆìŒ)
            None: Observerê°€ ì£½ì—ˆê±°ë‚˜ ì˜¤ë¥˜ ë°œìƒ (í´ë§ fallback í•„ìš”)
        """
        try:
            if not self._switch_to_frame_path(driver, frame_path):
                return None
            result = driver.execute_script(
                """
                if (!window.__subtitleBuffer) return null;
                var buf = window.__subtitleBuffer;
                window.__subtitleBuffer = [];
                return buf;
                """
            )
            if result is None:
                return None  # Observerê°€ ì—†ìŒ
            return result if isinstance(result, list) else []
        except Exception as e:
            logger.debug("Observer ë²„í¼ ìˆ˜ì§‘ ì˜¤ë¥˜: %s", e)
            return None
        finally:
            try:
                driver.switch_to.default_content()
            except Exception:
                pass

    def _activate_subtitle(self, driver) -> bool:
        """ìë§‰ ë ˆì´ì–´ í™œì„±í™” - ë‹¤ì–‘í•œ ë°©ë²• ì‹œë„

        Returns:
            bool: í™œì„±í™” ì„±ê³µ ì—¬ë¶€
        """
        activation_scripts = [
            # ë°©ë²• 1: layerSubtit í•¨ìˆ˜ í˜¸ì¶œ
            "if(typeof layerSubtit==='function'){layerSubtit(); return true;} return false;",
            # ë°©ë²• 2: ìë§‰ ë²„íŠ¼ í´ë¦­
            "var btn=document.querySelector('.btn_subtit'); if(btn){btn.click(); return true;} return false;",
            "var btn=document.querySelector('#btnSubtit'); if(btn){btn.click(); return true;} return false;",
            # ë°©ë²• 3: AI ìë§‰ ë²„íŠ¼
            "var btn=document.querySelector('[data-action=\\'subtitle\\']'); if(btn){btn.click(); return true;} return false;",
            # ë°©ë²• 4: ìë§‰ ë ˆì´ì–´ ì§ì ‘ í‘œì‹œ
            "var layer=document.querySelector('#viewSubtit'); if(layer){layer.style.display='block'; return true;} return false;",
        ]

        activated = False
        for idx, script in enumerate(activation_scripts, start=1):
            try:
                result = driver.execute_script(script)
                if result:
                    logger.info(
                        "ìë§‰ í™œì„±í™” ì„±ê³µ (step=%s/%s): %s...",
                        idx,
                        len(activation_scripts),
                        script[:50],
                    )
                    activated = True
                    break
                self.stop_event.wait(timeout=0.5)
            except Exception as e:
                logger.debug(f"ìë§‰ í™œì„±í™” ìŠ¤í¬ë¦½íŠ¸ ì‹¤íŒ¨: {e}")

        # ì¶”ê°€ ëŒ€ê¸° - ìë§‰ ë ˆì´ì–´ ë¡œë”© (ì¢…ë£Œ ì‹ í˜¸ì— ì¦‰ì‹œ ë°˜ì‘)
        self.stop_event.wait(timeout=2.0)
        return activated

    def _find_subtitle_selector(self, driver) -> str:
        """ì‚¬ìš© ê°€ëŠ¥í•œ ìë§‰ ì…€ë ‰í„° ìë™ ê°ì§€

        Returns:
            str: ì°¾ì€ ì…€ë ‰í„°. ì°¾ì§€ ëª»í•˜ë©´ ë¹ˆ ë¬¸ìì—´.
        """
        # ìš°ì„ ìˆœìœ„ëŒ€ë¡œ ì…€ë ‰í„° í™•ì¸
        selectors = [
            "#viewSubtit .smi_word:last-child",
            "#viewSubtit .smi_word",
            "#viewSubtit .incont",
            "#viewSubtit span",
            "#viewSubtit",
            ".subtitle_area",
            ".ai_subtitle",
            "[class*='subtitle']",
        ]

        text, matched_selector, found = self._read_subtitle_text_by_selectors(
            driver, selectors
        )
        if found and matched_selector:
            if text and len(text) > 2:
                logger.info(f"ìë§‰ ì…€ë ‰í„° ë°œê²¬: {matched_selector}")
            return matched_selector
        return ""

    def _clear_message_queue(self) -> None:
        """ë©”ì‹œì§€ í ë¹„ìš°ê¸° (ì¤‘ì§€/ì¬ì‹œì‘ ì•ˆì •ì„±ìš©)"""
        try:
            while True:
                self.message_queue.get_nowait()
        except queue.Empty:
            pass

    def _process_subtitle_segments(self, data) -> None:
        """ì„¸ê·¸ë¨¼íŠ¸ í˜•íƒœë¡œ ë“¤ì–´ì˜¨ ìë§‰ ë°ì´í„°ë¥¼ ì²˜ë¦¬í•œë‹¤."""
        if not data:
            return
        if isinstance(data, (list, tuple)):
            for item in data:
                if item:
                    prepared = self._prepare_preview_raw(item)
                    if prepared:
                        self._process_raw_text(prepared)
            return
        if isinstance(data, str):
            prepared = self._prepare_preview_raw(data)
            if prepared:
                self._process_raw_text(prepared)

    def _prepare_preview_raw(self, raw):
        """preview ì…ë ¥ì„ ì •ê·œí™”/ê²Œì´íŒ…í•˜ì—¬ core ì•Œê³ ë¦¬ì¦˜ì— ì „ë‹¬í• ì§€ ê²°ì •í•œë‹¤."""
        if raw is None:
            return None

        normalized = utils.clean_text_display(str(raw)).strip()
        if not normalized:
            return None

        raw_compact = utils.compact_subtitle_text(normalized)
        if not raw_compact:
            return None

        def accept(text: str):
            self._preview_desync_count = 0
            self._preview_ambiguous_skip_count = 0
            self._last_good_raw_compact = raw_compact
            return text

        suffix = self._trailing_suffix
        if not suffix:
            return accept(normalized)

        first_pos = raw_compact.find(suffix)
        if first_pos < 0:
            # 1ì°¨ fallback: ì§ì „ raw ëŒ€ë¹„ ì¦ë¶„(delta)ë§Œ ì¶”ì¶œ
            delta = self._extract_stream_delta(normalized, self._last_raw_text)
            if isinstance(delta, str):
                delta = delta.strip()
                if len(delta) >= 1 and len(delta) < len(normalized):
                    return accept(delta)
                if not delta:
                    pass

            # 2ì°¨ fallback: ìµœê·¼ í™•ì • íˆìŠ¤í† ë¦¬ tail anchorë¡œ ì¦ë¶„ ì¶”ì¶œ
            history_anchor = self._confirmed_history_compact_tail(
                max_entries=8, max_compact_len=3000
            )
            incremental = self._slice_incremental_part(
                normalized,
                history_anchor,
                raw_compact,
                min_anchor=20,
                min_overlap=8,
            )
            if isinstance(incremental, str):
                incremental = incremental.strip()
                if len(incremental) >= 1 and len(incremental) < len(normalized):
                    return accept(incremental)
                if not incremental:
                    return None

            self._preview_desync_count += 1
            if self._preview_desync_count >= self._preview_resync_threshold:
                logger.warning(
                    "preview suffix desync reset: count=%s", self._preview_desync_count
                )
                self._soft_resync()
                return accept(normalized)
            return None

        last_pos = raw_compact.rfind(suffix)
        if first_pos != last_pos:
            predicted_append = len(raw_compact) - (first_pos + len(suffix))
            large_append_threshold = max(200, len(raw_compact) // 3)
            if predicted_append > large_append_threshold:
                incremental = self._slice_incremental_part(
                    normalized,
                    suffix,
                    raw_compact,
                    min_anchor=20,
                    min_overlap=8,
                )
                if isinstance(incremental, str):
                    incremental = incremental.strip()
                    if len(incremental) >= 1 and len(incremental) < len(normalized):
                        return accept(incremental)
                    if not incremental:
                        return None

                self._preview_ambiguous_skip_count += 1
                if (
                    self._preview_ambiguous_skip_count
                    >= self._preview_ambiguous_resync_threshold
                ):
                    logger.warning(
                        "preview ambiguous suffix reset: count=%s, predicted_append=%s",
                        self._preview_ambiguous_skip_count,
                        predicted_append,
                    )
                    self._soft_resync()
                    return accept(normalized)
                return None

        return accept(normalized)

    def _drain_pending_previews(
        self, max_items: int = 2000, requeue_others: bool = True
    ) -> None:
        """íì— ë‚¨ì€ preview/segments ë©”ì‹œì§€ë¥¼ ì†Œì§„í•´ ë§ˆì§€ë§‰ ìë§‰ ëˆ„ë½ì„ ì¤„ì¸ë‹¤."""
        drained = 0
        pending = []
        try:
            while drained < max_items:
                msg_type, data = self.message_queue.get_nowait()
                if msg_type == "preview":
                    prepared = self._prepare_preview_raw(data)
                    if prepared:
                        self._process_raw_text(prepared)
                    elif data:
                        forced = utils.clean_text_display(str(data)).strip()
                        if forced:
                            self._process_raw_text(forced)
                elif msg_type == "subtitle_segments":
                    self._process_subtitle_segments(data)
                else:
                    pending.append((msg_type, data))
                drained += 1
        except queue.Empty:
            pass

        if drained >= max_items:
            logger.warning("preview í ì†Œì§„ ì œí•œ ë„ë‹¬: max_items=%s", max_items)

        if requeue_others and pending:
            for item in pending:
                self.message_queue.put(item)

    def _finalize_pending_subtitle(self) -> None:
        """ë‚¨ì•„ìˆëŠ” ë²„í¼ë¥¼ í™•ì • ì²˜ë¦¬ (ì¢…ë£Œ/ì¤‘ì§€ ì•ˆì „ì„±ìš©)."""
        if self.last_subtitle:
            self._finalize_subtitle(self.last_subtitle)
            self.last_subtitle = ""
            self._stream_start_time = None
            self._clear_preview()

    def _reset_stream_state_after_subtitle_change(
        self, keep_history_from_subtitles: bool
    ) -> None:
        """ì™¸ë¶€ í¸ì§‘/ë¡œë“œ í›„ ìŠ¤íŠ¸ë¦¬ë° íŒŒì´í”„ë¼ì¸ ìƒíƒœë¥¼ ì¬ë™ê¸°í™”í•œë‹¤."""
        self.last_subtitle = ""
        self._stream_start_time = None
        self._last_raw_text = ""
        self._last_processed_raw = ""
        self._preview_desync_count = 0
        self._preview_ambiguous_skip_count = 0
        self._last_good_raw_compact = ""
        self.finalize_timer.stop()
        self._clear_preview()

        if keep_history_from_subtitles:
            self._soft_resync()
        else:
            self._confirmed_compact = ""
            self._trailing_suffix = ""

    def _replace_subtitles_and_refresh(
        self, new_subtitles: list, keep_history_from_subtitles: bool | None = None
    ) -> None:
        """ìë§‰ ì „ì²´ êµì²´ í›„ ìƒíƒœ/í†µê³„/UIë¥¼ ì¼ê´€ë˜ê²Œ ê°±ì‹ í•œë‹¤."""
        if keep_history_from_subtitles is None:
            keep_history_from_subtitles = bool(new_subtitles)

        with self.subtitle_lock:
            self.subtitles = list(new_subtitles)

        self._reset_stream_state_after_subtitle_change(
            keep_history_from_subtitles=keep_history_from_subtitles
        )
        self._rebuild_stats_cache()
        self._refresh_text(force_full=True)
        self._update_count_label()

    def _check_finalize(self):
        """í˜„ì¬ ìŠ¤íŠ¸ë¦¬ë° ì•Œê³ ë¦¬ì¦˜ì—ì„œëŠ” finalize íƒ€ì´ë¨¸ë¥¼ ì‚¬ìš©í•˜ì§€ ì•ŠëŠ”ë‹¤."""
        if self.finalize_timer.isActive():
            self.finalize_timer.stop()

    # ========== ë©”ì‹œì§€ í ì²˜ë¦¬ ==========

    def _process_message_queue(self):
        """ë©”ì‹œì§€ í ì²˜ë¦¬ (100msë§ˆë‹¤ í˜¸ì¶œ) - ì˜ˆì™¸ ì²˜ë¦¬ ê°•í™”"""
        try:
            # ìµœëŒ€ 10ê°œ ë©”ì‹œì§€ê¹Œì§€ ì²˜ë¦¬ (ë¬´í•œ ë£¨í”„ ë°©ì§€)
            for _ in range(10):
                try:
                    msg_type, data = self.message_queue.get_nowait()
                    self._handle_message(msg_type, data)
                except queue.Empty:
                    break
        except Exception as e:
            logger.error(f"í ì²˜ë¦¬ ì˜¤ë¥˜: {e}")

    def _handle_message(self, msg_type, data):
        """ê°œë³„ ë©”ì‹œì§€ ì²˜ë¦¬"""
        try:
            if self._is_stopping and msg_type not in (
                "db_task_result",
                "db_task_error",
                "session_save_done",
                "session_save_failed",
                "session_load_done",
                "session_load_failed",
                "session_load_json_error",
            ):
                return
            if msg_type == "status":
                status_text = str(data)[:200]
                if status_text != self._last_status_message:
                    self.status_label.setText(status_text)
                    self._last_status_message = status_text

            elif msg_type == "toast":
                # ë‹¤ë¥¸ ìŠ¤ë ˆë“œì—ì„œ ì˜¨ UI ì•Œë¦¼ì€ Queueë¡œ ì „ë‹¬ë°›ì•„ UI ìŠ¤ë ˆë“œì—ì„œ ì²˜ë¦¬í•œë‹¤.
                if isinstance(data, dict):
                    message = data.get("message", "")
                    toast_type = data.get("toast_type", "info")
                    duration = data.get("duration", 3000)
                else:
                    message = data
                    toast_type = "info"
                    duration = 3000
                try:
                    duration = int(duration) if duration is not None else 3000
                except Exception:
                    duration = 3000
                self._show_toast(str(message), str(toast_type), duration)

            elif msg_type == "preview":
                prepared = self._prepare_preview_raw(data)
                if prepared:
                    self._process_raw_text(prepared)

            elif msg_type == "subtitle_reset":
                # ë°œì–¸ì ì „í™˜ìœ¼ë¡œ ìë§‰ ì˜ì—­ì´ í´ë¦¬ì–´ë¨ â†’ ì™„ì „ ë¦¬ì…‹
                # (ìë§‰ ì˜ì—­ì´ ë¹ˆ ìƒíƒœì´ë¯€ë¡œ ì¤‘ë³µ ìœ ì… ìœ„í—˜ ì—†ìŒ)
                logger.info("subtitle_reset ê°ì§€: %s", data)
                # ì´ì „ ë°œì–¸ìì˜ ë§ˆì§€ë§‰ ë²„í¼ í™•ì •
                if self.last_subtitle:
                    self._finalize_subtitle(self.last_subtitle)
                    self.last_subtitle = ""
                    self._stream_start_time = None
                self._confirmed_compact = ""
                self._trailing_suffix = ""
                self._last_raw_text = ""
                self._last_processed_raw = ""
                self._preview_desync_count = 0
                self._preview_ambiguous_skip_count = 0

            elif msg_type == "keepalive":
                self._handle_keepalive(data)

            elif msg_type == "error":
                self.progress.hide()
                self._reset_ui()
                self._update_tray_status("âšª ëŒ€ê¸° ì¤‘")
                self._update_connection_status("disconnected")
                self._clear_preview()
                QMessageBox.critical(self, "ì˜¤ë¥˜", str(data))

            elif msg_type == "finished":
                if self.last_subtitle:
                    self._finalize_subtitle(self.last_subtitle)
                    self.last_subtitle = ""
                    self._stream_start_time = None
                self.finalize_timer.stop()
                self._clear_preview()

                self._refresh_text()
                self._reset_ui()
                self._update_tray_status("âšª ëŒ€ê¸° ì¤‘")
                self._update_connection_status("disconnected")

                # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ í†µê³„ ê³„ì‚°
                with self.subtitle_lock:
                    subtitle_count = len(self.subtitles)
                total_chars = self._cached_total_chars
                self.status_label.setText(
                    f"ì™„ë£Œ - {subtitle_count}ë¬¸ì¥, {total_chars:,}ì"
                )

            elif msg_type == "session_save_done":
                self._session_save_in_progress = False
                info = data if isinstance(data, dict) else {}
                saved_count = int(info.get("saved_count", 0) or 0)
                db_saved = bool(info.get("db_saved", False))
                db_error = str(info.get("db_error", "") or "").strip()
                if db_saved:
                    self._set_status(
                        f"ì„¸ì…˜ ì €ì¥ ì™„ë£Œ ({saved_count}ê°œ, DB ì €ì¥ í¬í•¨)", "success"
                    )
                    self._show_toast("ì„¸ì…˜ ì €ì¥ ì™„ë£Œ! (JSON + DB)", "success")
                else:
                    self._set_status(f"ì„¸ì…˜ ì €ì¥ ì™„ë£Œ ({saved_count}ê°œ)", "success")
                    if db_error:
                        self._show_toast(
                            "ì„¸ì…˜ ì €ì¥ ì™„ë£Œ (DB ì €ì¥ì€ ì‹¤íŒ¨)", "warning", 3500
                        )
                        logger.warning("ì„¸ì…˜ ì €ì¥: DB ì €ì¥ ì‹¤íŒ¨ - %s", db_error)
                    else:
                        self._show_toast("ì„¸ì…˜ ì €ì¥ ì™„ë£Œ!", "success")

            elif msg_type == "session_save_failed":
                self._session_save_in_progress = False
                err = data.get("error") if isinstance(data, dict) else str(data)
                self._set_status(f"ì„¸ì…˜ ì €ì¥ ì‹¤íŒ¨: {err}", "error")
                QMessageBox.critical(self, "ì˜¤ë¥˜", f"ì„¸ì…˜ ì €ì¥ ì‹¤íŒ¨: {err}")

            elif msg_type == "session_load_done":
                self._session_load_in_progress = False
                payload = data if isinstance(data, dict) else {}
                session_version = str(payload.get("version", "unknown"))
                source_path = str(payload.get("path", ""))
                source_url = str(payload.get("url", "") or "")
                loaded_subtitles = payload.get("subtitles", [])
                skipped_items = int(payload.get("skipped", 0) or 0)

                if session_version != Config.VERSION:
                    reply = QMessageBox.question(
                        self,
                        "ë²„ì „ ë¶ˆì¼ì¹˜",
                        f"ì„¸ì…˜ ë²„ì „({session_version})ì´ í˜„ì¬ ë²„ì „({Config.VERSION})ê³¼ ë‹¤ë¦…ë‹ˆë‹¤.\n"
                        "ê³„ì† ë¶ˆëŸ¬ì˜¤ì‹œê² ìŠµë‹ˆê¹Œ?",
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    )
                    if reply == QMessageBox.StandardButton.No:
                        self._set_status("ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì·¨ì†Œë¨ (ë²„ì „ ë¶ˆì¼ì¹˜)", "warning")
                        return

                self._replace_subtitles_and_refresh(
                    loaded_subtitles, keep_history_from_subtitles=bool(loaded_subtitles)
                )
                if source_url:
                    self.url_combo.setCurrentText(source_url)

                summary = f"ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì™„ë£Œ! {len(self.subtitles)}ê°œ ë¬¸ì¥"
                if skipped_items > 0:
                    summary += f" (ì†ìƒ í•­ëª© {skipped_items}ê°œ ì œì™¸)"
                self._set_status(summary, "success")
                self._show_toast(summary, "success")
                if source_path:
                    logger.info("ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì™„ë£Œ: %s", source_path)

            elif msg_type == "session_load_json_error":
                self._session_load_in_progress = False
                info = data if isinstance(data, dict) else {}
                path = str(info.get("path", ""))
                err = str(info.get("error", "JSON íŒŒì‹± ì˜¤ë¥˜"))
                reply = QMessageBox.question(
                    self,
                    "íŒŒì¼ ì†ìƒ",
                    f"ì„¸ì…˜ íŒŒì¼ì´ ì†ìƒë˜ì—ˆìŠµë‹ˆë‹¤ (JSON ì˜¤ë¥˜).\nìœ„ì¹˜: {path}\nì˜¤ë¥˜: {err}\n\n"
                    "ë°±ì—… í´ë”ë¥¼ ì—´ì–´ ë³µêµ¬ë¥¼ ì‹œë„í•˜ì‹œê² ìŠµë‹ˆê¹Œ?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                )
                if reply == QMessageBox.StandardButton.Yes:
                    backup_path = os.path.abspath(Config.BACKUP_DIR)
                    if os.name == "nt":
                        os.startfile(backup_path)
                self._set_status("ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì‹¤íŒ¨ (JSON ì˜¤ë¥˜)", "error")

            elif msg_type == "session_load_failed":
                self._session_load_in_progress = False
                err = data.get("error") if isinstance(data, dict) else str(data)
                self._set_status(f"ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì‹¤íŒ¨: {err}", "error")
                QMessageBox.critical(self, "ì˜¤ë¥˜", f"ë¶ˆëŸ¬ì˜¤ê¸° ì‹¤íŒ¨: {err}")

            elif msg_type == "subtitle_not_found":
                # ìë§‰ ìš”ì†Œë¥¼ ì°¾ì§€ ëª»í–ˆì„ ë•Œ ì‚¬ìš©ì ì•ˆë‚´
                self.progress.hide()
                self._reset_ui()
                self._update_tray_status("âšª ëŒ€ê¸° ì¤‘")
                self._clear_preview()

                # ìƒì„¸ ì•ˆë‚´ ë‹¤ì´ì–¼ë¡œê·¸
                msg_box = QMessageBox(self)
                msg_box.setWindowTitle("ìë§‰ì„ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤")
                msg_box.setIcon(QMessageBox.Icon.Warning)
                msg_box.setText(str(data))
                msg_box.setStandardButtons(
                    QMessageBox.StandardButton.Open | QMessageBox.StandardButton.Ok
                )
                msg_box.button(QMessageBox.StandardButton.Open).setText(
                    "ğŸŒ ì‚¬ì´íŠ¸ ì—´ê¸°"
                )
                msg_box.button(QMessageBox.StandardButton.Ok).setText("í™•ì¸")

                result = msg_box.exec()

                # ì‚¬ì´íŠ¸ ì—´ê¸° ë²„íŠ¼ í´ë¦­ ì‹œ ë¸Œë¼ìš°ì €ì—ì„œ ì—´ê¸°
                if result == QMessageBox.StandardButton.Open:
                    import webbrowser

                    webbrowser.open("https://assembly.webcast.go.kr")

            # ì—°ê²° ìƒíƒœ ì—…ë°ì´íŠ¸ (#30)
            elif msg_type == "connection_status":
                status = data.get("status", "disconnected")
                latency = data.get("latency")
                self._update_connection_status(status, latency)

            # ì¬ì—°ê²° ì‹œë„ (#31)
            elif msg_type == "reconnecting":
                self.reconnect_attempts = data.get("attempt", 0)
                self._update_connection_status("reconnecting")
                self._show_toast(
                    f"ì¬ì—°ê²° ì‹œë„ ì¤‘... ({self.reconnect_attempts}/{Config.MAX_RECONNECT_ATTEMPTS})",
                    "warning",
                    2000,
                )

            # ì¬ì—°ê²° ì„±ê³µ (#31)
            elif msg_type == "reconnected":
                self.reconnect_attempts = 0
                self._update_connection_status("connected")
                self._show_toast("ì¬ì—°ê²° ì„±ê³µ!", "success", 2000)

            elif msg_type == "hwp_save_failed":
                error = data.get("error") if isinstance(data, dict) else data
                self._handle_hwp_save_failure(error)

            elif msg_type == "db_task_result":
                task_name = data.get("task") if isinstance(data, dict) else None
                result = data.get("result") if isinstance(data, dict) else None
                context = data.get("context") if isinstance(data, dict) else {}
                if task_name:
                    self._db_tasks_inflight.discard(task_name)
                    self._handle_db_task_result(task_name, result, context or {})

            elif msg_type == "db_task_error":
                task_name = data.get("task") if isinstance(data, dict) else "db_unknown"
                error = data.get("error") if isinstance(data, dict) else str(data)
                context = data.get("context") if isinstance(data, dict) else {}
                self._db_tasks_inflight.discard(task_name)
                self._handle_db_task_error(task_name, error, context or {})

        except Exception as e:
            logger.error(f"ë©”ì‹œì§€ ì²˜ë¦¬ ì˜¤ë¥˜ ({msg_type}): {e}")

    def _handle_keepalive(self, raw: str) -> None:
        """ë™ì¼ ìë§‰ì´ ìœ ì§€ë  ë•Œ ë§ˆì§€ë§‰ ì—”íŠ¸ë¦¬ì˜ end_timeì„ ê°±ì‹ í•œë‹¤."""
        if not raw:
            return
        # í˜„ì¬ ì§„í–‰ ì¤‘ì¸ ë²„í¼ê°€ ìˆìœ¼ë©´ í™•ì • ì´í›„ì—ë§Œ ê°±ì‹ 
        if self.last_subtitle:
            return
        raw_compact = utils.compact_subtitle_text(raw)
        if not raw_compact:
            return

        with self.subtitle_lock:
            if not self.subtitles:
                return
            last_entry = self.subtitles[-1]
            last_compact = utils.compact_subtitle_text(last_entry.text)
            if raw_compact != last_compact:
                return
            last_entry.end_time = datetime.now()

    def _should_merge_entry(
        self, last_entry: SubtitleEntry, new_text: str, now: datetime
    ) -> bool:
        """ê¸°ì¡´ ìë§‰ì— ì´ì–´ë¶™ì—¬ë„ ë˜ëŠ”ì§€ íŒë‹¨ (ë‹¨ìˆœ ì‹œê°„/ê¸¸ì´ ì²´í¬)"""
        if not last_entry or not new_text:
            return False

        # ë§ˆì§€ë§‰ ìë§‰ê³¼ ì‹œê°„ ì°¨ì´ê°€ ë„ˆë¬´ í¬ë©´ ë¶„ë¦¬
        if (
            last_entry.end_time
            and (now - last_entry.end_time).total_seconds() > Config.ENTRY_MERGE_MAX_GAP
        ):
            return False

        # ë„ˆë¬´ ê¸¸ì–´ì§€ë©´ ë¶„ë¦¬
        if len(last_entry.text) + 1 + len(new_text) > Config.ENTRY_MERGE_MAX_CHARS:
            return False

        return True

    def _confirmed_history_compact_tail(
        self, max_entries: int = 10, max_compact_len: int = 3000
    ) -> str:
        """ìµœê·¼ í™•ì • ìë§‰ë“¤ì˜ compact tail ë¬¸ìì—´(ê²¹ì¹¨/ì¤‘ë³µ ì œê±°ìš©)"""
        with self.subtitle_lock:
            if not self.subtitles:
                return ""
            tail_entries = self.subtitles[-max_entries:]
            combined = " ".join(e.text for e in tail_entries if e and e.text)

        compact = utils.compact_subtitle_text(combined)
        if max_compact_len > 0 and len(compact) > max_compact_len:
            compact = compact[-max_compact_len:]
        return compact

    def _slice_incremental_part(
        self,
        raw: str,
        anchor_compact: str,
        raw_compact: str,
        min_anchor: int = 12,
        min_overlap: int = 4,
    ):
        """anchor ê¸°ì¤€ìœ¼ë¡œ rawì—ì„œ ìƒˆë¡œ ì¶”ê°€ëœ ë¶€ë¶„ì„ ë°˜í™˜í•œë‹¤.

        ë°˜í™˜ê°’:
          - None: ê²¹ì¹¨ ì—†ìŒ (ë¬¸ë§¥ ì „í™˜ìœ¼ë¡œ ì²˜ë¦¬)
          - "": ê²¹ì¹¨ì€ ìˆìœ¼ë‚˜ ìƒˆ ë‚´ìš© ì—†ìŒ
          - str: ìƒˆë¡œ ì¶”ê°€ëœ í…ìŠ¤íŠ¸
        """
        if not raw or not anchor_compact or not raw_compact:
            return None

        # anchorê°€ ê¸¸ë©´, raw ë‚´ë¶€ì—ì„œ ë§ˆì§€ë§‰ìœ¼ë¡œ ë“±ì¥í•˜ëŠ” ìœ„ì¹˜ë¥¼ ê¸°ì¤€ìœ¼ë¡œ ìŠ¬ë¼ì´ìŠ¤
        if len(anchor_compact) >= min_anchor:
            idx = raw_compact.rfind(anchor_compact)
            if idx != -1:
                return utils.slice_from_compact_index(
                    raw, idx + len(anchor_compact)
                ).strip()

        # fallback: suffix-prefix overlap
        overlap_len = utils.find_compact_suffix_prefix_overlap(
            anchor_compact, raw_compact, min_overlap=min_overlap
        )
        if overlap_len > 0:
            return utils.slice_from_compact_index(raw, overlap_len).strip()

        return None

    def _extract_stream_delta(self, raw: str, last_raw: str):
        """ì´ì „ raw ëŒ€ë¹„ ìƒˆë¡œ ì¶”ê°€ëœ ë¶€ë¶„(delta)ì„ ì¶”ì¶œí•œë‹¤.

        Returns:
            str: ìƒˆë¡œ ì¶”ê°€ëœ í…ìŠ¤íŠ¸ (ì—†ìœ¼ë©´ "")
            None: ë¬¸ë§¥ ì „í™˜(ìƒˆ ë¬¸ì¥)ìœ¼ë¡œ íŒë‹¨
        """
        if not last_raw:
            return None

        raw_compact = utils.compact_subtitle_text(raw)
        last_compact = utils.compact_subtitle_text(last_raw)
        if not raw_compact or not last_compact:
            return ""

        if raw_compact == last_compact:
            return ""

        if last_compact in raw_compact:
            idx = raw_compact.rfind(last_compact)
            start = idx + len(last_compact)
            if start <= 0 or start >= len(raw_compact):
                return ""
            return utils.slice_from_compact_index(raw, start).strip()

        if utils.is_continuation_text(last_raw, raw):
            return ""

        return None

    def _process_raw_text(self, raw):
        """ê¸€ë¡œë²Œ íˆìŠ¤í† ë¦¬ + Suffix ë§¤ì¹­ (#GlobalHistorySuffix)

        ì´ ì„¸ì…˜ì—ì„œ í™•ì •ëœ ëª¨ë“  í…ìŠ¤íŠ¸ë¥¼ ëˆ„ì í•˜ê³ ,
        ìƒˆ rawì—ì„œ íˆìŠ¤í† ë¦¬ì˜ ë§ˆì§€ë§‰ ë¶€ë¶„(suffix) ì´í›„ë§Œ ì¶”ì¶œí•©ë‹ˆë‹¤.

        DOM ë£¨í•‘ì— ì™„ì „íˆ ë©´ì—­ì…ë‹ˆë‹¤.
        """
        if not raw:
            return

        raw = raw.strip()

        # ë™ì¼í•œ rawë©´ ë¬´ì‹œ (ë¹ ë¥¸ ê²½ë¡œ)
        if raw == self._last_raw_text:
            return
        self._last_raw_text = raw

        # rawë¥¼ compact (ê³µë°± ì œê±°)
        raw_compact = utils.compact_subtitle_text(raw)

        if not raw_compact:
            return

        # ìƒˆë¡œìš´ ë¶€ë¶„ ì¶”ì¶œ
        new_part = self._extract_new_part(raw, raw_compact)

        if not new_part or len(new_part.strip()) < 3:
            # ìƒˆ ë‚´ìš© ì—†ìŒ
            return

        new_part = new_part.strip()

        # ì½”ì–´ ì•Œê³ ë¦¬ì¦˜ ê²°ê³¼ë¥¼ ìµœê·¼ í™•ì • ìë§‰ ê¸°ì¤€ìœ¼ë¡œ í•œ ë²ˆ ë” ì •ì œí•´
        # ëŒ€ëŸ‰ ë°˜ë³µ ë¸”ë¡ì´ ê·¸ëŒ€ë¡œ ëˆ„ì ë˜ëŠ” ê²ƒì„ ë°©ì§€í•œë‹¤.
        with self.subtitle_lock:
            last_text = self.subtitles[-1].text if self.subtitles else ""

        if last_text:
            refined_part = utils.get_word_diff(last_text, new_part)
            if not refined_part or len(refined_part.strip()) < 3:
                return
            new_part = refined_part.strip()

        new_compact = utils.compact_subtitle_text(new_part)
        recent_compact_tail = self._confirmed_history_compact_tail(
            max_entries=12, max_compact_len=5000
        )
        if (
            len(new_compact) >= 20
            and recent_compact_tail
            and new_compact in recent_compact_tail
        ):
            return

        # íˆìŠ¤í† ë¦¬ì— ì¶”ê°€
        self._confirmed_compact += new_compact

        # suffix ê°±ì‹ 
        if len(self._confirmed_compact) >= self._suffix_length:
            self._trailing_suffix = self._confirmed_compact[-self._suffix_length :]
        else:
            self._trailing_suffix = self._confirmed_compact

        # ìë§‰ì— ì¶”ê°€
        self._add_text_to_subtitles(new_part)
        self._last_processed_raw = raw

    def _extract_new_part(self, raw: str, raw_compact: str) -> str:
        """íˆìŠ¤í† ë¦¬ì˜ suffix ì´í›„ì˜ ìƒˆ ë¶€ë¶„ë§Œ ì¶”ì¶œ

        Args:
            raw: ì›ë³¸ í…ìŠ¤íŠ¸
            raw_compact: ê³µë°± ì œê±°ëœ í…ìŠ¤íŠ¸

        Returns:
            ìƒˆë¡œìš´ ë¶€ë¶„ (ì›ë³¸ í˜•íƒœë¡œ)
        """
        # íˆìŠ¤í† ë¦¬ê°€ ì—†ìœ¼ë©´ ì „ì²´ê°€ ìƒˆë¡œìš´ ê²ƒ
        if not self._trailing_suffix:
            return raw

        # suffixê°€ raw_compactì— ìˆëŠ”ì§€ ê²€ìƒ‰ (rfind: ë§ˆì§€ë§‰ ìœ„ì¹˜ ê¸°ì¤€ìœ¼ë¡œ ê³¼ì‰ ì¶”ì¶œ ë°©ì§€)
        pos = raw_compact.rfind(self._trailing_suffix)

        if pos >= 0:
            # suffix ì´í›„ ë¶€ë¶„ë§Œ ë°˜í™˜
            start_idx = pos + len(self._trailing_suffix)
            if start_idx >= len(raw_compact):
                return ""  # suffix ì´í›„ì— ë‚´ìš© ì—†ìŒ
            return utils.slice_from_compact_index(raw, start_idx)

        # ì •ë§ ìƒˆë¡œìš´ ë¬¸ë§¥ - ì „ì²´ ë°˜í™˜
        return raw

    def _soft_resync(self) -> None:
        """ì „ì²´ ë¦¬ì…‹ ëŒ€ì‹ , ìµœê·¼ í™•ì • ìë§‰ì—ì„œ íˆìŠ¤í† ë¦¬ë¥¼ ì¬êµ¬ì„±í•œë‹¤.

        desync/ambiguous ë¦¬ì…‹ ì‹œ _confirmed_compactë¥¼ ì™„ì „ ì´ˆê¸°í™”í•˜ë©´
        ì´ë¯¸ ì²˜ë¦¬ëœ í…ìŠ¤íŠ¸ê°€ ì¬ìœ ì…ë˜ì–´ ëŒ€ëŸ‰ ì¤‘ë³µì´ ë°œìƒí•  ìˆ˜ ìˆë‹¤.
        ìµœê·¼ ìë§‰ì˜ compactë¥¼ ê¸°ë°˜ìœ¼ë¡œ suffixë¥¼ ë³µì›í•˜ë©´ ì´ë¥¼ ë°©ì§€í•œë‹¤.
        """
        with self.subtitle_lock:
            if self.subtitles:
                # ìµœê·¼ 5ê°œ ìë§‰ì˜ compactë¡œ íˆìŠ¤í† ë¦¬ ì¬êµ¬ì„±
                recent = " ".join(
                    e.text for e in self.subtitles[-5:] if e and e.text
                )
                self._confirmed_compact = utils.compact_subtitle_text(recent)
                if len(self._confirmed_compact) >= self._suffix_length:
                    self._trailing_suffix = self._confirmed_compact[
                        -self._suffix_length :
                    ]
                else:
                    self._trailing_suffix = self._confirmed_compact
                logger.info(
                    "ì†Œí”„íŠ¸ ë¦¬ì…‹: suffix=%s",
                    self._trailing_suffix[-20:] if self._trailing_suffix else "(empty)",
                )
            else:
                # ìë§‰ì´ ì—†ìœ¼ë©´ ì–´ì©” ìˆ˜ ì—†ì´ ì „ì²´ ë¦¬ì…‹
                self._confirmed_compact = ""
                self._trailing_suffix = ""
                logger.info("ì†Œí”„íŠ¸ ë¦¬ì…‹: ìë§‰ ì—†ìŒ, ì „ì²´ ë¦¬ì…‹")

    def _find_overlap(self, suffix: str, text: str) -> int:
        """suffixì˜ ë’·ë¶€ë¶„ê³¼ textì˜ ì•ë¶€ë¶„ì´ ê²¹ì¹˜ëŠ” ê¸¸ì´ ë°˜í™˜"""
        max_overlap = min(len(suffix), len(text))
        for i in range(max_overlap, 0, -1):
            if suffix[-i:] == text[:i]:
                return i
        return 0

    def _add_text_to_subtitles(self, text: str) -> None:
        """í™•ì •ëœ í…ìŠ¤íŠ¸ë¥¼ SubtitleEntryë¡œ ë³€í™˜í•˜ì—¬ ì €ì¥"""
        if not text or len(text) < 3:
            return

        with self.subtitle_lock:
            # ë§ˆì§€ë§‰ ì—”íŠ¸ë¦¬ì— ì´ì–´ë¶™ì¼ì§€, ìƒˆ ì—”íŠ¸ë¦¬ ë§Œë“¤ì§€ ê²°ì •
            if self.subtitles:
                last_entry = self.subtitles[-1]
                # [FIX] end_timeì´ Noneì¸ ê²½ìš° ë°©ì–´ ì²˜ë¦¬ (reflow í›„ ë°œìƒ ê°€ëŠ¥)
                can_append = False
                if last_entry.end_time is not None:
                    elapsed = (datetime.now() - last_entry.end_time).total_seconds()
                    if elapsed < 5.0 and len(last_entry.text) + len(text) < 300:
                        can_append = True

                if can_append:
                    old_chars = last_entry.char_count
                    old_words = last_entry.word_count
                    last_entry.update_text(
                        self._join_stream_text(last_entry.text, text)
                    )
                    last_entry.end_time = datetime.now()
                    self._cached_total_chars += last_entry.char_count - old_chars
                    self._cached_total_words += last_entry.word_count - old_words
                else:
                    # ìƒˆ ì—”íŠ¸ë¦¬
                    entry = SubtitleEntry(text)
                    entry.start_time = datetime.now()
                    entry.end_time = datetime.now()
                    self.subtitles.append(entry)
                    self._cached_total_chars += entry.char_count
                    self._cached_total_words += entry.word_count
            else:
                # ì²« ì—”íŠ¸ë¦¬
                entry = SubtitleEntry(text)
                entry.start_time = datetime.now()
                entry.end_time = datetime.now()
                self.subtitles.append(entry)
                self._cached_total_chars += entry.char_count
                self._cached_total_words += entry.word_count

        # í‚¤ì›Œë“œ ì•Œë¦¼ í™•ì¸
        self._check_keyword_alert(text)

        # ì¹´ìš´íŠ¸ ë¼ë²¨ ì—…ë°ì´íŠ¸
        self._update_count_label()

        # UI ê°±ì‹ 
        self._refresh_text(force_full=False)

        # ì‹¤ì‹œê°„ ì €ì¥
        if self.realtime_file:
            try:
                timestamp = datetime.now().strftime("%H:%M:%S")
                self.realtime_file.write(f"[{timestamp}] {text}\n")
                self.realtime_file.flush()
                self._realtime_error_count = 0
            except IOError as e:
                self._realtime_error_count = (
                    getattr(self, "_realtime_error_count", 0) + 1
                )
                if self._realtime_error_count == 3:
                    logger.error(f"ì‹¤ì‹œê°„ ì €ì¥ ì—°ì† ì‹¤íŒ¨: {e}")

    def _set_preview_text(self, text: str) -> None:
        """ë¯¸ë¦¬ë³´ê¸° í…ìŠ¤íŠ¸ í‘œì‹œ/ìˆ¨ê¹€"""
        if not hasattr(self, "preview_frame"):
            return
        content = (text or "").strip()
        if content:
            self.preview_label.setText(content)
            self.preview_frame.show()
        else:
            self.preview_label.setText("")
            self.preview_frame.hide()

    def _clear_preview(self) -> None:
        """ë¯¸ë¦¬ë³´ê¸° ì´ˆê¸°í™”"""
        self._set_preview_text("")

    def _update_preview(self, raw: str) -> None:
        """ë¯¸ë¦¬ë³´ê¸° ì—…ë°ì´íŠ¸ (ë³„ë„ UI)"""
        self._set_preview_text(raw)

    def _render_subtitles(self, force_full: bool = False) -> None:
        """Render subtitles with incremental updates when possible."""
        scrollbar = self.subtitle_text.verticalScrollBar()
        preserve_scroll = self._user_scrolled_up and scrollbar is not None
        saved_scroll = None
        if preserve_scroll:
            saved_scroll = scrollbar.value()
            scrollbar.blockSignals(True)

        # Snapshot for rendering
        with self.subtitle_lock:
            subtitles_copy = list(self.subtitles)

        total_count = len(subtitles_copy)

        # ì„±ëŠ¥ ìµœì í™”: ëŒ€ëŸ‰ ìë§‰ ì‹œ ìµœê·¼ í•­ëª©ë§Œ ë Œë”ë§ (#4)
        render_offset = 0
        if total_count > Config.MAX_RENDER_ENTRIES:
            render_offset = total_count - Config.MAX_RENDER_ENTRIES
            subtitles_copy = subtitles_copy[render_offset:]

        visible_count = len(subtitles_copy)
        last_text = subtitles_copy[-1].text if subtitles_copy else ""

        show_ts = self.timestamp_action.isChecked()

        # ë§ˆì§€ë§‰ ì¶œë ¥ëœ íƒ€ì„ìŠ¤íƒ¬í”„ (í¬ì†Œ íƒ€ì„ìŠ¤íƒ¬í”„ìš©)
        # ë Œë”ë§ ì‹œì‘ ì‹œ ì´ˆê¸°í™”
        if not hasattr(self, "_last_printed_ts"):
            self._last_printed_ts = None

        previous_visible_count = max(
            0, self._last_rendered_count - getattr(self, "_last_render_offset", 0)
        )
        offset_changed = render_offset != getattr(self, "_last_render_offset", 0)
        show_ts_changed = (
            self._last_render_show_ts is not None and show_ts != self._last_render_show_ts
        )

        needs_full_render = (
            force_full
            or offset_changed
            or show_ts_changed
            or (total_count < self._last_rendered_count)
            or (previous_visible_count > visible_count)
            or (
                total_count == self._last_rendered_count
                and last_text != self._last_rendered_last_text
            )
        )

        if needs_full_render:
            self.subtitle_text.clear()
            self._last_printed_ts = None  # í’€ ë Œë”ë§ ì‹œ ì´ˆê¸°í™”
            cursor = self.subtitle_text.textCursor()

            for i, entry in enumerate(subtitles_copy):
                same_second = False
                if i > 0:
                    prev_entry = subtitles_copy[i - 1]
                    same_second = entry.timestamp.replace(
                        microsecond=0
                    ) == prev_entry.timestamp.replace(microsecond=0)
                    cursor.insertText(" " if same_second else "\n\n")

                if show_ts:
                    # 1ë¶„ ê°„ê²© íƒ€ì„ìŠ¤íƒ¬í”„ ë¡œì§
                    should_print = False
                    if not same_second:
                        if self._last_printed_ts is None:
                            should_print = True
                        elif (
                            entry.timestamp - self._last_printed_ts
                        ).total_seconds() >= 60:
                            should_print = True

                    if should_print:
                        cursor.insertText(
                            f"[{entry.timestamp.strftime('%H:%M:%S')}] ",
                            self._timestamp_fmt,
                        )
                        self._last_printed_ts = entry.timestamp

                self._insert_highlighted_text(cursor, entry.text)

        else:
            cursor = self.subtitle_text.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.End)

            start_local_idx = min(previous_visible_count, visible_count)
            for local_idx in range(start_local_idx, visible_count):
                entry = subtitles_copy[local_idx]
                same_second = False
                if local_idx > 0:
                    prev_entry = subtitles_copy[local_idx - 1]
                    same_second = entry.timestamp.replace(
                        microsecond=0
                    ) == prev_entry.timestamp.replace(microsecond=0)
                    cursor.insertText(" " if same_second else "\n\n")

                if show_ts:
                    # 1ë¶„ ê°„ê²© íƒ€ì„ìŠ¤íƒ¬í”„ ë¡œì§ (ì¦ë¶„)
                    should_print = False
                    if not same_second:
                        if self._last_printed_ts is None:
                            should_print = True
                        elif (
                            entry.timestamp - self._last_printed_ts
                        ).total_seconds() >= 60:
                            should_print = True

                    if should_print:
                        cursor.insertText(
                            f"[{entry.timestamp.strftime('%H:%M:%S')}] ",
                            self._timestamp_fmt,
                        )
                        self._last_printed_ts = entry.timestamp

                self._insert_highlighted_text(cursor, entry.text)

        self._last_rendered_count = total_count
        self._last_rendered_last_text = last_text
        self._last_render_offset = render_offset
        self._last_render_show_ts = show_ts

        if self.auto_scroll_check.isChecked() and not self._user_scrolled_up:
            self.subtitle_text.moveCursor(QTextCursor.MoveOperation.End)

        if preserve_scroll:
            scrollbar.setValue(min(saved_scroll, scrollbar.maximum()))
            scrollbar.blockSignals(False)

    def _on_scroll_changed(self) -> None:
        """ìŠ¤í¬ë¡¤ë°” ìœ„ì¹˜ ë³€ê²½ ê°ì§€ - ìŠ¤ë§ˆíŠ¸ ìŠ¤í¬ë¡¤ìš©"""
        scrollbar = self.subtitle_text.verticalScrollBar()
        # ìŠ¤í¬ë¡¤ì´ ë§¨ ì•„ë˜ì—ì„œ ì¼ì • ê±°ë¦¬ ì´ë‚´ë©´ ìë™ ìŠ¤í¬ë¡¤ í™œì„±í™”
        at_bottom = (
            scrollbar.value() >= scrollbar.maximum() - Config.SCROLL_BOTTOM_THRESHOLD
        )

        if at_bottom:
            self._user_scrolled_up = False
            if hasattr(self, "scroll_to_bottom_btn"):
                self.scroll_to_bottom_btn.hide()
        else:
            # ì‚¬ìš©ìê°€ ìœ„ë¡œ ìŠ¤í¬ë¡¤í•œ ê²½ìš°
            self._user_scrolled_up = True
            # ì¶”ì¶œ ì¤‘ì¼ ë•Œë§Œ ë²„íŠ¼ í‘œì‹œ
            if self.is_running and hasattr(self, "scroll_to_bottom_btn"):
                self.scroll_to_bottom_btn.show()

    def _scroll_to_bottom(self) -> None:
        """ë§¨ ì•„ë˜ë¡œ ìŠ¤í¬ë¡¤í•˜ê³  ìë™ ìŠ¤í¬ë¡¤ ì¬ê°œ"""
        self._user_scrolled_up = False
        self.subtitle_text.moveCursor(QTextCursor.MoveOperation.End)
        if hasattr(self, "scroll_to_bottom_btn"):
            self.scroll_to_bottom_btn.hide()

    def _toggle_stats_panel(self) -> None:
        """í†µê³„ íŒ¨ë„ ìˆ¨ê¸°ê¸°/ë³´ì´ê¸°"""
        if hasattr(self, "stats_group") and hasattr(self, "toggle_stats_btn"):
            is_visible = self.stats_group.isVisible()
            self.stats_group.setVisible(not is_visible)

            if is_visible:
                self.toggle_stats_btn.setText("ğŸ“Š í†µê³„ ë³´ì´ê¸°")
                # ìë§‰ ì˜ì—­ì´ ì „ì²´ ë„ˆë¹„ë¥¼ ì‚¬ìš©í•˜ë„ë¡
                self.main_splitter.setSizes([1080, 0])
            else:
                self.toggle_stats_btn.setText("ğŸ“Š í†µê³„ ìˆ¨ê¸°ê¸°")
                self.main_splitter.setSizes([860, 220])

    def _refresh_text(self, force_full: bool = False) -> None:
        """í™•ì •ëœ ìë§‰ë§Œ í‘œì‹œ (ì§„í–‰ ì¤‘ì¸ ìë§‰ ì—†ìŒ)"""
        self._render_subtitles(force_full=force_full)

    def _refresh_text_full(self, *_args) -> None:
        """ë Œë” ìºì‹œë¥¼ ë¬´ì‹œí•˜ê³  ì „ì²´ ë‹¤ì‹œ ê·¸ë¦¬ê¸°"""
        self._refresh_text(force_full=True)

    def _finalize_subtitle(self, text):
        """ìë§‰ í™•ì • - ë‹¨ì–´ ë‹¨ìœ„ ì¦ë¶„ ë³‘í•© (Incremental Word Accumulation)

        [Thread Safety] subtitle_lockìœ¼ë¡œ self.subtitles ì ‘ê·¼ ë³´í˜¸
        """
        if not text:
            return

        # [Fix] ì¤‘ë³µ ì²˜ë¦¬ ë°©ì§€: _process_raw_textì—ì„œ ì´ë¯¸ ì²˜ë¦¬ëœ í…ìŠ¤íŠ¸ë©´ ê±´ë„ˆëœ€
        # _last_processed_rawê°€ ë™ì¼í•˜ë©´ ì´ë¯¸ _process_raw_textì—ì„œ ìë§‰ì´ ì¶”ê°€ë¨
        if text == self._last_processed_raw:
            return

        # [Fix] í™•ì •ëœ ìë§‰ì„ í”„ë¡œì„¸ì„œì— ë“±ë¡ (#SubtitleRepetitionFix)
        self.subtitle_processor.add_confirmed(text)

        # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ì ‘ê·¼
        with self.subtitle_lock:
            # ì´ì „ì— í™•ì •ëœ ìë§‰ì´ ìˆìœ¼ë©´, ê²¹ì¹˜ëŠ” ë¶€ë¶„ ì œê±°
            if self.subtitles:
                last_text = self.subtitles[-1].text

                # [Smart Merge] ë‹¨ì–´ ë‹¨ìœ„ ê²¹ì¹¨ ë¶„ì„í•˜ì—¬ ìƒˆë¡œìš´ ë‹¨ì–´ë§Œ ì¶”ì¶œ
                new_part = utils.get_word_diff(last_text, text)

                if new_part:
                    # [Length Check] ë¬¸ì¥ì´ ë„ˆë¬´ ê¸¸ì–´ì§€ë©´ ê°•ì œë¡œ ëŠê³  ì‹œ íƒ€ì„ìŠ¤íƒ¬í”„ ìƒì„±
                    current_len = len(last_text)
                    new_len = len(new_part)

                    # ê¸°ì¡´ ê¸¸ì´ì— ìƒˆ ë‚´ìš©ì„ ë”í–ˆì„ ë•Œ ìµœëŒ€ ê¸¸ì´ë¥¼ ì´ˆê³¼í•˜ë©´ ë¶„ë¦¬
                    if current_len + new_len > Config.STREAM_SUBTITLE_MAX_LENGTH:
                        entry = SubtitleEntry(new_part)
                        entry.start_time = datetime.now()
                        entry.end_time = datetime.now()
                        self.subtitles.append(entry)
                        # í†µê³„ ìºì‹œ ê°±ì‹ 
                        self._cached_total_chars += entry.char_count
                        self._cached_total_words += entry.word_count

                        if self.realtime_file:
                            try:
                                # ë¶„ë¦¬ëœ ìƒˆ ë¬¸ì¥ìœ¼ë¡œ ì €ì¥
                                timestamp = entry.timestamp.strftime("%H:%M:%S")
                                self.realtime_file.write(f"[{timestamp}] {new_part}\n")
                                self.realtime_file.flush()
                            except IOError as e:
                                logger.warning(f"ì‹¤ì‹œê°„ ì €ì¥ ì“°ê¸° ì˜¤ë¥˜: {e}")
                    else:
                        # ìƒˆë¡œìš´ ë‚´ìš©ì´ ìˆìœ¼ë©´ ì´ì–´ë¶™ì´ê¸° - update_text()ë¡œ ìºì‹œ ê°±ì‹  í¬í•¨
                        old_chars = self.subtitles[-1].char_count
                        old_words = self.subtitles[-1].word_count
                        self.subtitles[-1].update_text(
                            self._join_stream_text(self.subtitles[-1].text, new_part)
                        )
                        self.subtitles[-1].end_time = datetime.now()
                        # í†µê³„ ìºì‹œ ê°±ì‹ 
                        self._cached_total_chars += (
                            self.subtitles[-1].char_count - old_chars
                        )
                        self._cached_total_words += (
                            self.subtitles[-1].word_count - old_words
                        )

                        # ì‹¤ì‹œê°„ ì €ì¥
                        if self.realtime_file:
                            try:
                                # + ê¸°í˜¸ë¡œ ì´ì–´ë¶™ì—¬ì§„ ë‚´ìš©ì„ì„ í‘œì‹œ
                                self.realtime_file.write(f"+ {new_part}\n")
                                self.realtime_file.flush()
                            except IOError as e:
                                logger.warning(f"ì‹¤ì‹œê°„ ì €ì¥ ì“°ê¸° ì˜¤ë¥˜: {e}")
                    return
                else:
                    # ê²¹ì¹˜ëŠ” ë‚´ìš©ë§Œ ìˆê³  ìƒˆë¡œìš´ ë‚´ìš©ì´ ì—†ìœ¼ë©´ (ë¶€ë¶„ì§‘í•©)
                    # ì‹œê°„ë§Œ ê°±ì‹ í•˜ê³  ì¢…ë£Œ
                    self.subtitles[-1].end_time = datetime.now()
                    return

            # ì²« ë²ˆì§¸ ìë§‰ ë˜ëŠ” ì™„ì „íˆ ìƒˆë¡œìš´ ë¬¸ì¥
            entry = SubtitleEntry(text)
            entry.start_time = datetime.now()
            entry.end_time = datetime.now()

            self.subtitles.append(entry)
            # í†µê³„ ìºì‹œ ê°±ì‹ 
            self._cached_total_chars += entry.char_count
            self._cached_total_words += entry.word_count

        # í‚¤ì›Œë“œ ì•Œë¦¼ í™•ì¸ (ë½ ë°–ì—ì„œ - UI ì‘ì—…)
        self._check_keyword_alert(text)

        # ì¹´ìš´íŠ¸ ë¼ë²¨ ì—…ë°ì´íŠ¸
        self._update_count_label()

        # ì‹¤ì‹œê°„ ì €ì¥
        if self.realtime_file:
            try:
                timestamp = entry.timestamp.strftime("%H:%M:%S")
                self.realtime_file.write(f"[{timestamp}] {text}\n")
                self.realtime_file.flush()
            except IOError as e:
                logger.warning(f"ì‹¤ì‹œê°„ ì €ì¥ ì“°ê¸° ì˜¤ë¥˜: {e}")

        self._refresh_text(force_full=False)

    def _insert_highlighted_text(self, cursor, text):
        """í…ìŠ¤íŠ¸ì—ì„œ í‚¤ì›Œë“œë§Œ í•˜ì´ë¼ì´íŠ¸ (ì„±ëŠ¥ ìµœì í™”: ìºì‹œëœ íŒ¨í„´/í¬ë§· ì‚¬ìš©)"""
        # í‚¤ì›Œë“œ ìºì‹œê°€ ë¹„ì–´ìˆìœ¼ë©´ ì¼ë°˜ í…ìŠ¤íŠ¸ë¡œ ì‚½ì…
        if not self._keyword_pattern:
            cursor.insertText(text, self._normal_fmt)
            return

        # ìºì‹œëœ íŒ¨í„´ìœ¼ë¡œ ë¶„í• 
        parts = self._keyword_pattern.split(text)

        for part in parts:
            if not part:  # ë¹ˆ ë¬¸ìì—´ ê±´ë„ˆë›°ê¸°
                continue

            if part.lower() in self._keywords_lower_set:
                # í‚¤ì›Œë“œ: ìºì‹œëœ í•˜ì´ë¼ì´íŠ¸ í¬ë§· ì‚¬ìš©
                cursor.insertText(part, self._highlight_fmt)
            else:
                # ì¼ë°˜ í…ìŠ¤íŠ¸: ìºì‹œëœ ì¼ë°˜ í¬ë§· ì‚¬ìš©
                cursor.insertText(part, self._normal_fmt)

    def _rebuild_keyword_cache(
        self, keywords: list, update_settings: bool = True, refresh: bool = True
    ) -> None:
        """í•˜ì´ë¼ì´íŠ¸ í‚¤ì›Œë“œ ìºì‹œ ì¬êµ¬ì„±"""
        cleaned = [k.strip() for k in keywords if k and k.strip()]
        self.keywords = cleaned
        self._keywords_lower_set = {k.lower() for k in cleaned}

        if cleaned:
            pattern = "|".join(re.escape(k) for k in cleaned)
            try:
                self._keyword_pattern = re.compile(f"({pattern})", re.IGNORECASE)
            except re.error:
                self._keyword_pattern = None
        else:
            self._keyword_pattern = None

        if update_settings:
            self.settings.setValue("highlight_keywords", ", ".join(self.keywords))

        if refresh and hasattr(self, "subtitle_text"):
            self._refresh_text(force_full=True)

    def _update_keyword_cache(self):
        """í‚¤ì›Œë“œ íŒ¨í„´ ìºì‹œ ì—…ë°ì´íŠ¸ (ë””ë°”ìš´ì‹± ì ìš©)"""
        # ë””ë°”ìš´ì‹±: ì´ì „ íƒ€ì´ë¨¸ ì·¨ì†Œ
        if (
            hasattr(self, "_keyword_debounce_timer")
            and self._keyword_debounce_timer.isActive()
        ):
            self._keyword_debounce_timer.stop()

        def do_update():
            self._perform_keyword_cache_update()

        # 300ms í›„ ì‹¤í–‰
        self._keyword_debounce_timer = QTimer(self)
        self._keyword_debounce_timer.setSingleShot(True)
        self._keyword_debounce_timer.timeout.connect(do_update)
        self._keyword_debounce_timer.start(300)

    def _perform_keyword_cache_update(self):
        """ì‹¤ì œ í‚¤ì›Œë“œ ìºì‹œ ì—…ë°ì´íŠ¸ ë¡œì§"""
        try:
            if hasattr(self, "keyword_input"):
                raw_text = self.keyword_input.text()
            else:
                raw_text = ", ".join(self.keywords)

            keywords = [k.strip() for k in raw_text.split(",") if k.strip()]
            self._rebuild_keyword_cache(keywords, update_settings=True, refresh=True)
        except Exception as e:
            logger.error(f"í‚¤ì›Œë“œ ìºì‹œ ì—…ë°ì´íŠ¸ ì˜¤ë¥˜: {e}")

    # ========== í€µ ì•¡ì…˜ ==========

    def _copy_to_clipboard(self) -> None:
        """ìë§‰ ì „ì²´ë¥¼ í´ë¦½ë³´ë“œì— ë³µì‚¬"""
        with self.subtitle_lock:
            if not self.subtitles:
                self._show_toast("ë³µì‚¬í•  ìë§‰ì´ ì—†ìŠµë‹ˆë‹¤", "warning")
                return

            text = "\n".join(s.text for s in self.subtitles)

        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        self._show_toast(f"ğŸ“‹ {len(self.subtitles)}ê°œ ìë§‰ ë³µì‚¬ë¨", "success")

    def _clear_subtitles(self) -> None:
        """ìë§‰ ëª©ë¡ ì´ˆê¸°í™”"""
        with self.subtitle_lock:
            count = len(self.subtitles)
        if not count:
            self._show_toast("ì§€ìš¸ ìë§‰ì´ ì—†ìŠµë‹ˆë‹¤", "warning")
            return

        reply = QMessageBox.question(
            self,
            "ìë§‰ ì§€ìš°ê¸°",
            f"í˜„ì¬ {count}ê°œì˜ ìë§‰ì„ ëª¨ë‘ ì§€ìš°ì‹œê² ìŠµë‹ˆê¹Œ?\nì´ ì‘ì—…ì€ ë˜ëŒë¦´ ìˆ˜ ì—†ìŠµë‹ˆë‹¤.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )

        if reply != QMessageBox.StandardButton.Yes:
            return

        self._replace_subtitles_and_refresh([], keep_history_from_subtitles=False)
        self._show_toast(f"ğŸ—‘ï¸ {count}ê°œ ìë§‰ ì‚­ì œë¨", "success")

    def _toggle_theme_from_button(self) -> None:
        """íˆ´ë°” ë²„íŠ¼ì—ì„œ í…Œë§ˆ ì „í™˜"""
        self._toggle_theme()
        # ë²„íŠ¼ ì•„ì´ì½˜ ì—…ë°ì´íŠ¸
        if hasattr(self, "theme_toggle_btn"):
            self.theme_toggle_btn.setText("ğŸŒ™" if self.is_dark_theme else "â˜€ï¸")

    # ========== í†µê³„ ==========

    def _update_stats(self):
        """í†µê³„ ì—…ë°ì´íŠ¸ (ì„±ëŠ¥ ìµœì í™”: ìºì‹œëœ í†µê³„ ì‚¬ìš©)"""
        if self.start_time:
            elapsed = int(time.time() - self.start_time)
            h, r = divmod(elapsed, 3600)
            m, s = divmod(r, 60)
            self.stat_time.setText(f"â±ï¸ ì‹¤í–‰ ì‹œê°„: {h:02d}:{m:02d}:{s:02d}")

            # ìºì‹œëœ í†µê³„ ì‚¬ìš© (ìŠ¤ë ˆë“œ ì•ˆì „)
            with self.subtitle_lock:
                subtitle_count = len(self.subtitles)

            # ìºì‹œëœ ê°’ ì§ì ‘ ì‚¬ìš© (ë§¤ë²ˆ ì¬ê³„ì‚° ëŒ€ì‹ )
            total_chars = self._cached_total_chars
            total_words = self._cached_total_words

            self.stat_chars.setText(f"ğŸ“ ê¸€ì ìˆ˜: {total_chars:,}")
            self.stat_words.setText(f"ğŸ“– ë‹¨ì–´ ìˆ˜: {total_words:,}")
            self.stat_sents.setText(f"ğŸ’¬ ë¬¸ì¥ ìˆ˜: {subtitle_count}")

            if elapsed > 0:
                cpm = int(total_chars / (elapsed / 60))
                self.stat_cpm.setText(f"âš¡ ë¶„ë‹¹ ê¸€ì: {cpm}")

    # ========== ê²€ìƒ‰ ==========

    def _show_search(self):
        self.search_frame.show()
        self.search_input.setFocus()
        self.search_input.selectAll()

    def _hide_search(self):
        self.search_frame.hide()
        self._refresh_text(force_full=True)

    def _do_search(self):
        query = self.search_input.text()
        if not query:
            return

        query_l = query.lower()
        text_l = self.subtitle_text.toPlainText().lower()
        self.search_matches = []

        start = 0
        while True:
            idx = text_l.find(query_l, start)
            if idx == -1:
                break
            self.search_matches.append(idx)
            start = idx + 1

        self.search_idx = 0
        self.search_count.setText(f"{len(self.search_matches)}ê°œ")

        if self.search_matches:
            self._highlight_search(0)

    def _nav_search(self, delta):
        if not self.search_matches:
            return

        self.search_idx = (self.search_idx + delta) % len(self.search_matches)
        self._highlight_search(self.search_idx)

    def _highlight_search(self, idx):
        if not self.search_matches:
            return

        pos = self.search_matches[idx]
        query = self.search_input.text()

        cursor = self.subtitle_text.textCursor()
        cursor.setPosition(pos)
        cursor.setPosition(pos + len(query), QTextCursor.MoveMode.KeepAnchor)

        self.subtitle_text.setTextCursor(cursor)
        self.subtitle_text.ensureCursorVisible()

        self.search_count.setText(f"{idx + 1}/{len(self.search_matches)}")

    # ========== í‚¤ì›Œë“œ ==========

    def _set_keywords(self):
        """í•˜ì´ë¼ì´íŠ¸ í‚¤ì›Œë“œ ì„¤ì •"""
        current = ", ".join(self.keywords)
        text, ok = QInputDialog.getText(
            self,
            "í•˜ì´ë¼ì´íŠ¸ í‚¤ì›Œë“œ ì„¤ì •",
            "í•˜ì´ë¼ì´íŠ¸í•  í‚¤ì›Œë“œ (ì‰¼í‘œë¡œ êµ¬ë¶„):",
            text=current,
        )

        if ok:
            keywords = [k.strip() for k in text.split(",") if k.strip()]
            if hasattr(self, "keyword_input"):
                self.keyword_input.blockSignals(True)
                self.keyword_input.setText(", ".join(keywords))
                self.keyword_input.blockSignals(False)
            self._rebuild_keyword_cache(keywords, update_settings=True, refresh=True)
            self._show_toast(f"í•˜ì´ë¼ì´íŠ¸ í‚¤ì›Œë“œ {len(keywords)}ê°œ ì„¤ì •ë¨", "success")

    def _rebuild_alert_keyword_cache(
        self, keywords: list, update_settings: bool = True
    ) -> None:
        """ì•Œë¦¼ í‚¤ì›Œë“œ ìºì‹œë¥¼ ì¬êµ¬ì„±í•œë‹¤."""
        cleaned = [k.strip() for k in keywords if k and k.strip()]
        self.alert_keywords = cleaned
        self._alert_keywords_cache = [(k, k.lower()) for k in cleaned]
        if update_settings:
            self.settings.setValue("alert_keywords", ", ".join(cleaned))

    def _set_alert_keywords(self):
        """ì•Œë¦¼ í‚¤ì›Œë“œ ì„¤ì • - í•´ë‹¹ í‚¤ì›Œë“œ ê°ì§€ ì‹œ í† ìŠ¤íŠ¸ ì•Œë¦¼"""
        current = ", ".join(self.alert_keywords)
        text, ok = QInputDialog.getText(
            self,
            "ì•Œë¦¼ í‚¤ì›Œë“œ ì„¤ì •",
            "ì•Œë¦¼ì„ ë°›ì„ í‚¤ì›Œë“œ (ì‰¼í‘œë¡œ êµ¬ë¶„):\nì˜ˆ: ë²•ì•ˆ, ì˜ê²°, í†µê³¼",
            text=current,
        )

        if ok:
            self._rebuild_alert_keyword_cache(
                [k.strip() for k in text.split(",") if k.strip()],
                update_settings=True,
            )
            self._show_toast(
                f"ì•Œë¦¼ í‚¤ì›Œë“œ {len(self.alert_keywords)}ê°œ ì„¤ì •ë¨", "success"
            )

    def _check_keyword_alert(self, text: str):
        """í‚¤ì›Œë“œ í¬í•¨ ì‹œ ì•Œë¦¼ í‘œì‹œ"""
        if not self._alert_keywords_cache:
            return

        text_lower = text.lower()
        for original, keyword_lower in self._alert_keywords_cache:
            if keyword_lower and keyword_lower in text_lower:
                self._show_toast(f"ğŸ”” í‚¤ì›Œë“œ ê°ì§€: {original}", "warning", 5000)
                break  # í•œ ë²ˆë§Œ ì•Œë¦¼

    # ========== ìë™ ë°±ì—… ==========

    def _auto_backup(self):
        """ìë™ ë°±ì—… ì‹¤í–‰"""
        if not self.subtitles:
            return

        # UI ìŠ¤ë ˆë“œê°€ ë©ˆì¶”ì§€ ì•Šë„ë¡(ê¸´ ì„¸ì…˜/ëŒ€ìš©ëŸ‰) íŒŒì¼ I/OëŠ” ë°±ê·¸ë¼ìš´ë“œì—ì„œ ì²˜ë¦¬
        if not self._auto_backup_lock.acquire(blocking=False):
            return  # ì´ë¯¸ ë°±ì—… ì¤‘

        try:
            backup_dir = Path(Config.BACKUP_DIR)
            backup_dir.mkdir(exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_file = backup_dir / f"backup_{timestamp}.json"

            # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ìŠ¤ëƒ…ìƒ· ìƒì„± (UI ìŠ¤ë ˆë“œ)
            with self.subtitle_lock:
                subtitles_copy = [s.to_dict() for s in self.subtitles]

            data = {
                "version": Config.VERSION,
                "created": datetime.now().isoformat(),
                "url": self._get_current_url(),
                "subtitles": subtitles_copy,
            }

        except Exception as e:
            try:
                self._auto_backup_lock.release()
            except Exception:
                pass
            logger.error(f"ìë™ ë°±ì—… ì¤€ë¹„ ì˜¤ë¥˜: {e}")
            return

        def write_backup():
            try:
                utils.atomic_write_json(
                    backup_file,
                    data,
                    ensure_ascii=False,
                    indent=2,
                )

                # ì˜¤ë˜ëœ ë°±ì—… ì‚­ì œ (ìµœëŒ€ ê°œìˆ˜ ìœ ì§€)
                self._cleanup_old_backups()

                logger.info(f"ìë™ ë°±ì—… ì™„ë£Œ: {backup_file}")
            except Exception as e:
                logger.error(f"ìë™ ë°±ì—… ì˜¤ë¥˜: {e}")
            finally:
                try:
                    self._auto_backup_lock.release()
                except Exception:
                    pass

        threading.Thread(target=write_backup, daemon=True).start()

    def _cleanup_old_backups(self):
        """ì˜¤ë˜ëœ ë°±ì—… íŒŒì¼ ì •ë¦¬"""
        try:
            backup_dir = Path(Config.BACKUP_DIR)
            backups = sorted(backup_dir.glob("backup_*.json"), reverse=True)

            # ìµœëŒ€ ê°œìˆ˜ ì´ˆê³¼ë¶„ ì‚­ì œ
            for old_backup in backups[Config.MAX_BACKUP_COUNT :]:
                old_backup.unlink()
                logger.debug(f"ì˜¤ë˜ëœ ë°±ì—… ì‚­ì œ: {old_backup}")
        except Exception as e:
            logger.warning(f"ë°±ì—… ì •ë¦¬ ì¤‘ ì˜¤ë¥˜: {e}")

    # ========== íŒŒì¼ ì €ì¥ ==========

    def _save_in_background(
        self, save_func, path: str, success_msg: str, error_prefix: str
    ):
        """ë°±ê·¸ë¼ìš´ë“œì—ì„œ íŒŒì¼ ì €ì¥ (ìë§‰ ìˆ˜ì§‘ ì¤‘ë‹¨ ì—†ì´)

        Args:
            save_func: ì‹¤ì œ ì €ì¥ì„ ìˆ˜í–‰í•˜ëŠ” í•¨ìˆ˜ (pathë¥¼ ì¸ìë¡œ ë°›ìŒ)
            path: ì €ì¥í•  íŒŒì¼ ê²½ë¡œ
            success_msg: ì„±ê³µ ì‹œ í† ìŠ¤íŠ¸ ë©”ì‹œì§€
            error_prefix: ì‹¤íŒ¨ ì‹œ ì—ëŸ¬ ë©”ì‹œì§€ ì ‘ë‘ì–´
        """

        def background_save():
            try:
                save_func(path)
                # UI ìŠ¤ë ˆë“œë¡œ ì•ˆì „í•˜ê²Œ ì „ë‹¬ (Queue ê¸°ë°˜)
                self.message_queue.put(
                    ("toast", {"message": success_msg, "toast_type": "success"})
                )
            except Exception as e:
                logger.error(f"{error_prefix}: {e}")
                self.message_queue.put(
                    (
                        "toast",
                        {
                            "message": f"{error_prefix}: {e}",
                            "toast_type": "error",
                            "duration": 5000,
                        },
                    )
                )

        # ë°±ê·¸ë¼ìš´ë“œ ìŠ¤ë ˆë“œì—ì„œ ì €ì¥ ì‹¤í–‰
        save_thread = threading.Thread(target=background_save, daemon=True)
        save_thread.start()

        # ì €ì¥ ì‹œì‘ ì•Œë¦¼ (ì¦‰ì‹œ)
        self._show_toast(f"ğŸ’¾ ì €ì¥ ì¤‘... ({Path(path).name})", "info", 1500)

    def _get_accumulated_text(self):
        with self.subtitle_lock:
            return "\n".join(s.text for s in self.subtitles)

    def _export_stats(self):
        """ìë§‰ í†µê³„ ë‚´ë³´ë‚´ê¸°"""
        if not self.subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ë‚´ë³´ë‚¼ ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ìŠ¤ëƒ…ìƒ· ìƒì„±
        with self.subtitle_lock:
            subtitles_snapshot = list(self.subtitles)

        # í†µê³„ ê³„ì‚°
        total_chars = sum(len(s.text) for s in subtitles_snapshot)
        total_words = sum(len(s.text.split()) for s in subtitles_snapshot)

        # ì‹œê°„ëŒ€ë³„ í†µê³„
        hour_counts = {}
        for entry in subtitles_snapshot:
            hour = entry.timestamp.hour
            hour_counts[hour] = hour_counts.get(hour, 0) + 1

        # ê°€ì¥ ê¸´/ì§§ì€ ë¬¸ì¥
        longest = max(subtitles_snapshot, key=lambda s: len(s.text))
        shortest = min(subtitles_snapshot, key=lambda s: len(s.text))

        # íŒŒì¼ ì €ì¥
        filename = f"ìë§‰í†µê³„_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        path, _ = QFileDialog.getSaveFileName(
            self, "í†µê³„ ë‚´ë³´ë‚´ê¸°", filename, "í…ìŠ¤íŠ¸ (*.txt)"
        )

        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write("=" * 50 + "\n")
                    f.write("        ğŸ›ï¸ êµ­íšŒ ìë§‰ í†µê³„ ë³´ê³ ì„œ\n")
                    f.write("=" * 50 + "\n\n")

                    f.write(
                        f"ìƒì„± ì¼ì‹œ: {datetime.now().strftime('%Yë…„ %mì›” %dì¼ %H:%M:%S')}\n\n"
                    )

                    f.write("ğŸ“Š ê¸°ë³¸ í†µê³„\n")
                    f.write("-" * 30 + "\n")
                    f.write(f"  ì´ ë¬¸ì¥ ìˆ˜: {len(subtitles_snapshot):,}ê°œ\n")
                    f.write(f"  ì´ ê¸€ì ìˆ˜: {total_chars:,}ì\n")
                    f.write(f"  ì´ ë‹¨ì–´ ìˆ˜: {total_words:,}ê°œ\n")
                    f.write(
                        f"  í‰ê·  ë¬¸ì¥ ê¸¸ì´: {total_chars / len(subtitles_snapshot):.1f}ì\n"
                    )
                    f.write(
                        f"  í‰ê·  ë‹¨ì–´ ìˆ˜: {total_words / len(subtitles_snapshot):.1f}ê°œ\n\n"
                    )

                    f.write("ğŸ“ ë¬¸ì¥ ë¶„ì„\n")
                    f.write("-" * 30 + "\n")
                    f.write(f"  ê°€ì¥ ê¸´ ë¬¸ì¥: {len(longest.text)}ì\n")
                    f.write(
                        f'    "{longest.text[:50]}{"..." if len(longest.text) > 50 else ""}"\n'
                    )
                    f.write(f"  ê°€ì¥ ì§§ì€ ë¬¸ì¥: {len(shortest.text)}ì\n")
                    f.write(f'    "{shortest.text}"\n\n')

                    if hour_counts:
                        f.write("â° ì‹œê°„ëŒ€ë³„ ë¶„í¬\n")
                        f.write("-" * 30 + "\n")
                        for h in sorted(hour_counts.keys()):
                            bar = "â–ˆ" * min(hour_counts[h] // 2, 20)
                            f.write(f"  {h:02d}ì‹œ: {bar} {hour_counts[h]}ê°œ\n")
                        f.write("\n")

                    # í•˜ì´ë¼ì´íŠ¸ í‚¤ì›Œë“œ í†µê³„
                    if self.keywords:
                        f.write("ğŸ” í‚¤ì›Œë“œ ë¹ˆë„\n")
                        f.write("-" * 30 + "\n")
                        all_text = " ".join(s.text for s in subtitles_snapshot).lower()
                        for kw in self.keywords:
                            count = all_text.count(kw.lower())
                            if count > 0:
                                f.write(f"  {kw}: {count}íšŒ\n")
                        f.write("\n")

                    f.write("=" * 50 + "\n")

                self._show_toast("í†µê³„ ë‚´ë³´ë‚´ê¸° ì™„ë£Œ!", "success")
            except Exception as e:
                QMessageBox.critical(self, "ì˜¤ë¥˜", f"í†µê³„ ì €ì¥ ì‹¤íŒ¨: {e}")

    def _generate_smart_filename(self, extension: str) -> str:
        """URLê³¼ í˜„ì¬ ì‹œê°„ ê¸°ë°˜ ìŠ¤ë§ˆíŠ¸ íŒŒì¼ëª… ìƒì„± (#28)"""
        # ìœ„ì›íšŒëª… ì¶”ì¶œ (í˜„ì¬ URLì—ì„œ ìë™ ê°ì§€)
        current_url = self._get_current_url()
        committee_name = self._autodetect_tag(current_url)
        return utils.generate_filename(committee_name, extension)

    def _save_txt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ì €ì¥í•  ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        filename = self._generate_smart_filename("txt")
        path, _ = QFileDialog.getSaveFileName(
            self, "TXT ì €ì¥", filename, "í…ìŠ¤íŠ¸ (*.txt)"
        )

        if path:
            # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ìŠ¤ëƒ…ìƒ· ìƒì„± (UI ìŠ¤ë ˆë“œì—ì„œ)
            with self.subtitle_lock:
                subtitles_snapshot = [
                    (entry.timestamp, entry.text) for entry in self.subtitles
                ]

            # ì‹¤ì œ ì €ì¥ í•¨ìˆ˜ ì •ì˜
            def do_save(filepath):
                with open(filepath, "w", encoding="utf-8-sig") as f:
                    last_printed_ts = None
                    for i, (timestamp, text) in enumerate(subtitles_snapshot):
                        # ë©”ì¸ ìœˆë„ìš°ì™€ ë™ì¼: 1ë¶„ ê°„ê²©ìœ¼ë¡œ íƒ€ì„ìŠ¤íƒ¬í”„ í‘œì‹œ
                        should_print_ts = False
                        if last_printed_ts is None:
                            should_print_ts = True
                        elif (timestamp - last_printed_ts).total_seconds() >= 60:
                            should_print_ts = True

                        if should_print_ts:
                            f.write(f"[{timestamp.strftime('%H:%M:%S')}] {text}\n")
                            last_printed_ts = timestamp
                        else:
                            f.write(f"{text}\n")

            # ë°±ê·¸ë¼ìš´ë“œì—ì„œ ì €ì¥ ì‹¤í–‰
            self._save_in_background(do_save, path, "TXT ì €ì¥ ì™„ë£Œ!", "TXT ì €ì¥ ì‹¤íŒ¨")

    def _save_srt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ì €ì¥í•  ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        filename = self._generate_smart_filename("srt")
        path, _ = QFileDialog.getSaveFileName(
            self, "SRT ì €ì¥", filename, "SubRip (*.srt)"
        )

        if path:
            # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ìŠ¤ëƒ…ìƒ· ìƒì„± (UI ìŠ¤ë ˆë“œì—ì„œ)
            with self.subtitle_lock:
                subtitles_snapshot = [
                    (entry.start_time, entry.end_time, entry.timestamp, entry.text)
                    for entry in self.subtitles
                ]

            def do_save(filepath):
                with open(filepath, "w", encoding="utf-8") as f:
                    for i, (start_time, end_time, timestamp, text) in enumerate(
                        subtitles_snapshot, 1
                    ):
                        if start_time and end_time:
                            # [Fix] ë°€ë¦¬ì´ˆ ì •ë°€ë„ ê°œì„ 
                            start = f"{start_time.strftime('%H:%M:%S')},{start_time.microsecond // 1000:03d}"
                            end = f"{end_time.strftime('%H:%M:%S')},{end_time.microsecond // 1000:03d}"
                        else:
                            start = f"{timestamp.strftime('%H:%M:%S')},{timestamp.microsecond // 1000:03d}"
                            fallback_end = timestamp + timedelta(seconds=3)
                            end = f"{fallback_end.strftime('%H:%M:%S')},{fallback_end.microsecond // 1000:03d}"
                        f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

            self._save_in_background(do_save, path, "SRT ì €ì¥ ì™„ë£Œ!", "SRT ì €ì¥ ì‹¤íŒ¨")

    def _save_vtt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ì €ì¥í•  ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        filename = self._generate_smart_filename("vtt")
        path, _ = QFileDialog.getSaveFileName(
            self, "VTT ì €ì¥", filename, "WebVTT (*.vtt)"
        )

        if path:
            with self.subtitle_lock:
                subtitles_snapshot = [
                    (entry.start_time, entry.end_time, entry.timestamp, entry.text)
                    for entry in self.subtitles
                ]

            def do_save(filepath):
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write("WEBVTT\n\n")
                    for i, (start_time, end_time, timestamp, text) in enumerate(
                        subtitles_snapshot, 1
                    ):
                        if start_time and end_time:
                            # [Fix] ë°€ë¦¬ì´ˆ ì •ë°€ë„ ê°œì„ 
                            start = f"{start_time.strftime('%H:%M:%S')}.{start_time.microsecond // 1000:03d}"
                            end = f"{end_time.strftime('%H:%M:%S')}.{end_time.microsecond // 1000:03d}"
                        else:
                            start = f"{timestamp.strftime('%H:%M:%S')}.{timestamp.microsecond // 1000:03d}"
                            fallback_end = timestamp + timedelta(seconds=3)
                            end = f"{fallback_end.strftime('%H:%M:%S')}.{fallback_end.microsecond // 1000:03d}"
                        f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

            self._save_in_background(do_save, path, "VTT ì €ì¥ ì™„ë£Œ!", "VTT ì €ì¥ ì‹¤íŒ¨")

    def _save_docx(self):
        """DOCX (Word) íŒŒì¼ë¡œ ì €ì¥"""
        if not self.subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ì €ì¥í•  ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.warning(
                self,
                "ë¼ì´ë¸ŒëŸ¬ë¦¬ í•„ìš”",
                "DOCX ì €ì¥ì„ ìœ„í•´ python-docx ë¼ì´ë¸ŒëŸ¬ë¦¬ê°€ í•„ìš”í•©ë‹ˆë‹¤.\n\n"
                "ì„¤ì¹˜: pip install python-docx",
            )
            return

        filename = self._generate_smart_filename("docx")
        path, _ = QFileDialog.getSaveFileName(
            self, "DOCX ì €ì¥", filename, "Word ë¬¸ì„œ (*.docx)"
        )

        if path:
            # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ìŠ¤ëƒ…ìƒ· ìƒì„±
            with self.subtitle_lock:
                subtitles_snapshot = [
                    (entry.timestamp, entry.text) for entry in self.subtitles
                ]

            generated_at = datetime.now().strftime("%Yë…„ %mì›” %dì¼ %H:%M:%S")
            total_chars = sum(len(text) for _, text in subtitles_snapshot)

            def do_save(filepath):
                doc = Document()

                # ì œëª©
                title = doc.add_heading("êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ìë§‰", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # ìƒì„± ì¼ì‹œ
                doc.add_paragraph(f"ìƒì„± ì¼ì‹œ: {generated_at}")
                doc.add_paragraph()

                # ìë§‰ ë‚´ìš© - 1ë¶„ ê°„ê²©ìœ¼ë¡œ íƒ€ì„ìŠ¤íƒ¬í”„ í‘œì‹œ
                last_printed_ts = None
                for timestamp, text in subtitles_snapshot:
                    should_print_ts = False
                    if last_printed_ts is None:
                        should_print_ts = True
                    elif (timestamp - last_printed_ts).total_seconds() >= 60:
                        should_print_ts = True

                    p = doc.add_paragraph()
                    if should_print_ts:
                        ts = timestamp.strftime("%H:%M:%S")
                        run = p.add_run(f"[{ts}] ")
                        run.font.size = Pt(9)
                        run.font.color.rgb = None  # ê¸°ë³¸ ìƒ‰ìƒ
                        last_printed_ts = timestamp
                    p.add_run(text)

                # í†µê³„
                doc.add_paragraph()
                doc.add_paragraph(
                    f"ì´ {len(subtitles_snapshot)}ë¬¸ì¥, {total_chars:,}ì"
                )

                doc.save(filepath)

            self._save_in_background(do_save, path, "DOCX ì €ì¥ ì™„ë£Œ!", "DOCX ì €ì¥ ì‹¤íŒ¨")

    def _save_hwp(self):
        """HWP íŒŒì¼ë¡œ ì €ì¥ (Hancom Office í•„ìš”)"""
        if not self.subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ì €ì¥í•  ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        try:
            import win32com.client
        except ImportError:
            QMessageBox.warning(
                self,
                "ë¼ì´ë¸ŒëŸ¬ë¦¬ í•„ìš”",
                "HWP ì €ì¥ì„ ìœ„í•´ pywin32 ë¼ì´ë¸ŒëŸ¬ë¦¬ê°€ í•„ìš”í•©ë‹ˆë‹¤.\n\n"
                "ì„¤ì¹˜: pip install pywin32\n\n"
                "RTF í˜•ì‹ìœ¼ë¡œ ì €ì¥ì„ ì‹œë„í•©ë‹ˆë‹¤.",
            )
            self._save_rtf()
            return

        # ì €ì¥ ëŒ€í™”ìƒì
        filename = f"êµ­íšŒìë§‰_{datetime.now().strftime('%Y%m%d_%H%M%S')}.hwp"
        path, _ = QFileDialog.getSaveFileName(
            self, "HWP ì €ì¥", filename, "HWP ë¬¸ì„œ (*.hwp)"
        )

        if not path:
            return

        # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ìŠ¤ëƒ…ìƒ· ìƒì„±
        with self.subtitle_lock:
            subtitles_snapshot = [
                (entry.timestamp, entry.text) for entry in self.subtitles
            ]

        generated_at = datetime.now().strftime("%Yë…„ %mì›” %dì¼ %H:%M:%S")
        total_chars = sum(len(text) for _, text in subtitles_snapshot)

        def do_save(filepath):
            hwp = None
            pythoncom = None
            try:
                try:
                    import pythoncom

                    pythoncom.CoInitialize()
                except Exception:
                    pythoncom = None

                # win32com.client.dynamic.Dispatch ì‚¬ìš©ìœ¼ë¡œ ìºì‹œ ë¬¸ì œ íšŒí”¼
                hwp = win32com.client.dynamic.Dispatch("HWPFrame.HwpObject")
                hwp.XHwpWindows.Item(0).Visible = True
                hwp.RegisterModule(
                    "FilePathCheckDLL", "SecurityModule"
                )  # ë³´ì•ˆ ëª¨ë“ˆ ë“±ë¡ ì‹œë„

                # ìƒˆ ë¬¸ì„œ ìƒì„±
                hwp.HAction.Run("FileNew")

                # ì œëª© ì…ë ¥
                hwp.HAction.Run("CharShapeBold")
                hwp.HAction.Run("ParagraphShapeAlignCenter")
                hwp.HAction.GetDefault("InsertText", hwp.HParameterSet.HInsertText.HSet)
                hwp.HParameterSet.HInsertText.Text = "êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ìë§‰\r\n"
                hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)

                hwp.HAction.Run("CharShapeBold")
                hwp.HAction.Run("ParagraphShapeAlignLeft")

                # ìƒì„± ì¼ì‹œ
                hwp.HParameterSet.HInsertText.Text = (
                    f"ìƒì„± ì¼ì‹œ: {generated_at}\r\n\r\n"
                )
                hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)

                # ìë§‰ ë‚´ìš© - 1ë¶„ ê°„ê²©ìœ¼ë¡œ íƒ€ì„ìŠ¤íƒ¬í”„ í‘œì‹œ
                last_printed_ts = None
                for timestamp, text in subtitles_snapshot:
                    should_print_ts = False
                    if last_printed_ts is None:
                        should_print_ts = True
                    elif (timestamp - last_printed_ts).total_seconds() >= 60:
                        should_print_ts = True

                    if should_print_ts:
                        ts = timestamp.strftime("%H:%M:%S")
                        hwp.HParameterSet.HInsertText.Text = f"[{ts}] {text}\r\n"
                        last_printed_ts = timestamp
                    else:
                        hwp.HParameterSet.HInsertText.Text = f"{text}\r\n"
                    hwp.HAction.Execute(
                        "InsertText", hwp.HParameterSet.HInsertText.HSet
                    )

                # í†µê³„
                hwp.HParameterSet.HInsertText.Text = (
                    f"\r\nì´ {len(subtitles_snapshot)}ë¬¸ì¥, {total_chars:,}ì\r\n"
                )
                hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)

                # FileSaveAs_S ì•¡ì…˜ ì‚¬ìš©
                hwp.HAction.GetDefault(
                    "FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet
                )
                hwp.HParameterSet.HFileOpenSave.filename = filepath
                hwp.HParameterSet.HFileOpenSave.Format = "HWP"
                hwp.HAction.Execute(
                    "FileSaveAs_S", hwp.HParameterSet.HFileOpenSave.HSet
                )
            finally:
                if hwp:
                    try:
                        hwp.Quit()
                    except Exception:
                        pass
                if pythoncom:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass

        def do_save_with_error(filepath):
            last_error = None
            for attempt in range(2):
                try:
                    do_save(filepath)
                    saved_path = Path(filepath)
                    if not saved_path.exists() or saved_path.stat().st_size <= 0:
                        raise RuntimeError("ì €ì¥ëœ íŒŒì¼ì´ í™•ì¸ë˜ì§€ ì•ŠìŠµë‹ˆë‹¤.")
                    return
                except Exception as e:
                    last_error = e
                    logger.warning(f"HWP ì €ì¥ ì¬ì‹œë„ ì‹¤íŒ¨ ({attempt + 1}/2): {e}")
                    time.sleep(1)
            self.message_queue.put(("hwp_save_failed", {"error": str(last_error)}))
            raise last_error

        self._save_in_background(
            do_save_with_error, path, "HWP ì €ì¥ ì™„ë£Œ!", "HWP ì €ì¥ ì‹¤íŒ¨"
        )

    def _handle_hwp_save_failure(self, error) -> None:
        """HWP ì €ì¥ ì‹¤íŒ¨ ì‹œ ëŒ€ì²´ ì €ì¥ ì•ˆë‚´"""
        error_msg = str(error).lower()
        logger.error(f"HWP ì €ì¥ ì‹¤íŒ¨: {error}")

        # ê¶Œí•œ ë¬¸ì œ íŒíŠ¸ ì œê³µ
        if "access denied" in error_msg or "ê¶Œí•œ" in str(error):
            advice = (
                "\n\nê´€ë¦¬ì ê¶Œí•œìœ¼ë¡œ ì‹¤í–‰í•˜ê±°ë‚˜ í•œê¸€ í”„ë¡œê·¸ë¨ì„ ë¨¼ì € ì‹¤í–‰í•´ ë³´ì„¸ìš”."
            )
        elif "server execution failed" in error_msg:
            advice = "\n\ní•œê¸€ í”„ë¡œê·¸ë¨ì´ ì‘ë‹µí•˜ì§€ ì•ŠìŠµë‹ˆë‹¤. í•œê¸€ì„ ì¢…ë£Œí•˜ê³  ë‹¤ì‹œ ì‹œë„í•˜ì„¸ìš”."
        else:
            advice = ""

        # ì‚¬ìš©ìì—ê²Œ ëŒ€ì²´ ì €ì¥ ë°©ì‹ ì œì•ˆ
        reply = QMessageBox.question(
            self,
            "HWP ì €ì¥ ì‹¤íŒ¨",
            f"í•œê¸€ íŒŒì¼ ì €ì¥ ì¤‘ ì˜¤ë¥˜ê°€ ë°œìƒí–ˆìŠµë‹ˆë‹¤: {error}{advice}\n\n"
            "ëŒ€ì²´ í˜•ì‹ìœ¼ë¡œ ì €ì¥í•˜ì‹œê² ìŠµë‹ˆê¹Œ?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )

        if reply == QMessageBox.StandardButton.Yes:
            items = ["RTF (í•œê¸€ í˜¸í™˜)", "DOCX (Word)", "TXT (í…ìŠ¤íŠ¸)"]
            item, ok = QInputDialog.getItem(
                self, "í˜•ì‹ ì„ íƒ", "ì €ì¥ í˜•ì‹:", items, 0, False
            )
            if ok and item:
                if "RTF" in item:
                    self._save_rtf()
                elif "DOCX" in item:
                    self._save_docx()
                else:
                    self._save_txt()

    def _rtf_encode(self, text: str) -> str:
        """ìœ ë‹ˆì½”ë“œ ë¬¸ìë¥¼ RTF í˜•ì‹ìœ¼ë¡œ ì¸ì½”ë”© (íŠ¹ìˆ˜ë¬¸ì ì´ìŠ¤ì¼€ì´í”„)"""
        result = []
        for char in text:
            if char == "\\":
                result.append("\\\\")
            elif char == "{":
                result.append("\\{")
            elif char == "}":
                result.append("\\}")
            else:
                result.append(char)
        return "".join(result)

    def _save_rtf(self):
        """RTF íŒŒì¼ë¡œ ì €ì¥ (HWPì—ì„œ ì—´ê¸° ê°€ëŠ¥)"""
        if not self.subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ì €ì¥í•  ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        filename = self._generate_smart_filename("rtf")
        path, _ = QFileDialog.getSaveFileName(
            self, "RTF ì €ì¥", filename, "RTF ë¬¸ì„œ (*.rtf)"
        )

        if path:
            try:
                # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ìŠ¤ëƒ…ìƒ· ìƒì„±
                with self.subtitle_lock:
                    subtitles_snapshot = list(self.subtitles)

                with open(path, "wb") as f:  # ë°”ì´ë„ˆë¦¬ ëª¨ë“œë¡œ ë³€ê²½
                    # RTF í—¤ë” (ìœ ë‹ˆì½”ë“œ ì§€ì›)
                    f.write(b"{\\rtf1\\ansi\\ansicpg949\\deff0")
                    f.write(
                        b"{\\fonttbl{\\f0\\fnil\\fcharset129 \\'b8\\'c0\\'c0\\'ba \\'b0\\'ed\\'b5\\'f1;}}"
                    )
                    f.write(
                        b"{\\colortbl;\\red0\\green0\\blue0;\\red128\\green128\\blue128;}"
                    )
                    f.write(b"\n")

                    # ì œëª©
                    title = self._rtf_encode("êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ìë§‰")
                    f.write(b"\\pard\\qc\\b\\fs28 ")
                    f.write(title.encode("cp949", errors="replace"))
                    f.write(b"\\b0\\par\n")

                    # ìƒì„± ì¼ì‹œ
                    date_str = self._rtf_encode(
                        f"ìƒì„± ì¼ì‹œ: {datetime.now().strftime('%Yë…„ %mì›” %dì¼ %H:%M:%S')}"
                    )
                    f.write(b"\\pard\\ql\\fs20 ")
                    f.write(date_str.encode("cp949", errors="replace"))
                    f.write(b"\\par\\par\n")

                    # ìë§‰ ë‚´ìš©
                    for entry in subtitles_snapshot:
                        timestamp = entry.timestamp.strftime("%H:%M:%S")
                        text = self._rtf_encode(entry.text)
                        f.write(b"\\cf2[")
                        f.write(timestamp.encode("cp949", errors="replace"))
                        f.write(b"]\\cf1 ")
                        f.write(text.encode("cp949", errors="replace"))
                        f.write(b"\\par\n")

                    # í†µê³„
                    total_chars = sum(len(s.text) for s in subtitles_snapshot)
                    stats = self._rtf_encode(
                        f"ì´ {len(subtitles_snapshot)}ë¬¸ì¥, {total_chars:,}ì"
                    )
                    f.write(b"\\par\\fs18 ")
                    f.write(stats.encode("cp949", errors="replace"))
                    f.write(b"\\par}")

                QMessageBox.information(
                    self,
                    "ì„±ê³µ",
                    f"RTF ì €ì¥ ì™„ë£Œ!\n\níŒŒì¼: {path}\n\nì´ íŒŒì¼ì€ í•œê¸€(HWP)ì—ì„œ ì—´ ìˆ˜ ìˆìŠµë‹ˆë‹¤.",
                )

            except Exception as e:
                QMessageBox.critical(self, "ì˜¤ë¥˜", f"RTF ì €ì¥ ì‹¤íŒ¨: {e}")

    # ========== ì„¸ì…˜ ==========

    def _save_session(self):
        with self.subtitle_lock:
            has_subtitles = bool(self.subtitles)
        if not has_subtitles:
            QMessageBox.warning(self, "ì•Œë¦¼", "ì €ì¥í•  ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤.")
            return

        if self._session_save_in_progress:
            self._show_toast("ì´ë¯¸ ì„¸ì…˜ ì €ì¥ì´ ì§„í–‰ ì¤‘ì…ë‹ˆë‹¤.", "info")
            return

        filename = (
            f"{Config.SESSION_DIR}/ì„¸ì…˜_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        )
        path, _ = QFileDialog.getSaveFileName(
            self, "ì„¸ì…˜ ì €ì¥", filename, "JSON (*.json)"
        )

        if not path:
            return

        try:
            Path(path).parent.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            QMessageBox.critical(self, "ì˜¤ë¥˜", f"ì„¸ì…˜ ì €ì¥ ê²½ë¡œ ì¤€ë¹„ ì‹¤íŒ¨: {e}")
            return

        self._session_save_in_progress = True
        self._set_status("ì„¸ì…˜ ì €ì¥ ì¤‘...", "running")

        current_url = self._get_current_url()
        committee_name = self._autodetect_tag(current_url) or ""
        duration = int(time.time() - self.start_time) if self.start_time else 0

        def background_save():
            try:
                with self.subtitle_lock:
                    snapshot = [
                        (s.text, s.timestamp, s.start_time, s.end_time)
                        for s in self.subtitles
                    ]

                subtitles_copy = [
                    {
                        "text": text,
                        "timestamp": timestamp.isoformat() if timestamp else None,
                        "start_time": start_time.isoformat() if start_time else None,
                        "end_time": end_time.isoformat() if end_time else None,
                    }
                    for text, timestamp, start_time, end_time in snapshot
                ]

                data = {
                    "version": Config.VERSION,
                    "created": datetime.now().isoformat(),
                    "url": current_url,
                    "committee_name": committee_name,
                    "subtitles": subtitles_copy,
                }

                utils.atomic_write_json(
                    path,
                    data,
                    ensure_ascii=False,
                    indent=2,
                )

                db_saved = False
                db_error = ""
                if self.db:
                    try:
                        db_data = {
                            "url": current_url,
                            "committee_name": committee_name,
                            "subtitles": subtitles_copy,
                            "version": Config.VERSION,
                            "duration_seconds": duration,
                        }
                        self.db.save_session(db_data)
                        db_saved = True
                    except Exception as db_exc:
                        db_error = str(db_exc)

                self.message_queue.put(
                    (
                        "session_save_done",
                        {
                            "path": path,
                            "saved_count": len(subtitles_copy),
                            "db_saved": db_saved,
                            "db_error": db_error,
                        },
                    )
                )
            except Exception as e:
                logger.error(f"ì„¸ì…˜ ì €ì¥ ì˜¤ë¥˜: {e}")
                self.message_queue.put(
                    ("session_save_failed", {"path": path, "error": str(e)})
                )

        threading.Thread(
            target=background_save, daemon=True, name="SessionSaveWorker"
        ).start()
        self._show_toast(f"ğŸ’¾ ì„¸ì…˜ ì €ì¥ ì‹œì‘: {Path(path).name}", "info", 1500)

    def _load_session(self):
        if self._session_load_in_progress:
            self._show_toast("ì´ë¯¸ ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸°ê°€ ì§„í–‰ ì¤‘ì…ë‹ˆë‹¤.", "info")
            return

        path, _ = QFileDialog.getOpenFileName(
            self, "ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸°", f"{Config.SESSION_DIR}/", "JSON (*.json)"
        )

        if not path:
            return

        self._session_load_in_progress = True
        self._set_status("ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì¤‘...", "running")

        def background_load():
            try:
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                except json.JSONDecodeError as json_err:
                    self.message_queue.put(
                        (
                            "session_load_json_error",
                            {"path": path, "error": str(json_err)},
                        )
                    )
                    return

                session_version = data.get("version", "unknown")
                new_subtitles = []
                skipped = 0
                for item in data.get("subtitles", []):
                    try:
                        new_subtitles.append(SubtitleEntry.from_dict(item))
                    except ValueError as e:
                        logger.warning(f"ì†ìƒëœ ìë§‰ í•­ëª© ê±´ë„ˆëœ€: {e}")
                        skipped += 1

                self.message_queue.put(
                    (
                        "session_load_done",
                        {
                            "path": path,
                            "version": session_version,
                            "url": data.get("url", ""),
                            "subtitles": new_subtitles,
                            "skipped": skipped,
                        },
                    )
                )
            except Exception as e:
                logger.error(f"ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì˜¤ë¥˜: {e}")
                self.message_queue.put(
                    ("session_load_failed", {"path": path, "error": str(e)})
                )

        threading.Thread(
            target=background_load, daemon=True, name="SessionLoadWorker"
        ).start()
        self._show_toast(f"ğŸ“‚ ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì‹œì‘: {Path(path).name}", "info", 1500)

    # ========== ìœ í‹¸ë¦¬í‹° ==========
    def _clear_text(self):
        if not self.subtitles:
            return

        reply = QMessageBox.question(
            self,
            "í™•ì¸",
            "ëª¨ë“  ë‚´ìš©ì„ ì§€ìš°ì‹œê² ìŠµë‹ˆê¹Œ?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )

        if reply == QMessageBox.StandardButton.Yes:
            self._replace_subtitles_and_refresh(
                [], keep_history_from_subtitles=False
            )
            self.status_label.setText("ë‚´ìš© ì‚­ì œë¨")

    def _edit_subtitle(self):
        """ì„ íƒí•œ ìë§‰ í¸ì§‘"""
        if not self.subtitles:
            self._show_toast("í¸ì§‘í•  ìë§‰ì´ ì—†ìŠµë‹ˆë‹¤.", "warning")
            return
        if self.is_running:
            self._show_toast(
                "ì¶”ì¶œ ì¤‘ì—ëŠ” í¸ì§‘ì´ ë¶ˆì•ˆì •í•  ìˆ˜ ìˆìŠµë‹ˆë‹¤. ë¨¼ì € ì¤‘ì§€í•˜ì„¸ìš”.", "warning"
            )
            return

        # ìë§‰ ëª©ë¡ ë‹¤ì´ì–¼ë¡œê·¸
        dialog = QDialog(self)
        dialog.setWindowTitle("ìë§‰ í¸ì§‘")
        dialog.setMinimumSize(600, 400)

        layout = QVBoxLayout(dialog)

        # ì•ˆë‚´ ë¼ë²¨
        info_label = QLabel("í¸ì§‘í•  ìë§‰ì„ ì„ íƒí•˜ì„¸ìš”:")
        layout.addWidget(info_label)

        # ìë§‰ ëª©ë¡
        list_widget = QListWidget()
        for i, entry in enumerate(self.subtitles):
            timestamp = entry.timestamp.strftime("%H:%M:%S")
            text_preview = (
                entry.text[:50] + "..." if len(entry.text) > 50 else entry.text
            )
            list_widget.addItem(f"[{timestamp}] {text_preview}")
        layout.addWidget(list_widget)

        # í¸ì§‘ ì˜ì—­
        edit_label = QLabel("ìë§‰ ë‚´ìš©:")
        layout.addWidget(edit_label)

        edit_text = QTextEdit()
        edit_text.setMaximumHeight(100)
        layout.addWidget(edit_text)

        # ì„ íƒ ì‹œ ë‚´ìš© ë¡œë“œ
        def on_selection_changed():
            idx = list_widget.currentRow()
            if 0 <= idx < len(self.subtitles):
                edit_text.setText(self.subtitles[idx].text)

        list_widget.currentRowChanged.connect(on_selection_changed)

        # ë²„íŠ¼
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Save
            | QDialogButtonBox.StandardButton.Cancel
        )

        def save_edit():
            idx = list_widget.currentRow()
            if 0 <= idx < len(self.subtitles):
                new_text = edit_text.toPlainText().strip()
                if new_text:
                    with self.subtitle_lock:
                        entry = self.subtitles[idx]
                        old_chars = entry.char_count
                        old_words = entry.word_count
                        entry.update_text(new_text)
                        self._cached_total_chars += entry.char_count - old_chars
                        self._cached_total_words += entry.word_count - old_words
                    self._refresh_text(force_full=True)
                    self._update_count_label()
                    self._show_toast("ìë§‰ì´ ìˆ˜ì •ë˜ì—ˆìŠµë‹ˆë‹¤.", "success")
                    dialog.accept()
                else:
                    QMessageBox.warning(dialog, "ì•Œë¦¼", "ìë§‰ ë‚´ìš©ì„ ì…ë ¥í•´ì£¼ì„¸ìš”.")
            else:
                QMessageBox.warning(dialog, "ì•Œë¦¼", "í¸ì§‘í•  ìë§‰ì„ ì„ íƒí•´ì£¼ì„¸ìš”.")

        buttons.accepted.connect(save_edit)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        # ì²« ë²ˆì§¸ í•­ëª© ì„ íƒ
        if list_widget.count() > 0:
            list_widget.setCurrentRow(0)

        dialog.exec()

    def _delete_subtitle(self):
        """ì„ íƒí•œ ìë§‰ ì‚­ì œ"""
        if not self.subtitles:
            self._show_toast("ì‚­ì œí•  ìë§‰ì´ ì—†ìŠµë‹ˆë‹¤.", "warning")
            return
        if self.is_running:
            self._show_toast(
                "ì¶”ì¶œ ì¤‘ì—ëŠ” ì‚­ì œê°€ ë¶ˆì•ˆì •í•  ìˆ˜ ìˆìŠµë‹ˆë‹¤. ë¨¼ì € ì¤‘ì§€í•˜ì„¸ìš”.", "warning"
            )
            return

        # ìë§‰ ëª©ë¡ ë‹¤ì´ì–¼ë¡œê·¸
        dialog = QDialog(self)
        dialog.setWindowTitle("ìë§‰ ì‚­ì œ")
        dialog.setMinimumSize(600, 400)

        layout = QVBoxLayout(dialog)

        # ì•ˆë‚´ ë¼ë²¨
        info_label = QLabel("ì‚­ì œí•  ìë§‰ì„ ì„ íƒí•˜ì„¸ìš” (ë‹¤ì¤‘ ì„ íƒ ê°€ëŠ¥):")
        layout.addWidget(info_label)

        # ìë§‰ ëª©ë¡ (ë‹¤ì¤‘ ì„ íƒ ê°€ëŠ¥)
        list_widget = QListWidget()
        list_widget.setSelectionMode(QAbstractItemView.SelectionMode.MultiSelection)

        for i, entry in enumerate(self.subtitles):
            timestamp = entry.timestamp.strftime("%H:%M:%S")
            text_preview = (
                entry.text[:50] + "..." if len(entry.text) > 50 else entry.text
            )
            list_widget.addItem(f"[{timestamp}] {text_preview}")
        layout.addWidget(list_widget)

        # ë²„íŠ¼
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.button(QDialogButtonBox.StandardButton.Ok).setText("ì‚­ì œ")

        def delete_selected():
            selected_rows = sorted(
                [i.row() for i in list_widget.selectedIndexes()], reverse=True
            )
            if not selected_rows:
                QMessageBox.warning(dialog, "ì•Œë¦¼", "ì‚­ì œí•  ìë§‰ì„ ì„ íƒí•´ì£¼ì„¸ìš”.")
                return

            reply = QMessageBox.question(
                dialog,
                "í™•ì¸",
                f"ì„ íƒí•œ {len(selected_rows)}ê°œì˜ ìë§‰ì„ ì‚­ì œí•˜ì‹œê² ìŠµë‹ˆê¹Œ?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )

            if reply == QMessageBox.StandardButton.Yes:
                with self.subtitle_lock:
                    for row in selected_rows:
                        entry = self.subtitles[row]
                        self._cached_total_chars -= entry.char_count
                        self._cached_total_words -= entry.word_count
                        del self.subtitles[row]
                self._refresh_text(force_full=True)
                self._update_count_label()
                self._show_toast(
                    f"{len(selected_rows)}ê°œ ìë§‰ì´ ì‚­ì œë˜ì—ˆìŠµë‹ˆë‹¤.", "success"
                )
                dialog.accept()

        buttons.accepted.connect(delete_selected)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        dialog.exec()

    # ========== ë„ì›€ë§ ==========

    def _show_guide(self):
        """ì‚¬ìš©ë²• ê°€ì´ë“œ í‘œì‹œ"""
        guide = """
<h2>ğŸ›ï¸ ì‚¬ìš©ë²• ê°€ì´ë“œ</h2>

<h3>ğŸ“‹ ê¸°ë³¸ ì‚¬ìš©ë²•</h3>
<ol>
<li><b>URL ì…ë ¥</b> - êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ í˜ì´ì§€ URLì„ ì…ë ¥í•©ë‹ˆë‹¤</li>
<li><b>ì„ íƒì í™•ì¸</b> - ê¸°ë³¸ê°’ì„ ì‚¬ìš©í•˜ê±°ë‚˜ ìˆ˜ì •í•©ë‹ˆë‹¤</li>
<li><b>ì˜µì…˜ ì„¤ì •</b>
    <ul>
    <li>ìë™ ìŠ¤í¬ë¡¤: ìƒˆ ìë§‰ ìë™ ë”°ë¼ê°€ê¸°</li>
    <li>ì‹¤ì‹œê°„ ì €ì¥: ìë§‰ ì‹¤ì‹œê°„ íŒŒì¼ ì €ì¥</li>
    <li>í—¤ë“œë¦¬ìŠ¤ ëª¨ë“œ: ë¸Œë¼ìš°ì € ì°½ ìˆ¨ê¸°ê³  ì‹¤í–‰</li>
    </ul>
</li>
<li><b>ì‹œì‘</b> ë²„íŠ¼ í´ë¦­ (ë˜ëŠ” F5)</li>
<li>ìë§‰ ì¶”ì¶œ ì™„ë£Œ í›„ <b>íŒŒì¼ ì €ì¥</b></li>
</ol>

<h3>âŒ¨ï¸ ì£¼ìš” ë‹¨ì¶•í‚¤</h3>
<table>
<tr><td><b>F5</b></td><td>ì‹œì‘</td></tr>
<tr><td><b>Escape</b></td><td>ì¤‘ì§€</td></tr>
<tr><td><b>Ctrl+F</b></td><td>ê²€ìƒ‰</td></tr>
<tr><td><b>F3</b></td><td>ë‹¤ìŒ ê²€ìƒ‰</td></tr>
<tr><td><b>Ctrl+T</b></td><td>í…Œë§ˆ ì „í™˜</td></tr>
<tr><td><b>Ctrl+S</b></td><td>TXT ì €ì¥</td></tr>
</table>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("ì‚¬ìš©ë²• ê°€ì´ë“œ")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(guide)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()

    def _show_shortcuts(self):
        """í‚¤ë³´ë“œ ë‹¨ì¶•í‚¤ ëª©ë¡ í‘œì‹œ"""
        shortcuts = """
<h2>âŒ¨ï¸ í‚¤ë³´ë“œ ë‹¨ì¶•í‚¤</h2>

<h3>ğŸ“‹ ê¸°ë³¸ ì¡°ì‘</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>ë‹¨ì¶•í‚¤</th><th>ê¸°ëŠ¥</th></tr>
<tr><td><b>F5</b></td><td>ì¶”ì¶œ ì‹œì‘</td></tr>
<tr><td><b>Escape</b></td><td>ì¶”ì¶œ ì¤‘ì§€</td></tr>
<tr><td><b>Ctrl+Q</b></td><td>í”„ë¡œê·¸ë¨ ì¢…ë£Œ</td></tr>
</table>

<h3>ğŸ” ê²€ìƒ‰ ë° í¸ì§‘</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>ë‹¨ì¶•í‚¤</th><th>ê¸°ëŠ¥</th></tr>
<tr><td><b>Ctrl+F</b></td><td>ê²€ìƒ‰ì°½ ì—´ê¸°</td></tr>
<tr><td><b>F3</b></td><td>ë‹¤ìŒ ê²€ìƒ‰ ê²°ê³¼</td></tr>
<tr><td><b>Shift+F3</b></td><td>ì´ì „ ê²€ìƒ‰ ê²°ê³¼</td></tr>
<tr><td><b>Ctrl+E</b></td><td>ìë§‰ í¸ì§‘</td></tr>
<tr><td><b>Delete</b></td><td>ìë§‰ ì‚­ì œ</td></tr>
<tr><td><b>Ctrl+C</b></td><td>í´ë¦½ë³´ë“œ ë³µì‚¬</td></tr>
</table>

<h3>ğŸ’¾ ì €ì¥</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>ë‹¨ì¶•í‚¤</th><th>ê¸°ëŠ¥</th></tr>
<tr><td><b>Ctrl+S</b></td><td>TXT ì €ì¥</td></tr>
<tr><td><b>Ctrl+Shift+S</b></td><td>ì„¸ì…˜ ì €ì¥</td></tr>
<tr><td><b>Ctrl+O</b></td><td>ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸°</td></tr>
</table>

<h3>ğŸ¨ ë³´ê¸°</h3>
<table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
<tr style="background-color: #f0f0f0;"><th>ë‹¨ì¶•í‚¤</th><th>ê¸°ëŠ¥</th></tr>
<tr><td><b>Ctrl+T</b></td><td>í…Œë§ˆ ì „í™˜</td></tr>
<tr><td><b>Ctrl++</b></td><td>ê¸€ì í¬ê¸° í‚¤ìš°ê¸°</td></tr>
<tr><td><b>Ctrl+-</b></td><td>ê¸€ì í¬ê¸° ì¤„ì´ê¸°</td></tr>
<tr><td><b>F1</b></td><td>ì‚¬ìš©ë²• ê°€ì´ë“œ</td></tr>
</table>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("í‚¤ë³´ë“œ ë‹¨ì¶•í‚¤")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(shortcuts)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()

    def _show_features(self):
        """ê¸°ëŠ¥ ì†Œê°œ í‘œì‹œ"""
        features = """
<h2>âœ¨ ê¸°ëŠ¥ ì†Œê°œ</h2>

<h3>ğŸ¯ ì‹¤ì‹œê°„ ìë§‰ ì¶”ì¶œ</h3>
<p>êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ì›¹ì‚¬ì´íŠ¸ì˜ AI ìë§‰ì„ ì‹¤ì‹œê°„ìœ¼ë¡œ ìº¡ì²˜í•©ë‹ˆë‹¤.<br>
2ì´ˆ ë™ì•ˆ ìë§‰ì´ ë³€ê²½ë˜ì§€ ì•Šìœ¼ë©´ ìë™ìœ¼ë¡œ í™•ì •ë©ë‹ˆë‹¤.</p>

<h3>ğŸ’¾ ë‹¤ì–‘í•œ ì €ì¥ í˜•ì‹</h3>
<ul>
<li><b>TXT</b> - ì¼ë°˜ í…ìŠ¤íŠ¸</li>
<li><b>SRT</b> - ìë§‰ íŒŒì¼ í˜•ì‹</li>
<li><b>VTT</b> - WebVTT ìë§‰ í˜•ì‹</li>
<li><b>DOCX</b> - Word ë¬¸ì„œ</li>
</ul>

<h3>ğŸ” ê²€ìƒ‰ ë° í•˜ì´ë¼ì´íŠ¸</h3>
<ul>
<li><b>ì‹¤ì‹œê°„ ê²€ìƒ‰</b> - Ctrl+Fë¡œ ìë§‰ ë‚´ í…ìŠ¤íŠ¸ ê²€ìƒ‰</li>
<li><b>í‚¤ì›Œë“œ í•˜ì´ë¼ì´íŠ¸</b> - íŠ¹ì • ë‹¨ì–´ ê°•ì¡°</li>
</ul>

<h3>âš™ï¸ í—¤ë“œë¦¬ìŠ¤ ëª¨ë“œ (ì¸í„°ë„·ì°½ ìˆ¨ê¹€)</h3>
<p>ë¸Œë¼ìš°ì € ì°½ì„ ìˆ¨ê¸°ê³  ë°±ê·¸ë¼ìš´ë“œì—ì„œ ì‹¤í–‰í•©ë‹ˆë‹¤.<br>
ìë§‰ ì¶”ì¶œ ì¤‘ ë‹¤ë¥¸ ì‘ì—…ì„ í•  ìˆ˜ ìˆìŠµë‹ˆë‹¤.</p>

<h3>ğŸ“Š í†µê³„ íŒ¨ë„</h3>
<p>ì‹¤í–‰ ì‹œê°„, ê¸€ì ìˆ˜, ë‹¨ì–´ ìˆ˜, ë¶„ë‹¹ ê¸€ì ìˆ˜ë¥¼ í‘œì‹œí•©ë‹ˆë‹¤.</p>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("ê¸°ëŠ¥ ì†Œê°œ")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(features)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()

    def _show_about(self):
        """í”„ë¡œê·¸ë¨ ì •ë³´ í‘œì‹œ"""
        about = f"""
<h2>ğŸ›ï¸ êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ìë§‰ ì¶”ì¶œê¸°</h2>
<p><b>ë²„ì „:</b> {Config.VERSION}</p>
<p><b>ì„¤ëª…:</b> êµ­íšŒ ì˜ì‚¬ì¤‘ê³„ ì›¹ì‚¬ì´íŠ¸ì—ì„œ ì‹¤ì‹œê°„ AI ìë§‰ì„<br>
ìë™ìœ¼ë¡œ ì¶”ì¶œí•˜ê³  ì €ì¥í•˜ëŠ” í”„ë¡œê·¸ë¨ì…ë‹ˆë‹¤.</p>

<h3>ğŸ“¦ í•„ìš” ë¼ì´ë¸ŒëŸ¬ë¦¬</h3>
<ul>
<li>PyQt6</li>
<li>selenium</li>
<li>python-docx (DOCX ì €ì¥ìš©)</li>
</ul>

<p><b>Â© 2024-2025</b></p>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("ì •ë³´")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(about)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()

    # ========== ë°ì´í„°ë² ì´ìŠ¤ ê¸°ëŠ¥ (#26) ==========

    def _set_db_history_dialog_busy(self, busy: bool, message: str = "") -> None:
        """DB íˆìŠ¤í† ë¦¬ ë‹¤ì´ì–¼ë¡œê·¸ì˜ ë²„íŠ¼/ëª©ë¡ ìƒíƒœë¥¼ í† ê¸€í•œë‹¤."""
        state = self._db_history_dialog_state or {}
        load_btn = state.get("load_btn")
        delete_btn = state.get("delete_btn")
        close_btn = state.get("close_btn")
        list_widget = state.get("list_widget")
        status_label = state.get("status_label")

        for btn in (load_btn, delete_btn, close_btn):
            if btn is not None:
                btn.setEnabled(not busy)
        if list_widget is not None:
            list_widget.setEnabled(not busy)

        if status_label is not None:
            if busy:
                status_label.setText(message or "ì²˜ë¦¬ ì¤‘ì…ë‹ˆë‹¤...")
                status_label.show()
            elif message:
                status_label.setText(message)
                status_label.show()
            else:
                status_label.hide()

    def _clear_db_history_dialog_state(self) -> None:
        """í™œì„± DB íˆìŠ¤í† ë¦¬ ë‹¤ì´ì–¼ë¡œê·¸ ìƒíƒœë¥¼ ì •ë¦¬í•œë‹¤."""
        self._db_history_dialog_state = None

    def _run_db_task(
        self,
        task_name: str,
        worker,
        context: dict | None = None,
        loading_text: str = "",
    ) -> bool:
        """DB ì‘ì—…ì„ ë°±ê·¸ë¼ìš´ë“œ ìŠ¤ë ˆë“œì—ì„œ ì‹¤í–‰í•˜ê³  ê²°ê³¼ë¥¼ íë¡œ ì „ë‹¬í•œë‹¤."""
        if not self.db:
            QMessageBox.warning(self, "ì•Œë¦¼", "ë°ì´í„°ë² ì´ìŠ¤ê°€ ì´ˆê¸°í™”ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤.")
            return False

        if task_name in self._db_tasks_inflight:
            self._show_toast("ì´ë¯¸ ê°™ì€ DB ì‘ì—…ì´ ì‹¤í–‰ ì¤‘ì…ë‹ˆë‹¤.", "info")
            return False

        self._db_tasks_inflight.add(task_name)
        if loading_text:
            self._set_status(loading_text, "running")

        payload_context = dict(context or {})

        def _work():
            try:
                result = worker()
                self.message_queue.put(
                    (
                        "db_task_result",
                        {
                            "task": task_name,
                            "result": result,
                            "context": payload_context,
                        },
                    )
                )
            except Exception as e:
                logger.exception("DB ì‘ì—… ì‹¤íŒ¨ (%s)", task_name)
                self.message_queue.put(
                    (
                        "db_task_error",
                        {"task": task_name, "error": str(e), "context": payload_context},
                    )
                )

        threading.Thread(target=_work, daemon=True, name=f"DBTask-{task_name}").start()
        return True

    def _handle_db_task_result(
        self, task_name: str, result: Any, context: dict | None = None
    ) -> None:
        """DB ë¹„ë™ê¸° ì‘ì—… ì™„ë£Œ ì²˜ë¦¬ (UI ìŠ¤ë ˆë“œ)."""
        context = context or {}
        if task_name == "db_history_list":
            sessions = result if isinstance(result, list) else []
            if not sessions:
                QMessageBox.information(self, "ì„¸ì…˜ íˆìŠ¤í† ë¦¬", "ì €ì¥ëœ ì„¸ì…˜ì´ ì—†ìŠµë‹ˆë‹¤.")
                self._set_status("ì„¸ì…˜ íˆìŠ¤í† ë¦¬ ì—†ìŒ", "info")
                return
            self._open_db_history_dialog(sessions)
            self._set_status("ì„¸ì…˜ íˆìŠ¤í† ë¦¬ ì¡°íšŒ ì™„ë£Œ", "success")
            return

        if task_name == "db_search":
            query = str(context.get("query", "")).strip()
            results = result if isinstance(result, list) else []
            if not results:
                QMessageBox.information(
                    self, "ê²€ìƒ‰ ê²°ê³¼", f"'{query}'ì— ëŒ€í•œ ê²€ìƒ‰ ê²°ê³¼ê°€ ì—†ìŠµë‹ˆë‹¤."
                )
                self._set_status("ìë§‰ ê²€ìƒ‰ ê²°ê³¼ ì—†ìŒ", "info")
                return
            self._show_db_search_results(query, results)
            self._set_status(f"ìë§‰ ê²€ìƒ‰ ì™„ë£Œ ({len(results)}ê±´)", "success")
            return

        if task_name == "db_stats":
            stats = result if isinstance(result, dict) else {}
            self._show_db_stats_dialog(stats)
            self._set_status("DB í†µê³„ ì¡°íšŒ ì™„ë£Œ", "success")
            return

        if task_name == "db_history_load_selected":
            self._set_db_history_dialog_busy(False)
            payload = result if isinstance(result, dict) else {}
            if not payload.get("ok"):
                self._set_status("ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì‹¤íŒ¨", "error")
                self._show_toast("ì„¸ì…˜ì„ ë¶ˆëŸ¬ì˜¤ì§€ ëª»í–ˆìŠµë‹ˆë‹¤.", "error")
                return

            new_subtitles = payload.get("subtitles", [])
            skipped = int(payload.get("skipped", 0) or 0)
            self._replace_subtitles_and_refresh(new_subtitles)

            message = f"ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ê¸° ì™„ë£Œ! {len(new_subtitles)}ê°œ ë¬¸ì¥"
            if skipped > 0:
                message += f" (ì†ìƒ í•­ëª© {skipped}ê°œ ì œì™¸)"
            self._show_toast(message, "success")
            self._set_status(message, "success")

            state = self._db_history_dialog_state or {}
            dialog = state.get("dialog")
            if dialog is not None:
                dialog.accept()
            return

        if task_name == "db_history_delete_selected":
            self._set_db_history_dialog_busy(False)
            deleted = bool(result)
            if not deleted:
                self._set_status("ì„¸ì…˜ ì‚­ì œ ì‹¤íŒ¨", "error")
                self._show_toast("ì„¸ì…˜ ì‚­ì œ ì‹¤íŒ¨", "error")
                return

            state = self._db_history_dialog_state or {}
            sessions = state.get("sessions")
            list_widget = state.get("list_widget")
            session_id = context.get("session_id")

            remove_idx = None
            if isinstance(sessions, list):
                for i, item in enumerate(sessions):
                    if item.get("id") == session_id:
                        remove_idx = i
                        break
            if (
                remove_idx is not None
                and list_widget is not None
                and isinstance(sessions, list)
            ):
                list_widget.takeItem(remove_idx)
                sessions.pop(remove_idx)

            self._show_toast("ì„¸ì…˜ ì‚­ì œë¨", "info")
            self._set_status("ì„¸ì…˜ ì‚­ì œ ì™„ë£Œ", "success")
            return

    def _handle_db_task_error(
        self, task_name: str, error: str, context: dict | None = None
    ) -> None:
        """DB ë¹„ë™ê¸° ì‘ì—… ì‹¤íŒ¨ ì²˜ë¦¬ (UI ìŠ¤ë ˆë“œ)."""
        context = context or {}
        if task_name in ("db_history_load_selected", "db_history_delete_selected"):
            self._set_db_history_dialog_busy(False)
        query_hint = str(context.get("query", "")).strip()
        message = f"DB ì‘ì—… ì‹¤íŒ¨ ({task_name}): {error}"
        if task_name == "db_search" and query_hint:
            message = f"ê²€ìƒ‰ ì‹¤íŒ¨ ('{query_hint}'): {error}"
        self._set_status(message, "error")
        QMessageBox.warning(self, "ë°ì´í„°ë² ì´ìŠ¤ ì˜¤ë¥˜", message)

    def _show_db_history(self):
        """ì„¸ì…˜ íˆìŠ¤í† ë¦¬ ë‹¤ì´ì–¼ë¡œê·¸ í‘œì‹œ"""
        self._run_db_task(
            "db_history_list",
            worker=lambda: self.db.list_sessions(limit=50),
            loading_text="DB ì„¸ì…˜ íˆìŠ¤í† ë¦¬ ì¡°íšŒ ì¤‘...",
        )

    def _open_db_history_dialog(self, sessions: list[dict]) -> None:
        dialog = QDialog(self)
        dialog.setWindowTitle("ğŸ“‹ ì„¸ì…˜ íˆìŠ¤í† ë¦¬")
        dialog.setMinimumSize(700, 500)

        layout = QVBoxLayout(dialog)

        # ì„¸ì…˜ ëª©ë¡
        list_widget = QListWidget()
        for s in sessions:
            created = s.get("created_at", "")[:19] if s.get("created_at") else ""
            committee = s.get("committee_name") or "ì•Œ ìˆ˜ ì—†ìŒ"
            subtitles = s.get("total_subtitles", 0)
            chars = s.get("total_characters", 0)
            list_widget.addItem(
                f"[{created}] {committee} - {subtitles}ë¬¸ì¥, {chars:,}ì"
            )

        layout.addWidget(list_widget)

        status_label = QLabel("")
        status_label.hide()
        layout.addWidget(status_label)

        # ë²„íŠ¼
        btn_layout = QHBoxLayout()

        load_btn = QPushButton("ë¶ˆëŸ¬ì˜¤ê¸°")

        def load_selected():
            idx = list_widget.currentRow()
            if idx < 0 or idx >= len(sessions):
                return

            session_id = sessions[idx].get("id")
            if not session_id:
                self._show_toast("ìœ íš¨í•œ ì„¸ì…˜ IDê°€ ì—†ìŠµë‹ˆë‹¤.", "warning")
                return

            def worker(sid=session_id):
                session_data = self.db.load_session(sid)
                if not session_data:
                    return {"ok": False, "subtitles": [], "skipped": 0}
                new_subtitles = []
                skipped = 0
                for item in session_data.get("subtitles", []):
                    try:
                        new_subtitles.append(SubtitleEntry.from_dict(item))
                    except ValueError as e:
                        logger.warning(f"DB ìë§‰ í•­ëª© ê±´ë„ˆëœ€: {e}")
                        skipped += 1
                return {"ok": True, "subtitles": new_subtitles, "skipped": skipped}

            started = self._run_db_task(
                "db_history_load_selected",
                worker=worker,
                context={"session_id": session_id, "row": idx},
                loading_text="DB ì„¸ì…˜ ë¶ˆëŸ¬ì˜¤ëŠ” ì¤‘...",
            )
            if started:
                self._set_db_history_dialog_busy(
                    True, "ì„¸ì…˜ì„ ë¶ˆëŸ¬ì˜¤ëŠ” ì¤‘ì…ë‹ˆë‹¤..."
                )

        load_btn.clicked.connect(load_selected)
        btn_layout.addWidget(load_btn)

        delete_btn = QPushButton("ì‚­ì œ")

        def delete_selected():
            idx = list_widget.currentRow()
            if idx < 0 or idx >= len(sessions):
                return

            session_id = sessions[idx].get("id")
            if not session_id:
                self._show_toast("ìœ íš¨í•œ ì„¸ì…˜ IDê°€ ì—†ìŠµë‹ˆë‹¤.", "warning")
                return

            reply = QMessageBox.question(
                dialog,
                "ì‚­ì œ í™•ì¸",
                "ì„ íƒí•œ ì„¸ì…˜ì„ ì‚­ì œí•˜ì‹œê² ìŠµë‹ˆê¹Œ?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if reply != QMessageBox.StandardButton.Yes:
                return

            started = self._run_db_task(
                "db_history_delete_selected",
                worker=lambda sid=session_id: self.db.delete_session(sid),
                context={"session_id": session_id, "row": idx},
                loading_text="DB ì„¸ì…˜ ì‚­ì œ ì¤‘...",
            )
            if started:
                self._set_db_history_dialog_busy(True, "ì„¸ì…˜ì„ ì‚­ì œí•˜ëŠ” ì¤‘ì…ë‹ˆë‹¤...")

        delete_btn.clicked.connect(delete_selected)
        btn_layout.addWidget(delete_btn)

        close_btn = QPushButton("ë‹«ê¸°")
        close_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(close_btn)

        layout.addLayout(btn_layout)

        self._db_history_dialog_state = {
            "dialog": dialog,
            "sessions": sessions,
            "list_widget": list_widget,
            "status_label": status_label,
            "load_btn": load_btn,
            "delete_btn": delete_btn,
            "close_btn": close_btn,
        }
        dialog.finished.connect(lambda *_: self._clear_db_history_dialog_state())
        dialog.exec()

    def _show_db_search(self):
        """ìë§‰ í†µí•© ê²€ìƒ‰ ë‹¤ì´ì–¼ë¡œê·¸"""
        if not self.db:
            QMessageBox.warning(self, "ì•Œë¦¼", "ë°ì´í„°ë² ì´ìŠ¤ê°€ ì´ˆê¸°í™”ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤.")
            return

        query, ok = QInputDialog.getText(self, "ìë§‰ ê²€ìƒ‰", "ê²€ìƒ‰ì–´:")
        query = query.strip() if ok and query else ""
        if not query:
            return

        self._run_db_task(
            "db_search",
            worker=lambda q=query: self.db.search_subtitles(q),
            context={"query": query},
            loading_text=f"DB ìë§‰ ê²€ìƒ‰ ì¤‘... ({query[:15]})",
        )

    def _show_db_search_results(self, query: str, results: list[dict]) -> None:
        dialog = QDialog(self)
        dialog.setWindowTitle(f"ğŸ” ê²€ìƒ‰ ê²°ê³¼ - '{query}'")
        dialog.setMinimumSize(700, 500)

        layout = QVBoxLayout(dialog)
        layout.addWidget(QLabel(f"ì´ {len(results)}ê°œ ê²°ê³¼"))

        list_widget = QListWidget()
        for r in results:
            created = r.get("created_at", "")[:10] if r.get("created_at") else ""
            committee = r.get("committee_name") or ""
            text = r.get("text", "")[:100]
            list_widget.addItem(f"[{created}] {committee}: {text}")

        layout.addWidget(list_widget)

        close_btn = QPushButton("ë‹«ê¸°")
        close_btn.clicked.connect(dialog.reject)
        layout.addWidget(close_btn)

        dialog.exec()

    def _show_db_stats(self):
        """ë°ì´í„°ë² ì´ìŠ¤ ì „ì²´ í†µê³„"""
        self._run_db_task(
            "db_stats",
            worker=lambda: self.db.get_statistics(),
            loading_text="DB í†µê³„ ì¡°íšŒ ì¤‘...",
        )

    def _show_db_stats_dialog(self, stats: dict) -> None:
        msg = f"""
<h2>ğŸ“Š ë°ì´í„°ë² ì´ìŠ¤ í†µê³„</h2>
<table>
<tr><td><b>ì´ ì„¸ì…˜ ìˆ˜:</b></td><td>{stats.get("total_sessions", 0):,}ê°œ</td></tr>
<tr><td><b>ì´ ìë§‰ ìˆ˜:</b></td><td>{stats.get("total_subtitles", 0):,}ê°œ</td></tr>
<tr><td><b>ì´ ê¸€ì ìˆ˜:</b></td><td>{stats.get("total_characters", 0):,}ì</td></tr>
<tr><td><b>ì´ ë…¹í™” ì‹œê°„:</b></td><td>{stats.get("total_duration_hours", 0):.1f}ì‹œê°„</td></tr>
</table>
"""
        QMessageBox.information(self, "ë°ì´í„°ë² ì´ìŠ¤ í†µê³„", msg)

    # ========== ìë§‰ ë³‘í•© ê¸°ëŠ¥ (#20) ==========

    def _show_merge_dialog(self):
        """ìë§‰ ë³‘í•© ë‹¤ì´ì–¼ë¡œê·¸"""
        dialog = QDialog(self)
        dialog.setWindowTitle("ğŸ“ ìë§‰ ë³‘í•©")
        dialog.setMinimumSize(600, 500)

        layout = QVBoxLayout(dialog)

        # íŒŒì¼ ëª©ë¡
        layout.addWidget(QLabel("ë³‘í•©í•  ì„¸ì…˜ íŒŒì¼ì„ ì¶”ê°€í•˜ì„¸ìš”:"))
        file_list = QListWidget()
        file_list.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        layout.addWidget(file_list)

        file_paths = []

        # íŒŒì¼ ì¶”ê°€/ì œê±° ë²„íŠ¼
        file_btn_layout = QHBoxLayout()

        add_btn = QPushButton("â• íŒŒì¼ ì¶”ê°€")

        def add_files():
            paths, _ = QFileDialog.getOpenFileNames(
                dialog, "ì„¸ì…˜ íŒŒì¼ ì„ íƒ", f"{Config.SESSION_DIR}/", "JSON íŒŒì¼ (*.json)"
            )
            for path in paths:
                if path not in file_paths:
                    file_paths.append(path)
                    file_list.addItem(Path(path).name)

        add_btn.clicked.connect(add_files)
        file_btn_layout.addWidget(add_btn)

        remove_btn = QPushButton("â– ì„ íƒ ì œê±°")

        def remove_files():
            for item in file_list.selectedItems():
                idx = file_list.row(item)
                file_list.takeItem(idx)
                if idx < len(file_paths):
                    file_paths.pop(idx)

        remove_btn.clicked.connect(remove_files)
        file_btn_layout.addWidget(remove_btn)

        layout.addLayout(file_btn_layout)

        # ì˜µì…˜
        options_layout = QHBoxLayout()
        remove_dup_check = QCheckBox("ì¤‘ë³µ ìë§‰ ì œê±°")
        remove_dup_check.setChecked(True)
        sort_check = QCheckBox("ì‹œê°„ìˆœ ì •ë ¬")
        sort_check.setChecked(True)
        options_layout.addWidget(remove_dup_check)
        options_layout.addWidget(sort_check)
        options_layout.addStretch()
        layout.addLayout(options_layout)

        # ë²„íŠ¼
        btn_layout = QHBoxLayout()

        merge_btn = QPushButton("ë³‘í•© ì‹¤í–‰")

        def do_merge():
            if len(file_paths) < 2:
                QMessageBox.warning(dialog, "ì•Œë¦¼", "2ê°œ ì´ìƒì˜ íŒŒì¼ì„ ì„ íƒí•˜ì„¸ìš”.")
                return

            # ê¸°ì¡´ ìë§‰ ì²˜ë¦¬ ì˜µì…˜ í™•ì¸
            existing_subtitles = None
            if self.subtitles:
                reply = QMessageBox.question(
                    dialog,
                    "ê¸°ì¡´ ìë§‰ ì²˜ë¦¬",
                    f"í˜„ì¬ {len(self.subtitles)}ê°œì˜ ìë§‰ì´ ìˆìŠµë‹ˆë‹¤.\n\n"
                    "ê¸°ì¡´ ìë§‰ì„ ë³‘í•© ê²°ê³¼ì— í¬í•¨í•˜ì‹œê² ìŠµë‹ˆê¹Œ?\n"
                    "(Yes: í¬í•¨í•˜ì—¬ ë³‘í•© / No: ê¸°ì¡´ ìë§‰ ë¬´ì‹œí•˜ê³  íŒŒì¼ë“¤ë§Œ ë³‘í•©)",
                    QMessageBox.StandardButton.Yes
                    | QMessageBox.StandardButton.No
                    | QMessageBox.StandardButton.Cancel,
                )
                if reply == QMessageBox.StandardButton.Cancel:
                    return

                if reply == QMessageBox.StandardButton.Yes:
                    with self.subtitle_lock:
                        existing_subtitles = list(self.subtitles)

            merged = self._merge_sessions(
                file_paths,
                remove_duplicates=remove_dup_check.isChecked(),
                sort_by_time=sort_check.isChecked(),
                existing_subtitles=existing_subtitles,
            )

            if merged:
                self._replace_subtitles_and_refresh(merged)
                self._show_toast(f"ë³‘í•© ì™„ë£Œ! {len(merged)}ê°œ ë¬¸ì¥", "success")
                dialog.accept()

        merge_btn.clicked.connect(do_merge)
        btn_layout.addWidget(merge_btn)

        cancel_btn = QPushButton("ì·¨ì†Œ")
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(cancel_btn)

        layout.addLayout(btn_layout)
        dialog.exec()

    def _clean_newlines(self):
        """ì¤„ë„˜ê¹€ ì •ë¦¬: ë¬¸ì¥ ë¶€í˜¸ ë¶„ë¦¬ ë° ë³‘í•© (ìŠ¤ë§ˆíŠ¸ ë¦¬í”Œë¡œìš°)"""
        # ìë§‰ì´ ì—†ëŠ” ê²½ìš° ì²˜ë¦¬
        if not self.subtitles:
            self._show_toast("ì •ë¦¬í•  ìë§‰ì´ ì—†ìŠµë‹ˆë‹¤.", "warning")
            return

        # ì‚¬ìš©ì í™•ì¸
        reply = QMessageBox.question(
            self,
            "ì¤„ë„˜ê¹€ ì •ë¦¬ (Smart Reflow)",
            "ìë§‰ ì¬ì •ë ¬ì„ ìˆ˜í–‰í•˜ì‹œê² ìŠµë‹ˆê¹Œ?\n\n"
            "ê¸°ëŠ¥:\n"
            "1. í…ìŠ¤íŠ¸ ë‚´ íƒ€ì„ìŠ¤íƒ¬í”„([HH:MM:SS])ë¥¼ ê°ì§€í•˜ì—¬ ë¶„ë¦¬\n"
            "2. ë¬¸ì¥ ë¶€í˜¸(. ? !) ê¸°ì¤€ìœ¼ë¡œ ì¤„ ë°”ê¿ˆ\n"
            "3. ëŠì–´ì§„ ë¬¸ì¥ ë³‘í•©\n\n"
            "(ì£¼ì˜: ë˜ëŒë¦¬ê¸°ëŠ” ì§€ì›ë˜ì§€ ì•ŠìŠµë‹ˆë‹¤.)",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply == QMessageBox.StandardButton.No:
            return

        # ë³‘í•©/ë¶„ë¦¬ ë¡œì§ (utils ëª¨ë“ˆ ìœ„ì„)
        try:
            old_count = len(self.subtitles)

            # ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ë³µì‚¬ë³¸ ìƒì„± í›„ ì²˜ë¦¬
            with self.subtitle_lock:
                current_subs = list(self.subtitles)

            # ì²˜ë¦¬ (ì‹œê°„ ì†Œìš” ê°€ëŠ¥ì„± ìˆìœ¼ë¯€ë¡œ ë½ ë°–ì—ì„œ ìˆ˜í–‰ ê¶Œì¥í•˜ì§€ë§Œ, ë°ì´í„°ê°€ í¬ì§€ ì•Šì•„ ì¦‰ì‹œ ìˆ˜í–‰)
            new_subtitles = utils.reflow_subtitles(current_subs)

            if not new_subtitles:
                return

            # ì›ë³¸ê³¼ ê°œìˆ˜ ì°¨ì´ í™•ì¸
            logger.info(f"ìŠ¤ë§ˆíŠ¸ ë¦¬í”Œë¡œìš°: {old_count} -> {len(new_subtitles)}")

            self._replace_subtitles_and_refresh(new_subtitles)

            # ê²°ê³¼ ì•Œë¦¼
            self._show_toast(f"ì •ë¦¬ ì™„ë£Œ! ({len(new_subtitles)}ê°œ ë¬¸ì¥)", "success")

        except Exception as e:
            logger.error(f"ë¦¬í”Œë¡œìš° ì¤‘ ì˜¤ë¥˜: {e}")
            self._show_toast(f"ì˜¤ë¥˜ ë°œìƒ: {e}", "error")

    def _merge_sessions(
        self,
        file_paths: list,
        remove_duplicates: bool = True,
        sort_by_time: bool = True,
        existing_subtitles: list = None,
    ) -> list:
        """ì—¬ëŸ¬ ì„¸ì…˜ íŒŒì¼ì„ ë³‘í•©

        Args:
            file_paths: ë³‘í•©í•  íŒŒì¼ ê²½ë¡œ ëª©ë¡
            remove_duplicates: ì¤‘ë³µ ìë§‰ ì œê±° ì—¬ë¶€
            sort_by_time: ì‹œê°„ìˆœ ì •ë ¬ ì—¬ë¶€
            existing_subtitles: ê¸°ì¡´ ìë§‰ ë¦¬ìŠ¤íŠ¸ (ì„ íƒ)

        Returns:
            List[SubtitleEntry]: ë³‘í•©ëœ ìë§‰ ëª©ë¡
        """
        all_entries = []

        # ê¸°ì¡´ ìë§‰ ì¶”ê°€
        if existing_subtitles:
            all_entries.extend(existing_subtitles)

        for path in file_paths:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)

                for item in data.get("subtitles", []):
                    try:
                        entry = SubtitleEntry.from_dict(item)
                        all_entries.append(entry)
                    except ValueError as e:
                        logger.warning(f"ìë§‰ í•­ëª© ê±´ë„ˆëœ€ ({path}): {e}")

            except Exception as e:
                logger.warning(f"íŒŒì¼ ë¡œë“œ ì‹¤íŒ¨ ({path}): {e}")

        # ì‹œê°„ìˆœ ì •ë ¬
        if sort_by_time:
            all_entries.sort(key=lambda e: e.timestamp)

        # ì¤‘ë³µ ì œê±°
        if remove_duplicates:
            seen = set()
            unique_entries = []
            for entry in all_entries:
                text_normalized = entry.text.strip().lower()
                if text_normalized not in seen:
                    seen.add(text_normalized)
                    unique_entries.append(entry)
            all_entries = unique_entries

        logger.info(f"ë³‘í•© ì™„ë£Œ: {len(all_entries)}ê°œ ìë§‰")
        return all_entries

    def closeEvent(self, event):
        # íŠ¸ë ˆì´ ìµœì†Œí™” ëª¨ë“œ
        if self.minimize_to_tray and self.tray_icon.isVisible():
            self.hide()
            self.tray_icon.showMessage(
                Config.APP_NAME,
                "í”„ë¡œê·¸ë¨ì´ íŠ¸ë ˆì´ë¡œ ìµœì†Œí™”ë˜ì—ˆìŠµë‹ˆë‹¤.\níŠ¸ë ˆì´ ì•„ì´ì½˜ì„ ë”ë¸”í´ë¦­í•˜ì—¬ ë‹¤ì‹œ ì—´ ìˆ˜ ìˆìŠµë‹ˆë‹¤.",
                QSystemTrayIcon.MessageIcon.Information,
                2000,
            )
            event.ignore()
            return

        # ì¶”ì¶œ ì¤‘ì´ë©´ ì¢…ë£Œ í™•ì¸ í›„ ë¨¼ì € ì•ˆì „ ì¤‘ì§€(í ì†Œì§„/ë§ˆì§€ë§‰ ìë§‰ í™•ì •)
        if self.is_running:
            reply = QMessageBox.question(
                self,
                "ì¢…ë£Œ",
                "ì¶”ì¶œ ì¤‘ì…ë‹ˆë‹¤. ì¢…ë£Œí•˜ì‹œê² ìŠµë‹ˆê¹Œ?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if reply == QMessageBox.StandardButton.No:
                event.ignore()
                return
            self._stop(for_app_exit=True)

        with self.subtitle_lock:
            subtitle_count = len(self.subtitles)

        # ì €ì¥í•˜ì§€ ì•Šì€ ìë§‰ì´ ìˆìœ¼ë©´ í™•ì¸
        if subtitle_count:
            reply = QMessageBox.question(
                self,
                "ì¢…ë£Œ í™•ì¸",
                f"ì €ì¥í•˜ì§€ ì•Šì€ ìë§‰ {subtitle_count}ê°œê°€ ìˆìŠµë‹ˆë‹¤.\n\n"
                "ì €ì¥í•˜ì‹œê² ìŠµë‹ˆê¹Œ?",
                QMessageBox.StandardButton.Save
                | QMessageBox.StandardButton.Discard
                | QMessageBox.StandardButton.Cancel,
            )
            if reply == QMessageBox.StandardButton.Cancel:
                event.ignore()
                return
            elif reply == QMessageBox.StandardButton.Save:
                # ì €ì¥ ëŒ€í™”ìƒìì—ì„œ ì·¨ì†Œí•˜ë©´ ì¢…ë£Œë„ ì·¨ì†Œ
                filename = f"êµ­íšŒìë§‰_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                path, _ = QFileDialog.getSaveFileName(
                    self, "TXT ì €ì¥", filename, "í…ìŠ¤íŠ¸ (*.txt)"
                )
                if not path:  # ì‚¬ìš©ìê°€ ì·¨ì†Œí•œ ê²½ìš°
                    event.ignore()
                    return
                # íŒŒì¼ ì €ì¥ (ìŠ¤ë ˆë“œ ì•ˆì „í•˜ê²Œ ìë§‰ ë³µì‚¬)
                try:
                    with self.subtitle_lock:
                        subtitles_snapshot = list(self.subtitles)
                    with open(path, "w", encoding="utf-8") as f:
                        for entry in subtitles_snapshot:
                            f.write(
                                f"[{entry.timestamp.strftime('%H:%M:%S')}] {entry.text}\n"
                            )
                except Exception as e:
                    QMessageBox.critical(self, "ì˜¤ë¥˜", f"ì €ì¥ ì‹¤íŒ¨: {e}")
                    event.ignore()
                    return

        # ì°½ ìœ„ì¹˜/í¬ê¸° ì €ì¥
        self.settings.setValue("geometry", self.saveGeometry())
        self.settings.setValue("windowState", self.saveState())

        # íƒ€ì´ë¨¸ ì •ë¦¬
        self.queue_timer.stop()
        self.stats_timer.stop()
        self.finalize_timer.stop()
        self.backup_timer.stop()

        if self.realtime_file:
            try:
                self.realtime_file.close()
            except Exception as e:
                logger.debug(f"íŒŒì¼ ë‹«ê¸° ì˜¤ë¥˜: {e}")
            self.realtime_file = None

        # closeEvent ì‹œì—ëŠ” ì‹¤í–‰ ì—¬ë¶€ì™€ ë¬´ê´€í•˜ê²Œ ë“œë¼ì´ë²„ ì •ë¦¬ë¥¼ ì‹œë„í•œë‹¤.
        if self.driver:
            driver = self.driver
            self.driver = None
            self._force_quit_driver_with_timeout(
                driver, timeout=2.0, source="close_event_idle"
            )
        self._cleanup_detached_drivers_with_timeout(timeout=2.0)

        # ë°ì´í„°ë² ì´ìŠ¤ ì—°ê²° ì •ë¦¬
        if getattr(self, "db", None):
            try:
                self.db.close_all()
            except Exception as e:
                logger.debug(f"DB ì—°ê²° ì¢…ë£Œ ì˜¤ë¥˜: {e}")

        logger.info("í”„ë¡œê·¸ë¨ ì¢…ë£Œ")
        event.accept()

    def _toggle_top_header(self):
        """ìƒë‹¨ ì˜ì—­(í—¤ë”/íˆ´ë°”) í‘œì‹œ/ìˆ¨ê¹€ í† ê¸€"""
        if self.top_header_container.isVisible():
            # ì ‘ê¸°: í—¤ë” ìˆ¨ê¹€ & ì„¤ì • ê·¸ë£¹ ì ‘ê¸°
            self.top_header_container.hide()
            self.settings_group.set_collapsed(True)
            self.toggle_header_btn.setText("ğŸ”½ ìƒë‹¨ í¼ì¹˜ê¸°")
        else:
            # í¼ì¹˜ê¸°: í—¤ë” ë³´ì„ & ì„¤ì • ê·¸ë£¹ í¼ì¹˜ê¸°
            self.top_header_container.show()
            self.settings_group.set_collapsed(False)
            self.toggle_header_btn.setText("ğŸ”¼ ìƒë‹¨ ì ‘ê¸°")
