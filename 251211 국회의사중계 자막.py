# -*- coding: utf-8 -*-
"""
국회 의사중계 자막 추출기 v10.0
- 참조 코드 기반 안정적 구조
- PyQt6 모던 UI
- 추가 기능: 타임스탬프, 실시간 저장, 검색, 테마, URL 히스토리, SRT/VTT 내보내기, 통계, 키워드 하이라이트, 세션 저장
"""

import sys
import os
import time
import threading
import queue
import re
import json
import logging
from datetime import datetime, timedelta
from pathlib import Path


# ============================================================
# 로깅 설정
# ============================================================

def setup_logging():
    """로깅 시스템 초기화 - 파일 및 콘솔 출력"""
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    
    log_file = log_dir / f"subtitle_{datetime.now().strftime('%Y%m%d')}.log"
    
    # 로거 설정
    logger = logging.getLogger("SubtitleExtractor")
    logger.setLevel(logging.DEBUG)
    
    # 이미 핸들러가 있으면 추가하지 않음
    if logger.handlers:
        return logger
    
    # 파일 핸들러 (DEBUG 레벨)
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_format = logging.Formatter(
        '%(asctime)s [%(levelname)s] %(funcName)s: %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(file_format)
    
    # 콘솔 핸들러 (INFO 레벨)
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_format = logging.Formatter('[%(levelname)s] %(message)s')
    console_handler.setFormatter(console_format)
    
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    return logger


# 전역 로거 초기화
logger = setup_logging()

try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QLabel, QPushButton, QLineEdit, QTextEdit, QComboBox, QCheckBox,
        QFrame, QProgressBar, QMessageBox, QFileDialog, QSizePolicy,
        QGroupBox, QGridLayout, QDialog, QDialogButtonBox, QListWidget,
        QSplitter, QMenu, QMenuBar, QInputDialog
    )
    from PyQt6.QtCore import Qt, QTimer, QSettings
    from PyQt6.QtGui import (
        QFont, QColor, QTextCursor, QTextCharFormat, QAction,
        QShortcut, QKeySequence, QPalette
    )
except ImportError:
    print("PyQt6 필요: pip install PyQt6")
    sys.exit(1)

try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.common.exceptions import WebDriverException, NoSuchElementException, StaleElementReferenceException
    from selenium.webdriver.chrome.options import Options
except ImportError:
    print("selenium 필요: pip install selenium")
    sys.exit(1)


# ============================================================
# 상수 정의
# ============================================================

class Config:
    """프로그램 설정 상수"""
    VERSION = "10.1"  # 디버깅/리팩토링 버전
    
    # 타이밍 상수 (초)
    SUBTITLE_FINALIZE_DELAY = 2.0      # 자막 확정까지 대기 시간
    FINALIZE_CHECK_INTERVAL = 500      # 자막 확정 체크 간격 (ms)
    QUEUE_PROCESS_INTERVAL = 100       # 메시지 큐 처리 간격 (ms)
    STATS_UPDATE_INTERVAL = 1000       # 통계 업데이트 간격 (ms)
    SUBTITLE_CHECK_INTERVAL = 0.2      # 자막 확인 간격 (초)
    THREAD_STOP_TIMEOUT = 3            # 스레드 종료 대기 시간 (초)
    PAGE_LOAD_WAIT = 3                 # 페이지 로딩 대기 시간 (초)
    WEBDRIVER_WAIT_TIMEOUT = 20        # WebDriver 대기 타임아웃 (초)
    SCRIPT_DELAY = 0.5                 # 스크립트 실행 후 대기 (초)
    
    # 기본 CSS 선택자
    DEFAULT_SELECTORS = [
        "#viewSubtit .incont",
        "#viewSubtit",
        ".subtitle_area"
    ]
    
    # 기본 URL
    DEFAULT_URL = "https://assembly.webcast.go.kr/main/player.asp"


# ============================================================
# 테마 정의
# ============================================================

DARK_THEME = """
QMainWindow, QWidget {
    background-color: #1a1a2e;
    color: #eaeaea;
    font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
}
QLabel { 
    color: #eaeaea; 
}
QLabel#headerLabel {
    font-size: 18px;
    font-weight: bold;
    color: #4fc3f7;
    padding: 10px;
}
QPushButton {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #0f3460, stop:1 #0a2540);
    color: white;
    border: none;
    padding: 10px 20px;
    border-radius: 8px;
    font-weight: bold;
    font-size: 12px;
}
QPushButton:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1a4a7a, stop:1 #0f3460);
}
QPushButton:pressed {
    background: #0a2540;
}
QPushButton:disabled { 
    background-color: #333; 
    color: #666; 
}
QPushButton#startBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #00c853, stop:1 #00a86b);
    font-size: 14px;
    min-width: 120px;
}
QPushButton#startBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #00e676, stop:1 #00c853);
}
QPushButton#stopBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ff5252, stop:1 #e94560);
    font-size: 14px;
    min-width: 120px;
}
QPushButton#stopBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ff8a80, stop:1 #ff5252);
}
QLineEdit, QComboBox {
    background-color: #16213e;
    border: 2px solid #0f3460;
    border-radius: 8px;
    padding: 10px;
    color: #eaeaea;
    font-size: 13px;
}
QLineEdit:focus, QComboBox:focus {
    border: 2px solid #4fc3f7;
}
QTextEdit {
    background-color: #0d1421;
    border: 2px solid #0f3460;
    border-radius: 10px;
    padding: 15px;
    color: #eaeaea;
    font-size: 14px;
    line-height: 1.5;
}
QGroupBox {
    border: 2px solid #0f3460;
    border-radius: 10px;
    margin-top: 15px;
    padding-top: 15px;
    color: #4fc3f7;
    font-weight: bold;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 15px;
    padding: 0 8px;
    color: #4fc3f7;
}
QCheckBox {
    spacing: 8px;
    color: #eaeaea;
}
QCheckBox::indicator {
    width: 18px;
    height: 18px;
    border-radius: 4px;
    border: 2px solid #0f3460;
    background: #16213e;
}
QCheckBox::indicator:checked {
    background: #4fc3f7;
    border-color: #4fc3f7;
}
QCheckBox::indicator:hover {
    border-color: #4fc3f7;
}
QProgressBar {
    background-color: #16213e;
    border-radius: 6px;
    height: 8px;
    text-align: center;
}
QProgressBar::chunk { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #4fc3f7, stop:1 #e94560);
    border-radius: 6px; 
}
QMenuBar { 
    background-color: #16213e; 
    color: #eaeaea; 
    padding: 5px;
}
QMenuBar::item { 
    padding: 8px 15px;
    border-radius: 5px;
}
QMenuBar::item:selected { 
    background-color: #0f3460; 
}
QMenu { 
    background-color: #16213e; 
    color: #eaeaea; 
    border: 2px solid #0f3460;
    border-radius: 8px;
    padding: 5px;
}
QMenu::item { 
    padding: 8px 25px;
    border-radius: 5px;
}
QMenu::item:selected { 
    background-color: #0f3460; 
}
QScrollBar:vertical {
    background: #0d1421;
    width: 12px;
    border-radius: 6px;
}
QScrollBar::handle:vertical {
    background: #0f3460;
    border-radius: 6px;
    min-height: 30px;
}
QScrollBar::handle:vertical:hover {
    background: #1a4a7a;
}
QSplitter::handle {
    background: #0f3460;
}
"""

LIGHT_THEME = """
QMainWindow, QWidget {
    background-color: #f8f9fa;
    color: #333333;
    font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
}
QLabel { 
    color: #333333; 
}
QLabel#headerLabel {
    font-size: 18px;
    font-weight: bold;
    color: #1976d2;
    padding: 10px;
}
QPushButton {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #5a9fd4, stop:1 #4a90d9);
    color: white;
    border: none;
    padding: 10px 20px;
    border-radius: 8px;
    font-weight: bold;
    font-size: 12px;
}
QPushButton:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #6ab0e5, stop:1 #5a9fd4);
}
QPushButton:disabled { 
    background-color: #cccccc; 
    color: #888888; 
}
QPushButton#startBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #4caf50, stop:1 #27ae60);
    font-size: 14px;
    min-width: 120px;
}
QPushButton#startBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #66bb6a, stop:1 #4caf50);
}
QPushButton#stopBtn { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ef5350, stop:1 #e74c3c);
    font-size: 14px;
    min-width: 120px;
}
QPushButton#stopBtn:hover { 
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ff7043, stop:1 #ef5350);
}
QLineEdit, QComboBox {
    background-color: #ffffff;
    border: 2px solid #e0e0e0;
    border-radius: 8px;
    padding: 10px;
    color: #333333;
    font-size: 13px;
}
QLineEdit:focus, QComboBox:focus {
    border: 2px solid #1976d2;
}
QTextEdit {
    background-color: #ffffff;
    border: 2px solid #e0e0e0;
    border-radius: 10px;
    padding: 15px;
    color: #333333;
    font-size: 14px;
}
QGroupBox {
    border: 2px solid #e0e0e0;
    border-radius: 10px;
    margin-top: 15px;
    padding-top: 15px;
    color: #1976d2;
    font-weight: bold;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 15px;
    padding: 0 8px;
    color: #1976d2;
}
QCheckBox::indicator {
    width: 18px;
    height: 18px;
    border-radius: 4px;
    border: 2px solid #e0e0e0;
    background: #ffffff;
}
QCheckBox::indicator:checked {
    background: #1976d2;
    border-color: #1976d2;
}
QProgressBar {
    background-color: #e0e0e0;
    border-radius: 6px;
    height: 8px;
}
QProgressBar::chunk { 
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #1976d2, stop:1 #4caf50);
    border-radius: 6px; 
}
QMenuBar { 
    background-color: #ffffff; 
    color: #333333;
    padding: 5px;
}
QMenuBar::item { 
    padding: 8px 15px;
    border-radius: 5px;
}
QMenuBar::item:selected { 
    background-color: #e3f2fd; 
}
QMenu { 
    background-color: #ffffff; 
    color: #333333;
    border: 2px solid #e0e0e0;
    border-radius: 8px;
}
QMenu::item { 
    padding: 8px 25px;
    border-radius: 5px;
}
QMenu::item:selected { 
    background-color: #e3f2fd; 
}
QScrollBar:vertical {
    background: #f0f0f0;
    width: 12px;
    border-radius: 6px;
}
QScrollBar::handle:vertical {
    background: #c0c0c0;
    border-radius: 6px;
    min-height: 30px;
}
QScrollBar::handle:vertical:hover {
    background: #a0a0a0;
}
"""


# ============================================================
# 자막 데이터 클래스
# ============================================================

class SubtitleEntry:
    """자막 항목"""
    def __init__(self, text, timestamp=None):
        self.text = text
        self.timestamp = timestamp or datetime.now()
        self.start_time = None  # SRT용
        self.end_time = None    # SRT용


# ============================================================
# 메인 윈도우
# ============================================================

class MainWindow(QMainWindow):
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"국회 의사중계 자막 추출기 v{Config.VERSION}")
        self.setMinimumSize(1100, 750)
        self.resize(1200, 800)
        
        # 설정
        self.settings = QSettings("AssemblySubtitle", "Extractor")
        self.is_dark_theme = self.settings.value("dark_theme", True, type=bool)
        
        # 메시지 큐
        self.message_queue = queue.Queue()
        
        # 상태
        self.worker = None
        self.driver = None
        self.is_running = False
        self.stop_event = threading.Event()  # 스레드 안전한 종료 시그널
        self.start_time = None
        self.last_subtitle = ""
        
        # 자막 데이터 (타임스탬프 포함)
        self.subtitles = []  # List[SubtitleEntry]
        self.keywords = []   # 하이라이트 키워드
        self.last_update_time = 0  # 마지막 자막 업데이트 시간
        
        # 실시간 저장
        self.realtime_file = None
        
        # URL 히스토리
        self.url_history = self._load_url_history()
        
        # UI 생성
        self._create_menu()
        self._create_ui()
        self._apply_theme()
        self._setup_shortcuts()
        
        # 타이머
        self.queue_timer = QTimer(self)
        self.queue_timer.timeout.connect(self._process_message_queue)
        self.queue_timer.start(Config.QUEUE_PROCESS_INTERVAL)
        
        self.stats_timer = QTimer(self)
        self.stats_timer.timeout.connect(self._update_stats)
        
        # 자막 확정 타이머
        self.finalize_timer = QTimer(self)
        self.finalize_timer.timeout.connect(self._check_finalize)
        
        # 디렉토리 생성
        Path("sessions").mkdir(exist_ok=True)
        Path("realtime_output").mkdir(exist_ok=True)
    
    def _create_menu(self):
        menubar = self.menuBar()
        
        # 파일 메뉴
        file_menu = menubar.addMenu("파일")
        
        save_txt = QAction("TXT 저장", self)
        save_txt.setShortcut("Ctrl+S")
        save_txt.triggered.connect(self._save_txt)
        file_menu.addAction(save_txt)
        
        save_srt = QAction("SRT 저장", self)
        save_srt.triggered.connect(self._save_srt)
        file_menu.addAction(save_srt)
        
        save_vtt = QAction("VTT 저장", self)
        save_vtt.triggered.connect(self._save_vtt)
        file_menu.addAction(save_vtt)
        
        save_docx = QAction("DOCX 저장 (Word)", self)
        save_docx.triggered.connect(self._save_docx)
        file_menu.addAction(save_docx)
        
        file_menu.addSeparator()
        
        save_session = QAction("세션 저장", self)
        save_session.setShortcut("Ctrl+Shift+S")
        save_session.triggered.connect(self._save_session)
        file_menu.addAction(save_session)
        
        load_session = QAction("세션 불러오기", self)
        load_session.setShortcut("Ctrl+O")
        load_session.triggered.connect(self._load_session)
        file_menu.addAction(load_session)
        
        file_menu.addSeparator()
        
        exit_action = QAction("종료", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # 편집 메뉴
        edit_menu = menubar.addMenu("편집")
        
        search_action = QAction("검색", self)
        search_action.setShortcut("Ctrl+F")
        search_action.triggered.connect(self._show_search)
        edit_menu.addAction(search_action)
        
        keyword_action = QAction("키워드 설정", self)
        keyword_action.triggered.connect(self._set_keywords)
        edit_menu.addAction(keyword_action)
        
        edit_menu.addSeparator()
        
        copy_action = QAction("클립보드 복사", self)
        copy_action.setShortcut("Ctrl+C")
        copy_action.triggered.connect(self._copy_to_clipboard)
        edit_menu.addAction(copy_action)
        
        clear_action = QAction("내용 지우기", self)
        clear_action.triggered.connect(self._clear_text)
        edit_menu.addAction(clear_action)
        
        # 보기 메뉴
        view_menu = menubar.addMenu("보기")
        
        self.theme_action = QAction("라이트 테마" if self.is_dark_theme else "다크 테마", self)
        self.theme_action.setShortcut("Ctrl+T")
        self.theme_action.triggered.connect(self._toggle_theme)
        view_menu.addAction(self.theme_action)
        
        self.timestamp_action = QAction("타임스탬프 표시", self)
        self.timestamp_action.setCheckable(True)
        self.timestamp_action.setChecked(True)
        self.timestamp_action.triggered.connect(self._refresh_text)
        view_menu.addAction(self.timestamp_action)
        
        # 도움말 메뉴
        help_menu = menubar.addMenu("도움말")
        
        guide_action = QAction("사용법 가이드", self)
        guide_action.setShortcut("F1")
        guide_action.triggered.connect(self._show_guide)
        help_menu.addAction(guide_action)
        
        features_action = QAction("기능 소개", self)
        features_action.triggered.connect(self._show_features)
        help_menu.addAction(features_action)
        
        help_menu.addSeparator()
        
        about_action = QAction("정보", self)
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)
    
    def _create_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        
        layout = QVBoxLayout(central)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # === 헤더 ===
        header = QLabel("🏛️ 국회 의사중계 자막 추출기")
        header.setObjectName("headerLabel")
        header.setFont(QFont("맑은 고딕", 18, QFont.Weight.Bold))
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(header)
        
        # === URL/설정 영역 ===
        settings_group = QGroupBox("⚙️ 설정")
        settings_layout = QGridLayout(settings_group)
        settings_layout.setSpacing(10)
        
        # URL
        url_label = QLabel("📌 URL:")
        url_label.setToolTip("국회 의사중계 웹사이트 URL을 입력하세요")
        settings_layout.addWidget(url_label, 0, 0)
        
        url_layout = QHBoxLayout()
        self.url_combo = QComboBox()
        self.url_combo.setEditable(True)
        self.url_combo.setToolTip("국회 의사중계 웹사이트 URL\n최근 사용한 URL이 자동 저장됩니다\n태그가 있으면 [태그] 형태로 표시됩니다")
        
        # URL 히스토리 로드 및 콤보박스 초기화
        for url, tag in self.url_history.items():
            if tag:
                self.url_combo.addItem(f"[{tag}] {url}", url)
            else:
                self.url_combo.addItem(url, url)
        
        if not self.url_history:
            self.url_combo.addItem("https://assembly.webcast.go.kr/main/player.asp")
        
        url_layout.addWidget(self.url_combo, 1)
        
        # 태그 버튼
        self.tag_btn = QPushButton("🏷️ 태그")
        self.tag_btn.setToolTip("현재 URL에 태그 추가/편집\n예: 본회의, 법사위, 상임위")
        self.tag_btn.setFixedWidth(80)
        self.tag_btn.clicked.connect(self._edit_url_tag)
        url_layout.addWidget(self.tag_btn)
        
        settings_layout.addLayout(url_layout, 0, 1)
        
        # 선택자
        selector_label = QLabel("🔍 선택자:")
        selector_label.setToolTip("자막 요소의 CSS 선택자")
        settings_layout.addWidget(selector_label, 1, 0)
        self.selector_combo = QComboBox()
        self.selector_combo.setEditable(True)
        self.selector_combo.addItems(["#viewSubtit .incont", "#viewSubtit", ".subtitle_area"])
        self.selector_combo.setToolTip("자막 텍스트가 표시되는 HTML 요소의 CSS 선택자\n기본값을 사용하거나 직접 입력하세요")
        settings_layout.addWidget(self.selector_combo, 1, 1)
        
        # 옵션
        options_layout = QHBoxLayout()
        options_layout.setSpacing(20)
        
        self.auto_scroll_check = QCheckBox("📜 자동 스크롤")
        self.auto_scroll_check.setChecked(True)
        self.auto_scroll_check.setToolTip("새 자막이 추가될 때 자동으로 맨 아래로 스크롤합니다")
        
        self.realtime_save_check = QCheckBox("💾 실시간 저장")
        self.realtime_save_check.setChecked(False)
        self.realtime_save_check.setToolTip("자막이 확정될 때마다 자동으로 파일에 저장합니다\n저장 위치: realtime_output 폴더")
        
        self.headless_check = QCheckBox("🔇 헤드리스 모드 (인터넷창 숨김)")
        self.headless_check.setChecked(False)
        self.headless_check.setToolTip("Chrome 브라우저 창을 숨기고 백그라운드에서 실행합니다.\n자막 추출 중 다른 작업을 할 수 있습니다.")
        options_layout.addWidget(self.auto_scroll_check)
        options_layout.addWidget(self.realtime_save_check)
        options_layout.addWidget(self.headless_check)
        options_layout.addStretch()
        settings_layout.addLayout(options_layout, 2, 0, 1, 2)
        
        layout.addWidget(settings_group)
        
        # === 컨트롤 버튼 ===
        btn_layout = QHBoxLayout()
        
        self.start_btn = QPushButton("▶  시작")
        self.start_btn.setObjectName("startBtn")
        self.start_btn.clicked.connect(self._start)
        
        self.stop_btn = QPushButton("⏹  중지")
        self.stop_btn.setObjectName("stopBtn")
        self.stop_btn.setEnabled(False)
        self.stop_btn.clicked.connect(self._stop)
        
        btn_layout.addWidget(self.start_btn)
        btn_layout.addWidget(self.stop_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        # === 진행 표시 ===
        self.progress = QProgressBar()
        self.progress.setMaximum(0)
        self.progress.hide()
        layout.addWidget(self.progress)
        
        # === 메인 콘텐츠 (자막 + 통계) ===
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # 자막 텍스트
        self.subtitle_text = QTextEdit()
        self.subtitle_text.setReadOnly(True)
        splitter.addWidget(self.subtitle_text)
        
        # 통계 패널
        stats_group = QGroupBox("통계")
        stats_layout = QVBoxLayout(stats_group)
        
        self.stat_time = QLabel("⏱️ 실행 시간: 00:00:00")
        self.stat_chars = QLabel("📝 글자 수: 0")
        self.stat_words = QLabel("📖 단어 수: 0")
        self.stat_sents = QLabel("💬 문장 수: 0")
        self.stat_cpm = QLabel("⚡ 분당 글자: 0")
        
        for label in [self.stat_time, self.stat_chars, self.stat_words, self.stat_sents, self.stat_cpm]:
            label.setFont(QFont("맑은 고딕", 10))
            stats_layout.addWidget(label)
        
        stats_layout.addStretch()
        stats_group.setFixedWidth(180)
        splitter.addWidget(stats_group)
        
        splitter.setSizes([900, 180])
        layout.addWidget(splitter)
        
        # === 검색바 (숨김) ===
        self.search_frame = QFrame()
        search_layout = QHBoxLayout(self.search_frame)
        search_layout.setContentsMargins(0, 0, 0, 0)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("검색어 입력...")
        self.search_input.returnPressed.connect(self._do_search)
        
        search_prev = QPushButton("◀")
        search_prev.setFixedWidth(40)
        search_prev.clicked.connect(lambda: self._nav_search(-1))
        
        search_next = QPushButton("▶")
        search_next.setFixedWidth(40)
        search_next.clicked.connect(lambda: self._nav_search(1))
        
        search_close = QPushButton("✕")
        search_close.setFixedWidth(40)
        search_close.clicked.connect(self._hide_search)
        
        self.search_count = QLabel("")
        
        search_layout.addWidget(self.search_input)
        search_layout.addWidget(self.search_count)
        search_layout.addWidget(search_prev)
        search_layout.addWidget(search_next)
        
        # 검색바 항상 표시 (숨기기 버튼 제거)
        layout.addWidget(self.search_frame)
        
        # === 상태바 ===
        self.status_label = QLabel("대기 중")
        layout.addWidget(self.status_label)
        
        # 검색 상태
        self.search_matches = []
        self.search_idx = 0
    
    def _apply_theme(self):
        self.setStyleSheet(DARK_THEME if self.is_dark_theme else LIGHT_THEME)
        self.theme_action.setText("라이트 테마" if self.is_dark_theme else "다크 테마")
    
    def _toggle_theme(self):
        self.is_dark_theme = not self.is_dark_theme
        self.settings.setValue("dark_theme", self.is_dark_theme)
        self._apply_theme()
    
    def _setup_shortcuts(self):
        QShortcut(QKeySequence("F5"), self, self._start)
        QShortcut(QKeySequence("Escape"), self, self._stop)
        QShortcut(QKeySequence("F3"), self, lambda: self._nav_search(1))
        QShortcut(QKeySequence("Shift+F3"), self, lambda: self._nav_search(-1))
    
    # ========== URL 히스토리 (태그 지원) ==========
    
    def _load_url_history(self):
        """URL 히스토리 로드 - {url: tag} 형태"""
        try:
            if Path("url_history.json").exists():
                with open("url_history.json", 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # dict 형태인지 확인 (새로운 형식)
                    if isinstance(data, dict):
                        return data
                    # 이전 list 형태면 dict로 변환
                    elif isinstance(data, list):
                        return {url: "" for url in data}
        except Exception as e:
            logger.warning(f"URL 히스토리 로드 오류: {e}")
        return {}
    
    def _save_url_history(self):
        """URL 히스토리 저장"""
        try:
            if not isinstance(self.url_history, dict):
                self.url_history = {}
            with open("url_history.json", 'w', encoding='utf-8') as f:
                json.dump(self.url_history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"URL 히스토리 저장 오류: {e}")
    
    def _add_to_history(self, url, tag=""):
        """URL 히스토리에 추가"""
        if not isinstance(self.url_history, dict):
            self.url_history = {}
        
        # 기존 태그가 있으면 유지
        if url in self.url_history and not tag:
            tag = self.url_history[url]
        
        self.url_history[url] = tag
        self._save_url_history()
        self._refresh_url_combo()
    
    def _refresh_url_combo(self):
        """URL 콤보박스 새로고침"""
        current_text = self.url_combo.currentText()
        self.url_combo.clear()
        
        for url, tag in self.url_history.items():
            if tag:
                self.url_combo.addItem(f"[{tag}] {url}", url)
            else:
                self.url_combo.addItem(url, url)
        
        # 기본 URL 추가
        if not self.url_history:
            self.url_combo.addItem("https://assembly.webcast.go.kr/main/player.asp")
        
        # 이전 텍스트 복원
        if current_text:
            self.url_combo.setCurrentText(current_text)
    
    def _get_current_url(self):
        """현재 선택된 URL 반환 (태그 제거)"""
        # userData에 원본 URL이 저장되어 있으면 그것 사용
        data = self.url_combo.currentData()
        if data:
            return data
        
        # 없으면 텍스트에서 추출
        text = self.url_combo.currentText()
        if text.startswith("[") and "] " in text:
            return text.split("] ", 1)[1]
        return text
    
    def _edit_url_tag(self):
        """현재 URL의 태그 편집"""
        url = self._get_current_url()
        if not url or not url.startswith(('http://', 'https://')):
            QMessageBox.warning(self, "알림", "태그를 지정할 URL을 먼저 선택하세요.")
            return
        
        current_tag = self.url_history.get(url, "")
        tag, ok = QInputDialog.getText(
            self, "URL 태그 설정",
            f"URL: {url[:50]}...\n\n태그 입력 (예: 본회의, 법사위, 상임위):",
            text=current_tag
        )
        
        if ok:
            self._add_to_history(url, tag.strip())
            QMessageBox.information(self, "성공", f"태그가 설정되었습니다: [{tag}]" if tag else "태그가 제거되었습니다.")
    
    # ========== 추출 제어 ==========
    
    def _start(self):
        if self.is_running:
            return
        
        url = self._get_current_url().strip()
        selector = self.selector_combo.currentText().strip()
        
        if not url or not selector:
            QMessageBox.warning(self, "오류", "URL과 선택자를 입력하세요.")
            return
        
        if not url.startswith(('http://', 'https://')):
            QMessageBox.warning(self, "오류", "올바른 URL을 입력하세요.")
            return
        
        try:
            # 히스토리에 추가
            self._add_to_history(url)
            
            # 초기화
            self.subtitle_text.clear()
            self.subtitles = []
            self.last_subtitle = ""
            self.start_time = time.time()
            
            # 큐 비우기
            while not self.message_queue.empty():
                try:
                    self.message_queue.get_nowait()
                except queue.Empty:
                    break
            
            # 실시간 저장 설정 (예외 처리 추가)
            self.realtime_file = None
            if self.realtime_save_check.isChecked():
                try:
                    Path("realtime_output").mkdir(exist_ok=True)
                    filename = f"realtime_output/자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                    self.realtime_file = open(filename, 'w', encoding='utf-8')
                    self.status_label.setText(f"실시간 저장: {filename}")
                except Exception as e:
                    logger.error(f"실시간 저장 파일 생성 오류: {e}")
            
            # UI 업데이트
            self.is_running = True
            self.stop_event.clear()  # 종료 이벤트 초기화
            self.start_btn.setEnabled(False)
            self.stop_btn.setEnabled(True)
            self.url_combo.setEnabled(False)
            self.selector_combo.setEnabled(False)
            self.progress.show()
            
            self.status_label.setText("Chrome 브라우저 시작 중...")
            
            # UI 값을 시작 시점에 복사 (스레드 안전성)
            headless = self.headless_check.isChecked()
            
            # 워커 시작
            self.worker = threading.Thread(
                target=self._extraction_worker,
                args=(url, selector, headless),
                daemon=True
            )
            self.worker.start()
            
            self.stats_timer.start(1000)
        
        except Exception as e:
            logger.exception(f"시작 오류: {e}")
            self._reset_ui()
            QMessageBox.critical(self, "오류", f"시작 중 오류 발생: {e}")
    
    def _stop(self):
        if not self.is_running:
            return
        
        self.is_running = False
        self.stop_event.set()  # 워커 스레드에 종료 신호
        self.status_label.setText("중지 중...")
        
        # 워커 스레드 종료 대기 (최대 3초)
        if self.worker and self.worker.is_alive():
            self.worker.join(timeout=3)
            if self.worker.is_alive():
                logger.warning("워커 스레드가 시간 내에 종료되지 않음")
        
        # 마지막 자막 저장
        if self.last_subtitle:
            self._finalize_subtitle(self.last_subtitle)
            self.last_subtitle = ""
        
        # 실시간 저장 종료
        if self.realtime_file:
            self.realtime_file.close()
            self.realtime_file = None
        
        # WebDriver 종료
        if self.driver:
            try:
                self.driver.quit()
            except Exception as e:
                logger.debug(f"WebDriver 종료 중 오류 (무시됨): {e}")
            self.driver = None
        
        self._reset_ui()
        self.status_label.setText("중지됨")
    
    def _reset_ui(self):
        self.is_running = False
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.url_combo.setEnabled(True)
        self.selector_combo.setEnabled(True)
        self.progress.hide()
        self.stats_timer.stop()
    
    # ========== 워커 스레드 ==========
    
    def _extraction_worker(self, url, selector, headless):
        """자막 추출 워커 스레드
        
        Args:
            url: 대상 웹사이트 URL
            selector: 자막 요소 CSS 선택자
            headless: 헤드리스 모드 여부 (시작 시점에 복사됨)
        """
        driver = None
        
        try:
            options = Options()
            options.add_argument("--log-level=3")
            options.add_argument("--disable-blink-features=AutomationControlled")
            options.add_experimental_option("excludeSwitches", ["enable-logging", "enable-automation"])
            options.add_experimental_option('useAutomationExtension', False)
            
            # 헤드리스 모드
            if headless:
                options.add_argument("--headless=new")
                options.add_argument("--window-size=1280,720")
                self.message_queue.put(("status", "헤드리스 모드로 시작 중..."))
            
            try:
                driver = webdriver.Chrome(options=options)
                self.driver = driver
                self.message_queue.put(("status", "Chrome 시작 완료"))
            except Exception as e:
                self.message_queue.put(("error", f"Chrome 오류: {e}"))
                return
            
            self.message_queue.put(("status", "페이지 로딩 중..."))
            driver.get(url)
            time.sleep(3)
            
            self.message_queue.put(("status", "AI 자막 활성화 중..."))
            self._activate_subtitle(driver)
            
            self.message_queue.put(("status", "자막 요소 검색 중..."))
            wait = WebDriverWait(driver, 20)
            
            found = False
            for sel in [selector, "#viewSubtit .incont", "#viewSubtit", ".subtitle_area"]:
                try:
                    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, sel)))
                    self.message_queue.put(("status", f"자막 요소 찾음: {sel}"))
                    found = True
                    break
                except Exception:
                    continue
            
            if not found:
                self.message_queue.put(("error", "자막 요소를 찾을 수 없습니다."))
                return
            
            self.message_queue.put(("status", "자막 모니터링 중"))
            
            last_check = time.time()
            # stop_event 사용으로 더 빠른 종료 응답
            while not self.stop_event.is_set():
                try:
                    now = time.time()
                    if now - last_check >= 0.2:
                        try:
                            text = driver.find_element(By.CSS_SELECTOR, selector).text.strip()
                        except (NoSuchElementException, StaleElementReferenceException):
                            text = ""
                        
                        text = self._clean_text(text)
                        if text and text != self.last_subtitle:
                            self.message_queue.put(("preview", text))
                        
                        last_check = now
                    
                    # stop_event 대기 (0.05초, 즉시 응답 가능)
                    self.stop_event.wait(timeout=0.05)
                except Exception as e:
                    if not self.stop_event.is_set():
                        logger.warning(f"모니터링 중 오류: {e}")
                    time.sleep(0.5)
        
        except Exception as e:
            if not self.stop_event.is_set():
                self.message_queue.put(("error", str(e)))
        
        finally:
            if driver:
                try:
                    driver.quit()
                except Exception as e:
                    logger.debug(f"워커 WebDriver 종료 중 오류: {e}")
                self.driver = None
            self.message_queue.put(("finished", ""))
    
    def _activate_subtitle(self, driver):
        scripts = [
            "if(typeof layerSubtit==='function'){layerSubtit();}",
            "document.querySelector('.btn_subtit')?.click();",
            "document.querySelector('#btnSubtit')?.click();",
        ]
        for script in scripts:
            try:
                driver.execute_script(script)
                time.sleep(0.5)
            except Exception:
                pass  # 자막 활성화 스크립트 실패는 무시
    
    def _clean_text(self, text):
        text = re.sub(r'\b\d{4}년\b', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    # ========== 메시지 큐 처리 ==========
    
    def _process_message_queue(self):
        """메시지 큐 처리 (100ms마다 호출) - 예외 처리 강화"""
        try:
            # 최대 10개 메시지까지 처리 (무한 루프 방지)
            for _ in range(10):
                try:
                    msg_type, data = self.message_queue.get_nowait()
                    self._handle_message(msg_type, data)
                except queue.Empty:
                    break
        except Exception as e:
            logger.error(f"큐 처리 오류: {e}")
    
    def _handle_message(self, msg_type, data):
        """개별 메시지 처리"""
        try:
            if msg_type == "status":
                self.status_label.setText(str(data)[:200])
            
            elif msg_type == "preview":
                self._process_raw_text(data)
            
            elif msg_type == "error":
                self.progress.hide()
                self._reset_ui()
                QMessageBox.critical(self, "오류", str(data))
            
            elif msg_type == "finished":
                if self.last_subtitle:
                    self._finalize_subtitle(self.last_subtitle)
                    self.last_subtitle = ""
                
                self._refresh_text()
                self._reset_ui()
                
                total_chars = sum(len(s.text) for s in self.subtitles)
                self.status_label.setText(f"완료 - {len(self.subtitles)}문장, {total_chars:,}자")
        
        except Exception as e:
            logger.error(f"메시지 처리 오류 ({msg_type}): {e}")
    
    def _process_raw_text(self, raw):
        if not raw or raw == self.last_subtitle:
            return
        
        # 시간 기반 확정: 자막이 변경될 때마다 시간 기록
        # 2초 동안 변경 없으면 확정
        self.last_update_time = time.time()
        self.last_subtitle = raw
        
        # 확정 타이머 시작 (아직 안돌고 있으면)
        if not self.finalize_timer.isActive():
            self.finalize_timer.start(500)
        
        # 화면에는 현재 자막만 표시 (누적 없음, 미리보기만)
        self._update_preview(raw)
    
    def _check_finalize(self):
        """2초 동안 변경 없으면 현재 자막 확정"""
        if not self.last_subtitle:
            return
        
        elapsed = time.time() - self.last_update_time
        if elapsed >= 2.0:  # 2초 동안 변경 없음
            self._finalize_subtitle(self.last_subtitle)
            self.last_subtitle = ""
            self.finalize_timer.stop()
            self._refresh_text()
    
    def _update_preview(self, raw):
        """미리보기 업데이트 - 확정 자막 + 현재 자막 모두 표시"""
        self.subtitle_text.clear()
        cursor = self.subtitle_text.textCursor()
        
        show_ts = self.timestamp_action.isChecked()
        
        # 확정된 자막들 표시
        for i, entry in enumerate(self.subtitles):
            if i > 0:
                cursor.insertText("\n\n")
            
            # 타임스탬프
            if show_ts:
                fmt = QTextCharFormat()
                fmt.setForeground(QColor("#888888"))
                cursor.insertText(f"[{entry.timestamp.strftime('%H:%M:%S')}] ", fmt)
            
            # 텍스트 (키워드 하이라이트)
            self._insert_highlighted_text(cursor, entry.text)
        
        # 현재 진행 중인 자막 (회색, 진행 중 표시)
        if raw:
            if self.subtitles:
                cursor.insertText("\n\n")
            
            fmt = QTextCharFormat()
            fmt.setForeground(QColor("#aaaaaa"))
            cursor.insertText("⏳ ", fmt)
            cursor.insertText(raw)
        
        if self.auto_scroll_check.isChecked():
            self.subtitle_text.moveCursor(QTextCursor.MoveOperation.End)
    
    def _finalize_subtitle(self, text):
        """자막 확정 - 이전 확정 자막과 겹치는 부분 제거 후 이어붙이기"""
        if not text:
            return
        
        # 이전에 확정된 자막이 있으면, 겹치는 부분 제거
        if self.subtitles:
            last_text = self.subtitles[-1].text
            
            # 새 텍스트가 이전 텍스트로 시작하면 -> 새로운 부분만 추가
            if text.startswith(last_text):
                new_part = text[len(last_text):].strip()
                if new_part:
                    # 이전 자막에 이어붙이기
                    self.subtitles[-1].text = text
                    self.subtitles[-1].end_time = datetime.now()
                    
                    # 실시간 저장 (덮어쓰기 대신 업데이트)
                    if self.realtime_file:
                        try:
                            timestamp = self.subtitles[-1].timestamp.strftime('%H:%M:%S')
                            self.realtime_file.write(f"[{timestamp}] +{new_part}\n")
                            self.realtime_file.flush()
                        except IOError as e:
                            logger.warning(f"실시간 저장 쓰기 오류: {e}")
                return
            
            # 텍스트 겹침 확인 (앞부분이 잘려서 시작하는 경우)
            # 새 텍스트의 앞부분이 이전 텍스트의 끝부분과 겹치는지 확인
            for overlap_len in range(min(len(last_text), len(text)), 10, -1):
                if last_text.endswith(text[:overlap_len]):
                    # 겹치는 부분을 제외하고 새 부분만 이어붙이기
                    new_part = text[overlap_len:].strip()
                    if new_part:
                        self.subtitles[-1].text += " " + new_part
                        self.subtitles[-1].end_time = datetime.now()
                        
                        if self.realtime_file:
                            try:
                                self.realtime_file.write(f"+ {new_part}\n")
                                self.realtime_file.flush()
                            except IOError as e:
                                logger.warning(f"실시간 저장 쓰기 오류: {e}")
                    return
        
        # 첫 번째 자막 또는 완전히 새로운 문장
        entry = SubtitleEntry(text)
        entry.start_time = datetime.now()
        entry.end_time = datetime.now()
        
        self.subtitles.append(entry)
        
        # 실시간 저장
        if self.realtime_file:
            try:
                timestamp = entry.timestamp.strftime('%H:%M:%S')
                self.realtime_file.write(f"[{timestamp}] {text}\n")
                self.realtime_file.flush()
            except IOError as e:
                logger.warning(f"실시간 저장 쓰기 오류: {e}")
        
        self._refresh_text()
    
    def _refresh_text(self):
        self.subtitle_text.clear()
        cursor = self.subtitle_text.textCursor()
        
        show_ts = self.timestamp_action.isChecked()
        
        for i, entry in enumerate(self.subtitles):
            if i > 0:
                cursor.insertText("\n\n")
            
            # 타임스탬프
            if show_ts:
                fmt = QTextCharFormat()
                fmt.setForeground(QColor("#888888"))
                cursor.insertText(f"[{entry.timestamp.strftime('%H:%M:%S')}] ", fmt)
            
            # 텍스트 (키워드 하이라이트)
            self._insert_highlighted_text(cursor, entry.text)
        
        if self.auto_scroll_check.isChecked():
            self.subtitle_text.moveCursor(QTextCursor.MoveOperation.End)
    
    def _insert_highlighted_text(self, cursor, text):
        """텍스트에서 키워드만 하이라이트"""
        # 빈 키워드 필터링
        valid_keywords = [k.strip() for k in self.keywords if k and k.strip()]
        
        if not valid_keywords:
            cursor.insertText(text)
            return
        
        # 키워드 하이라이트 (해당 단어만)
        pattern = '|'.join(re.escape(k) for k in valid_keywords)
        parts = re.split(f'({pattern})', text, flags=re.IGNORECASE)
        
        # 키워드 소문자 리스트 (비교용)
        keywords_lower = [k.lower() for k in valid_keywords]
        
        for part in parts:
            if not part:  # 빈 문자열 건너뛰기
                continue
            
            if part.lower() in keywords_lower:
                # 키워드만 하이라이트 (눈에 띄는 스타일)
                fmt = QTextCharFormat()
                fmt.setBackground(QColor("#ffd700"))  # 골드 배경
                fmt.setForeground(QColor("#000000"))  # 검정 글자
                fmt.setFontWeight(QFont.Weight.Bold)  # 볼드
                cursor.insertText(part, fmt)
            else:
                # 일반 텍스트 (기본 포맷으로)
                fmt = QTextCharFormat()
                cursor.insertText(part, fmt)
    
    # ========== 통계 ==========
    
    def _update_stats(self):
        if self.start_time:
            elapsed = int(time.time() - self.start_time)
            h, r = divmod(elapsed, 3600)
            m, s = divmod(r, 60)
            self.stat_time.setText(f"⏱️ 실행 시간: {h:02d}:{m:02d}:{s:02d}")
            
            total_chars = sum(len(s.text) for s in self.subtitles)
            total_words = sum(len(s.text.split()) for s in self.subtitles)
            
            self.stat_chars.setText(f"📝 글자 수: {total_chars:,}")
            self.stat_words.setText(f"📖 단어 수: {total_words:,}")
            self.stat_sents.setText(f"💬 문장 수: {len(self.subtitles)}")
            
            if elapsed > 0:
                cpm = int(total_chars / (elapsed / 60))
                self.stat_cpm.setText(f"⚡ 분당 글자: {cpm}")
    
    # ========== 검색 ==========
    
    def _show_search(self):
        self.search_frame.show()
        self.search_input.setFocus()
        self.search_input.selectAll()
    
    def _hide_search(self):
        self.search_frame.hide()
        self._refresh_text()
    
    def _do_search(self):
        query = self.search_input.text()
        if not query:
            return
        
        text = self.subtitle_text.toPlainText()
        self.search_matches = []
        
        start = 0
        while True:
            idx = text.lower().find(query.lower(), start)
            if idx == -1:
                break
            self.search_matches.append(idx)
            start = idx + 1
        
        self.search_idx = 0
        self.search_count.setText(f"{len(self.search_matches)}개")
        
        if self.search_matches:
            self._highlight_search(0)
    
    def _nav_search(self, delta):
        if not self.search_matches:
            return
        
        self.search_idx = (self.search_idx + delta) % len(self.search_matches)
        self._highlight_search(self.search_idx)
    
    def _highlight_search(self, idx):
        if not self.search_matches:
            return
        
        pos = self.search_matches[idx]
        query = self.search_input.text()
        
        cursor = self.subtitle_text.textCursor()
        cursor.setPosition(pos)
        cursor.setPosition(pos + len(query), QTextCursor.MoveMode.KeepAnchor)
        
        self.subtitle_text.setTextCursor(cursor)
        self.subtitle_text.ensureCursorVisible()
        
        self.search_count.setText(f"{idx + 1}/{len(self.search_matches)}")
    
    # ========== 키워드 ==========
    
    def _set_keywords(self):
        current = ", ".join(self.keywords)
        text, ok = QInputDialog.getText(
            self, "키워드 설정",
            "하이라이트할 키워드 (쉼표로 구분):",
            text=current
        )
        
        if ok:
            self.keywords = [k.strip() for k in text.split(",") if k.strip()]
            self._refresh_text()
    
    # ========== 파일 저장 ==========
    
    def _get_accumulated_text(self):
        return "\n".join(s.text for s in self.subtitles)
    
    def _save_txt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        path, _ = QFileDialog.getSaveFileName(self, "TXT 저장", filename, "텍스트 (*.txt)")
        
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                for entry in self.subtitles:
                    f.write(f"[{entry.timestamp.strftime('%H:%M:%S')}] {entry.text}\n")
            QMessageBox.information(self, "성공", "저장 완료!")
    
    def _save_srt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.srt"
        path, _ = QFileDialog.getSaveFileName(self, "SRT 저장", filename, "SubRip (*.srt)")
        
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                for i, entry in enumerate(self.subtitles, 1):
                    start = entry.start_time.strftime('%H:%M:%S,000')
                    end = entry.end_time.strftime('%H:%M:%S,000')
                    f.write(f"{i}\n{start} --> {end}\n{entry.text}\n\n")
            QMessageBox.information(self, "성공", "SRT 저장 완료!")
    
    def _save_vtt(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.vtt"
        path, _ = QFileDialog.getSaveFileName(self, "VTT 저장", filename, "WebVTT (*.vtt)")
        
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                f.write("WEBVTT\n\n")
                for i, entry in enumerate(self.subtitles, 1):
                    start = entry.start_time.strftime('%H:%M:%S.000')
                    end = entry.end_time.strftime('%H:%M:%S.000')
                    f.write(f"{i}\n{start} --> {end}\n{entry.text}\n\n")
            QMessageBox.information(self, "성공", "VTT 저장 완료!")
    
    def _save_docx(self):
        """DOCX (Word) 파일로 저장"""
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.warning(
                self, "라이브러리 필요",
                "DOCX 저장을 위해 python-docx 라이브러리가 필요합니다.\n\n"
                "설치: pip install python-docx"
            )
            return
        
        filename = f"국회자막_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path, _ = QFileDialog.getSaveFileName(self, "DOCX 저장", filename, "Word 문서 (*.docx)")
        
        if path:
            try:
                doc = Document()
                
                # 제목
                title = doc.add_heading("국회 의사중계 자막", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 생성 일시
                doc.add_paragraph(f"생성 일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}")
                doc.add_paragraph()
                
                # 자막 내용
                for entry in self.subtitles:
                    timestamp = entry.timestamp.strftime('%H:%M:%S')
                    p = doc.add_paragraph()
                    run = p.add_run(f"[{timestamp}] ")
                    run.font.size = Pt(9)
                    run.font.color.rgb = None  # 기본 색상
                    p.add_run(entry.text)
                
                # 통계
                doc.add_paragraph()
                total_chars = sum(len(s.text) for s in self.subtitles)
                doc.add_paragraph(f"총 {len(self.subtitles)}문장, {total_chars:,}자")
                
                doc.save(path)
                QMessageBox.information(self, "성공", f"DOCX 저장 완료!\n\n파일: {path}")
            
            except Exception as e:
                QMessageBox.critical(self, "오류", f"DOCX 저장 실패: {e}")
    
    # ========== 세션 ==========
    
    def _save_session(self):
        if not self.subtitles:
            QMessageBox.warning(self, "알림", "저장할 내용이 없습니다.")
            return
        
        filename = f"sessions/세션_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        path, _ = QFileDialog.getSaveFileName(self, "세션 저장", filename, "JSON (*.json)")
        
        if path:
            data = {
                'version': self.VERSION,
                'created': datetime.now().isoformat(),
                'url': self.url_combo.currentText(),
                'subtitles': [
                    {
                        'text': s.text,
                        'timestamp': s.timestamp.isoformat(),
                        'start_time': s.start_time.isoformat() if s.start_time else None,
                        'end_time': s.end_time.isoformat() if s.end_time else None,
                    }
                    for s in self.subtitles
                ]
            }
            
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            QMessageBox.information(self, "성공", "세션 저장 완료!")
    
    def _load_session(self):
        path, _ = QFileDialog.getOpenFileName(self, "세션 불러오기", "sessions/", "JSON (*.json)")
        
        if path:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                self.subtitles = []
                for item in data.get('subtitles', []):
                    entry = SubtitleEntry(item['text'])
                    entry.timestamp = datetime.fromisoformat(item['timestamp'])
                    if item.get('start_time'):
                        entry.start_time = datetime.fromisoformat(item['start_time'])
                    if item.get('end_time'):
                        entry.end_time = datetime.fromisoformat(item['end_time'])
                    self.subtitles.append(entry)
                
                if data.get('url'):
                    self.url_combo.setCurrentText(data['url'])
                
                self._refresh_text()
                QMessageBox.information(self, "성공", f"세션 불러오기 완료!\n{len(self.subtitles)}개 문장")
            
            except Exception as e:
                QMessageBox.critical(self, "오류", f"불러오기 실패: {e}")
    
    # ========== 유틸리티 ==========
    
    def _copy_to_clipboard(self):
        text = self._get_accumulated_text()
        if not text:
            QMessageBox.warning(self, "알림", "복사할 내용이 없습니다.")
            return
        
        QApplication.clipboard().setText(text)
        QMessageBox.information(self, "성공", f"클립보드에 복사되었습니다.\n{len(text):,}자")
    
    def _clear_text(self):
        if not self.subtitles:
            return
        
        reply = QMessageBox.question(
            self, "확인", "모든 내용을 지우시겠습니까?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.subtitles = []
            self.last_subtitle = ""
            self.subtitle_text.clear()
            self.status_label.setText("내용 삭제됨")
    
    # ========== 도움말 ==========
    
    def _show_guide(self):
        """사용법 가이드 표시"""
        guide = """
<h2>🏛️ 사용법 가이드</h2>

<h3>📋 기본 사용법</h3>
<ol>
<li><b>URL 입력</b> - 국회 의사중계 페이지 URL을 입력합니다</li>
<li><b>선택자 확인</b> - 기본값을 사용하거나 수정합니다</li>
<li><b>옵션 설정</b>
    <ul>
    <li>자동 스크롤: 새 자막 자동 따라가기</li>
    <li>실시간 저장: 자막 실시간 파일 저장</li>
    <li>헤드리스 모드: 브라우저 창 숨기고 실행</li>
    </ul>
</li>
<li><b>시작</b> 버튼 클릭 (또는 F5)</li>
<li>자막 추출 완료 후 <b>파일 저장</b></li>
</ol>

<h3>⌨️ 주요 단축키</h3>
<table>
<tr><td><b>F5</b></td><td>시작</td></tr>
<tr><td><b>Escape</b></td><td>중지</td></tr>
<tr><td><b>Ctrl+F</b></td><td>검색</td></tr>
<tr><td><b>F3</b></td><td>다음 검색</td></tr>
<tr><td><b>Ctrl+T</b></td><td>테마 전환</td></tr>
<tr><td><b>Ctrl+S</b></td><td>TXT 저장</td></tr>
</table>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("사용법 가이드")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(guide)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()
    
    def _show_features(self):
        """기능 소개 표시"""
        features = """
<h2>✨ 기능 소개</h2>

<h3>🎯 실시간 자막 추출</h3>
<p>국회 의사중계 웹사이트의 AI 자막을 실시간으로 캡처합니다.<br>
2초 동안 자막이 변경되지 않으면 자동으로 확정됩니다.</p>

<h3>💾 다양한 저장 형식</h3>
<ul>
<li><b>TXT</b> - 일반 텍스트</li>
<li><b>SRT</b> - 자막 파일 형식</li>
<li><b>VTT</b> - WebVTT 자막 형식</li>
<li><b>DOCX</b> - Word 문서</li>
</ul>

<h3>🔍 검색 및 하이라이트</h3>
<ul>
<li><b>실시간 검색</b> - Ctrl+F로 자막 내 텍스트 검색</li>
<li><b>키워드 하이라이트</b> - 특정 단어 강조</li>
</ul>

<h3>⚙️ 헤드리스 모드 (인터넷창 숨김)</h3>
<p>브라우저 창을 숨기고 백그라운드에서 실행합니다.<br>
자막 추출 중 다른 작업을 할 수 있습니다.</p>

<h3>📊 통계 패널</h3>
<p>실행 시간, 글자 수, 단어 수, 분당 글자 수를 표시합니다.</p>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("기능 소개")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(features)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()
    
    def _show_about(self):
        """프로그램 정보 표시"""
        about = f"""
<h2>🏛️ 국회 의사중계 자막 추출기</h2>
<p><b>버전:</b> {self.VERSION}</p>
<p><b>설명:</b> 국회 의사중계 웹사이트에서 실시간 AI 자막을<br>
자동으로 추출하고 저장하는 프로그램입니다.</p>

<h3>📦 필요 라이브러리</h3>
<ul>
<li>PyQt6</li>
<li>selenium</li>
<li>python-docx (DOCX 저장용)</li>
</ul>

<p><b>© 2024</b></p>
"""
        msg = QMessageBox(self)
        msg.setWindowTitle("정보")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(about)
        msg.setIcon(QMessageBox.Icon.Information)
        msg.exec()
    
    def closeEvent(self, event):
        if self.is_running:
            reply = QMessageBox.question(
                self, "종료", "추출 중입니다. 종료하시겠습니까?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                event.ignore()
                return
        
        # 타이머 정리
        self.queue_timer.stop()
        self.stats_timer.stop()
        self.finalize_timer.stop()
        
        self.is_running = False
        if self.realtime_file:
            try:
                self.realtime_file.close()
            except Exception as e:
                logger.debug(f"파일 닫기 오류: {e}")
        if self.driver:
            try:
                self.driver.quit()
            except Exception as e:
                logger.debug(f"WebDriver 종료 오류: {e}")
        
        logger.info("프로그램 종료")
        event.accept()


def main():
    """메인 함수 - 예외 처리 강화"""
    try:
        # IDLE 호환성을 위한 이벤트 루프 체크
        app = QApplication.instance()
        if app is None:
            app = QApplication(sys.argv)
        
        app.setStyle("Fusion")
        app.setFont(QFont("맑은 고딕", 10))
        
        window = MainWindow()
        window.show()
        
        # IDLE에서 실행 시 exec() 대신 processEvents 사용
        if hasattr(sys, 'ps1'):  # IDLE/인터프리터 환경
            while window.isVisible():
                app.processEvents()
                time.sleep(0.01)
        else:
            sys.exit(app.exec())
    
    except Exception as e:
        logger.exception(f"프로그램 오류: {e}")
        QMessageBox.critical(None, "오류", f"프로그램 실행 중 오류 발생:\n{e}")


if __name__ == '__main__':
    main()
